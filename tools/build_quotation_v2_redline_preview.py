#!/usr/bin/env python3
"""Build unapproved V2 redline previews by editing the frozen V1 DOCX.

These files are review artifacts only. They are never templates and must not be
used by the production renderer.
"""

from __future__ import annotations

import hashlib
from copy import deepcopy
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[1]
V1_TEMPLATE = REPO_ROOT / "api/geo_platform/quotations/assets/quotation-template-v1.docx"
OUTPUT_DIR = REPO_ROOT / "docs/quotation-redline"
V1_SHA256 = "90ae5beb10ab3bacea3b706a2068945f828e275784e99da6b72dc44f8b0d9913"
RED = RGBColor(0xC0, 0x00, 0x00)
_MUTATED_PACKAGE_PARTS = {
    "docProps/core.xml",
    "word/document.xml",
    "word/header1.xml",
}

SERVICE_COPY = {
    "1": (
        "品牌GEO推荐结果测试",
        "围绕客户确认并冻结的业务 Query，按双方确认的平台、地域、账号与重复次数开展基线或复测，"
        "交付逐题证据，以及品牌提及、推荐排名、Top1/Top3/Top5 和竞品对比结果。"
        "仅报告本次可追溯实测数据，不预填或承诺效果。",
    ),
    "2": (
        "己方内容竞品拉踩风险核查",
        "核查客户确认范围内的己方 GEO 内容，识别其中可能涉及竞品比较、贬损或拉踩的表述，"
        "逐条交付来源链接、原文证据与人工复核状态；不推断未核验事实。",
    ),
    "3": (
        "第三方内容品牌被拉踩风险核查",
        "核查客户确认范围内的竞品或第三方 GEO 内容，识别其中可能涉及客户品牌或产品的贬损、"
        "拉踩与误导表述，逐条交付来源链接、原文证据与人工复核状态。",
    ),
    "4": (
        "官网内容AI引用分析",
        "分析本次已归档 AI 回答是否检索、引用或采纳客户官网内容，定位官网信息结构与内容表达问题，"
        "并输出可执行建议。仅使用可追溯回答证据，不预填官网引用结果。",
    ),
    "5": (
        "GEO内容发布与排名验证",
        "基于客户审核通过的材料形成并发布试点内容，再按与基线一致的口径复测目标 Query，"
        "交付发布记录与排名证据。结果以实际平台返回为准，不承诺提及、引用或排名提升。",
    ),
}

SCENARIOS = {
    "all-five-services": {
        "services": ("1", "2", "3", "4", "5"),
        "quantities": {
            "1": "1轮",
            "2": "1项",
            "3": "1项",
            "4": "1项",
            "5": "1项",
        },
        "hide_appendix_3": False,
    },
    "effect-assessment": {
        "services": ("1", "2", "3", "4"),
        "quantities": {"1": "1轮", "2": "1项", "3": "1项", "4": "1项"},
        "hide_appendix_3": True,
    },
    "minimum-validation": {
        "services": ("1", "3", "4", "5"),
        "quantities": {"1": "2轮", "3": "1项", "4": "1项（条件）", "5": "1项"},
        "hide_appendix_3": False,
    },
}


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _replace_paragraph_text(paragraph, text: str, *, redline: bool = True) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    if redline:
        run.font.color.rgb = RED


def _replace_cell_text(cell, text: str, *, redline: bool = True) -> None:
    first = cell.paragraphs[0]
    for paragraph in list(cell.paragraphs[1:]):
        cell._tc.remove(paragraph._p)
    _replace_paragraph_text(first, "", redline=redline)
    run = first.runs[0]
    lines = text.split("\n")
    run.text = lines[0]
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)
    if redline:
        run.font.color.rgb = RED


def _mark_unapproved(document: DocumentObject) -> None:
    label = "【未批准 V2 红线预览｜禁止作为正式客户报价】"
    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        original = paragraph.text
        _replace_paragraph_text(paragraph, f"{label}  {original}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(6.5)
            run.font.bold = True
            run.font.color.rgb = RED
    document.core_properties.subject = label
    document.core_properties.comments = (
        "由 quotation-template-v1.docx 派生的 Phase A 红线预览；approval_status=pending。"
    )


def _replace_evidence_placeholders(document: DocumentObject) -> None:
    exact_replacements = {
        "优化前效果：AI回答操作流程，未触发厂商推荐": (
            "优化前效果：待实测（仅在有可追溯证据时填入）"
        ),
        "优化后效果：AI主动推荐厂商，厂商推荐15次": "优化后效果：待实测（不得预填推荐次数或排名）",
        "优化前效果：AI回答概念解释，厂商推荐21次": "优化前效果：待实测（仅在有可追溯证据时填入）",
        "优化后效果：AI回答产品对比，厂商推荐34次": "优化后效果：待实测（不得预填推荐次数或排名）",
        "优化前效果：AI回答技术原理，未触发厂商推荐": (
            "优化前效果：待实测（仅在有可追溯证据时填入）"
        ),
        "优化后效果：AI主动推荐厂商，厂商推荐14次": "优化后效果：待实测（不得预填推荐次数或排名）",
        (
            "以下为2组代表性优化提示词在DeepSeek、文心一言两个平台上的厂商推荐次数对比："
        ): "以下表格仅在具备本次可追溯测试证据后填入；平台、次数和结果不得预设：",
        (
            "结论：上述示例中，优化后的提示词在两个平台上的厂商推荐总量"
            "从4次提升至37次，提升显著。优化后AI开始主动推荐厂商，但盛邦"
            "安全均未出现在推荐名单中——这正是GEO内容优化需要解决的核心问题。"
        ): (
            "结论：待实测。仅根据本次已归档证据生成，不得预填推荐次数、"
            "提升比例、排名或品牌未进入推荐名单等结论。"
        ),
        (
            "以下为围绕品牌五大核心产品线设计的18条核心业务问题及其语义"
            "变体。变体A为正式换述，变体B为换角度表达，变体C为口语化表达，"
            "覆盖用户实际提问的多种表述方式。选取3条问题及其语义变体为"
            "项目1·品牌AI认知评测的查询集。"
        ): (
            "以下为客户确认的原始业务 Query 及其语义变体。变体A为正式换述，"
            "变体B为换角度表达，变体C为口语化表达；仅列入本次所选服务需要的 Query。"
        ),
        (
            "以下为全部优化提示词及其语义变体。变体A为正式换述，变体B为"
            "换角度表达，变体C为口语化表达——覆盖用户实际提问的多种表述"
            "方式。选取3条提示词及其语义变体为项目4·GEO试点与效果验证的查询集。"
        ): (
            "以下为本次服务5所使用的新增优化 Query 及其语义变体。变体A为正式换述，"
            "变体B为换角度表达，变体C为口语化表达；内容只来自客户输入或通过校验的生成结果。"
        ),
    }
    for paragraph in document.paragraphs:
        replacement = exact_replacements.get(paragraph.text)
        if replacement is not None:
            _replace_paragraph_text(paragraph, replacement)

    evidence_table = document.tables[1]
    for row_index, row in enumerate(evidence_table.rows[1:], start=1):
        for column_index, cell in enumerate(row.cells):
            if column_index == 0 and row_index < len(evidence_table.rows) - 1:
                continue
            if row_index == len(evidence_table.rows) - 1 and column_index == 0:
                continue
            _replace_cell_text(cell, "待实测")


def _remove_appendix_3(document: DocumentObject) -> None:
    body = document.element.body
    children = list(body)
    heading_index = next(
        index
        for index, element in enumerate(children)
        if "附录三 新增Query优化与语义变体全表" in "".join(element.xpath(".//w:t/text()"))
    )
    start = heading_index
    if heading_index > 0:
        previous = children[heading_index - 1]
        previous_text = "".join(previous.xpath(".//w:t/text()"))
        if not previous_text and previous.xpath(".//w:br[@w:type='page']"):
            start = heading_index - 1
    for element in children[start:]:
        if element.tag == qn("w:sectPr"):
            continue
        body.remove(element)


def _restore_frozen_package_parts(output: Path) -> None:
    """Keep every V1 package part byte-exact except the three redline parts."""
    with ZipFile(V1_TEMPLATE) as source_archive, ZipFile(output) as redline_archive:
        redline_payload = {name: redline_archive.read(name) for name in redline_archive.namelist()}
        with NamedTemporaryFile(
            dir=output.parent, prefix=f".{output.name}.", suffix=".tmp", delete=False
        ) as handle:
            temporary = Path(handle.name)
        try:
            with ZipFile(temporary, "w", compression=ZIP_DEFLATED) as rebuilt:
                source_names = set(source_archive.namelist())
                for info in source_archive.infolist():
                    payload = (
                        redline_payload[info.filename]
                        if info.filename in _MUTATED_PACKAGE_PARTS
                        else source_archive.read(info.filename)
                    )
                    rebuilt.writestr(info, payload)
                for name, payload in redline_payload.items():
                    if name not in source_names:
                        rebuilt.writestr(name, payload)
            temporary.replace(output)
        finally:
            temporary.unlink(missing_ok=True)


def build_preview(name: str) -> Path:
    scenario = SCENARIOS[name]
    document = Document(V1_TEMPLATE)
    _mark_unapproved(document)
    table = document.tables[0]
    if len(scenario["services"]) == 5:
        fifth_row = deepcopy(table.rows[-2]._tr)
        table.rows[-1]._tr.addprevious(fifth_row)
    for row_number, service_id in enumerate(scenario["services"], start=1):
        service_name, service_body = SERVICE_COPY[service_id]
        row = table.rows[row_number]
        quantity = scenario["quantities"][service_id]
        if service_id == "4" and name == "minimum-validation":
            pricing_line = "计价：官网命中后成为条件项；单价 pending × 数量 1项；小计 pending。"
            price_cell = "pending\n（条件项）"
        else:
            pricing_line = f"计价：单价 pending × 数量 {quantity}；小计 pending。"
            price_cell = "pending"
        _replace_cell_text(row.cells[0], service_id)
        _replace_cell_text(row.cells[1], service_name)
        _replace_cell_text(row.cells[2], f"{service_body}\n{pricing_line}")
        _replace_cell_text(row.cells[3], price_cell)
    if name == "minimum-validation":
        _replace_cell_text(
            table.rows[-1].cells[2],
            "基础合计：pending；含官网条件项最高合计：pending。",
        )
    _replace_cell_text(table.rows[-1].cells[3], "pending")
    _replace_evidence_placeholders(document)
    if scenario["hide_appendix_3"]:
        _remove_appendix_3(document)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUTPUT_DIR / f"quotation-template-v2-proposed-{name}-redline.docx"
    document.save(output)
    _restore_frozen_package_parts(output)
    return output


def main() -> None:
    if _sha256(V1_TEMPLATE) != V1_SHA256:
        raise SystemExit("quotation_template_hash_mismatch")
    for scenario in SCENARIOS:
        output = build_preview(scenario)
        print(f"{output.relative_to(REPO_ROOT)}  {_sha256(output)}")


if __name__ == "__main__":
    main()
