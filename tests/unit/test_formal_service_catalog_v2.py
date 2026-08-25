from __future__ import annotations

import json
from contextlib import contextmanager
from datetime import UTC, date, datetime, timedelta
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from typing import Any

import pytest
from docx import Document
from geo_platform.reports import formal_production as production
from geo_platform.reports import formal_review_service2_outbound, formal_review_service5
from geo_platform.reports.formal_production import (
    LEGACY_SERVICE_CATALOG,
    QUOTATION_SERVICE_CATALOG,
    FormalProductionIncomplete,
    FormalProductionInvalid,
    FormalProductionRequest,
    FormalReportProductionService,
    FormalWindow,
    request_contract,
)
from geo_platform.reports.formal_production_router import FormalProductionCreate
from pydantic import ValidationError

from domain.reporting.formal_review_service2_outbound_docx import (
    render_outbound_disparagement_docx,
)
from domain.reporting.formal_review_service5_docx import render_publishing_pilot_docx


def _request(
    services: tuple[int, ...],
    *,
    catalog: str = QUOTATION_SERVICE_CATALOG,
    sop_project_pub_id: str | None = None,
    before_window: FormalWindow | None = None,
    after_window: FormalWindow | None = None,
) -> FormalProductionRequest:
    return FormalProductionRequest(
        pub_id="frp_unit",
        tenant_pub_id="ten_unit",
        project_pub_id="prj_unit",
        services=services,
        window=FormalWindow(date(2026, 8, 1), date(2026, 8, 20)),
        document_status="internal_review",
        candidate_group_strategy="preregistered_scope_v1",
        frozen_at=datetime(2026, 8, 20, tzinfo=UTC),
        created_by_pub_id="usr_unit",
        request_hash="a" * 64,
        before_window=before_window,
        after_window=after_window,
        service_catalog_version=catalog,
        sop_project_pub_id=sop_project_pub_id,
    )


def _document_text(payload: bytes) -> tuple[str, str]:
    document = Document(BytesIO(payload))
    text = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text.extend(paragraph.text for paragraph in cell.paragraphs)
    return document.core_properties.title or "", "\n".join(text)


def test_catalog_registries_have_stable_non_overlapping_semantics() -> None:
    assert {
        number: adapter.service_code
        for number, adapter in production.FORMAL_REPORT_REGISTRY.items()
    } == production.LEGACY_SERVICE_CODES
    assert {
        number: adapter.service_code
        for number, adapter in production.QUOTATION_FORMAL_REPORT_REGISTRY.items()
    } == production.QUOTATION_SERVICE_CODES
    assert set(production.FORMAL_REPORT_REGISTRY) == {1, 2, 3, 4}
    assert set(production.QUOTATION_FORMAL_REPORT_REGISTRY) == {1, 2, 3, 4, 5}
    assert type(production.FORMAL_REPORT_REGISTRY[2]) is not type(
        production.QUOTATION_FORMAL_REPORT_REGISTRY[2]
    )
    assert type(production.FORMAL_REPORT_REGISTRY[4]) is not type(
        production.QUOTATION_FORMAL_REPORT_REGISTRY[5]
    )


def test_create_model_rejects_service5_under_default_legacy_catalog() -> None:
    fields = {
        "project_pub_id": "prj_unit",
        "services": [5],
        "window_start": date(2026, 8, 1),
        "window_end": date(2026, 8, 20),
        "prepared_by": "项目组",
        "prepared_date": date(2026, 8, 20),
    }
    historical = FormalProductionCreate.model_validate(fields | {"services": [1, 4]})
    assert historical.root.service_catalog_version == LEGACY_SERVICE_CATALOG
    assert historical.root.services == [1, 4]

    with pytest.raises(ValidationError):
        FormalProductionCreate.model_validate(fields)
    with pytest.raises(ValidationError):
        FormalProductionCreate.model_validate(
            fields | {"service_catalog_version": LEGACY_SERVICE_CATALOG}
        )

    current = FormalProductionCreate.model_validate(
        fields
        | {
            "service_catalog_version": QUOTATION_SERVICE_CATALOG,
            "sop_project_pub_id": "spr_unit",
            "before_window": {"start": "2026-08-01", "end": "2026-08-08"},
            "after_window": {"start": "2026-08-12", "end": "2026-08-20"},
        }
    )
    assert current.root.services == [5]


def test_generated_create_contract_binds_service_range_to_catalog() -> None:
    root = Path(__file__).parents[2]
    openapi = json.loads((root / "contracts" / "openapi.json").read_text(encoding="utf-8"))
    schemas = openapi["components"]["schemas"]
    create_schema = schemas["FormalProductionCreate"]

    assert len(create_schema["anyOf"]) == 2
    branch_names = {branch["$ref"].rsplit("/", 1)[-1] for branch in create_schema["anyOf"]}
    assert branch_names == {
        "LegacyFormalProductionCreate",
        "QuotationFormalProductionCreate",
    }
    legacy = schemas["LegacyFormalProductionCreate"]
    quotation = schemas["QuotationFormalProductionCreate"]
    assert legacy["properties"]["services"]["items"]["enum"] == [1, 2, 3, 4]
    assert "service_catalog_version" not in legacy.get("required", [])
    assert legacy["properties"]["service_catalog_version"]["const"] == LEGACY_SERVICE_CATALOG
    assert quotation["properties"]["services"]["items"]["enum"] == [1, 2, 3, 4, 5]
    assert "service_catalog_version" in quotation["required"]
    assert quotation["properties"]["service_catalog_version"]["const"] == QUOTATION_SERVICE_CATALOG

    generated = (root / "packages" / "api-client" / "src" / "schema.generated.ts").read_text(
        encoding="utf-8"
    )
    assert (
        'FormalProductionCreate: components["schemas"]["LegacyFormalProductionCreate"]' in generated
    )
    assert '| components["schemas"]["QuotationFormalProductionCreate"];' in generated


def test_service1_v2_renderer_and_publication_qa_share_the_new_title(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    captured: dict[str, Any] = {}

    def renderer(facts: dict[str, Any], **kwargs: Any) -> bytes:
        del facts
        captured.update(kwargs)
        return b"docx"

    monkeypatch.setattr(
        production,
        "import_module",
        lambda name: (
            SimpleNamespace(render_service1_delivery_docx=renderer)
            if name == "domain.reporting.formal_service1_delivery_docx"
            else pytest.fail(name)
        ),
    )
    adapter = production.QUOTATION_FORMAL_REPORT_REGISTRY[1]
    assert (
        adapter.render(
            {"_formal_evidence_assets": {}},
            blob_loader=lambda *_args: b"",
        )
        == b"docx"
    )
    request = _request((1,))

    assert captured["service_number"] == 1
    assert captured["report_title"] == "AI 推荐排名效果测试报告"
    assert captured["report_subtitle"] == "服务 1 · AI 推荐排名效果与 API/手机端差异"
    assert production._service1_report_title(request, {}) == "AI 推荐排名效果测试报告"
    assert (
        production._service1_report_title(
            _request((1,), catalog=LEGACY_SERVICE_CATALOG),
            {"service1": {"delivery_v3": {"scope": {"scope_label": "网空线三类资产治理场景"}}}},
        )
        == "网空线三类资产治理场景品牌 GEO 推荐结果评测报告"
    )


def test_legacy_contract_keeps_the_pre_catalog_hash() -> None:
    contract = request_contract(
        project_pub_id="prj_unit",
        services=[1, 2, 3],
        window=FormalWindow(date(2026, 7, 1), date(2026, 7, 31)),
        document_status="pre_formal",
        candidate_group_strategy="evidence_completeness_v1",
        before_window=None,
        after_window=None,
    )

    assert "service_catalog_version" not in contract
    assert "sop_project_pub_id" not in contract
    assert production._canonical_hash(contract) == (
        "c640a828b4db84eb61a0ffd134a3a0763306deddd4fb1d580899a177efcf07cf"
    )
    with pytest.raises(FormalProductionInvalid, match="invalid_services"):
        request_contract(
            project_pub_id="prj_unit",
            services=[5],
            window=FormalWindow(date(2026, 7, 1), date(2026, 7, 31)),
            document_status="pre_formal",
            candidate_group_strategy="evidence_completeness_v1",
            before_window=None,
            after_window=None,
        )


def test_quotation_contract_applies_sop_and_service5_window_invariants() -> None:
    window = FormalWindow(date(2026, 8, 1), date(2026, 8, 20))
    before = FormalWindow(date(2026, 8, 1), date(2026, 8, 8))
    after = FormalWindow(date(2026, 8, 12), date(2026, 8, 20))
    common = {
        "project_pub_id": "prj_unit",
        "window": window,
        "document_status": "pre_formal",
        "candidate_group_strategy": "evidence_completeness_v1",
        "service_catalog_version": QUOTATION_SERVICE_CATALOG,
    }
    service2 = request_contract(
        **common,
        services=[2],
        before_window=None,
        after_window=None,
    )
    with pytest.raises(FormalProductionInvalid, match="service5_windows_required"):
        request_contract(
            **common,
            services=[5],
            before_window=None,
            after_window=None,
            sop_project_pub_id="spr_unit",
        )
    with pytest.raises(FormalProductionInvalid, match="sop_project_required"):
        request_contract(
            **common,
            services=[5],
            before_window=before,
            after_window=after,
        )
    with pytest.raises(FormalProductionInvalid, match="sop_project_not_applicable"):
        request_contract(
            **common,
            services=[1, 3, 4],
            before_window=None,
            after_window=None,
            sop_project_pub_id="spr_unit",
        )

    service4 = request_contract(
        **common,
        services=[4],
        before_window=None,
        after_window=None,
    )
    complete = request_contract(
        **common,
        services=[1, 2, 3, 4, 5],
        before_window=before,
        after_window=after,
        sop_project_pub_id="spr_unit",
    )
    assert service4["services"] == [4]
    assert service2["services"] == [2]
    assert "sop_project_pub_id" not in service2
    assert complete["service_catalog_version"] == QUOTATION_SERVICE_CATALOG
    assert complete["sop_project_pub_id"] == "spr_unit"


def test_service2_cutover_rehydrates_old_own_content_and_new_all_u_requests() -> None:
    window = FormalWindow(date(2026, 8, 1), date(2026, 8, 20))
    common = {
        "project_pub_id": "prj_unit",
        "services": [2],
        "window": window,
        "document_status": "pre_formal",
        "candidate_group_strategy": "evidence_completeness_v1",
        "before_window": None,
        "after_window": None,
        "service_catalog_version": QUOTATION_SERVICE_CATALOG,
    }
    old_contract = request_contract(
        **common,
        sop_project_pub_id="spr_legacy",
        legacy_service2_sop_compat=True,
    )
    new_contract = request_contract(**common)
    base_row = {
        "pub_id": "frp_unit",
        "tenant_pub_id": "tnt_unit",
        "project_pub_id": "prj_unit",
        "services": [2],
        "window_start": window.start,
        "window_end": window.end,
        "document_status": "pre_formal",
        "candidate_group_strategy": "evidence_completeness_v1",
        "before_start": None,
        "before_end": None,
        "after_start": None,
        "after_end": None,
        "document_governance": {},
        "service_catalog_version": QUOTATION_SERVICE_CATALOG,
        "frozen_at": datetime(2026, 8, 24, tzinfo=UTC),
        "created_by_pub_id": "usr_unit",
    }
    service = FormalReportProductionService(
        dsn="postgresql://unused",
        evidence=SimpleNamespace(),  # type: ignore[arg-type]
    )

    legacy = service._request_from_row(
        {
            **base_row,
            "sop_project_pub_id": "spr_legacy",
            "request_hash": production._canonical_hash(old_contract),
            "created_at": production.SERVICE2_ALL_U_EFFECTIVE_AT - timedelta(seconds=1),
        }
    )
    current = service._request_from_row(
        {
            **base_row,
            "sop_project_pub_id": None,
            "request_hash": production._canonical_hash(new_contract),
            "created_at": production.SERVICE2_ALL_U_EFFECTIVE_AT + timedelta(seconds=1),
        }
    )

    assert legacy.service2_semantics == "own_content_v1"
    assert type(production._adapter_for(legacy, 2)) is production._LegacyOutboundService2Adapter
    assert current.service2_semantics == "all_u_v2"
    assert production._adapter_for(current, 2) is production.QUOTATION_FORMAL_REPORT_REGISTRY[2]


def test_sop_binding_accepts_any_project_brand_or_alias(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        production,
        "_platform_project_brand_names",
        lambda *_args: frozenset({"第一品牌", "第二品牌", "second-brand"}),
    )
    request = _request((2,), sop_project_pub_id="spr_unit")

    production._assert_sop_brand_binding(
        request,
        ["第二品牌", "Second-Brand"],
        dsn="postgresql://unused",
    )
    with pytest.raises(FormalProductionInvalid, match="sop_project_brand_mismatch"):
        production._assert_sop_brand_binding(
            request,
            ["无关品牌"],
            dsn="postgresql://unused",
        )


def test_outbound_service_requires_every_expected_window_and_bounds_publications(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    judgment_rows: list[dict[str, Any]] = [
        {
            "article_version_pub_id": "sav_unit",
            "version_no": 1,
            "title": "安全能力说明",
            "body": "盛邦安全与奇安信的能力比较。",
            "body_sha256": "a" * 64,
            "publication_ready": True,
            "version_created_at": datetime(2026, 8, 5, tzinfo=UTC),
            "judgment_pub_id": "dpj_1",
            "subject_brand": "盛邦安全",
            "target_brand": "奇安信",
            "attitude": "neutral",
            "disparagement": False,
            "evidence_quote": "",
            "confidence": 0.9,
            "method": "llm",
            "model": "unit",
            "prompt_version": "disparage-v2",
            "judgment_status": "ok",
            "window_hash": "window-1",
            "judgment_created_at": datetime(2026, 8, 6, tzinfo=UTC),
            "factcheck_verdict": None,
            "factcheck_summary": None,
            "factcheck_source_url": None,
        }
    ]
    publication_rows = [
        {
            "article_version_pub_id": "sav_unit",
            "platform": "客户官网",
            "public_url": "https://example.test/article",
            "published_at": datetime(2026, 8, 7, tzinfo=UTC),
            "status": "public",
        }
    ]
    statements: list[tuple[str, object]] = []

    class Result:
        def __init__(self, value: object) -> None:
            self.value = value

        def fetchone(self) -> object:
            return self.value

        def fetchall(self) -> object:
            return self.value

    class Connection:
        calls = 0

        def execute(self, statement: str, parameters: object) -> Result:
            self.calls += 1
            statements.append((" ".join(statement.split()), parameters))
            if self.calls == 1:
                return Result(
                    {
                        "pub_id": "spr_unit",
                        "name": "盛邦安全内容核查",
                        "brand_standard_name": "盛邦安全",
                        "brand_profile": {"competitors": ["奇安信"]},
                    }
                )
            if self.calls == 2:
                return Result(judgment_rows)
            return Result(publication_rows)

    @contextmanager
    def connection(*_args: object, **_kwargs: object):
        yield Connection()

    expected_windows = [
        SimpleNamespace(window_hash="window-1", target_brand="奇安信"),
        SimpleNamespace(window_hash="window-2", target_brand="奇安信"),
    ]
    monkeypatch.setattr(formal_review_service2_outbound, "tenant_connection", connection)
    monkeypatch.setattr(
        formal_review_service2_outbound,
        "extract_windows",
        lambda **_kwargs: expected_windows,
    )
    monkeypatch.setattr(
        formal_review_service2_outbound,
        "dedupe_windows",
        lambda windows: windows,
    )

    incomplete = formal_review_service2_outbound.build_outbound_disparagement_facts(
        dsn="postgresql://unused",
        tenant_pub_id="ten_unit",
        sop_project_pub_id="spr_unit",
        start=date(2026, 8, 1),
        end=date(2026, 8, 20),
        generated_at=datetime(2026, 8, 20, tzinfo=UTC),
    )
    assert incomplete["scope"]["expected_windows"] == 2
    assert incomplete["scope"]["completed_windows"] == 1
    assert incomplete["scope"]["judged_content_versions"] == 0
    assert incomplete["evidence_gate"]["status"] == "insufficient"
    publication_query = next(row for row in statements if "FROM sop.publication" in row[0])
    assert "publication.published_at::date BETWEEN %s AND %s" in publication_query[0]
    assert publication_query[1] == (
        "ten_unit",
        "spr_unit",
        date(2026, 8, 1),
        date(2026, 8, 20),
    )

    judgment_rows.append(
        judgment_rows[0]
        | {
            "judgment_pub_id": "dpj_2",
            "window_hash": "window-2",
        }
    )
    complete = formal_review_service2_outbound.build_outbound_disparagement_facts(
        dsn="postgresql://unused",
        tenant_pub_id="ten_unit",
        sop_project_pub_id="spr_unit",
        start=date(2026, 8, 1),
        end=date(2026, 8, 20),
        generated_at=datetime(2026, 8, 20, tzinfo=UTC),
    )
    assert complete["scope"]["completed_windows"] == 2
    assert complete["scope"]["judged_content_versions"] == 1
    assert complete["evidence_gate"]["status"] == "ready"


def test_service5_adapter_uses_the_bound_sop_brand_for_comparison(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    before = FormalWindow(date(2026, 8, 1), date(2026, 8, 8))
    after = FormalWindow(date(2026, 8, 12), date(2026, 8, 20))
    request = _request(
        (5,),
        sop_project_pub_id="spr_unit",
        before_window=before,
        after_window=after,
    )
    received: dict[str, Any] = {}

    def publishing(**kwargs: Any) -> dict[str, Any]:
        del kwargs
        return {
            "schema_version": "formal-publishing-evidence-v1",
            "target_brand": "第二品牌",
            "brand_aliases": ["Second-Brand"],
        }

    def comparison(**kwargs: Any) -> dict[str, Any]:
        received.update(kwargs)
        return {"target_brand": kwargs["target_brand"]}

    def strategy(**kwargs: Any) -> dict[str, Any]:
        received["strategy"] = kwargs
        return {
            "schema_version": "formal-uvw-content-strategy-v1",
            "status": "insufficient",
            "cohort_counts": {},
            "feature_comparison": {},
            "recommendations": [],
        }

    def importer(name: str) -> object:
        if name.endswith("formal_review_service5"):
            return SimpleNamespace(
                build_publishing_evidence=publishing,
                load_uvw_content_strategy_evidence=strategy,
            )
        if name.endswith("formal_review_service4"):
            return SimpleNamespace(build_service4_review_facts=comparison)
        raise AssertionError(name)

    monkeypatch.setattr(production, "import_module", importer)
    monkeypatch.setattr(production, "_assert_sop_brand_binding", lambda *_args, **_kwargs: None)
    adapter = production.QUOTATION_FORMAL_REPORT_REGISTRY[5]
    facts = adapter.build(
        production.FormalBuildContext(
            dsn="postgresql://unused",
            request=request,
            blob_loader=lambda *_args: b"",
        )
    )

    assert received["target_brand"] == "第二品牌"
    assert facts["publication_evidence"]["target_brand"] == "第二品牌"
    assert received["strategy"]["project_pub_id"] == "prj_unit"
    assert facts["uvw_content_strategy"]["status"] == "insufficient"


def test_service5_missing_uvw_analysis_stays_unknown(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class Result:
        def fetchone(self) -> None:
            return None

    class Connection:
        def execute(self, *_args: object, **_kwargs: object) -> Result:
            return Result()

    @contextmanager
    def connection(*_args: object, **_kwargs: object):
        yield Connection()

    monkeypatch.setattr(formal_review_service5, "tenant_connection", connection)
    facts = formal_review_service5.load_uvw_content_strategy_evidence(
        dsn="postgresql://unused",
        tenant_pub_id="ten_unit",
        project_pub_id="prj_unit",
        window_start=date(2026, 8, 1),
        window_end=date(2026, 8, 20),
    )

    assert facts["status"] == "insufficient"
    assert facts["reasons"] == ["analysis_missing"]
    assert facts["cohort_counts"] == {}


def test_service5_publication_boundary_is_strict(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    statements: list[str] = []
    rows = [
        {
            "platform": "客户官网",
            "title": "有效干预",
            "public_url": "https://example.test/valid",
            "published_at": datetime(2026, 8, 10, 8, tzinfo=UTC),
            "public_checked_at": datetime(2026, 8, 10, 9, tzinfo=UTC),
            "public_http_status": 200,
            "status": "public",
            "body_sha256": "a" * 64,
            "version_no": 1,
            "publication_ready": True,
            "has_publication_evidence": True,
            "has_approved_distribution": True,
            "has_publication_attribution": False,
        },
        {
            "platform": "客户官网",
            "title": "后测首日发布",
            "public_url": "https://example.test/ambiguous",
            "published_at": datetime(2026, 8, 12, 0, tzinfo=UTC),
            "public_checked_at": datetime(2026, 8, 12, 1, tzinfo=UTC),
            "public_http_status": 200,
            "status": "public",
            "body_sha256": "b" * 64,
            "version_no": 2,
            "publication_ready": True,
            "has_publication_evidence": True,
            "has_approved_distribution": True,
            "has_publication_attribution": True,
        },
    ]

    class Result:
        def __init__(self, value: object) -> None:
            self.value = value

        def fetchone(self) -> object:
            return self.value

        def fetchall(self) -> object:
            return self.value

    class Connection:
        calls = 0

        def execute(self, *_args: object, **_kwargs: object) -> Result:
            statements.append(" ".join(str(_args[0]).split()))
            self.calls += 1
            if self.calls == 1:
                return Result(
                    {
                        "pub_id": "spr_unit",
                        "name": "盛邦安全最小验证",
                        "brand_standard_name": "盛邦安全",
                        "brand_profile": {"aliases": ["盛邦"]},
                    }
                )
            return Result(rows)

    @contextmanager
    def connection(*_args: object, **_kwargs: object):
        yield Connection()

    monkeypatch.setattr(formal_review_service5, "tenant_connection", connection)
    facts = formal_review_service5.build_publishing_evidence(
        dsn="postgresql://unused",
        tenant_pub_id="ten_unit",
        sop_project_pub_id="spr_unit",
        before_end=date(2026, 8, 8),
        after_start=date(2026, 8, 12),
        window_start=date(2026, 8, 1),
        window_end=date(2026, 8, 20),
        generated_at=datetime(2026, 8, 20, tzinfo=UTC),
    )

    assert facts["brand_aliases"] == ("盛邦",)
    assert facts["summary"]["between_measurement_arms"] == 1
    assert facts["publications"][0]["between_measurement_arms"] is True
    assert facts["publications"][1]["between_measurement_arms"] is False
    assert facts["publication_window"]["required_intervention_before"] == "2026-08-12"
    assert facts["evidence_gate"]["status"] == "ready"
    publication_query = next(
        statement for statement in statements if "FROM sop.publication publication" in statement
    )
    assert "attribution.sop_publication_pub_id=publication.pub_id" in publication_query
    assert "attribution.public_url=publication.public_url" in publication_query

    rows[0]["has_approved_distribution"] = False
    unapproved = formal_review_service5.build_publishing_evidence(
        dsn="postgresql://unused",
        tenant_pub_id="ten_unit",
        sop_project_pub_id="spr_unit",
        before_end=date(2026, 8, 8),
        after_start=date(2026, 8, 12),
        window_start=date(2026, 8, 1),
        window_end=date(2026, 8, 20),
        generated_at=datetime(2026, 8, 20, tzinfo=UTC),
    )
    assert unapproved["evidence_gate"]["status"] == "insufficient"
    assert "approved_distribution_missing" in unapproved["evidence_gate"]["reasons"]


def test_new_service2_and_service5_render_separate_docx_products() -> None:
    generated_at = datetime(2026, 8, 20, tzinfo=UTC)
    governance = {
        "version": "V1.0",
        "prepared_by": "项目组",
        "prepared_date": "2026-08-20",
    }
    service2 = render_outbound_disparagement_docx(
        {
            "schema_version": "formal-outbound-disparagement-v1",
            "service_code": "outbound_disparagement_audit",
            "project_name": "盛邦安全 GEO 项目",
            "target_brand": "盛邦安全",
            "window": {"start": "2026-08-01", "end": "2026-08-20"},
            "generated_at": generated_at,
            "document_status": "internal_review",
            "document_governance": governance,
            "scope": {
                "finalized_content_versions": 1,
                "judged_content_versions": 1,
                "judgment_coverage_complete": True,
                "validation_failures": 0,
                "risk_cases": 0,
            },
            "content_versions": [
                {
                    "title": "安全能力说明",
                    "version_no": 1,
                    "body_sha256": "a" * 64,
                    "judgments": 1,
                    "publications": [],
                }
            ],
            "cases": [],
            "limitations": [],
        }
    )
    service5 = render_publishing_pilot_docx(
        {
            "schema_version": "service4-formal-review-v2",
            "service_code": "content_publishing_pilot",
            "project_name": "盛邦安全 GEO 项目",
            "target_brand": "盛邦安全",
            "window": {"start": "2026-08-01", "end": "2026-08-20"},
            "windows": {
                "before": {"start": "2026-08-01", "end": "2026-08-08"},
                "after": {"start": "2026-08-12", "end": "2026-08-20"},
            },
            "generated_at": generated_at,
            "document_status": "internal_review",
            "document_governance": governance,
            "publication_evidence": {
                "schema_version": "formal-publishing-evidence-v1",
                "summary": {
                    "publications": 1,
                    "between_measurement_arms": 1,
                    "evidence_complete": True,
                },
                "publications": [],
                "causal_boundary": "只能报告描述性关联，不能证明因果。",
            },
            "comparability": {
                "status": "comparable",
                "checks": [{"label": "同问题矩阵", "passed": True, "detail": "一致"}],
            },
            "metrics": [
                {
                    "label": "目标品牌提及率",
                    "before": 0.2,
                    "after": 0.3,
                    "absolute_change": 0.1,
                    "unit": "rate",
                    "before_n": 10,
                    "after_n": 10,
                    "stability": "descriptive",
                }
            ],
            "arms": {"before": {}, "after": {}},
        }
    )

    service2_title, service2_text = _document_text(service2)
    service5_title, service5_text = _document_text(service5)
    assert service2_title == "主动拉踩内容核查报告"
    assert "GEO-S2-V10-INTERNAL-20260820" in service2_text
    assert service5_title == "内容发布与排名提升试点报告"
    assert "GEO-S5-V10-INTERNAL-20260820" in service5_text
    assert "不承诺发帖一定提升排名" in service5_text
    assert service2_title not in service5_text


def test_public_response_derives_catalog_and_output_code_for_history() -> None:
    now = datetime(2026, 8, 20, tzinfo=UTC)
    row = {
        "pub_id": "frp_unit",
        "project_pub_id": "prj_unit",
        "services": [2],
        "status": "awaiting_review",
        "document_status": "internal_review",
        "window_start": date(2026, 8, 1),
        "window_end": date(2026, 8, 20),
        "before_start": None,
        "before_end": None,
        "after_start": None,
        "after_end": None,
        "candidate_group_strategy": "preregistered_scope_v1",
        "document_governance": {},
        "workflow_id": "formal-report/ten_unit/frp_unit",
        "created_at": now,
        "updated_at": now,
    }
    output = {
        "production_pub_id": "frp_unit",
        "service_number": 2,
        "report_pub_id": "rpt_unit",
        "report_version_pub_id": "rpv_unit",
        "fact_snapshot_hash": "a" * 64,
        "format": None,
    }
    service = FormalReportProductionService(
        dsn="postgresql://unused",
        evidence=SimpleNamespace(store=SimpleNamespace()),
    )

    historical = service._public_row(row, outputs=[output])
    current = service._public_row(
        row
        | {
            "services": [5],
            "service_catalog_version": QUOTATION_SERVICE_CATALOG,
            "sop_project_pub_id": "spr_unit",
        },
        outputs=[output | {"service_number": 5}],
    )
    assert historical["service_catalog_version"] == LEGACY_SERVICE_CATALOG
    assert historical["outputs"][0]["service_code"] == "legacy_content_ecosystem_risk"
    assert current["outputs"][0]["service_code"] == "content_publishing_pilot"
    with pytest.raises(FormalProductionIncomplete, match="service_catalog_mismatch"):
        service._public_row(
            row | {"services": [5]},
            outputs=[output | {"service_number": 5}],
        )
    with pytest.raises(FormalProductionIncomplete, match="service_catalog_mismatch"):
        service._public_row(
            row,
            outputs=[output | {"service_number": 3}],
        )


def test_catalog_migration_keeps_versioned_database_invariants() -> None:
    migration = (
        Path(__file__).parents[2]
        / "migrations"
        / "versions"
        / "s06_0031_formal_service_catalog_v2.py"
    ).read_text(encoding="utf-8")
    merge = (
        Path(__file__).parents[2]
        / "migrations"
        / "versions"
        / "s06_0033_merge_source_and_formal_catalog.py"
    ).read_text(encoding="utf-8")

    assert 'revision: str = "s06_0031_formal_catalog"' in migration
    assert 'down_revision: str | Sequence[str] | None = "s06_catalog_0001"' in migration
    assert "WHEN service_catalog_version='quotation_services_v2' THEN" in migration
    assert "services <@ ARRAY[1,2,3,4,5]::SMALLINT[]" in migration
    assert "services <@ ARRAY[1,2,3,4]::SMALLINT[]" in migration
    assert "services && ARRAY[2,5]::SMALLINT[]" in migration
    assert "before_end < after_start AND 5 = ANY(services)" in migration
    assert "formal_output_service_catalog_trg" in migration
    assert "NEW.service_number = ANY(parent_services)" in migration
    assert "FOR NO KEY UPDATE" in migration
    assert "legacy formal output service must be between 1 and 4" in migration
    assert "formal_production_catalog_immutable_trg" in migration
    assert "cannot downgrade while quotation_services_v2 productions exist" in migration
    assert '"s06_0032",' in merge
    assert '"s06_0031_formal_catalog",' in merge
    assert "target ``s06_0031_formal_catalog`` explicitly" in merge
