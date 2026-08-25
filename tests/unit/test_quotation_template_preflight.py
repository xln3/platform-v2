from __future__ import annotations

import hashlib
import json
import shutil
import subprocess
import sys
from pathlib import Path
from zipfile import ZipFile

import yaml
from docx import Document

ROOT = Path(__file__).resolve().parents[2]
ASSETS = ROOT / "api/geo_platform/quotations/assets"
V1 = ASSETS / "quotation-template-v1.docx"
MANIFEST = ASSETS / "quotation-template-v1.yaml"
STRUCTURE = ASSETS / "quotation-template-v1.structure.json"
PREFLIGHT = ROOT / "tools/quotation_template_preflight.py"
V1_SHA256 = "90ae5beb10ab3bacea3b706a2068945f828e275784e99da6b72dc44f8b0d9913"


def _run_preflight(*args: str, script: Path = PREFLIGHT) -> tuple[int, dict[str, object]]:
    result = subprocess.run(
        [sys.executable, str(script), *args],
        cwd=script.parents[1],
        check=False,
        capture_output=True,
        text=True,
    )
    return result.returncode, json.loads(result.stdout)


def test_v1_template_identity_and_manifest_are_frozen() -> None:
    manifest = yaml.safe_load(MANIFEST.read_text(encoding="utf-8"))
    assert V1.is_file()
    assert hashlib.sha256(V1.read_bytes()).hexdigest() == V1_SHA256
    assert manifest["template_id"] == "geo-quotation-v1"
    assert manifest["template_version"] == "v1"
    assert manifest["status"] == "approved"
    assert manifest["approval_status"] == "approved"
    assert manifest["sha256"] == V1_SHA256
    assert manifest["production_use"]["status"] == "blocked_pending_v2_approval"


def test_v1_structure_contract_matches_approved_docx() -> None:
    contract = json.loads(STRUCTURE.read_text(encoding="utf-8"))
    document = Document(V1)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    assert contract["template_sha256"] == V1_SHA256
    assert len(document.sections) == contract["document"]["section_count"] == 1
    assert len(document.paragraphs) == contract["document"]["paragraph_count"] == 200
    assert len(document.tables) == contract["document"]["table_count"] == 2
    section = document.sections[0]
    assert section.page_width.twips == 11906
    assert section.page_height.twips == 16838
    assert section.top_margin.twips == section.bottom_margin.twips == 1020
    assert section.left_margin.twips == section.right_margin.twips == 850

    assert paragraphs[0] == contract["anchors"]["title"] == "GEO验证服务报价单"
    main_table = document.tables[0]
    assert len(main_table.rows) == contract["main_table"]["rows"] == 6
    assert len(main_table.columns) == contract["main_table"]["columns"] == 4
    assert [cell.text for cell in main_table.rows[0].cells] == contract["main_table"]["headers"]
    assert [main_table.cell(index, 1).text for index in range(1, 5)] == contract["main_table"][
        "service_names"
    ]
    assert all(term in paragraphs for term in contract["commercial_terms"])
    assert contract["anchors"]["signature"] in paragraphs
    assert section.header.paragraphs[0].text == contract["anchors"]["header"]

    appendix_positions = [
        paragraphs.index(appendix["title"]) for appendix in contract["appendices"]
    ]
    assert appendix_positions == sorted(appendix_positions)
    all_text = "\n".join(paragraphs)
    assert "服务输入、执行与交付说明" not in all_text
    assert "GEO 服务报价单" not in all_text
    with ZipFile(V1) as archive:
        footer = archive.read("word/footer1.xml").decode("utf-8")
    assert "PAGE  \\* MERGEFORMAT" in footer


def test_preflight_passes_v1_identity_but_fails_production_gate() -> None:
    code, payload = _run_preflight()
    assert code == 0
    assert payload == {
        "approval_status": "approved",
        "canonical_template": str(V1),
        "ok": True,
        "production_use_status": "blocked_pending_v2_approval",
        "template_id": "geo-quotation-v1",
        "template_sha256": V1_SHA256,
        "template_version": "v1",
    }

    code, payload = _run_preflight("--require-production")
    assert code == 1
    assert payload["code"] == "quotation_template_not_approved"


def test_preflight_fails_closed_for_missing_modified_unapproved_and_unknown(
    tmp_path: Path,
) -> None:
    code, payload = _run_preflight("--manifest", str(tmp_path / "missing.yaml"))
    assert code == 1
    assert payload["code"] == "quotation_template_missing"

    sandbox_root = tmp_path / "platform-v2"
    sandbox_tools = sandbox_root / "tools"
    sandbox_assets = sandbox_root / "api/geo_platform/quotations/assets"
    sandbox_tools.mkdir(parents=True)
    sandbox_assets.mkdir(parents=True)
    sandbox_script = sandbox_tools / PREFLIGHT.name
    copied_manifest = sandbox_assets / MANIFEST.name
    template = sandbox_assets / V1.name
    shutil.copyfile(PREFLIGHT, sandbox_script)
    shutil.copyfile(V1, template)
    manifest = yaml.safe_load(MANIFEST.read_text(encoding="utf-8"))
    copied_manifest.write_text(
        yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False), encoding="utf-8"
    )

    template.write_bytes(template.read_bytes() + b"modified")
    code, payload = _run_preflight(script=sandbox_script)
    assert code == 1
    assert payload["code"] == "quotation_template_hash_mismatch"

    manifest["sha256"] = hashlib.sha256(template.read_bytes()).hexdigest()
    copied_manifest.write_text(
        yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False), encoding="utf-8"
    )
    code, payload = _run_preflight(script=sandbox_script)
    assert code == 1
    assert payload["code"] == "quotation_template_hash_mismatch"

    shutil.copyfile(V1, template)
    manifest["sha256"] = V1_SHA256
    manifest["approval_status"] = "pending"
    copied_manifest.write_text(
        yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False), encoding="utf-8"
    )
    code, payload = _run_preflight(script=sandbox_script)
    assert code == 1
    assert payload["code"] == "quotation_template_not_approved"

    manifest["approval_status"] = "approved"
    manifest["schema_version"] = "unknown"
    copied_manifest.write_text(
        yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False), encoding="utf-8"
    )
    code, payload = _run_preflight(script=sandbox_script)
    assert code == 1
    assert payload["code"] == "quotation_template_version_unknown"
