"""报价单：XLSX 解析、LLM 内容契约、DOCX 格式与端到端编排。"""

from __future__ import annotations

from collections import Counter
from datetime import date
from io import BytesIO
from xml.etree import ElementTree
from xml.sax.saxutils import escape
from zipfile import ZIP_DEFLATED, ZipFile

import geo_platform.quotations.router as quotation_router
import pytest
from docx import Document
from docx.oxml.ns import qn
from fastapi.testclient import TestClient
from geo_platform.config import Settings
from geo_platform.identity.policy import Principal, Role, get_principal
from geo_platform.main import app
from geo_platform.quotations.generator import (
    QuotationGenerationFailed,
    plan_from_payload,
    selection_quotas,
)
from geo_platform.quotations.models import QuotationConfiguration, TargetQuery
from geo_platform.quotations.renderer import render_quotation_docx
from geo_platform.quotations.service import QuotationInputInvalid, generate_quotation
from geo_platform.quotations.xlsx import TargetWorkbookInvalid, parse_target_queries


def _xlsx(rows: list[dict[str, str]], *, extra_name: str | None = None) -> bytes:
    row_xml: list[str] = []
    for row_number, cells in enumerate(rows, start=1):
        rendered = "".join(
            f'<c r="{column}{row_number}" t="inlineStr"><is><t>{escape(value)}</t></is></c>'
            for column, value in cells.items()
        )
        row_xml.append(f'<row r="{row_number}">{rendered}</row>')
    worksheet = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f"<sheetData>{''.join(row_xml)}</sheetData></worksheet>"
    )
    workbook = """<?xml version="1.0" encoding="UTF-8"?>
    <workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"
      xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
      <sheets><sheet name="目标词" sheetId="1" r:id="rId1"/></sheets>
    </workbook>"""
    relationships = """<?xml version="1.0" encoding="UTF-8"?>
    <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
      <Relationship Id="rId1"
        Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet"
        Target="worksheets/sheet1.xml"/>
    </Relationships>"""
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        archive.writestr("xl/workbook.xml", workbook)
        archive.writestr("xl/_rels/workbook.xml.rels", relationships)
        archive.writestr("xl/worksheets/sheet1.xml", worksheet)
        if extra_name:
            archive.writestr(extra_name, "hostile")
    return output.getvalue()


def _queries() -> list[TargetQuery]:
    result: list[TargetQuery] = []
    index = 1
    for group in ("网空线", "防护线", "星网线", "检测线", "安服线"):
        for item in range(1, 5):
            result.append(
                TargetQuery(
                    query_id=f"Q{index:03d}",
                    group=group,
                    text=f"{group}目标问题{item}",
                    sheet="目标词",
                    row=index + 2,
                )
            )
            index += 1
    return result


def _configuration(
    package_code: str = "geo_effect_assessment",
    *,
    priced: bool = True,
    official_site_in_citations: bool | None = True,
    artifact_kind: str = "complete",
) -> QuotationConfiguration:
    if package_code == "geo_effect_assessment":
        rows = (
            ("ranking_test", 1, 2_000_000),
            ("outbound_disparagement_audit", 1, 800_000),
            ("inbound_disparagement_audit", 1, 1_200_000),
            ("official_site_audit", 1, 1_000_000),
        )
    else:
        rows = (
            ("ranking_test", 2, 2_000_000),
            ("inbound_disparagement_audit", 1, 1_200_000),
            *(
                (("official_site_audit", 1, 1_000_000),)
                if official_site_in_citations is not False
                else ()
            ),
            ("content_publishing_pilot", 1, 3_000_000),
        )
    return QuotationConfiguration.model_validate(
        {
            "package_code": package_code,
            "artifact_kind": artifact_kind,
            "website_url": "https://www.webray.com.cn",
            "official_site_in_citations": official_site_in_citations,
            "official_site_citation_url": (
                "https://www.webray.com.cn/cited-page"
                if package_code == "minimum_validation" and official_site_in_citations is True
                else ""
            ),
            "pricing_status": "priced" if priced else "pending",
            "service_quotes": [
                {
                    "service_code": code,
                    "quantity": quantity,
                    "unit_price_cents": price if priced else None,
                }
                for code, quantity, price in rows
            ],
        }
    )


def _llm_payload(queries: list[TargetQuery]) -> dict[str, object]:
    quotas = selection_quotas(queries)
    selected: list[dict[str, str]] = []
    for group, quota in quotas.items():
        for query in [row for row in queries if row.group == group][:quota]:
            selected.append(
                {
                    "source_id": query.query_id,
                    "variant_a": f"围绕{query.text}，国内主要产品与服务商有哪些？",
                    "variant_b": f"企业评估{query.text}相关方案时应关注哪些选型维度？",
                    "variant_c": f"想解决{query.text}，找哪家公司比较靠谱？",
                }
            )
    opportunities = [
        {
            "keyword": f"相邻业务机会词{index}",
            "optimized_query": f"相邻业务机会词{index}有哪些成熟产品和服务商？",
            "variant_a": f"国内提供相邻业务机会词{index}解决方案的主要厂商有哪些？",
            "variant_b": f"采购相邻业务机会词{index}方案时应比较哪些技术维度？",
            "variant_c": f"相邻业务机会词{index}这块找谁做比较靠谱？",
            "rewrite_rationale": "由概念了解转向产品选型、厂商推荐与方案比较意图",
        }
        for index in range(1, 17)
    ]
    return {
        "profile": {
            "category_label": "网络安全产品与服务",
            "sec_profile": "trust",
            "category_analysis": (
                "该类产品专业门槛较高，采购方在购买前难以仅凭参数判断实际效果，"
                "通常需要结合专业评测、适用场景、资质材料与可信第三方信息进行决策。"
            ),
            "intent_diagnosis": (
                "原始目标词兼有概念了解、问题解决和厂商选型意图。改写时保留核心业务语义，"
                "并明确产品、服务商、比较维度或应用场景，形成可验证的推荐型查询。"
            ),
        },
        "selected_queries": selected,
        "opportunities": opportunities,
        "sources": [{"title": "品牌官网", "url": "https://example.com/brand"}],
    }


def _workbook_for_queries(queries: list[TargetQuery]) -> bytes:
    rows: list[dict[str, str]] = [{"B": "产线"}]
    last_group = ""
    sequence = 1
    for query in queries:
        if query.group != last_group:
            rows.append({"A": "序号", "B": f"{query.group}-关键问题"})
            last_group = query.group
            sequence = 1
        rows.append({"A": str(sequence), "B": query.text})
        sequence += 1
    return _xlsx(rows)


def test_xlsx_parser_preserves_group_order_and_deduplicates() -> None:
    payload = _xlsx(
        [
            {"B": "产线"},
            {"A": "序号", "B": "网空线-关键问题"},
            {"A": "1", "B": "企业未知资产怎么排查"},
            {"A": "2", "B": "资产测绘平台推荐"},
            {"A": "序号", "B": "防护线-关键问题"},
            {"A": "1", "B": "老旧系统打不了补丁怎么办"},
            {"A": "2", "B": "资产测绘平台推荐"},
        ]
    )
    rows = parse_target_queries(payload)
    assert [(row.query_id, row.group, row.text) for row in rows] == [
        ("Q001", "网空线", "企业未知资产怎么排查"),
        ("Q002", "网空线", "资产测绘平台推荐"),
        ("Q003", "防护线", "老旧系统打不了补丁怎么办"),
    ]


def test_xlsx_parser_rejects_non_xlsx_and_archive_traversal() -> None:
    with pytest.raises(TargetWorkbookInvalid, match="xlsx_signature_invalid"):
        parse_target_queries(b"not-an-xlsx")
    with pytest.raises(TargetWorkbookInvalid, match="xlsx_archive_path_invalid"):
        parse_target_queries(_xlsx([{"A": "目标词"}, {"A": "测试问题"}], extra_name="../x"))


def test_service_rejects_secret_in_workbook_group_before_llm() -> None:
    workbook = _xlsx(
        [
            {"B": "产线"},
            {"A": "序号", "B": "access_token=topsecret-关键问题"},
            {"A": "1", "B": "正常测试问题"},
        ]
    )
    with pytest.raises(QuotationInputInvalid, match="target_words_contain_secret"):
        generate_quotation(
            brand_name="盛邦安全",
            configuration=_configuration(),
            workbook_payload=workbook,
            settings=Settings(),
        )


def test_selection_quota_matches_reference_balance() -> None:
    assert selection_quotas(_queries()) == {
        "网空线": 4,
        "防护线": 4,
        "星网线": 4,
        "检测线": 3,
        "安服线": 3,
    }


def test_package_pricing_keeps_service_one_baseline_and_retest_separate() -> None:
    legacy_payload = _configuration().model_dump(mode="json")
    legacy_payload.pop("artifact_kind")
    assert QuotationConfiguration.model_validate(legacy_payload).artifact_kind == "complete"

    minimum = _configuration("minimum_validation")
    assert minimum.service_codes == (
        "ranking_test",
        "inbound_disparagement_audit",
        "official_site_audit",
        "content_publishing_pilot",
    )
    assert minimum.service_quotes[0].quantity == 2
    assert minimum.service_quotes[0].subtotal_cents == 4_000_000
    assert minimum.total_price_cents == 9_200_000
    assert minimum.maximum_total_price_cents == 9_200_000

    conditional = _configuration(
        "minimum_validation",
        official_site_in_citations=None,
    )
    assert "official_site_audit" in conditional.service_codes
    assert conditional.total_price_cents == 8_200_000
    assert conditional.maximum_total_price_cents == 9_200_000
    conditional_document = Document(
        BytesIO(
            render_quotation_docx(
                brand_name="盛邦安全",
                quote_date=date(2026, 8, 20),
                configuration=conditional,
                plan=None,
            )
        )
    )
    conditional_text = "\n".join(
        [*(paragraph.text for paragraph in conditional_document.paragraphs)]
        + [
            cell.text
            for table in conditional_document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    assert "基础服务费合计（不含条件项）" in conditional_text
    assert "已确认服务费合计" not in conditional_text
    rendered = Document(
        BytesIO(
            render_quotation_docx(
                brand_name="盛邦安全",
                quote_date=date(2026, 8, 20),
                configuration=conditional,
                plan=None,
            )
        )
    )
    rendered_text = "\n".join(
        [*(paragraph.text for paragraph in rendered.paragraphs)]
        + [cell.text for table in rendered.tables for row in table.rows for cell in row.cells]
    )
    assert "触发后 ￥10,000.00" in rendered_text
    assert "基础服务费合计（不含条件项）" in rendered_text
    assert "￥82,000.00" in rendered_text
    assert "官网命中后最高服务费" in rendered_text
    assert "￥92,000.00" in rendered_text

    without_site = _configuration(
        "minimum_validation",
        official_site_in_citations=False,
    )
    assert "official_site_audit" not in without_site.service_codes
    assert without_site.total_price_cents == 8_200_000
    assert without_site.maximum_total_price_cents == 8_200_000


def test_package_rejects_browser_supplied_service_shape_or_missing_website() -> None:
    payload = _configuration().model_dump(mode="json")
    payload["service_quotes"] = payload["service_quotes"][:-1]
    with pytest.raises(ValueError, match="package_service_quantities_invalid"):
        QuotationConfiguration.model_validate(payload)

    payload = _configuration().model_dump(mode="json")
    payload["website_url"] = ""
    with pytest.raises(ValueError, match="official_website_required"):
        QuotationConfiguration.model_validate(payload)

    payload = _configuration("minimum_validation").model_dump(mode="json")
    payload["official_site_citation_url"] = "https://not-webray.example/evidence"
    with pytest.raises(ValueError, match="official_site_citation_evidence_required"):
        QuotationConfiguration.model_validate(payload)

    payload = {
        "package_code": "custom",
        "service_quotes": [
            {
                "service_code": "content_publishing_pilot",
                "quantity": 1,
                "unit_price_cents": 3_000_000,
            }
        ],
    }
    with pytest.raises(ValueError, match="publishing_requires_two_ranking_rounds"):
        QuotationConfiguration.model_validate(payload)

    payload = {
        "package_code": "custom",
        "official_site_in_citations": True,
        "official_site_citation_url": "https://www.webray.com.cn/cited-page",
        "service_quotes": [
            {
                "service_code": "ranking_test",
                "quantity": 1,
                "unit_price_cents": 2_000_000,
            }
        ],
    }
    with pytest.raises(ValueError, match="official_site_citation_evidence_not_applicable"):
        QuotationConfiguration.model_validate(payload)


def test_custom_two_round_ranking_without_publishing_keeps_quantity_in_sequence() -> None:
    configuration = QuotationConfiguration.model_validate(
        {
            "package_code": "custom",
            "service_quotes": [
                {
                    "service_code": "ranking_test",
                    "quantity": 2,
                    "unit_price_cents": 2_000_000,
                }
            ],
        }
    )
    payload = render_quotation_docx(
        brand_name="盛邦安全",
        quote_date=date(2026, 8, 20),
        configuration=configuration,
        plan=None,
    )
    document = Document(BytesIO(payload))
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "执行顺序：1 测试（2轮）" in body_text


def test_llm_payload_is_bound_back_to_uploaded_source_and_group_quota() -> None:
    queries = _queries()
    payload = _llm_payload(queries)
    plan = plan_from_payload(payload, queries=queries, model="fixture-model")
    source = {row.query_id: row for row in queries}
    assert len(plan.selected_queries) == 18
    assert Counter(row.group for row in plan.selected_queries) == selection_quotas(queries)
    assert all(row.original == source[row.source_id].text for row in plan.selected_queries)
    assert len(plan.opportunities) == 16
    assert plan.sources[0].url == "https://example.com/brand"


def test_llm_payload_rejects_unverified_effect_numbers() -> None:
    queries = _queries()
    payload = _llm_payload(queries)
    profile = payload["profile"]
    assert isinstance(profile, dict)
    profile["intent_diagnosis"] = "优化后厂商推荐提升37次，因此可以直接确认品牌表现已经明显改善。"
    with pytest.raises(QuotationGenerationFailed, match="llm_unverified_measurement_claim"):
        plan_from_payload(payload, queries=queries, model="fixture-model")


def test_legacy_renderer_marks_noncompliance_and_avoids_unsupported_measurements() -> None:
    queries = _queries()
    plan = plan_from_payload(_llm_payload(queries), queries=queries, model="fixture-model")
    payload = render_quotation_docx(
        brand_name="盛邦安全",
        quote_date=date(2026, 8, 10),
        configuration=_configuration(),
        plan=plan,
    )
    assert payload == render_quotation_docx(
        brand_name="盛邦安全",
        quote_date=date(2026, 8, 10),
        configuration=_configuration(),
        plan=plan,
    )
    assert payload.startswith(b"PK")
    document = Document(BytesIO(payload))
    assert document.core_properties.created is not None
    assert document.core_properties.created.date() == date(2026, 8, 10)
    assert document.core_properties.created.hour == 0
    assert len(document.sections) == 1
    section = document.sections[0]
    assert section.page_width is not None
    assert section.page_height is not None
    assert section.left_margin is not None
    assert section.top_margin is not None
    assert round(section.page_width.mm) == 210
    assert round(section.page_height.mm) == 297
    assert round(section.left_margin.mm) == 15
    assert round(section.top_margin.mm) == 18
    assert len(document.tables) == 1
    assert len(document.tables[0].rows) == 6
    assert len(document.tables[0].columns) == 6
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "盛邦安全" in body_text
    assert "附录一 服务输入、执行与交付说明" in body_text
    assert "附录二 原推广 Query 与变体构建说明" in body_text
    assert "新增 Query 优化与语义变体全表" not in body_text
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "￥20,000.00" in table_text
    assert "￥50,000.00" in table_text
    assert "厂商推荐15次" not in body_text
    document_text = f"{body_text}\n{table_text}"
    assert "主动拉踩内容核查" in document_text
    assert "被拉踩内容核查" in document_text
    assert "API 与豆包 App" in document_text
    with ZipFile(BytesIO(payload)) as archive:
        header = archive.read("word/header1.xml").decode()
        footer = archive.read("word/footer1.xml").decode()
        document_xml = archive.read("word/document.xml").decode()
    assert "本报价为保密商业资料" in header
    assert "非最终模板合规产物·仅供内部回归" in header
    assert " PAGE " in footer
    assert document_xml.count('w:type="page"') >= 2


def test_service_runs_one_pipeline_and_returns_integrity_metadata() -> None:
    queries = _queries()
    llm_payload = _llm_payload(queries)
    llm_payload["opportunities"] = []
    calls: list[dict[str, object]] = []

    def runner(
        *args: object, **kwargs: object
    ) -> tuple[dict[str, object], list[dict[str, str]], dict[str, int]]:
        calls.append(kwargs)
        return llm_payload, [], {"input_tokens": 10, "output_tokens": 20}

    result = generate_quotation(
        brand_name="盛邦安全",
        configuration=_configuration(),
        workbook_payload=_workbook_for_queries(queries),
        settings=Settings(research_llm_api_key="unit-test-key"),
        quote_date=date(2026, 8, 10),
        runner=runner,
    )
    assert len(calls) == 1
    assert (
        result.metadata.filename == "非最终模板合规产物-报价单-盛邦安全-GEO效果评测-20260810.docx"
    )
    assert result.metadata.artifact_kind == "complete"
    assert result.metadata.total_price_cents == 5_000_000
    assert result.metadata.maximum_total_price_cents == 5_000_000
    assert result.metadata.service_count == 4
    assert result.metadata.target_query_count == 20
    assert result.metadata.selected_query_count == 18
    assert result.metadata.opportunity_count == 0
    assert len(result.metadata.sha256) == 64
    assert result.payload.startswith(b"PK")


def test_quote_table_artifact_skips_llm_and_excludes_all_appendices() -> None:
    def runner(
        *_: object, **__: object
    ) -> tuple[dict[str, object], list[dict[str, str]], dict[str, int]]:
        raise AssertionError("quote_table_must_not_call_llm")

    result = generate_quotation(
        brand_name="盛邦安全",
        configuration=_configuration(artifact_kind="quote_table"),
        workbook_payload=_workbook_for_queries(_queries()),
        settings=Settings(),
        quote_date=date(2026, 8, 20),
        runner=runner,
    )
    assert result.plan is None
    assert result.metadata.artifact_kind == "quote_table"
    assert (
        result.metadata.filename
        == "非最终模板合规产物-报价单-盛邦安全-GEO效果评测-报价单表格-20260820.docx"
    )
    assert result.metadata.target_query_count == 20
    assert result.metadata.query_appendix_included is False
    document = Document(BytesIO(result.payload))
    assert document.core_properties.title == "盛邦安全 已开展 GEO · 效果评测报价单表格"
    assert document.core_properties.subject == "GEO 服务报价单表格"
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert len(document.tables) == 1
    assert "GEO 服务报价单" in text
    assert "非最终模板合规产物（仅供内部回归，禁止作为正式客户报价）" in text
    assert "附录一 服务输入、执行与交付说明" not in text
    assert "原推广 Query 与变体构建说明" not in text


def test_query_appendix_artifact_contains_only_query_sections() -> None:
    llm_payload = _llm_payload(_queries())
    llm_payload["opportunities"] = []

    def runner(
        *_: object, **__: object
    ) -> tuple[dict[str, object], list[dict[str, str]], dict[str, int]]:
        return llm_payload, [], {"input_tokens": 10, "output_tokens": 20}

    result = generate_quotation(
        brand_name="盛邦安全",
        configuration=_configuration(artifact_kind="query_appendix"),
        workbook_payload=_workbook_for_queries(_queries()),
        settings=Settings(research_llm_api_key="unit-test-key"),
        quote_date=date(2026, 8, 20),
        runner=runner,
    )
    assert result.plan is not None
    assert result.metadata.artifact_kind == "query_appendix"
    assert (
        result.metadata.filename
        == "非最终模板合规产物-报价单-盛邦安全-GEO效果评测-查询附件-20260820.docx"
    )
    assert result.metadata.query_appendix_included is True
    document = Document(BytesIO(result.payload))
    assert document.core_properties.title == "盛邦安全 已开展 GEO · 效果评测查询附件"
    assert document.core_properties.subject == "GEO 查询附件"
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert len(document.tables) == 0
    assert "盛邦安全 GEO 查询附件" in text
    assert "报价日期：2026-08-20" in text
    assert "附录一 原推广 Query 与变体构建说明" in text
    assert "GEO 服务报价单" not in text
    assert "非最终模板合规产物（仅供内部回归，禁止作为正式客户报价）" in text
    assert "服务输入、执行与交付说明" not in text
    with ZipFile(BytesIO(result.payload)) as archive:
        document_xml = archive.read("word/document.xml").decode()
    assert document_xml.count('w:type="page"') == 0
    root = ElementTree.fromstring(document_xml)
    variant_paragraphs = []
    for paragraph in root.iter(qn("w:p")):
        paragraph_text = "".join(node.text or "" for node in paragraph.iter(qn("w:t")))
        if paragraph_text.startswith(("A  ", "B  ")):
            variant_paragraphs.append(paragraph)
    assert len(variant_paragraphs) == 36
    assert all(
        paragraph.find(f"{qn('w:pPr')}/{qn('w:keepNext')}") is not None
        for paragraph in variant_paragraphs
    )


def test_query_appendix_artifact_fails_closed_without_source_or_query_service() -> None:
    with pytest.raises(QuotationInputInvalid, match="query_appendix_workbook_required"):
        generate_quotation(
            brand_name="盛邦安全",
            configuration=_configuration(artifact_kind="query_appendix"),
            workbook_payload=None,
            settings=Settings(),
        )

    configuration = QuotationConfiguration.model_validate(
        {
            "package_code": "custom",
            "artifact_kind": "query_appendix",
            "service_quotes": [
                {
                    "service_code": "inbound_disparagement_audit",
                    "quantity": 1,
                    "unit_price_cents": 1_200_000,
                }
            ],
        }
    )
    with pytest.raises(QuotationInputInvalid, match="query_appendix_service_required"):
        generate_quotation(
            brand_name="盛邦安全",
            configuration=configuration,
            workbook_payload=_workbook_for_queries(_queries()),
            settings=Settings(),
        )
    with pytest.raises(ValueError, match="query_appendix_plan_required"):
        render_quotation_docx(
            brand_name="盛邦安全",
            quote_date=date(2026, 8, 20),
            configuration=_configuration(artifact_kind="query_appendix"),
            plan=None,
        )


def test_service_generates_priced_business_quote_without_xlsx_or_llm() -> None:
    result = generate_quotation(
        brand_name="盛邦安全",
        configuration=_configuration("minimum_validation"),
        workbook_payload=None,
        settings=Settings(),
        quote_date=date(2026, 8, 20),
    )
    assert result.plan is None
    assert result.metadata.query_appendix_included is False
    assert result.metadata.target_query_count == 0
    assert result.metadata.selected_query_count == 0
    assert result.metadata.opportunity_count == 0
    assert result.metadata.total_price_cents == 9_200_000
    document = Document(BytesIO(result.payload))
    text = "\n".join(
        [*(paragraph.text for paragraph in document.paragraphs)]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "1 测试（基线） → 3 找被拉踩帖 → 4 官网分析 → 5 发帖提排名 → 1 测试（复测）" in text
    assert "￥92,000.00" in text
    assert "具体 Query 及语义变体将由我方提出候选，客户可补充并最终确认后冻结" in text


def test_generate_api_returns_named_no_store_docx(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    queries = _queries()
    llm_payload = _llm_payload(queries)
    llm_payload["opportunities"] = []

    def runner(
        *args: object, **kwargs: object
    ) -> tuple[dict[str, object], list[dict[str, str]], dict[str, int]]:
        return llm_payload, [], {"input_tokens": 10, "output_tokens": 20}

    generated = generate_quotation(
        brand_name="盛邦安全",
        configuration=_configuration(),
        workbook_payload=_workbook_for_queries(queries),
        settings=Settings(research_llm_api_key="unit-test-key"),
        quote_date=date(2026, 8, 10),
        runner=runner,
    )
    monkeypatch.setattr(quotation_router, "generate_quotation", lambda **_: generated)
    app.dependency_overrides[get_principal] = lambda: Principal(
        subject="operator@example.test",
        role=Role.OPERATOR,
        tenant_pub_id="ten_unit",
        user_pub_id="usr_unit",
    )
    try:
        with TestClient(app) as client:
            response = client.post(
                "/api/v2/quotations/generate",
                headers={"Origin": "http://127.0.0.1:45101"},
                data={
                    "brand_name": "盛邦安全",
                    "quote_date": "2026-08-10",
                    "quotation_config": _configuration().model_dump_json(),
                },
                files={
                    "target_words": (
                        "目标词.xlsx",
                        _workbook_for_queries(queries),
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )
                },
            )
    finally:
        app.dependency_overrides.pop(get_principal, None)

    assert response.status_code == 200
    assert response.content == generated.payload
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert response.headers["cache-control"] == "no-store"
    assert response.headers["x-quotation-target-query-count"] == "20"
    assert response.headers["x-quotation-package-code"] == "geo_effect_assessment"
    assert response.headers["x-quotation-artifact-kind"] == "complete"
    assert response.headers["x-quotation-service-count"] == "4"
    assert response.headers["x-quotation-total-cents"] == "5000000"
    assert response.headers["x-quotation-maximum-total-cents"] == "5000000"
    assert response.headers["x-quotation-query-appendix"] == "included"
    assert response.headers["x-quotation-template-compliance"] == "non-final-template"
    assert "UTF-8''" in response.headers["content-disposition"]
    exposed = {
        value.strip().lower()
        for value in response.headers["access-control-expose-headers"].split(",")
    }
    assert {
        "content-disposition",
        "x-quotation-artifact-kind",
        "x-quotation-total-cents",
        "x-quotation-maximum-total-cents",
        "x-quotation-sha256",
        "x-quotation-template-compliance",
    } <= exposed
