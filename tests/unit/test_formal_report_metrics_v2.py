from __future__ import annotations

import copy
from datetime import UTC, date, datetime
from io import BytesIO
from types import SimpleNamespace
from typing import Any

import pytest
from docx import Document
from geo_platform.reports import formal_production as production
from geo_platform.reports.formal_production import (
    QUOTATION_SERVICE_CATALOG,
    FormalProductionIncomplete,
    FormalProductionInvalid,
    FormalProductionRequest,
    FormalReportProductionService,
    FormalWindow,
    customer_fact_snapshot,
)
from geo_platform.reports.formal_production_router import FormalProductionCreate
from pydantic import ValidationError

from domain.reporting.formal_metric_snapshot_docx import (
    render_bound_metric_snapshot_docx,
)
from domain.reporting.metric_snapshot_binding import bind_metric_snapshot_set

HASH = "a" * 64
FILTERS = {"model": [], "region": [], "mode": []}


def _snapshot_document() -> dict[str, Any]:
    return {
        "schema_version": "metric-snapshot-set-v2",
        "snapshot_set_pub_id": "mss_report",
        "snapshot_set_hash": HASH,
        "project_pub_id": "prj_report",
        "state": "partial",
        "window": {"start": date(2026, 8, 1), "end": date(2026, 8, 27)},
        "filters": FILTERS,
        "aggregation_method": "query_macro",
        "metrics": [
            {
                "snapshot_pub_id": "msn_ready",
                "snapshot_hash": HASH,
                "focal_entity_id": "entity_target",
                "metric_name": "ai_recommendation_organic_mention_rate_v2",
                "metric_version": "2.0.0",
                "metric_definition_hash": HASH,
                "state": "ready",
                "value": 0.5,
                "observed_value": 0.5,
                "raw_numerator": 2,
                "raw_denominator": 4,
                "unique_query_count": 1,
                "coverage": {"semantic": 1},
                "contribution_set_hash": HASH,
                "query_contribution_set_hash": HASH,
                "design_contribution_set_hash": HASH,
            },
            {
                "snapshot_pub_id": "msn_experimental",
                "snapshot_hash": HASH,
                "focal_entity_id": "entity_target",
                "metric_name": "geo_visibility_index",
                "metric_version": "2.0.0",
                "metric_definition_hash": HASH,
                "state": "experimental",
                "value": None,
                "observed_value": 0.7,
                "raw_numerator": 7,
                "raw_denominator": 10,
                "unique_query_count": 1,
                "coverage": {"semantic": 1},
                "contribution_set_hash": HASH,
                "query_contribution_set_hash": HASH,
                "design_contribution_set_hash": HASH,
            },
        ],
    }


class SnapshotReader:
    def __init__(self, document: dict[str, Any]) -> None:
        self.document = document
        self.exports = 0

    def get_snapshot_set(self, *, tenant_pub_id: str, set_pub_id: str) -> dict[str, Any]:
        assert tenant_pub_id == "ten_report"
        assert set_pub_id == "mss_report"
        return copy.deepcopy(self.document)

    def export_bundle(self, *, tenant_pub_id: str, set_pub_id: str) -> dict[str, Any]:
        assert tenant_pub_id == "ten_report"
        assert set_pub_id == "mss_report"
        self.exports += 1
        return {
            "readme": [
                {
                    "snapshot_set_pub_id": "mss_report",
                    "snapshot_set_hash": HASH,
                    "project_pub_id": "prj_report",
                    "state": self.document["state"],
                    "window_start": "2026-08-01",
                    "window_end": "2026-08-27",
                    "filters": FILTERS,
                    "aggregation_method": "query_macro",
                }
            ],
            "metrics": [
                {
                    "pub_id": metric["snapshot_pub_id"],
                    "snapshot_hash": metric["snapshot_hash"],
                    "focal_entity_id": metric["focal_entity_id"],
                    "metric_name": metric["metric_name"],
                    "metric_version": metric["metric_version"],
                    "state": metric["state"],
                    "value": metric["value"],
                    "observed_value": metric["observed_value"],
                    "raw_numerator": metric["raw_numerator"],
                    "raw_denominator": metric["raw_denominator"],
                    "unique_query_count": metric["unique_query_count"],
                    "semantic_coverage": metric["coverage"]["semantic"],
                    "metric_definition_hash": metric["metric_definition_hash"],
                    "contribution_set_hash": metric["contribution_set_hash"],
                    "query_contribution_set_hash": metric["query_contribution_set_hash"],
                    "design_contribution_set_hash": metric["design_contribution_set_hash"],
                }
                for metric in self.document["metrics"]
            ],
            "hashes": [
                {
                    "object_type": "snapshot_set",
                    "object_pub_id": "mss_report",
                    "content_hash": HASH,
                },
                *[
                    {
                        "object_type": object_type,
                        "object_pub_id": snapshot_pub_id,
                        "content_hash": HASH,
                    }
                    for snapshot_pub_id in ("msn_ready", "msn_experimental")
                    for object_type in (
                        "snapshot_hash",
                        "contribution_set_hash",
                        "query_contribution_set_hash",
                        "design_contribution_set_hash",
                    )
                ],
            ],
            "queries": [
                {
                    "snapshot_pub_id": "msn_ready",
                    "query_key": "q-1",
                    "query_text": "推荐安全公司",
                    "query_weight": 1,
                    "query_numerator": 1,
                    "query_denominator": 1,
                    "query_value": 1,
                    "unknown_weight": 0,
                    "reason_codes": ["included"],
                    "contribution_hash": HASH,
                }
            ],
            "answers": [],
            "decisions": [],
            "events": [],
            "exclusions": [],
            "design_cells": [],
        }


def _request(document: dict[str, Any]) -> FormalProductionRequest:
    binding = bind_metric_snapshot_set(
        document,
        expected_project_pub_id="prj_report",
        expected_set_pub_id="mss_report",
        expected_set_hash=HASH,
        expected_window_start=date(2026, 8, 1),
        expected_window_end=date(2026, 8, 27),
        expected_filters=FILTERS,
    )
    return FormalProductionRequest(
        pub_id="frp_report",
        tenant_pub_id="ten_report",
        project_pub_id="prj_report",
        services=(1,),
        window=FormalWindow(date(2026, 8, 1), date(2026, 8, 27)),
        document_status="internal_review",
        candidate_group_strategy="preregistered_scope_v1",
        frozen_at=datetime(2026, 8, 27, tzinfo=UTC),
        created_by_pub_id="usr_report",
        request_hash=HASH,
        document_governance={
            "version": "V1.0",
            "prepared_by": "分析师",
            "prepared_date": "2026-08-27",
        },
        service_catalog_version=QUOTATION_SERVICE_CATALOG,
        metric_snapshot_set_pub_id="mss_report",
        metric_snapshot_set_hash=HASH,
        metric_snapshot_filters=FILTERS,
        metric_snapshot_dependency_hash=binding.dependency_hash,
    )


def test_new_official_request_requires_an_exact_metric_snapshot_binding() -> None:
    body = {
        "project_pub_id": "prj_report",
        "services": [1],
        "service_catalog_version": QUOTATION_SERVICE_CATALOG,
        "window_start": "2026-08-01",
        "window_end": "2026-08-27",
        "prepared_by": "分析师",
        "prepared_date": "2026-08-27",
    }
    with pytest.raises(ValidationError):
        FormalProductionCreate.model_validate(body)
    parsed = FormalProductionCreate.model_validate(
        body
        | {
            "metric_snapshot_set_pub_id": "mss_report",
            "metric_snapshot_set_hash": HASH,
            "metric_snapshot_filters": FILTERS,
        }
    )
    assert parsed.root.metric_snapshot_set_pub_id == "mss_report"


def test_v2_fact_freeze_never_calls_a_legacy_fact_builder(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document = _snapshot_document()
    reader = SnapshotReader(document)
    service = FormalReportProductionService(
        dsn="postgresql://unit",
        evidence=SimpleNamespace(store=SimpleNamespace()),
        metric_snapshots=reader,
    )

    def forbidden(*args: object, **kwargs: object) -> object:
        del args, kwargs
        raise AssertionError("legacy project metric recalculation was called")

    monkeypatch.setattr(production, "build_formal_review_facts", forbidden)
    monkeypatch.setattr(production, "enrich_service1_v2_facts", forbidden)
    facts, _ = service._build_fact_bundle(_request(document))

    projection = facts[1]["metric_snapshot"]
    assert projection["metric_snapshot_set_pub_id"] == "mss_report"
    assert projection["metric_snapshot_set_hash"] == HASH
    assert facts[1]["metric_query_contributions"][0]["contribution_hash"] == HASH
    assert reader.exports == 1
    customer = customer_fact_snapshot(facts[1])
    assert customer["metric_snapshot"]["metric_snapshot_set_pub_id"] == "mss_report"
    assert customer["metric_snapshot"]["metrics"][0]["snapshot_pub_id"] == "msn_ready"


def test_snapshot_binding_fails_closed_on_window_filter_or_hash_drift() -> None:
    document = _snapshot_document()
    service = FormalReportProductionService(
        dsn="postgresql://unit",
        evidence=SimpleNamespace(store=SimpleNamespace()),
        metric_snapshots=SnapshotReader(document),
    )
    with pytest.raises(FormalProductionInvalid, match="metric_snapshot_window_mismatch"):
        service.validate_metric_snapshot_binding(
            tenant_pub_id="ten_report",
            project_pub_id="prj_report",
            window=FormalWindow(date(2026, 8, 2), date(2026, 8, 27)),
            snapshot_set_pub_id="mss_report",
            snapshot_set_hash=HASH,
            filters=FILTERS,
        )
    with pytest.raises(FormalProductionInvalid, match="metric_snapshot_filters_mismatch"):
        service.validate_metric_snapshot_binding(
            tenant_pub_id="ten_report",
            project_pub_id="prj_report",
            window=FormalWindow(date(2026, 8, 1), date(2026, 8, 27)),
            snapshot_set_pub_id="mss_report",
            snapshot_set_hash=HASH,
            filters={"model": ["other"], "region": [], "mode": []},
        )
    with pytest.raises(FormalProductionInvalid, match="metric_snapshot_set_hash_mismatch"):
        service.validate_metric_snapshot_binding(
            tenant_pub_id="ten_report",
            project_pub_id="prj_report",
            window=FormalWindow(date(2026, 8, 1), date(2026, 8, 27)),
            snapshot_set_pub_id="mss_report",
            snapshot_set_hash="b" * 64,
            filters=FILTERS,
        )


def test_unavailable_or_failed_snapshot_set_reports_not_ready() -> None:
    class MissingReader(SnapshotReader):
        def get_snapshot_set(self, *, tenant_pub_id: str, set_pub_id: str) -> dict[str, Any]:
            del tenant_pub_id, set_pub_id
            raise LookupError("missing")

    missing_service = FormalReportProductionService(
        dsn="postgresql://unit",
        evidence=SimpleNamespace(store=SimpleNamespace()),
        metric_snapshots=MissingReader(_snapshot_document()),
    )
    with pytest.raises(FormalProductionInvalid, match="metric_snapshot_set_not_ready"):
        missing_service.validate_metric_snapshot_binding(
            tenant_pub_id="ten_report",
            project_pub_id="prj_report",
            window=FormalWindow(date(2026, 8, 1), date(2026, 8, 27)),
            snapshot_set_pub_id="mss_report",
            snapshot_set_hash=HASH,
            filters=FILTERS,
        )

    failed_document = _snapshot_document()
    failed_document["state"] = "failed"
    failed_service = FormalReportProductionService(
        dsn="postgresql://unit",
        evidence=SimpleNamespace(store=SimpleNamespace()),
        metric_snapshots=SnapshotReader(failed_document),
    )
    with pytest.raises(FormalProductionInvalid, match="metric_snapshot_set_not_ready"):
        failed_service.validate_metric_snapshot_binding(
            tenant_pub_id="ten_report",
            project_pub_id="prj_report",
            window=FormalWindow(date(2026, 8, 1), date(2026, 8, 27)),
            snapshot_set_pub_id="mss_report",
            snapshot_set_hash=HASH,
            filters=FILTERS,
        )


def test_report_rejects_an_export_whose_member_hash_manifest_is_incomplete() -> None:
    document = _snapshot_document()

    class TamperedReader(SnapshotReader):
        def export_bundle(self, *, tenant_pub_id: str, set_pub_id: str) -> dict[str, Any]:
            export = super().export_bundle(
                tenant_pub_id=tenant_pub_id,
                set_pub_id=set_pub_id,
            )
            export["hashes"] = [
                row
                for row in export["hashes"]
                if not (
                    row["object_type"] == "snapshot_hash"
                    and row["object_pub_id"] == "msn_experimental"
                )
            ]
            return export

    service = FormalReportProductionService(
        dsn="postgresql://unit",
        evidence=SimpleNamespace(store=SimpleNamespace()),
        metric_snapshots=TamperedReader(document),
    )
    with pytest.raises(FormalProductionIncomplete, match="metric_snapshot_export_hashes_invalid"):
        service._build_fact_bundle(_request(document))


@pytest.mark.parametrize(
    ("tamper", "error_code"),
    [
        ("filters", "metric_snapshot_export_binding_invalid"),
        ("metric_value", "metric_snapshot_export_members_invalid"),
        ("cross_set_answer", "metric_snapshot_export_member_scope_invalid"),
    ],
)
def test_report_rejects_export_scope_or_member_drift(tamper: str, error_code: str) -> None:
    document = _snapshot_document()

    class TamperedReader(SnapshotReader):
        def export_bundle(self, *, tenant_pub_id: str, set_pub_id: str) -> dict[str, Any]:
            export = super().export_bundle(
                tenant_pub_id=tenant_pub_id,
                set_pub_id=set_pub_id,
            )
            if tamper == "filters":
                export["readme"][0]["filters"] = {
                    "model": ["other"],
                    "region": [],
                    "mode": [],
                }
            elif tamper == "metric_value":
                export["metrics"][0]["value"] = 0.75
            else:
                export["answers"] = [{"snapshot_pub_id": "msn_other"}]
            return export

    service = FormalReportProductionService(
        dsn="postgresql://unit",
        evidence=SimpleNamespace(store=SimpleNamespace()),
        metric_snapshots=TamperedReader(document),
    )
    with pytest.raises(FormalProductionIncomplete, match=error_code):
        service._build_fact_bundle(_request(document))


def test_docx_only_displays_frozen_values_and_trace_hashes() -> None:
    document = _snapshot_document()
    binding = bind_metric_snapshot_set(
        document,
        expected_project_pub_id="prj_report",
        expected_set_pub_id="mss_report",
        expected_set_hash=HASH,
        expected_window_start=date(2026, 8, 1),
        expected_window_end=date(2026, 8, 27),
        expected_filters=FILTERS,
    )
    payload = render_bound_metric_snapshot_docx(
        title="V2 指标报告",
        service_number=1,
        binding=binding,
        query_contributions=[],
        document_status="内部审核稿",
        governance={"prepared_by": "分析师"},
    )
    rendered = Document(BytesIO(payload))
    text = "\n".join(
        [paragraph.text for paragraph in rendered.paragraphs]
        + [cell.text for table in rendered.tables for row in table.rows for cell in row.cells]
    )
    assert "mss_report" in text
    assert HASH in text
    assert "msn_ready" in text
    assert "中性 AI 推荐自然提及率" in text
    assert "查询数" in text
    assert "不形成正式数值（experimental）" in text
    assert "仅格式化上述不可变 Metrics V2 快照" in text
