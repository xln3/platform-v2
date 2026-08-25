from __future__ import annotations

import json
from contextlib import contextmanager
from datetime import UTC, date, datetime
from hashlib import sha256
from io import BytesIO
from typing import Any
from zipfile import ZipFile

import pytest
from docx import Document
from geo_platform.reports import formal_review_service2_source_corpus as builder
from PIL import Image

from domain.reporting.formal_review_service2_source_corpus_docx import (
    render_service2_source_corpus_docx,
)


def _manifest_facts() -> dict[str, Any]:
    return {
        "schema_version": "formal-service2-source-corpus-v2",
        "scope": {
            "project_pub_id": "prj_service2",
            "batch_pub_id": "s2b_service2",
            "run_pub_ids": ["run_a", "run_b"],
        },
        "coverage": {
            "expected_occurrences": 5,
            "materialized_items": 5,
            "distinct_urls": 3,
            "processing_states": {"processed": 4, "blocked": 1},
            "fetch_states": {"succeeded": 4, "blocked": 1},
            "entered_judgment": 4,
            "findings": 3,
            "reviewed_findings": 3,
            "eligible_cases": 2,
            "coverage_complete": True,
            "query_outcomes_complete": True,
            "query_coverage_complete": True,
        },
        "cases": [
            {
                "finding_pub_id": "s2f_case",
                "level": "L2b",
                "is_disparagement": True,
                "relation_direction": "target_degraded_peer_elevated",
                "textual_speaker": "页面叙述者",
                "target_entity": "目标品牌",
                "evidence_quote": "目标品牌被无锚点地置于次级位置。",
                "canonical_url": "https://example.com/case",
                "snapshot_pub_id": "snp_case",
                "snapshot_text_sha256": "a" * 64,
                "visual_evidence_pub_id": "evd_case",
                "visual_evidence_sha256": "c" * 64,
                "visual_page_snapshot_evidence_pub_id": "evd_page",
                "visual_page_snapshot_sha256": "d" * 64,
                "factcheck_verdict": "unverifiable",
                "factcheck_evidence": [],
                "factcheck_boundary": "当前公开材料不足，不能判断该陈述真假。",
                "publisher_attribution": {
                    "party": None,
                    "confidence": "unknown",
                    "evidence": [],
                },
                "commissioner_attribution": {
                    "party": None,
                    "confidence": "unknown",
                    "evidence": [],
                },
            }
        ],
        "evidence_pub_ids": ["evd_case", "evd_page"],
        "rendering_boundary": "frozen_facts_only_no_network_or_model",
    }


def _canonical_hash(value: object) -> str:
    payload = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return sha256(payload.encode()).hexdigest()


class _Result:
    def __init__(
        self,
        row: dict[str, Any] | None,
        rows: list[dict[str, Any]] | None = None,
    ) -> None:
        self.row = row
        self.rows = rows or []

    def fetchone(self) -> dict[str, Any] | None:
        return self.row

    def fetchall(self) -> list[dict[str, Any]]:
        return self.rows


class _Connection:
    def __init__(
        self,
        row: dict[str, Any] | None,
        asset_rows: list[dict[str, Any]] | None = None,
    ) -> None:
        self.row = row
        self.asset_rows = asset_rows or []
        self.statements: list[str] = []

    def execute(self, statement: str, _params: object) -> _Result:
        self.statements.append(statement)
        if "FROM evidence.evidence_asset" in statement:
            return _Result(None, self.asset_rows)
        return _Result(self.row)


def _document_text(payload: bytes) -> str:
    document = Document(BytesIO(payload))
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        values.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(values)


def _png() -> bytes:
    payload = BytesIO()
    Image.new("RGB", (32, 20), color=(255, 255, 255)).save(payload, format="PNG")
    return payload.getvalue()


def test_builder_reads_one_hash_verified_manifest_and_discloses_blocked_coverage(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    facts = _manifest_facts()
    row = {
        "manifest_pub_id": "s2m_service2",
        "revision": 1,
        "manifest_hash": _canonical_hash(facts),
        "facts": facts,
        "case_count": 1,
        "evidence_reference_count": 1,
        "created_at": datetime(2026, 8, 24, tzinfo=UTC),
        "batch_pub_id": "s2b_service2",
        "corpus_policy_version": "service2-all-u-occurrence-v1",
        "judgment_policy_version": "service2-entity-relation-v1",
        "project_name": "全 U 核查项目",
        "target_brand": "目标品牌",
    }
    connection = _Connection(
        row,
        [
            {
                "pub_id": "evd_case",
                "object_key": "cas/visual-case",
                "sha256": "c" * 64,
                "mime_type": "image/png",
                "kind": "service2_exact_quote_screenshot",
            },
            {
                "pub_id": "evd_page",
                "object_key": "cas/visual-page",
                "sha256": "d" * 64,
                "mime_type": "text/html",
                "kind": "service2_visual_page_snapshot",
            },
        ],
    )

    @contextmanager
    def fake_connection(*_args: object, **_kwargs: object) -> Any:
        yield connection

    monkeypatch.setattr(builder, "tenant_connection", fake_connection)
    result = builder.build_service2_source_corpus_facts(
        dsn="postgresql://unused",
        tenant_pub_id="tnt_service2",
        project_pub_id="prj_service2",
        start=date(2026, 8, 1),
        end=date(2026, 8, 24),
        generated_at=datetime(2026, 8, 24, tzinfo=UTC),
        manifest_pub_id="s2m_service2",
        manifest_hash=_canonical_hash(facts),
    )

    assert result["coverage"]["expected_occurrences"] == 5
    assert result["coverage"]["distinct_urls"] == 3
    assert result["evidence_gate"] == {
        "status": "insufficient",
        "reasons": ["source_or_evidence_coverage_incomplete"],
        "incomplete_processing_states": {"blocked": 1},
    }
    assert result["rendering_boundary"] == "frozen_facts_only_no_network_or_model"
    assert "service2_fact_manifest" in connection.statements[0]
    assert "AT TIME ZONE 'Asia/Shanghai'" in connection.statements[0]


@pytest.mark.parametrize(
    ("coverage_changes", "expected_status", "expected_reason"),
    [
        (
            {"processing_states": {"processed": 4, "partial": 1}},
            "insufficient",
            "source_or_evidence_coverage_incomplete",
        ),
        (
            {"query_coverage_complete": False, "failed_queries": 1},
            "insufficient",
            "failed_queries_require_retry",
        ),
        (
            {
                "expected_occurrences": 0,
                "materialized_items": 0,
                "processing_states": {"processed": 0},
            },
            "ready",
            None,
        ),
    ],
)
def test_formal_gate_blocks_partial_and_failed_queries_but_allows_complete_zero_results(
    monkeypatch: pytest.MonkeyPatch,
    coverage_changes: dict[str, object],
    expected_status: str,
    expected_reason: str | None,
) -> None:
    facts = _manifest_facts()
    facts["cases"] = []
    facts["evidence_pub_ids"] = []
    facts["coverage"] = {
        **facts["coverage"],
        "processing_states": {"processed": 5},
        **coverage_changes,
    }
    row = {
        "manifest_pub_id": "s2m_complete_gate",
        "revision": 1,
        "manifest_hash": _canonical_hash(facts),
        "facts": facts,
        "case_count": 0,
        "evidence_reference_count": 0,
        "created_at": datetime(2026, 8, 24, tzinfo=UTC),
        "batch_pub_id": "s2b_complete_gate",
        "corpus_policy_version": "service2-all-u-occurrence-v1",
        "judgment_policy_version": "service2-entity-relation-v1",
        "project_name": "全 U 核查项目",
        "target_brand": "目标品牌",
    }
    connection = _Connection(row)

    @contextmanager
    def fake_connection(*_args: object, **_kwargs: object) -> Any:
        yield connection

    monkeypatch.setattr(builder, "tenant_connection", fake_connection)
    result = builder.build_service2_source_corpus_facts(
        dsn="postgresql://unused",
        tenant_pub_id="tnt_service2",
        project_pub_id="prj_service2",
        start=date(2026, 8, 1),
        end=date(2026, 8, 24),
        generated_at=datetime(2026, 8, 24, tzinfo=UTC),
        manifest_pub_id="s2m_complete_gate",
        manifest_hash=_canonical_hash(facts),
    )

    assert result["evidence_gate"]["status"] == expected_status
    if expected_reason is None:
        assert result["evidence_gate"]["reasons"] == []
    else:
        assert expected_reason in result["evidence_gate"]["reasons"]


def test_builder_rejects_manifest_hash_mismatch(monkeypatch: pytest.MonkeyPatch) -> None:
    facts = _manifest_facts()
    connection = _Connection(
        {
            "manifest_pub_id": "s2m_service2",
            "revision": 1,
            "manifest_hash": "0" * 64,
            "facts": facts,
            "case_count": 1,
            "evidence_reference_count": 1,
            "created_at": datetime(2026, 8, 24, tzinfo=UTC),
            "batch_pub_id": "s2b_service2",
            "corpus_policy_version": "service2-all-u-occurrence-v1",
            "judgment_policy_version": "service2-entity-relation-v1",
            "project_name": "全 U 核查项目",
            "target_brand": "目标品牌",
        }
    )

    @contextmanager
    def fake_connection(*_args: object, **_kwargs: object) -> Any:
        yield connection

    monkeypatch.setattr(builder, "tenant_connection", fake_connection)
    with pytest.raises(ValueError, match="service2_frozen_manifest_integrity_failed"):
        builder.build_service2_source_corpus_facts(
            dsn="postgresql://unused",
            tenant_pub_id="tnt_service2",
            project_pub_id="prj_service2",
            start=date(2026, 8, 1),
            end=date(2026, 8, 24),
            generated_at=datetime(2026, 8, 24, tzinfo=UTC),
            manifest_pub_id="s2m_service2",
            manifest_hash="0" * 64,
        )


def test_retry_never_switches_from_bound_manifest_to_newer_same_window_manifest(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    b1_facts = _manifest_facts()
    b2_facts = {
        **_manifest_facts(),
        "scope": {**_manifest_facts()["scope"], "batch_pub_id": "s2b_newer"},
    }
    manifests = {
        "s2m_bound_b1": {
            "manifest_pub_id": "s2m_bound_b1",
            "revision": 1,
            "manifest_hash": _canonical_hash(b1_facts),
            "facts": b1_facts,
            "case_count": 1,
            "evidence_reference_count": 2,
            "created_at": datetime(2026, 8, 24, 1, tzinfo=UTC),
            "batch_pub_id": "s2b_service2",
            "corpus_policy_version": "service2-all-u-occurrence-v1",
            "judgment_policy_version": "service2-entity-relation-v1",
            "project_name": "全 U 核查项目",
            "target_brand": "目标品牌",
        },
        "s2m_newer_b2": {
            "manifest_pub_id": "s2m_newer_b2",
            "revision": 2,
            "manifest_hash": _canonical_hash(b2_facts),
            "facts": b2_facts,
            "case_count": 1,
            "evidence_reference_count": 2,
            "created_at": datetime(2026, 8, 24, 2, tzinfo=UTC),
            "batch_pub_id": "s2b_newer",
            "corpus_policy_version": "service2-all-u-occurrence-v1",
            "judgment_policy_version": "service2-entity-relation-v1",
            "project_name": "全 U 核查项目",
            "target_brand": "目标品牌",
        },
    }
    assets = [
        {
            "pub_id": "evd_case",
            "object_key": "cas/visual-case",
            "sha256": "c" * 64,
            "mime_type": "image/png",
            "kind": "service2_exact_quote_screenshot",
        },
        {
            "pub_id": "evd_page",
            "object_key": "cas/visual-page",
            "sha256": "d" * 64,
            "mime_type": "text/html",
            "kind": "service2_visual_page_snapshot",
        },
    ]

    class CompetitiveConnection:
        newest_manifest = "s2m_bound_b1"
        fail_first_asset_read = True
        selected_manifests: list[str] = []
        manifest_sql: list[str] = []

        @classmethod
        def execute(cls, statement: str, params: object) -> _Result:
            if "FROM evidence.evidence_asset" in statement:
                if cls.fail_first_asset_read:
                    cls.fail_first_asset_read = False
                    raise RuntimeError("injected first-attempt failure")
                return _Result(None, assets)
            cls.manifest_sql.append(statement)
            values = tuple(params)  # type: ignore[arg-type]
            # A regressed "latest manifest" query would have no exact identity and
            # therefore receive B2 after the first attempt.  The fixed query always
            # chooses the caller-bound identity at positions 1 and 2.
            selected = (
                str(values[1])
                if len(values) >= 3 and values[1] in manifests
                else cls.newest_manifest
            )
            cls.selected_manifests.append(selected)
            return _Result(manifests[selected])

    @contextmanager
    def fake_connection(*_args: object, **_kwargs: object) -> Any:
        yield CompetitiveConnection()

    monkeypatch.setattr(builder, "tenant_connection", fake_connection)
    kwargs = {
        "dsn": "postgresql://unused",
        "tenant_pub_id": "tnt_service2",
        "project_pub_id": "prj_service2",
        "start": date(2026, 8, 1),
        "end": date(2026, 8, 24),
        "generated_at": datetime(2026, 8, 24, tzinfo=UTC),
        "manifest_pub_id": "s2m_bound_b1",
        "manifest_hash": _canonical_hash(b1_facts),
    }

    with pytest.raises(RuntimeError, match="injected first-attempt failure"):
        builder.build_service2_source_corpus_facts(**kwargs)
    CompetitiveConnection.newest_manifest = "s2m_newer_b2"
    retried = builder.build_service2_source_corpus_facts(**kwargs)

    assert retried["manifest"]["manifest_pub_id"] == "s2m_bound_b1"
    assert retried["manifest"]["manifest_hash"] == _canonical_hash(b1_facts)
    assert CompetitiveConnection.selected_manifests == ["s2m_bound_b1", "s2m_bound_b1"]
    assert all("manifest.pub_id=%s" in sql for sql in CompetitiveConnection.manifest_sql)
    assert all("manifest.manifest_hash=%s" in sql for sql in CompetitiveConnection.manifest_sql)
    assert all("ORDER BY manifest." not in sql for sql in CompetitiveConnection.manifest_sql)


def test_docx_keeps_occurrence_and_url_counts_separate_and_unknown_attribution() -> None:
    facts = {
        **_manifest_facts(),
        "project_name": "全 U 核查项目",
        "target_brand": "目标品牌",
        "generated_at": datetime(2026, 8, 24, tzinfo=UTC),
        "window": {"start": "2026-08-01", "end": "2026-08-24"},
        "document_status": "internal_review",
        "document_governance": {
            "version": "V1.0",
            "prepared_by": "GEO 项目组",
            "prepared_date": "2026-08-24",
        },
        "manifest": {
            "batch_pub_id": "s2b_service2",
            "manifest_pub_id": "s2m_service2",
            "revision": 1,
            "manifest_hash": "b" * 64,
            "corpus_policy_version": "service2-all-u-occurrence-v1",
            "judgment_policy_version": "service2-entity-relation-v1",
        },
        "evidence_gate": {
            "status": "insufficient",
            "reasons": ["source_or_evidence_coverage_incomplete"],
            "incomplete_processing_states": {"blocked": 1},
        },
        "limitations": [
            "入池总体为全部 U occurrence；URL 抓取复用不缩小分母。",
            "publisher/commissioner 归属为 unknown 时不作委托归因。",
        ],
    }

    payload = render_service2_source_corpus_docx(
        facts,
        visual_screenshots={"s2f_case": _png()},
    )
    text = _document_text(payload)

    assert "主动拉踩内容核查报告" in text
    assert "全部 U 信源帖子实体—关系核查" in text
    assert "U occurrence\n5\n冻结范围总体" in text
    assert "distinct URL\n3\n仅用于抓取复用" in text
    assert "unknown（未作归因）" in text
    assert "当前公开材料不足，不能判断该陈述真假。" in text
    assert "service2-all-u-occurrence-v1" in text
    assert "b" * 64 in text
    assert "本稿不得被表述为完整无风险结论" in text
    with ZipFile(BytesIO(payload)) as archive:
        assert any(name.startswith("word/media/") for name in archive.namelist())
