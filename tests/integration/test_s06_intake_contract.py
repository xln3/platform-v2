"""S06 合同 docx 生成端点：模板填槽确定性 + unfilled 披露 + 租户隔离。

断言口径：
  * 全字段 profile+promo+trigger → 关键槽全部落值（甲方名/信用代码/Query/附件二条目）、
    真实性确认五条 □→☑、乙方栏与签署区不被误填（无错位）；
  * 空 profile → 不崩、空槽留模板、unfilled 计数=26（20 未填写 + 6 无数据源）；
  * 越租户 → 404。
"""

import io
import secrets
from collections.abc import Iterator
from typing import Any

import docx
import pytest
from fastapi.testclient import TestClient
from geo_platform.intake import models as intake_models
from geo_platform.main import app
from geo_platform.tenancy.database import SessionLocal
from geo_platform.tenancy.repository import TenantRepository
from sqlalchemy import select

_LICENSE_CODE = "91310000MA1FL0000A"  # 18 位 [0-9A-Z]
_CUSTOMER = "测试甲方保险科技有限公司"


def _bootstrap(client: TestClient, subject: str) -> tuple[str, dict[str, str]]:
    response = client.post(
        "/api/v2/identity/bootstrap",
        headers={"X-Bootstrap-Secret": "development-bootstrap"},
        json={"tenant_name": subject, "subject": subject, "display_name": "Admin"},
    )
    assert response.status_code == 201, response.text
    tenant = str(response.json()["tenant_pub_id"])
    return tenant, {
        "X-Tenant-Id": tenant,
        "X-Actor-Id": subject,
        "X-Actor-Role": "admin",
    }


def _create_member(
    client: TestClient, admin_headers: dict[str, str], tenant: str, role: str
) -> dict[str, str]:
    subject = f"{role}-" + secrets.token_hex(8)
    response = client.post(
        "/api/v2/identity/members",
        headers={**admin_headers, "Idempotency-Key": "member-" + secrets.token_hex(16)},
        json={"subject": subject, "display_name": role.title(), "role": role},
    )
    assert response.status_code == 201, response.text
    return {"X-Tenant-Id": tenant, "X-Actor-Id": subject, "X-Actor-Role": role}


def _idem() -> str:
    return "idem-" + secrets.token_hex(16)


def _create_project(client: TestClient, headers: dict[str, str]) -> str:
    response = client.post(
        "/api/v2/projects",
        headers={**headers, "Idempotency-Key": "project-" + secrets.token_hex(16)},
        json={"name": "合同测试项目", "customer_name": _CUSTOMER},
    )
    assert response.status_code == 201, response.text
    return str(response.json()["pub_id"])


@pytest.fixture()
def contract_env() -> Iterator[tuple[TestClient, str, dict[str, str], str]]:
    client = TestClient(app)
    tenant, admin_headers = _bootstrap(client, "contract-admin-" + secrets.token_hex(6))
    customer_headers = _create_member(client, admin_headers, tenant, "customer")
    project = _create_project(client, admin_headers)
    yield client, tenant, customer_headers, project


def _full_setup(client: TestClient, headers: dict[str, str], tenant: str, project: str) -> None:
    base = f"/api/v2/projects/{project}/intake"
    put = client.put(
        f"{base}/profile",
        headers=headers,
        json={
            "contact_person": "张三",
            "contact_info": "13800000000",
            "business_license_code": _LICENSE_CODE,
            "review_category": "A",
            "pre_review_required": True,
            "ad_review_no": "京药广审（文）第2024010001号",
            "ad_review_authority": "北京市药品监督管理局",
            "ad_review_expiry": "2027-12-31",
            "ad_review_doc_types": ["药品广告审查批准文号"],
            "goals": ["提升AI搜索曝光", "获取销售线索"],
            "audience_type": ["B2C个人消费者"],
            "platforms": ["豆包", "DeepSeek"],
            "regions": ["全国"],
            "selling_points": "百年品牌，OTC 品类市占率第一（中康数据 2025）",
            "trademarks": ["商标注册证第12345678号"],
            "licenses": [
                {"name": "药品生产许可证", "number": "京20160001", "expiry": "2027-01-01"}
            ],
            "evidence_links": ["https://example.com/test-report"],
            "truth_confirmed": True,
            "filler_name": "李四",
        },
    )
    assert put.status_code == 200, put.text
    promo = client.post(
        f"{base}/promos",
        headers={**headers, "Idempotency-Key": _idem()},
        json={
            "kind": "product",
            "payload": {"name": "清风感冒灵颗粒", "desc": "OTC 甲类，家庭常备感冒药"},
        },
    )
    assert promo.status_code == 201, promo.text
    trig = client.post(
        f"{base}/trigger-questions",
        headers={**headers, "Idempotency-Key": _idem()},
        json={"text": "感冒药哪个牌子见效快\n风寒感冒吃什么药\n感冒灵颗粒怎么选"},
    )
    assert trig.status_code == 201, trig.text
    # 把第三条翻成 claim_created（draft+claim_created 都应进合同）
    frozen_pub = str(trig.json()["items"][2]["pub_id"])
    with SessionLocal() as session:
        repository = TenantRepository(session, tenant)
        row = session.scalar(
            select(intake_models.IntakeTriggerQuestion).where(
                intake_models.IntakeTriggerQuestion.tenant_id == repository.tenant.id,
                intake_models.IntakeTriggerQuestion.pub_id == frozen_pub,
            )
        )
        assert row is not None
        row.status = "claim_created"
        session.commit()


def _all_text(document: Any) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_contract_docx_full_profile_fills_slots(
    contract_env: tuple[TestClient, str, dict[str, str], str],
) -> None:
    client, tenant, customer, project = contract_env
    _full_setup(client, customer, tenant, project)

    response = client.get(f"/api/v2/projects/{project}/intake/contract.docx", headers=customer)
    assert response.status_code == 200, response.text
    assert response.headers["Content-Type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "filename*=UTF-8''" in response.headers["Content-Disposition"]
    # 全字段后只剩 6 个无数据源槽（注册/通信地址、法定代表人、邮箱、所属行业、合同编号）
    assert response.headers["X-Contract-Unfilled-Count"] == "6"

    document = docx.Document(io.BytesIO(response.content))
    text = _all_text(document)

    # 甲方头部
    assert f"甲方（委托方）：{_CUSTOMER}" in text
    assert f"统一社会信用代码：{_LICENSE_CODE}" in text
    assert "联系人：张三" in text
    assert "联系电话：13800000000" in text
    # 无数据源槽留空
    assert "注册地址：" in text and f"注册地址：{_CUSTOMER}" not in text
    # 第一条验证对象
    assert "【清风感冒灵颗粒】" in text and "【】" not in text
    # Query（draft + claim_created 全部收录）
    assert "感冒药哪个牌子见效快" in text
    assert "风寒感冒吃什么药" in text
    assert "感冒灵颗粒怎么选" in text
    # 附件二表一
    assert "A类·法定前置审查" in text
    assert "京药广审（文）第2024010001号" in text
    assert "北京市药品监督管理局" in text
    assert "提升AI搜索曝光、获取销售线索" in text
    assert "豆包、DeepSeek" in text
    assert "百年品牌，OTC 品类市占率第一" in text
    assert "https://example.com/test-report" in text
    # 附件二表二
    assert "药品生产许可证" in text and "京20160001" in text
    assert "商标注册证第12345678号" in text
    assert "药品广告审查批准文号" in text
    # 真实性确认五条全勾
    assert text.count("☑") == 5
    # 填表人
    assert "填表人：李四" in text
    # 附件三甲方名称槽已填
    assert "【甲方公司名称】" not in text
    assert f"出具方：{_CUSTOMER}" in text
    assert f"承诺方（甲方）：{_CUSTOMER}（盖章）" in text
    # 合同编号无数据源 → 留空槽
    assert "合同编号：【 】" in text

    # 无错位：乙方栏与签署区保持模板空槽
    paragraphs = [p.text for p in document.paragraphs]
    b_idx = next(i for i, t in enumerate(paragraphs) if t.strip().startswith("乙方（服务方）"))
    b_block = paragraphs[b_idx : b_idx + 9]
    assert "联系人：" in b_block  # 乙方联系人仍为纯标签空槽
    assert not any("张三" in t for t in b_block)
    assert not any(_LICENSE_CODE in t for t in b_block)

    # 审计落库
    with SessionLocal() as session:
        repository = TenantRepository(session, tenant)
        from geo_platform.tenancy.models import AuditLog

        rows = session.scalars(
            select(AuditLog).where(
                AuditLog.tenant_id == repository.tenant.id,
                AuditLog.resource_type == "intake_contract",
            )
        ).all()
        assert len(rows) == 1
        assert rows[0].action.startswith("intake_contract.exported:")


def test_contract_docx_empty_profile_keeps_blank_slots(
    contract_env: tuple[TestClient, str, dict[str, str], str],
) -> None:
    client, _, customer, project = contract_env

    response = client.get(f"/api/v2/projects/{project}/intake/contract.docx", headers=customer)
    assert response.status_code == 200, response.text
    # 20 个有数据源但空值 + 6 个无数据源
    assert response.headers["X-Contract-Unfilled-Count"] == "26"

    document = docx.Document(io.BytesIO(response.content))
    text = _all_text(document)
    # 甲方名始终可填（Customer/Project 名称）
    assert f"甲方（委托方）：{_CUSTOMER}" in text
    # 关键空槽留模板、不崩不编造
    assert "统一社会信用代码：" in text
    assert "【】" in text  # 第一条验证对象空槽保留
    assert "【甲方公司名称】" not in text  # 附件三甲方名已填
    assert "合同编号：【 】" in text
    # 真实性确认未勾选（truth_confirmed 空 → 如实留 □）
    assert "☑" not in text
    assert "□ 本表所填信息及所附材料真实" in text
    # 附件二空值「未填写」口径
    assert "未填写" in text


def test_contract_docx_cross_tenant_404(
    contract_env: tuple[TestClient, str, dict[str, str], str],
) -> None:
    client, _, customer, project = contract_env
    _, other_admin = _bootstrap(client, "contract-other-" + secrets.token_hex(6))
    other_customer = _create_member(client, other_admin, other_admin["X-Tenant-Id"], "customer")
    response = client.get(
        f"/api/v2/projects/{project}/intake/contract.docx", headers=other_customer
    )
    assert response.status_code == 404
    # 本租户正常
    assert (
        client.get(f"/api/v2/projects/{project}/intake/contract.docx", headers=customer).status_code
        == 200
    )
