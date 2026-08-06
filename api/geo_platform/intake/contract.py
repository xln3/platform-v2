"""GEO 技术验证（MVP）服务合同 docx 渲染（模板填槽，确定性定位 + unfilled 披露）。

模板底本=assets/contract-template.docx（《GEO技术验证MVP服务合同》清洁交付 1 版，
14 条主合同 + 附件一须持证行业勾选 + 附件二客户信息收集表 + 附件三合规承诺函）。

填槽纪律（INV-32 零合成）：
  * 只写真库值（profile/promo/trigger/project/customer）；profile 无对应字段的槽
    （注册地址/通信地址/法定代表人/联系邮箱/所属行业/合同编号）一律留模板空槽并记
    unfilled「无数据源」，绝不编造；
  * 槽位一律按段落文本/run 文本/表格坐标+行标签双重确定性定位；定位不到 →
    unfilled「模板未定位」，绝不猜测错位填充；
  * 签署时才填的乙方栏/收款账户/盖章签字栏/填表日期不属于数据槽，不填也不计 unfilled；
  * Query 无独立枚举槽（模板第一条只有「5 个 Query（双方书面确认）」条款文字），
    全部收录问法（draft+claim_created，上限 _QUERY_CAP 条）填入附件二「期望的用户提问场景」格。
"""

import uuid
from pathlib import Path
from typing import Any

from sqlalchemy.orm import Session

from ..projects.models import Customer, Project
from . import models, service

_TEMPLATE = Path(__file__).resolve().parent / "assets" / "contract-template.docx"
# Query 上限：模板无 Query 枚举槽（见模块 docstring），取 intake 收录上限 20 条封顶。
_QUERY_CAP = 20

_REASON_NO_SOURCE = "无数据源"
_REASON_EMPTY = "未填写"
_REASON_NOT_FOUND = "模板未定位"
_REASON_NO_PROMO = "无推广内容"

_EMPTY = "未填写"

# 附件二表一（宣传内容与目标，13 行）行标签（确定性坐标校验用，cell0 前缀匹配）。
_TABLE1_LABELS = (
    "公司 / 品牌名称",
    "所属行业",
    "行业广告审查分类",
    "是否属于法定前置审查行业",
    "广告审查批准文号",
    "联系人及联系方式",
    "拟推广产品",
    "推广目标",
    "期望的用户提问场景",
    "目标 AI 平台",
    "重点地域",
    "核心卖点",
    "可公开引用的佐证材料",
)
# 附件二表二（资质，4 行）。
_TABLE2_LABELS = ("营业执照", "行业许可证", "商标", "广告审查批准文件")


# ── 底层填槽助手（全部确定性定位，失败只记 unfilled 不乱写）─────────────────────
def _para_texts(doc: Any) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs]


def _find_para(doc: Any, predicate: Any) -> Any | None:
    for p in doc.paragraphs:
        if predicate(p.text.strip()):
            return p
    return None


def _replace_in_runs(paragraph: Any, old: str, new: str, *, count: int = -1) -> bool:
    """在 paragraph 的 run 序列里替换 old（占位符都在单 run 内，模板分析已确认）。"""
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, count if count >= 0 else -1)
            return True
    return False


def _fill_header_slot(
    doc: Any,
    start: int,
    end: int,
    label: str,
    value: str | None,
    slot: str,
    unfilled: list[str],
) -> None:
    """甲方头部「标签：」段（标签独占一段，值追加为新 run）。scope=[start,end) 段落下标。"""
    value = (value or "").strip()
    for p in doc.paragraphs[start:end]:
        if p.text.strip() == label:
            if value:
                p.add_run(value)
            else:
                unfilled.append(f"{slot}（{_REASON_EMPTY}）")
            return
    unfilled.append(f"{slot}（{_REASON_NOT_FOUND}）")


def _note_no_source(
    doc: Any, start: int, end: int, label: str, slot: str, unfilled: list[str]
) -> None:
    """无数据源的槽：确认模板槽位存在后记「无数据源」，不存在记「模板未定位」；绝不填值。"""
    for p in doc.paragraphs[start:end]:
        if p.text.strip() == label:
            unfilled.append(f"{slot}（{_REASON_NO_SOURCE}）")
            return
    unfilled.append(f"{slot}（{_REASON_NOT_FOUND}）")


def _set_cell(cell: Any, text: str) -> None:
    """整格重写为填表值（保留首段首 run 字体格式；多行用软换行）。"""
    para = cell.paragraphs[0]
    lines = text.split("\n")
    if para.runs:
        keep = para.runs[0]
        keep.text = lines[0]
        for run in para.runs[1:]:
            run._element.getparent().remove(run._element)  # noqa: SLF001
        target = keep
    else:
        target = para.add_run(lines[0])
    for line in lines[1:]:
        target.add_break()
        target.add_text(line)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)  # noqa: SLF001


def _fill_table_row(
    table: Any,
    row: int,
    label: str,
    value: str | None,
    slot: str,
    unfilled: list[str],
    *,
    reason: str = _REASON_EMPTY,
) -> None:
    """按行坐标+行标签双重校验后填值；标签不符 → 模板未定位，绝不错位写。"""
    cell0 = table.rows[row].cells[0].text.strip()
    if not cell0.startswith(label):
        unfilled.append(f"{slot}（{_REASON_NOT_FOUND}）")
        return
    value = (value or "").strip()
    if not value:
        unfilled.append(f"{slot}（{reason}）")
        value = _EMPTY
    _set_cell(table.rows[row].cells[1], value)


def _table_by_header(doc: Any, header: str) -> Any | None:
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip().startswith(header):
            return table
    return None


# ── 值组装（与 service.build_profile_docx 同源口径）─────────────────────────────
def _yn(value: bool | None) -> str | None:
    if value is None:
        return None
    return "是" if value else "否"


def _promo_lines(promos: list[models.IntakePromo]) -> str:
    lines = []
    for promo in promos:
        d = promo.payload or {}
        name = str(d.get("name") or "").strip() or (
            "产品 / 服务" if promo.kind == "product" else "公司 / 品牌"
        )
        brief = str(d.get("desc") or d.get("advantage") or "").strip()
        lines.append(f"{name}：{brief}" if brief else name)
    return "\n".join(lines)


def _license_lines(licenses: list[dict[str, str]]) -> str:
    lines = []
    for lic in licenses:
        parts = []
        if (lic.get("name") or "").strip():
            parts.append(f"证照名称：{lic['name'].strip()}")
        if (lic.get("number") or "").strip():
            parts.append(f"编号：{lic['number'].strip()}")
        if (lic.get("expiry") or "").strip():
            parts.append(f"有效期至：{lic['expiry'].strip()}")
        if parts:
            lines.append("　".join(parts))
    return "\n".join(lines)


def _ad_review_line(profile: models.IntakeProfile) -> str | None:
    no = (profile.ad_review_no or "").strip()
    auth = (profile.ad_review_authority or "").strip()
    exp = (profile.ad_review_expiry or "").strip()
    if not (no or auth or exp):
        return None
    return f"批准文号：{no or '—'}　审查机关：{auth or '—'}　有效期至：{exp or '—'}（附扫描件）"


# ── 主渲染 ─────────────────────────────────────────────────────────────────
def render_contract_docx(
    session: Session,
    *,
    tenant_id: uuid.UUID,
    project: Project,
) -> tuple[bytes, list[str]]:
    """以模板填槽生成合同 docx。返回 (bytes, unfilled)；unfilled=「槽位（原因）」清单，
    覆盖：无数据源 / 有数据源但未填写 / 无推广内容 / 模板未定位 四类。"""
    import docx  # 延迟导入：重依赖，端点按需加载（同 service.build_profile_docx）

    profile = service.get_profile(session, tenant_id=tenant_id, project_id=project.id)
    promos = service.list_promos(session, tenant_id=tenant_id, project_id=project.id)
    triggers = service.list_triggers(session, tenant_id=tenant_id, project_id=project.id)[
        :_QUERY_CAP
    ]
    customer = session.get(Customer, project.customer_id)
    company = (
        customer.name if customer is not None and customer.tenant_id == tenant_id else project.name
    )

    doc = docx.Document(str(_TEMPLATE))
    unfilled: list[str] = []

    # ── 主合同头部：甲方信息栏（段落 scope=「甲方（委托方）：」..「乙方（服务方）：」）──
    texts = _para_texts(doc)

    def _is_a_head(t: str) -> bool:
        return t.startswith("甲方（委托方）")

    def _is_b_head(t: str) -> bool:
        return t.startswith("乙方（服务方）")

    a_idx = next((i for i, t in enumerate(texts) if _is_a_head(t)), None)
    b_idx = next((i for i, t in enumerate(texts) if _is_b_head(t)), None)
    if a_idx is None or b_idx is None or b_idx <= a_idx:
        for slot in (
            "甲方名称",
            "统一社会信用代码",
            "注册地址",
            "通信地址",
            "法定代表人",
            "联系人",
            "联系电话",
            "联系邮箱",
        ):
            unfilled.append(f"{slot}（{_REASON_NOT_FOUND}）")
    else:
        _fill_header_slot(doc, a_idx, b_idx, "甲方（委托方）：", company, "甲方名称", unfilled)
        _fill_header_slot(
            doc,
            a_idx,
            b_idx,
            "统一社会信用代码：",
            profile.business_license_code if profile else None,
            "统一社会信用代码",
            unfilled,
        )
        # profile 无地址/法定代表人/邮箱字段 → 留空槽 + 无数据源披露（INV-32，绝不编造）
        _note_no_source(doc, a_idx, b_idx, "注册地址：", "甲方注册地址", unfilled)
        _note_no_source(doc, a_idx, b_idx, "通信地址：", "甲方通信地址", unfilled)
        _note_no_source(doc, a_idx, b_idx, "法定代表人：", "甲方法定代表人", unfilled)
        _fill_header_slot(
            doc,
            a_idx,
            b_idx,
            "联系人：",
            profile.contact_person if profile else None,
            "联系人",
            unfilled,
        )
        _fill_header_slot(
            doc,
            a_idx,
            b_idx,
            "联系电话：",
            profile.contact_info if profile else None,
            "联系电话",
            unfilled,
        )
        _note_no_source(doc, a_idx, b_idx, "联系邮箱：", "甲方联系邮箱", unfilled)

    # ── 第一条：验证对象（【】占位 run；取首个 product promo，无则首个 promo）──
    target = next((p for p in promos if p.kind == "product"), promos[0] if promos else None)
    target_name = str((target.payload or {}).get("name") or "").strip() if target else ""
    para = _find_para(doc, lambda t: t.startswith("（验证对象）"))
    if para is None:
        unfilled.append(f"第一条·验证对象（{_REASON_NOT_FOUND}）")
    elif not target_name:
        unfilled.append(f"第一条·验证对象（{_REASON_NO_PROMO}）")
    elif not _replace_in_runs(para, "【】", f"【{target_name}】"):
        unfilled.append(f"第一条·验证对象（{_REASON_NOT_FOUND}）")

    # ── 附件三：甲方名称槽 + 合同编号（无数据源，留空披露）──
    hit = False
    for p in doc.paragraphs:
        if "【甲方公司名称】" in p.text:
            hit = _replace_in_runs(p, "【甲方公司名称】", company) or hit
    if not hit:
        unfilled.append(f"附件三·甲方名称（{_REASON_NOT_FOUND}）")
    if _find_para(doc, lambda t: "合同编号：【 】" in t) is not None:
        unfilled.append(f"附件三·合同编号（{_REASON_NO_SOURCE}）")
    else:
        unfilled.append(f"附件三·合同编号（{_REASON_NOT_FOUND}）")

    # ── 附件二表一：宣传内容与目标（行坐标+标签双校验，空值「未填写」）──
    table1 = _table_by_header(doc, _TABLE1_LABELS[0])
    if table1 is None or len(table1.rows) < len(_TABLE1_LABELS):
        for label in _TABLE1_LABELS:
            unfilled.append(f"附件二·{label}（{_REASON_NOT_FOUND}）")
    else:
        review_label = None
        if profile and profile.review_category:
            review_label = models.REVIEW_CATEGORY_LABELS.get(
                profile.review_category, profile.review_category
            )
        contact = None
        if profile and (
            (profile.contact_person or "").strip() or (profile.contact_info or "").strip()
        ):
            contact = (
                f"联系人：{(profile.contact_person or '').strip() or '—'}"
                f"　联系方式：{(profile.contact_info or '').strip() or '—'}"
            )
        values: list[tuple[str | None, str]] = [
            (company, _REASON_EMPTY),
            (None, _REASON_NO_SOURCE),  # 所属行业：profile 无 industry 字段
            (review_label, _REASON_EMPTY),
            (_yn(profile.pre_review_required) if profile else None, _REASON_EMPTY),
            (_ad_review_line(profile) if profile else None, _REASON_EMPTY),
            (contact, _REASON_EMPTY),
            (_promo_lines(promos), _REASON_EMPTY),
            (("、".join(profile.goals) if profile else "") or None, _REASON_EMPTY),
            ("\n".join(t.text for t in triggers) or None, _REASON_EMPTY),
            (("、".join(profile.platforms) if profile else "") or None, _REASON_EMPTY),
            (("、".join(profile.regions) if profile else "") or None, _REASON_EMPTY),
            ((profile.selling_points if profile else None), _REASON_EMPTY),
            (("\n".join(profile.evidence_links) if profile else "") or None, _REASON_EMPTY),
        ]
        for row, (label, (value, reason)) in enumerate(zip(_TABLE1_LABELS, values, strict=True)):
            _fill_table_row(table1, row, label, value, f"附件二·{label}", unfilled, reason=reason)

    # ── 附件二表二：资质 ──
    table2 = _table_by_header(doc, _TABLE2_LABELS[0])
    if table2 is None or len(table2.rows) < len(_TABLE2_LABELS):
        for label in _TABLE2_LABELS:
            unfilled.append(f"附件二·{label}（{_REASON_NOT_FOUND}）")
    else:
        code = (profile.business_license_code or "").strip() if profile else ""
        values2: list[str | None] = [
            (f"统一社会信用代码：{code}（附扫描件）" if code else None),
            (_license_lines(list(profile.licenses)) if profile else "") or None,
            (("、".join(profile.trademarks) if profile else "") or None),
            (("、".join(profile.ad_review_doc_types) if profile else "") or None),
        ]
        for row, (label, value) in enumerate(zip(_TABLE2_LABELS, values2, strict=True)):
            _fill_table_row(table2, row, label, value, f"附件二·{label}", unfilled)

    # ── 附件二：信息真实性确认（truth_confirmed=True 才五条全勾 □→☑；未确认如实留 □）──
    if profile and profile.truth_confirmed is True:
        for item in models.TRUTH_CONFIRM_ITEMS:
            para = _find_para(doc, lambda t, item=item: t.startswith("□") and item in t)
            if para is None or not _replace_in_runs(para, "□", "☑", count=1):
                unfilled.append(f"附件二·真实性确认勾选（{_REASON_NOT_FOUND}）")

    # ── 附件二：填表人（日期留签署时填写，不算数据槽）──
    filler = (profile.filler_name or "").strip() if profile else ""
    para = _find_para(doc, lambda t: t.startswith("填表人："))
    if para is None:
        unfilled.append(f"附件二·填表人（{_REASON_NOT_FOUND}）")
    elif not filler:
        unfilled.append(f"附件二·填表人（{_REASON_EMPTY}）")
    elif not _replace_in_runs(para, "＿" * 8, filler, count=1):
        unfilled.append(f"附件二·填表人（{_REASON_NOT_FOUND}）")

    return service.docx_bytes(doc), unfilled
