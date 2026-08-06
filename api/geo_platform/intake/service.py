"""Intake 存储层 + 调研预填 + docx 导出（execute-only，NO commit——调用方单点提交 + 审计）。

纪律（照搬旧 profile.py 语义）：
  * 全行 tenant 维隔离（fail-closed），跨租户一律查无（→ 404）；
  * prefilled 是调研预填 provenance：只有「该字段当前为空」才预填并标注，
    用户一经 API 写该字段即清标（预填≠用户确认）；
  * 期望问法同 project 同文本去重跳过（唯一索引兜底）。
"""

import datetime
import io
import re
import uuid
from typing import Any

from sqlalchemy import select
from sqlalchemy.orm import Session

from ..tenancy.ids import new_pub_id
from . import models

PREFILL_SOURCE_AI = "research:ai-live"

# 调研可预填的 profile 字段（合规亲笔项 contact_person/contact_info/filler_name/truth_confirmed
# 永远不在内——AI 只补公开可查信息）。
_PREFILL_SCALARS = (
    "website",
    "wechat",
    "douyin",
    "social_media",
    "business_license_code",
    "review_category",
    "selling_points",
    "ad_review_no",
    "ad_review_authority",
    "ad_review_expiry",
)
_PREFILL_LISTS = (
    "goals",
    "audience_type",
    "platforms",
    "regions",
    "ad_review_doc_types",
    "evidence_links",
)
_PREFILL_TEXTS = ("audience_desc",)


# ── profile ──────────────────────────────────────────────────────────────
def get_profile(
    session: Session, *, tenant_id: uuid.UUID, project_id: uuid.UUID
) -> models.IntakeProfile | None:
    return session.scalar(
        select(models.IntakeProfile).where(
            models.IntakeProfile.tenant_id == tenant_id,
            models.IntakeProfile.project_id == project_id,
        )
    )


def upsert_profile(
    session: Session,
    *,
    tenant_id: uuid.UUID,
    project_id: uuid.UUID,
    sets: dict[str, Any],
    clear_prefill: list[str],
) -> models.IntakeProfile:
    """建/改 profile（部分更新：sets 只含要写的列）；clear_prefill=用户本次显式写的 API 字段名
    （其预填标记清除——用户接手即用户数据）。列名由 schema 层白名单校验后传入。"""
    row = get_profile(session, tenant_id=tenant_id, project_id=project_id)
    if row is None:
        row = models.IntakeProfile(
            pub_id=new_pub_id("itp"),
            tenant_id=tenant_id,
            project_id=project_id,
            goals=[],
            audience_type=[],
            platforms=[],
            regions=[],
            trademarks=[],
            ad_review_doc_types=[],
            evidence_links=[],
            licenses=[],
            prefilled={},
        )
        session.add(row)
        session.flush()
    for key, value in sets.items():
        setattr(row, key, value)
    if clear_prefill:
        marks = dict(row.prefilled or {})
        if any(marks.pop(k, None) is not None for k in clear_prefill):
            row.prefilled = marks
    session.flush()
    return row


def prefill_from_research(
    session: Session,
    *,
    tenant_id: uuid.UUID,
    project_id: uuid.UUID,
    data: dict[str, Any],
    source: str = PREFILL_SOURCE_AI,
) -> list[str]:
    """调研预填：只填**当前为空**的字段，并记 prefilled[API名]=source（诚实标注）。

    已被用户/前次填写占据的字段一律跳过（预填绝不覆盖）；词表/格式 fail-closed 已在
    research._filter_vocab 完成，此处再兜底一次。返回实际预填的 API 字段名列表。"""
    row = upsert_profile(
        session, tenant_id=tenant_id, project_id=project_id, sets={}, clear_prefill=[]
    )
    marks = dict(row.prefilled or {})
    filled: list[str] = []
    for col in _PREFILL_SCALARS:
        value = data.get(col)
        if not isinstance(value, str) or not value.strip():
            continue
        value = value.strip()
        if col == "business_license_code" and not _license_code_ok(value):
            continue
        if col == "review_category" and value not in models.REVIEW_CATEGORIES:
            continue
        if getattr(row, col):  # 已有值 → 绝不覆盖
            continue
        setattr(row, col, value)
        marks[col] = source
        filled.append(col)
    for api_name in _PREFILL_LISTS:
        values = data.get(api_name)
        if not isinstance(values, list):
            continue
        vocab = models.PROFILE_LIST_FIELDS[api_name]
        vals = [
            v for v in values if isinstance(v, str) and v.strip() and (vocab is None or v in vocab)
        ]
        if not vals:
            continue
        if getattr(row, api_name):  # 已有值 → 绝不覆盖
            continue
        setattr(row, api_name, vals)
        marks[api_name] = source
        filled.append(api_name)
    for api_name in _PREFILL_TEXTS:
        value = data.get(api_name)
        if not isinstance(value, str) or not value.strip():
            continue
        if getattr(row, api_name):
            continue
        setattr(row, api_name, value.strip())
        marks[api_name] = source
        filled.append(api_name)
    if filled:
        row.prefilled = marks
        session.flush()
    return filled


def add_prefill_marks(
    session: Session,
    *,
    tenant_id: uuid.UUID,
    project_id: uuid.UUID,
    keys: tuple[str, ...],
    source: str = PREFILL_SOURCE_AI,
) -> list[str]:
    """给非 profile 列的 AI 草稿（promo/trigger_question）记 provenance 标记。只加不覆盖。"""
    row = upsert_profile(
        session, tenant_id=tenant_id, project_id=project_id, sets={}, clear_prefill=[]
    )
    marks = dict(row.prefilled or {})
    added = [k for k in keys if k not in marks]
    if added:
        for k in added:
            marks[k] = source
        row.prefilled = marks
        session.flush()
    return added


def clear_prefill_marks(
    session: Session,
    *,
    tenant_id: uuid.UUID,
    project_id: uuid.UUID,
    keys: tuple[str, ...],
) -> list[str]:
    """用户经 API 改动对应数据（promo/trigger）→ 清标（AI 草稿一经用户接手即用户数据）。"""
    row = get_profile(session, tenant_id=tenant_id, project_id=project_id)
    if row is None:
        return []
    marks = dict(row.prefilled or {})
    cleared = [k for k in keys if marks.pop(k, None) is not None]
    if cleared:
        row.prefilled = marks
        session.flush()
    return cleared


_LICENSE_CODE_RE = re.compile(r"[0-9A-Z]{18}")


def _license_code_ok(value: str) -> bool:
    return bool(_LICENSE_CODE_RE.fullmatch(value))


def apply_research_data(
    session: Session,
    *,
    tenant_id: uuid.UUID,
    project_id: uuid.UUID,
    data: dict[str, Any],
) -> dict[str, Any]:
    """调研结果落库（intake 运营端与 intake_form 免登录通道共用）：
    profile 空字段预填（带 provenance）→ 无 promo 时建 product≤3+company≤1 草稿 →
    问法收录 draft ≤20（去重跳过）。返回实际写入的清单（供审计/响应披露）。"""
    filled = prefill_from_research(session, tenant_id=tenant_id, project_id=project_id, data=data)
    promos_created: list[str] = []
    if not list_promos(session, tenant_id=tenant_id, project_id=project_id):
        for p in (data.get("products") or [])[:3]:
            if not isinstance(p, dict) or not (p.get("name") or "").strip():
                continue
            promo = create_promo(
                session,
                tenant_id=tenant_id,
                project_id=project_id,
                kind="product",
                payload={
                    "name": p["name"].strip(),
                    "category": p.get("category") or "",
                    "features": p.get("features") or [],
                    "desc": p.get("desc") or "",
                    "price": p.get("price") or "",
                },
            )
            promos_created.append(promo.pub_id)
        cb = data.get("company_brief")
        if isinstance(cb, dict):
            payload = {
                "name": cb.get("name") or "",
                "strength": cb.get("strength") or [],
                "advantage": cb.get("advantage") or "",
                "cases": cb.get("cases") or "",
                "data": cb.get("data") or "",
            }
            if any(payload.values()):
                promo = create_promo(
                    session,
                    tenant_id=tenant_id,
                    project_id=project_id,
                    kind="company",
                    payload=payload,
                )
                promos_created.append(promo.pub_id)
    if promos_created:
        add_prefill_marks(session, tenant_id=tenant_id, project_id=project_id, keys=("promos",))
    tq_text = data.get("trigger_questions") or ""
    tq_lines = [ln.strip() for ln in re.split(r"[\r\n]+", str(tq_text)) if ln.strip()][:20]
    triggers_created: list[str] = []
    triggers_skipped: list[str] = []
    if tq_lines:
        rows, triggers_skipped = create_trigger_questions(
            session, tenant_id=tenant_id, project_id=project_id, texts=tq_lines
        )
        triggers_created = [r.pub_id for r in rows]
    if triggers_created:
        add_prefill_marks(
            session, tenant_id=tenant_id, project_id=project_id, keys=("trigger_questions",)
        )
    return {
        "prefilled": filled,
        "promos_created": promos_created,
        "triggers_created": triggers_created,
        "triggers_skipped": triggers_skipped,
    }


# ── promo ────────────────────────────────────────────────────────────────
def create_promo(
    session: Session,
    *,
    tenant_id: uuid.UUID,
    project_id: uuid.UUID,
    kind: str,
    payload: dict[str, Any],
) -> models.IntakePromo:
    promo = models.IntakePromo(
        pub_id=new_pub_id("prm"),
        tenant_id=tenant_id,
        project_id=project_id,
        kind=kind,
        payload=payload,
    )
    session.add(promo)
    session.flush()
    return promo


def get_promo(
    session: Session, *, tenant_id: uuid.UUID, project_id: uuid.UUID, promo_pub_id: str
) -> models.IntakePromo | None:
    return session.scalar(
        select(models.IntakePromo).where(
            models.IntakePromo.tenant_id == tenant_id,
            models.IntakePromo.project_id == project_id,
            models.IntakePromo.pub_id == promo_pub_id,
        )
    )


def list_promos(
    session: Session, *, tenant_id: uuid.UUID, project_id: uuid.UUID
) -> list[models.IntakePromo]:
    return list(
        session.scalars(
            select(models.IntakePromo)
            .where(
                models.IntakePromo.tenant_id == tenant_id,
                models.IntakePromo.project_id == project_id,
            )
            .order_by(models.IntakePromo.created_at, models.IntakePromo.id)
        ).all()
    )


def delete_promo(session: Session, promo: models.IntakePromo) -> None:
    session.delete(promo)
    session.flush()


# ── trigger question ─────────────────────────────────────────────────────
def create_trigger_questions(
    session: Session,
    *,
    tenant_id: uuid.UUID,
    project_id: uuid.UUID,
    texts: list[str],
) -> tuple[list[models.IntakeTriggerQuestion], list[str]]:
    """批量收录问法；同 project 同文本去重跳过（唯一索引兜底）。返回 (新建行, 跳过的文本)。"""
    existing = {
        row.text for row in list_triggers(session, tenant_id=tenant_id, project_id=project_id)
    }
    created: list[models.IntakeTriggerQuestion] = []
    skipped: list[str] = []
    seen: set[str] = set()
    for text in texts:
        t = text.strip()
        if not t or t in seen:
            continue
        seen.add(t)
        if t in existing:
            skipped.append(t)
            continue
        row = models.IntakeTriggerQuestion(
            pub_id=new_pub_id("trq"),
            tenant_id=tenant_id,
            project_id=project_id,
            text=t,
            status="draft",
        )
        session.add(row)
        created.append(row)
        existing.add(t)
    if created:
        session.flush()
    return created, skipped


def get_trigger(
    session: Session, *, tenant_id: uuid.UUID, project_id: uuid.UUID, trigger_pub_id: str
) -> models.IntakeTriggerQuestion | None:
    return session.scalar(
        select(models.IntakeTriggerQuestion).where(
            models.IntakeTriggerQuestion.tenant_id == tenant_id,
            models.IntakeTriggerQuestion.project_id == project_id,
            models.IntakeTriggerQuestion.pub_id == trigger_pub_id,
        )
    )


def list_triggers(
    session: Session, *, tenant_id: uuid.UUID, project_id: uuid.UUID
) -> list[models.IntakeTriggerQuestion]:
    return list(
        session.scalars(
            select(models.IntakeTriggerQuestion)
            .where(
                models.IntakeTriggerQuestion.tenant_id == tenant_id,
                models.IntakeTriggerQuestion.project_id == project_id,
            )
            .order_by(models.IntakeTriggerQuestion.created_at, models.IntakeTriggerQuestion.id)
        ).all()
    )


# ══ 导出 Word（旧 _build_profile_docx 五节结构的服务端移植；python-docx 延迟导入）════
def _docx_table(doc: Any, rows: list[tuple[str, Any]]) -> Any:
    """两列表（标签/值）；空值如实写「未填写」（不补不猜）。"""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, val in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = (val or "").strip() if isinstance(val, str) else (val or "未填写")
        if not cells[1].text:
            cells[1].text = "未填写"
    return table


def build_profile_docx(
    company_name: str,
    profile: models.IntakeProfile | None,
    promos: list[models.IntakePromo],
    triggers: list[models.IntakeTriggerQuestion],
) -> tuple[Any, str]:
    """按五节结构生成客户信息表 docx：一、基础信息；二、推广对象；三、GEO 优化目标；
    四、意向 AI 平台；五、资质证明。数据=project+profile+promo+trigger 真库值。"""
    import docx  # 延迟导入：重依赖，端点按需加载

    today = datetime.datetime.now().strftime("%Y-%m-%d")
    doc = docx.Document()
    doc.add_heading("GEO 客户信息收集表", level=0)
    doc.add_paragraph("AI 搜索优化服务 · " + today)

    def _yn(v: bool | None) -> str:
        if v is None:
            return "未填写"
        return "是" if v else "否"

    review_label = None
    if profile and profile.review_category:
        review_label = models.REVIEW_CATEGORY_LABELS.get(
            profile.review_category, profile.review_category
        )
    doc.add_heading("一、基础信息", level=1)
    _docx_table(
        doc,
        [
            ("公司/品牌名称", company_name),
            ("行业广告审查分类", review_label),
            ("是否属于法定前置审查行业", _yn(None if not profile else profile.pre_review_required)),
            ("联系人", profile.contact_person if profile else None),
            ("联系方式", profile.contact_info if profile else None),
            ("官网", profile.website if profile else None),
            ("微信公众号/小程序", profile.wechat if profile else None),
            ("抖音/视频号", profile.douyin if profile else None),
            ("其他社媒", profile.social_media if profile else None),
        ],
    )

    doc.add_heading("二、推广对象", level=1)
    if not promos:
        doc.add_paragraph("未填写")
    for i, promo in enumerate(promos, 1):
        d = promo.payload or {}
        label = "产品 / 服务" if promo.kind == "product" else "公司 / 品牌"
        doc.add_paragraph(f"【推广对象 {i}】{d.get('name') or label}")

        def _joined(key: str, payload: dict[str, object] = d) -> str:
            value = payload.get(key)
            return "、".join(str(x) for x in value) if isinstance(value, list) else ""

        if promo.kind == "product":
            rows = [
                ("类型", label),
                ("产品/服务名称", d.get("name")),
                ("类别", d.get("category")),
                ("核心卖点", _joined("features")),
                ("功能与特点", d.get("desc")),
                ("价格区间", d.get("price")),
            ]
        else:
            rows = [
                ("类型", label),
                ("公司/品牌名称", d.get("name")),
                ("企业实力", _joined("strength")),
                ("核心差异化优势", d.get("advantage")),
                ("代表性成功案例", d.get("cases")),
                ("关键数据", d.get("data")),
            ]
        _docx_table(doc, rows)

    doc.add_heading("三、GEO 优化目标", level=1)
    _docx_table(
        doc,
        [
            ("核心推广目的", "、".join(profile.goals) if profile else ""),
            ("目标客户类型", "、".join(profile.audience_type) if profile else ""),
            ("决策人补充", profile.audience_desc if profile else None),
            ("期望触发的用户问法", "\n".join(t.text for t in triggers)),
            ("目标推广地域", "、".join(profile.regions) if profile else ""),
            ("核心卖点", profile.selling_points if profile else None),
            ("可公开引用的佐证材料", "\n".join(profile.evidence_links) if profile else ""),
        ],
    )

    doc.add_heading("四、意向 AI 平台", level=1)
    _docx_table(
        doc,
        [
            ("选定平台", "、".join(profile.platforms) if profile else ""),
        ],
    )

    licenses = profile.licenses if profile else []
    lic_text = "\n".join(
        f"{lic.get('name') or '（未命名）'}（编号 {lic.get('number') or '—'}"
        f"{'，有效期至 ' + lic['expiry'] if lic.get('expiry') else ''}）"
        for lic in licenses
    )
    ad_review = None
    if profile and (
        profile.ad_review_no or profile.ad_review_authority or profile.ad_review_expiry
    ):
        ad_review = (
            f"{profile.ad_review_no or '—'}（审查机关 {profile.ad_review_authority or '—'}"
            f"{'，有效期至 ' + profile.ad_review_expiry if profile.ad_review_expiry else ''}）"
        )
    doc.add_heading("五、资质证明", level=1)
    _docx_table(
        doc,
        [
            ("营业执照-统一社会信用代码", profile.business_license_code if profile else None),
            ("行业许可证", lic_text),
            ("商标/品牌权属证明", "、".join(profile.trademarks) if profile else ""),
            ("广告审查批准文号", ad_review),
            ("广告审查批准文件类型", "、".join(profile.ad_review_doc_types) if profile else ""),
            ("信息真实性确认", _yn(None if not profile else profile.truth_confirmed)),
            ("填表人", profile.filler_name if profile else None),
        ],
    )
    doc.add_paragraph("Powered by GEO Platform · AI 联网调研版")
    return doc, today


def docx_bytes(doc: Any) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
