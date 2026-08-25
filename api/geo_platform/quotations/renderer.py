"""Legacy GEO 报价 DOCX renderer（仅供内部回归）。

该 renderer 仍从空白 ``Document()`` 构造文档，并不符合用户最终批准模板。Phase A
保留它只为回归现有五服务业务逻辑；所有产物必须带醒目的非合规标识，禁止作为正式客户报价。
"""

from __future__ import annotations

from collections import OrderedDict
from collections.abc import Iterable, Sequence
from datetime import date, datetime, time
from io import BytesIO
from typing import Literal, cast
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .catalog import PACKAGE_BY_CODE, SERVICE_BY_CODE
from .models import (
    ExistingQueryVariants,
    OpportunityVariants,
    QuotationConfiguration,
    QuotationPlan,
)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
ISSUER_COMPANY = "北京硅基守望科技有限公司"
VALID_WORKING_DAYS = 30
TEMPLATE_COMPLIANCE = "non-final-template"
NON_FINAL_TEMPLATE_NOTICE = "非最终模板合规产物（仅供内部回归，禁止作为正式客户报价）"

_FONT = "宋体"
_BLACK = RGBColor(0x00, 0x00, 0x00)
_BODY_GRAY = RGBColor(0x44, 0x44, 0x44)
_MUTED_GRAY = RGBColor(0x80, 0x80, 0x80)

_SEC_LABELS = {
    "search": "搜索型产品或服务",
    "experience": "体验型产品或服务",
    "trust": "专业信任型产品或服务",
    "mixed": "复合决策型产品或服务",
}

_CHINESE_NUMERALS = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五"}


def _appendix_number(value: int) -> str:
    return _CHINESE_NUMERALS.get(value, str(value))


_TERMS = (
    "1. 本报价有效期为30个工作日，生效日期以报价日期为准。贵单位确认后，我司将以"
    "报价产品或服务的相关价格签订合同。",
    "2. 因AI平台算法更迭频繁，如项目在服务过程中因平台算法等调整无法按照既定方案"
    "执行，双方需签订补充协议对内容进行调整，并通过电话或邮件等方式进行确认。",
    "3. 因涉及行业及品类的专业性与合规性，相关产品或服务的内容基础材料均来源于"
    "品牌方，我方基于品牌方提供内容进行评测。",
    "4. 特定情况：（1）合作期间品牌方出现重大舆情事件，需重新评估合作及相关费用；"
    "（2）品牌方因自身原因需中途更新评测问题集的，需重新评估费用。",
    "5. 本报价单及附件内价格、方案、技术参数、商务条件均为我方保密商业信息。收件方"
    "仅可用于本次项目内部评审，不得向外部第三方、竞品、其他合作方进行披露、转发、"
    "摘抄。如因贵方泄密造成我方损失，我方有权取消本次报价效力并追究相关责任。"
    "本保密义务在报价失效后仍然持续有效。",
)


def _stable_docx_bytes(payload: bytes) -> bytes:
    """Normalize OOXML ZIP metadata so equal inputs produce equal document bytes and SHA-256."""
    output = BytesIO()
    with (
        ZipFile(BytesIO(payload)) as source,
        ZipFile(
            output,
            mode="w",
            compression=ZIP_DEFLATED,
            compresslevel=9,
        ) as target,
    ):
        for name in sorted(source.namelist()):
            info = ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = ZIP_DEFLATED
            info.create_system = 0
            info.external_attr = 0o600 << 16
            target.writestr(info, source.read(name))
    return output.getvalue()


def _set_east_asia_font(run: Run, name: str = _FONT) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)


def _style_run(
    run: Run,
    *,
    size: float,
    bold: bool = False,
    underline: bool = False,
    color: RGBColor = _BLACK,
) -> Run:
    _set_east_asia_font(run)
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline
    run.font.color.rgb = color
    return run


def _paragraph_spacing(
    paragraph: Paragraph,
    *,
    before: float = 0,
    after: float = 0,
    line: float = 1.0,
) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def _add_text_paragraph(
    document: DocumentObject,
    text: str,
    *,
    size: float = 10.5,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    before: float = 0,
    after: float = 0,
    line: float = 1.0,
    keep_with_next: bool = False,
    color: RGBColor = _BLACK,
) -> Paragraph:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    _paragraph_spacing(paragraph, before=before, after=after, line=line)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    _style_run(paragraph.add_run(text), size=size, bold=bold, color=color)
    return paragraph


def _add_heading(
    document: DocumentObject,
    text: str,
    *,
    level: Literal[1, 2, 3],
    centered: bool = False,
) -> Paragraph:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    sizes = {1: 14.0, 2: 13.0, 3: 10.5}
    _paragraph_spacing(
        paragraph,
        before=0 if level == 1 else 5,
        after=4 if level != 3 else 2,
        line=1.0,
    )
    paragraph.paragraph_format.keep_with_next = True
    _style_run(paragraph.add_run(text), size=sizes[level], bold=True)
    return paragraph


def _set_cell_margins(
    cell: _Cell, *, top: int = 0, left: int = 108, bottom: int = 0, right: int = 108
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _shade_cell(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _cant_split(row: _Row) -> None:
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _repeat_table_header(row: _Row) -> None:
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_column_widths(table: Table, widths_mm: Sequence[float]) -> None:
    table.autofit = False
    for column, width in zip(table.columns, widths_mm, strict=True):
        column.width = Mm(width)
        for cell in column.cells:
            cell.width = Mm(width)


def _cell_paragraph(
    cell: _Cell,
    text: str,
    *,
    size: float,
    bold: bool = False,
    underline: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    line: float = 1.0,
    after: float = 0,
) -> Paragraph:
    paragraph = (
        cell.paragraphs[0]
        if len(cell.paragraphs) == 1 and not cell.text
        else cast(Paragraph, cell.add_paragraph())
    )
    if align is not None:
        paragraph.alignment = align
    _paragraph_spacing(paragraph, after=after, line=line)
    _style_run(
        paragraph.add_run(text),
        size=size,
        bold=bold,
        underline=underline,
    )
    return paragraph


def _prepare_cell(cell: _Cell, *, vertical_center: bool = True) -> None:
    cell.text = ""
    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER if vertical_center else WD_CELL_VERTICAL_ALIGNMENT.TOP
    )
    _set_cell_margins(cell)


def _format_price(cents: int | None) -> str:
    if cents is None:
        return "待确认"
    return f"￥{cents / 100:,.2f}"


def _add_service_table(document: DocumentObject, configuration: QuotationConfiguration) -> None:
    has_pending_site_condition = (
        configuration.package_code == "minimum_validation"
        and configuration.official_site_in_citations is None
        and "official_site_audit" in configuration.service_codes
    )
    table = document.add_table(
        rows=len(configuration.service_quotes) + (3 if has_pending_site_condition else 2),
        cols=6,
    )
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_column_widths(table, (10, 35, 67, 24, 14, 30))

    for row in table.rows:
        _cant_split(row)
        for cell in row.cells:
            _prepare_cell(cell)

    headers = ("编号", "服务项目", "服务范围", "单价", "数量", "小计")
    _repeat_table_header(table.rows[0])
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        _shade_cell(cell, "F0F0F0")
        _cell_paragraph(
            cell,
            text,
            size=10.5,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=1.15,
        )

    for row_index, quote in enumerate(configuration.service_quotes, start=1):
        service = SERVICE_BY_CODE[quote.service_code]
        is_conditional = has_pending_site_condition and quote.service_code == "official_site_audit"
        row = table.rows[row_index]
        _cell_paragraph(
            row.cells[0],
            str(service.number),
            size=10.5,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=1.15,
        )
        _cell_paragraph(
            row.cells[1],
            f"{service.short_name}\n{service.name}",
            size=9.0,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=1.1,
        )
        content_cell = row.cells[2]
        content_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # 首页只呈现用于决策的服务摘要；完整执行范围统一放在附录一，避免报价表被长文撑散。
        _cell_paragraph(content_cell, service.summary, size=8.5, line=1.1)
        _cell_paragraph(
            row.cells[3],
            _format_price(quote.unit_price_cents),
            size=9.0,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=1.1,
        )
        _cell_paragraph(
            row.cells[4],
            f"{quote.quantity}{service.unit}" + ("\n条件" if is_conditional else ""),
            size=9.0,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=1.1,
        )
        _cell_paragraph(
            row.cells[5],
            (
                "触发后 " + _format_price(quote.subtotal_cents)
                if is_conditional
                else _format_price(quote.subtotal_cents)
            ),
            size=9.0,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=1.1,
        )

    total = table.rows[-2] if has_pending_site_condition else table.rows[-1]
    _cell_paragraph(
        total.cells[2],
        "基础服务费合计（不含条件项）" if has_pending_site_condition else "服务费合计",
        size=10.5,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        line=1.15,
    )
    if has_pending_site_condition:
        maximum = table.rows[-1]
        _cell_paragraph(
            maximum.cells[2],
            "官网命中后最高服务费",
            size=10.0,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            line=1.15,
        )
        _cell_paragraph(
            maximum.cells[5],
            _format_price(configuration.maximum_total_price_cents),
            size=10.0,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=1.15,
        )
    _cell_paragraph(
        total.cells[5],
        _format_price(configuration.total_price_cents),
        size=10.5,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line=1.15,
    )


def _configure_styles(document: DocumentObject) -> None:
    normal = document.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(10.5)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), _FONT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    for name, size in (("Heading 1", 14.0), ("Heading 2", 13.0), ("Heading 3", 10.5)):
        style = document.styles[name]
        style.font.name = _FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _BLACK
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), _FONT)
    bullet = document.styles["List Bullet"]
    bullet.font.name = _FONT
    bullet._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), _FONT)
    bullet.paragraph_format.space_after = Pt(0)


def _add_page_break(document: DocumentObject) -> None:
    paragraph = document.add_paragraph()
    _paragraph_spacing(paragraph, line=1.0)
    run = paragraph.add_run()
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run._r.append(page_break)


def _add_page_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    _style_run(run, size=9.0)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, placeholder, end):
        run._r.append(element)


def _configure_page(document: DocumentObject) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing(paragraph, line=1.0)
    _style_run(
        paragraph.add_run("【非最终模板合规产物·仅供内部回归】\n"),
        size=8.0,
        bold=True,
        color=RGBColor(0xC0, 0x00, 0x00),
    )
    _style_run(
        paragraph.add_run(
            "本报价为保密商业资料，未经出具方书面同意，禁止对外泄露、转发。"
            "报价仅供本次合作评估使用。"
        ),
        size=7.5,
        color=_MUTED_GRAY,
    )

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing(footer_paragraph, line=1.0)
    _add_page_field(footer_paragraph)


def _format_date(value: date) -> str:
    return f"{value.year}年{value.month}月{value.day}日"


def _execution_sequence(configuration: QuotationConfiguration) -> str:
    package = PACKAGE_BY_CODE[configuration.package_code]
    selected = frozenset(configuration.service_codes)
    quantities = {line.service_code: line.quantity for line in configuration.service_quotes}
    sequence = [code for code in package.execution_sequence if code in selected]
    if configuration.package_code == "custom":
        sequence = list(configuration.service_codes)
        if "content_publishing_pilot" in selected and quantities.get("ranking_test") == 2:
            sequence = [
                "ranking_test",
                *(code for code in sequence if code != "ranking_test"),
                "ranking_test",
            ]
    labels: list[str] = []
    ranking_seen = 0
    for code in sequence:
        service = SERVICE_BY_CODE[code]
        label = f"{service.number} {service.short_name}"
        if code == "ranking_test" and sequence.count(code) > 1:
            ranking_seen += 1
            label += "（基线）" if ranking_seen == 1 else "（复测）"
        elif (
            code == "ranking_test"
            and configuration.package_code == "custom"
            and quantities.get(code, 1) > 1
        ):
            label += f"（{quantities[code]}轮）"
        if (
            code == "official_site_audit"
            and configuration.package_code == "minimum_validation"
            and configuration.official_site_in_citations is None
        ):
            label += "（命中后）"
        labels.append(label)
    return " → ".join(labels)


def _add_cover(
    document: DocumentObject,
    brand_name: str,
    quote_date: date,
    configuration: QuotationConfiguration,
) -> None:
    package = PACKAGE_BY_CODE[configuration.package_code]
    _add_text_paragraph(
        document,
        "GEO 服务报价单",
        size=15.0,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=1,
    )
    _add_text_paragraph(
        document,
        package.name,
        size=11.0,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=1,
    )
    _add_text_paragraph(
        document,
        f"{brand_name}      报价日期：{_format_date(quote_date)}     "
        f"报价有效期：{VALID_WORKING_DAYS}个工作日",
        size=9.0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=1,
        line=1.5,
    )
    _add_text_paragraph(
        document,
        f"执行顺序：{_execution_sequence(configuration)}",
        size=8.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=2,
        line=1.1,
        color=_BODY_GRAY,
    )
    _add_service_table(document, configuration)
    if (
        configuration.package_code == "minimum_validation"
        and configuration.official_site_in_citations is None
    ):
        _add_text_paragraph(
            document,
            "条件说明：服务 4 的单价已单独列示，但暂不计入基础总价。"
            "首轮测试的引用 URL 命中客户官网后，才触发该项服务和费用。",
            size=8.5,
            bold=True,
            before=3,
            after=1,
            line=1.1,
        )
    if configuration.pricing_status == "pending":
        _add_text_paragraph(
            document,
            "样例说明：本文件中的服务单价与总价待商务确认，不构成正式价格承诺。",
            size=8.5,
            bold=True,
            before=3,
            after=1,
            line=1.1,
        )
    if configuration.commercial_note:
        _add_text_paragraph(
            document,
            "报价备注：" + configuration.commercial_note,
            size=8.5,
            before=2,
            after=1,
            line=1.1,
        )
    _add_text_paragraph(
        document,
        "商务口径：币种为人民币（CNY），文档版本为 V1.0。"
        "含税状态、发票类型、付款节点、交付周期和第三方代采费用如未在报价备注中明确，"
        "均待合同确认。",
        size=8.5,
        before=2,
        after=1,
        line=1.1,
    )
    for index, term in enumerate(_TERMS):
        _add_text_paragraph(
            document,
            term,
            size=9.0,
            color=_BODY_GRAY,
            before=6 if index == 0 else 0,
            after=1,
            line=1.15,
        )
    _add_text_paragraph(document, "", size=10.5, after=1, line=1.15)
    signature = document.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _paragraph_spacing(signature, after=10, line=1.15)
    _style_run(
        signature.add_run("报价人：__________________　　　　　　　　公司：" + ISSUER_COMPANY),
        size=10.5,
    )


def _add_service_delivery_appendix(
    document: DocumentObject,
    *,
    configuration: QuotationConfiguration,
    query_appendix_included: bool,
) -> None:
    package = PACKAGE_BY_CODE[configuration.package_code]
    _add_page_break(document)
    _add_heading(document, "附录一 服务输入、执行与交付说明", level=1, centered=True)
    _add_labelled_paragraph(document, "适用客户：", package.audience, size=9.0, after=2)
    _add_labelled_paragraph(document, "套餐目标：", package.summary, size=9.0, after=2)
    _add_labelled_paragraph(
        document,
        "执行顺序：",
        _execution_sequence(configuration),
        size=9.0,
        after=2,
    )
    _add_labelled_paragraph(
        document,
        "计价关系：",
        (
            "每项小计 = 本项单价 × 数量；基础总价不包含待触发的服务 4；"
            "官网命中后最高总价 = 基础总价 + 服务 4 小计。"
            if (
                configuration.package_code == "minimum_validation"
                and configuration.official_site_in_citations is None
            )
            else "每项小计 = 本项单价 × 数量；服务费总报价 = 全部已选服务小计之和。"
        )
        + "套餐决定预设服务组合、数量、条件项和执行顺序，不覆盖任何服务单价。",
        size=9.0,
        after=3,
    )
    if package.preconditions:
        _add_labelled_paragraph(
            document,
            "执行前置：",
            "",
            size=9.0,
            keep_with_next=True,
        )
        for item in package.preconditions:
            _add_bullet(document, "前置", item, size=8.5)
    if package.handoffs:
        _add_labelled_paragraph(
            document,
            "服务传递：",
            "",
            size=9.0,
            keep_with_next=True,
        )
        for item in package.handoffs:
            _add_bullet(document, "传递", item, size=8.5)
    _add_labelled_paragraph(
        document,
        "范围冻结：",
        "问题数/变体数/渠道/重复数、URL 或页面上限、发文篇数/媒体数与观察窗必须在"
        "商务备注或合同中确认。未确认时，本报价不表示验收工作量已冻结。",
        size=8.5,
        after=3,
    )
    if "ranking_test" in configuration.service_codes:
        _add_labelled_paragraph(
            document,
            "采集前置：",
            "开放 API 与豆包 App 两类渠道必须分别接通并验证。网页结果不得标记为 App 结果，"
            "网页内部请求不得标记为模型开放 API。前置门未通过时，双方应调整执行范围。",
            size=8.5,
            after=3,
        )
    if "official_site_audit" in configuration.service_codes:
        _add_labelled_paragraph(
            document,
            "官网范围：",
            configuration.website_url,
            size=9.0,
            after=3,
        )
    if configuration.official_site_citation_url:
        _add_labelled_paragraph(
            document,
            "官网命中证据 URL：",
            configuration.official_site_citation_url,
            size=8.5,
            after=3,
        )
    if not query_appendix_included and any(
        code in configuration.service_codes for code in ("ranking_test", "content_publishing_pilot")
    ):
        query_freeze_text = {
            "geo_effect_assessment": (
                "具体 Query 及语义变体将在客户提供既有 GEO 目标问题、补充并确认后冻结，"
                "不在报价阶段编造。"
            ),
            "minimum_validation": (
                "具体 Query 及语义变体将由我方提出候选，客户可补充并最终确认后冻结，"
                "不在报价阶段编造。"
            ),
            "custom": (
                "具体 Query 的提出方、补充方式及语义变体将在商务备注或合同中确认后冻结，"
                "不在报价阶段编造。"
            ),
        }[configuration.package_code]
        _add_text_paragraph(
            document,
            "本次未上传目标词 XLSX，因此报价单只确认服务范围、输入、交付物和价格状态。"
            + query_freeze_text,
            size=8.5,
            bold=True,
            after=4,
            line=1.1,
        )

    for quote in configuration.service_quotes:
        block_start = len(document.paragraphs)
        service = SERVICE_BY_CODE[quote.service_code]
        title = f"服务 {service.number}｜{service.short_name}：{service.name}"
        if quote.quantity > 1:
            title += f"（{quote.quantity}{service.unit}）"
        _add_heading(document, title, level=2)
        _appendix_body(document, service.summary, after=2)
        _add_labelled_paragraph(document, "客户需提供：", "", size=9.0, keep_with_next=True)
        for item in service.inputs:
            _add_bullet(document, "输入", item, size=8.5)
        _add_labelled_paragraph(document, "执行范围：", "", size=9.0, keep_with_next=True)
        for item in service.scope:
            _add_bullet(document, "执行", item, size=8.5)
        _add_labelled_paragraph(document, "交付物：", "", size=9.0, keep_with_next=True)
        for item in service.outputs:
            _add_bullet(document, "输出", item, size=8.5, after=1)
        # A service definition is one commercial unit. Keep its heading, inputs,
        # execution scope and deliverables together so a page break cannot detach
        # the acceptance-facing outputs from the service they belong to.
        block = document.paragraphs[block_start:]
        for paragraph in block[:-1]:
            paragraph.paragraph_format.keep_with_next = True


def _appendix_body(document: DocumentObject, text: str, *, after: float = 3) -> Paragraph:
    return _add_text_paragraph(document, text, size=10.5, after=after, line=1.2)


def _add_labelled_paragraph(
    document: DocumentObject,
    label: str,
    text: str,
    *,
    size: float = 9.0,
    after: float = 0,
    keep_with_next: bool = False,
) -> Paragraph:
    paragraph = document.add_paragraph()
    _paragraph_spacing(paragraph, after=after, line=1.0)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    _style_run(paragraph.add_run(label), size=size, bold=True)
    _style_run(paragraph.add_run(text), size=size)
    return paragraph


def _add_bullet(
    document: DocumentObject,
    label: str,
    text: str,
    *,
    size: float = 8.5,
    bold: bool = False,
    after: float = 0,
    keep_with_next: bool = False,
) -> Paragraph:
    paragraph = document.add_paragraph(style="List Bullet")
    _paragraph_spacing(paragraph, after=after, line=1.05)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    _style_run(paragraph.add_run(f"{label}  {text}"), size=size, bold=bold)
    return paragraph


def _add_validation_table(
    document: DocumentObject, opportunities: Sequence[OpportunityVariants]
) -> None:
    table = document.add_table(rows=1 + min(2, len(opportunities)), cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_column_widths(table, (49, 23, 23, 23, 23, 39))
    headers = (
        "拟新增目标词",
        "模型 API\n优化前",
        "模型 API\n优化后",
        "豆包 App\n优化前",
        "豆包 App\n优化后",
        "验证状态",
    )
    _repeat_table_header(table.rows[0])
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        _prepare_cell(cell)
        _shade_cell(cell, "F0F0F0")
        _cell_paragraph(
            cell,
            text,
            size=8.0,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=1.0,
        )
    for row, opportunity in zip(table.rows[1:], opportunities[:2], strict=True):
        _cant_split(row)
        for cell in row.cells:
            _prepare_cell(cell)
        values = (opportunity.keyword, "待实测", "待实测", "待实测", "待实测", "项目执行阶段验证")
        for cell, text in zip(row.cells, values, strict=True):
            _cell_paragraph(
                cell,
                text,
                size=7.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                line=1.0,
            )


def _add_appendix_one(
    document: DocumentObject,
    *,
    brand_name: str,
    plan: QuotationPlan,
    appendix_number: int,
    page_break_before: bool = True,
) -> None:
    if page_break_before:
        _add_page_break(document)
    _add_heading(
        document,
        f"附录{_appendix_number(appendix_number)} Query 优化方案说明",
        level=1,
        centered=True,
    )
    _add_heading(document, "一、优化目标", level=2)
    _appendix_body(
        document,
        '本附录所展示的"Query优化"工作，是基于客户提供的业务关键词，将其改写为更贴近'
        "真实用户在AI平台上的提问方式。消费者在使用AI搜索时，通常不会直接输入产品名称"
        "或技术术语，而是以自然语言描述自己的需求场景。",
    )
    _appendix_body(
        document,
        "因此，我们将客户提供的关键词转化为模拟用户真实提问的检索语句，使其能够触发"
        "AI平台给出包含厂商推荐的回答。同时为每条优化后的提问生成多个语义变体（正式换述、"
        "换角度、口语化），覆盖不同用户群体的表达习惯，确保评测结果能够反映品牌在真实"
        "检索场景中的可见性。",
    )

    _add_heading(document, "二、优化方法论", level=2)
    _appendix_body(
        document,
        "本方案基于消费心理学中的搜索品-体验品-信任品（SEC）经典理论（Nelson, 1970; "
        "Darby & Karni, 1973），结合我司自研的九维消费认知量化模型，对品牌所属品类进行"
        "消费心理画像分析，进而完善Query优化。",
    )
    _appendix_body(
        document,
        f"结合本次目标词特征，{plan.category_label}更接近{_SEC_LABELS[plan.sec_profile]}。"
        f"{plan.category_analysis}",
    )
    _appendix_body(document, f"基于该画像，{plan.intent_diagnosis}")

    _add_heading(document, "三、优化示例", level=2)
    for opportunity in plan.opportunities[:3]:
        _add_labelled_paragraph(
            document,
            "拟新增目标词：",
            opportunity.keyword,
            keep_with_next=True,
        )
        _add_bullet(
            document,
            "优化",
            opportunity.optimized_query,
            size=9.0,
            bold=True,
            keep_with_next=True,
        )
        _add_bullet(
            document,
            "改写目的",
            opportunity.rewrite_rationale,
            size=8.5,
            keep_with_next=True,
        )
        _add_bullet(
            document,
            "验证状态",
            "报价阶段未实测，实际效果以项目执行结果为准",
            size=8.5,
            after=2,
        )
    _add_text_paragraph(
        document,
        "*本节仅展示Query设计逻辑，不代表任何AI平台的实测推荐结果。",
        size=8.0,
        after=2,
        line=1.0,
    )

    _add_heading(document, "四、效果验证口径", level=2)
    _add_text_paragraph(
        document,
        "以下示例将在项目执行阶段，于相同平台、地域、账号与重复次数条件下开展优化前后"
        "对照测试；报价阶段不填入未经采样的推荐次数。",
        size=9.0,
        after=2,
        line=1.1,
    )
    _add_validation_table(document, plan.opportunities)
    _add_text_paragraph(
        document,
        f"结论：针对{brand_name}的实际优化效果，将以正式采样后的品牌提及率、推荐排名分布"
        "及厂商推荐次数为准。本报价单不将文案推演表述为实测提升数据。",
        size=9.0,
        bold=True,
        before=3,
        after=3,
        line=1.1,
    )
    _add_text_paragraph(
        document,
        "文献参考：Nelson, P. (1970). Information and Consumer Behavior. Journal of "
        "Political Economy, 78(2), 311–329.",
        size=8.0,
        color=_MUTED_GRAY,
        line=1.0,
    )
    _add_text_paragraph(
        document,
        "Darby, M. R., & Karni, E. (1973). Free Competition and the Optimal Amount of "
        "Fraud. Journal of Law and Economics, 16(1), 67–88.",
        size=8.0,
        line=1.0,
    )


def _selected_by_group(
    selected: Iterable[ExistingQueryVariants],
) -> OrderedDict[str, list[ExistingQueryVariants]]:
    ordered = sorted(selected, key=lambda row: int(row.source_id[1:]))
    groups: OrderedDict[str, list[ExistingQueryVariants]] = OrderedDict()
    for row in ordered:
        groups.setdefault(row.group, []).append(row)
    return groups


def _add_appendix_two(
    document: DocumentObject,
    plan: QuotationPlan,
    *,
    appendix_number: int,
    page_break_before: bool = True,
) -> None:
    if page_break_before:
        _add_page_break(document)
    _add_heading(
        document,
        f"附录{_appendix_number(appendix_number)} 原推广 Query 与变体构建说明",
        level=1,
        centered=True,
    )
    _add_heading(document, "核心检索问题库与语义变体", level=3)
    groups = _selected_by_group(plan.selected_queries)
    _add_text_paragraph(
        document,
        f"以下为围绕品牌{len(groups)}个核心业务方向设计的{len(plan.selected_queries)}条候选"
        "业务问题及其语义变体。变体A为正式换述，变体B为换角度表达，变体C为口语化表达。"
        "本次实际测试问题数和变体数必须在商务备注或合同中冻结；本附录的候选库数量不自动等于计费或验收数量。",
        size=9.0,
        after=2,
        line=1.1,
    )
    for group, rows in groups.items():
        _add_text_paragraph(
            document,
            group,
            size=10.5,
            bold=True,
            before=3,
            after=1,
            keep_with_next=True,
        )
        for row in rows:
            _add_text_paragraph(
                document,
                row.original,
                size=9.0,
                bold=True,
                after=0,
                line=1.0,
                keep_with_next=True,
            )
            _add_bullet(document, "A", row.variant_a, keep_with_next=True)
            _add_bullet(document, "B", row.variant_b, keep_with_next=True)
            _add_bullet(document, "C", row.variant_c, after=1)


def _add_appendix_three(
    document: DocumentObject,
    plan: QuotationPlan,
    *,
    appendix_number: int,
    page_break_before: bool = True,
) -> None:
    if page_break_before:
        _add_page_break(document)
    _add_heading(
        document,
        f"附录{_appendix_number(appendix_number)} 新增 Query 优化与语义变体全表",
        level=1,
        centered=True,
    )
    _add_text_paragraph(
        document,
        f"以下为{len(plan.opportunities)}条候选机会词、推荐型优化问句及其语义变体。"
        "变体A为正式换述，变体B为换角度表达，变体C为口语化表达——覆盖用户实际提问的"
        "多种表述方式。实际发布问题集必须由客户确认并在商务备注或合同中冻结；候选库数量不自动等于发文或验收数量。",
        size=9.0,
        after=2,
        line=1.1,
    )
    for row in plan.opportunities:
        _add_labelled_paragraph(
            document,
            "拟新增目标词：",
            row.keyword,
            size=9.0,
            keep_with_next=True,
        )
        _add_bullet(
            document,
            "优化",
            row.optimized_query,
            size=9.0,
            bold=True,
            keep_with_next=True,
        )
        _add_bullet(document, "A", row.variant_a, keep_with_next=True)
        _add_bullet(document, "B", row.variant_b, keep_with_next=True)
        _add_bullet(document, "C", row.variant_c, after=2)


def _add_query_appendices(
    document: DocumentObject,
    *,
    brand_name: str,
    configuration: QuotationConfiguration,
    plan: QuotationPlan,
    first_appendix_number: int,
    first_page_break: bool,
) -> None:
    """复用同一组 Query 章节，按制品位置决定首个分页和附录编号。"""
    appendix_number = first_appendix_number
    page_break_before = first_page_break
    if "ranking_test" in configuration.service_codes:
        _add_appendix_two(
            document,
            plan,
            appendix_number=appendix_number,
            page_break_before=page_break_before,
        )
        appendix_number += 1
        page_break_before = True
    if "content_publishing_pilot" in configuration.service_codes:
        _add_appendix_one(
            document,
            brand_name=brand_name,
            plan=plan,
            appendix_number=appendix_number,
            page_break_before=page_break_before,
        )
        appendix_number += 1
        _add_appendix_three(
            document,
            plan,
            appendix_number=appendix_number,
            page_break_before=True,
        )


def render_quotation_docx(
    *,
    brand_name: str,
    quote_date: date,
    configuration: QuotationConfiguration,
    plan: QuotationPlan | None,
) -> bytes:
    """用 legacy 空白文档版式渲染仅供内部回归的 DOCX bytes。"""
    document = Document()
    _configure_styles(document)
    _configure_page(document)
    package = PACKAGE_BY_CODE[configuration.package_code]
    artifact_title = {
        "complete": "报价单",
        "quote_table": "报价单表格",
        "query_appendix": "查询附件",
    }[configuration.artifact_kind]
    document.core_properties.title = f"{brand_name} {package.name}{artifact_title}"
    document.core_properties.subject = {
        "complete": "GEO 服务报价",
        "quote_table": "GEO 服务报价单表格",
        "query_appendix": "GEO 查询附件",
    }[configuration.artifact_kind]
    document.core_properties.author = ISSUER_COMPANY
    document.core_properties.comments = (
        f"{NON_FINAL_TEMPLATE_NOTICE}；由 legacy GEO 报价单生成服务生成；"
        "报价阶段未包含平台实测结果。"
    )
    # Core properties describe the quotation's effective document date, not the
    # wall-clock render time. Midnight keeps the date deterministic and avoids a
    # future local timestamp when OOXML serializes the naive value as UTC.
    document_timestamp = datetime.combine(quote_date, time())
    document.core_properties.created = document_timestamp
    document.core_properties.modified = document_timestamp

    _add_text_paragraph(
        document,
        NON_FINAL_TEMPLATE_NOTICE,
        size=10.5,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
        color=RGBColor(0xC0, 0x00, 0x00),
    )

    if configuration.artifact_kind == "quote_table":
        _add_cover(document, brand_name, quote_date, configuration)
    elif configuration.artifact_kind == "query_appendix":
        if plan is None:
            raise ValueError("query_appendix_plan_required")
        _add_heading(document, f"{brand_name} GEO 查询附件", level=1, centered=True)
        _add_text_paragraph(
            document,
            f"报价日期：{quote_date.isoformat()} · 套餐：{package.name}",
            size=9.0,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            after=5,
            color=_MUTED_GRAY,
        )
        _add_query_appendices(
            document,
            brand_name=brand_name,
            configuration=configuration,
            plan=plan,
            first_appendix_number=1,
            first_page_break=False,
        )
    else:
        _add_cover(document, brand_name, quote_date, configuration)
        _add_service_delivery_appendix(
            document,
            configuration=configuration,
            query_appendix_included=plan is not None,
        )
        if plan is not None:
            _add_query_appendices(
                document,
                brand_name=brand_name,
                configuration=configuration,
                plan=plan,
                first_appendix_number=2,
                first_page_break=True,
            )

    output = BytesIO()
    document.save(output)
    return _stable_docx_bytes(output.getvalue())
