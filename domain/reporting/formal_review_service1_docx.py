"""Client-oriented DOCX renderer for the service-1 V2 review delivery."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from domain.reporting.formal_review_docx import (
    FONT_BOLD,
    FONT_REGULAR,
    MUTED,
    NAVY,
    PALE_GREEN,
    WHITE,
    FormalDocument,
    _fmt_datetime,
    _fmt_percent,
    _fmt_ratio,
    _set_font,
    _shade,
    add_native_toc,
    build_report_code,
    is_formal_document,
)

SOURCE_SHARE_COLORS = (
    "1769AA",
    "29A3C6",
    "2B7A5A",
    "D99B2B",
    "8055A5",
    "D05B65",
    "5E86B3",
    "CCD8E2",
)


def _set_cell_width(cell: Any, width_mm: float) -> None:
    width_twips = int(width_mm * 1440 / 25.4)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Mm(width_mm)


def _table_layout_fixed(table: Any) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _style_specific_columns(table: Any, widths: tuple[float, ...]) -> None:
    _table_layout_fixed(table)
    for column_index, width in enumerate(widths):
        for cell in table.columns[column_index].cells:
            _set_cell_width(cell, width)


def _highlight_row(table: Any, row_index: int, *, fill: str = PALE_GREEN) -> None:
    if row_index >= len(table.rows):
        return
    for cell in table.rows[row_index].cells:
        _shade(cell, fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(NAVY)


def _caption(doc: FormalDocument, text: str) -> None:
    paragraph = doc.document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    _set_font(run, size=8)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(MUTED)


def _hyperlink(paragraph: Any, url: str, text: str = "打开原网页") -> None:
    """Add a bounded display label while preserving the complete URL target."""

    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Noto Sans CJK SC")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "14")
    language = OxmlElement("w:lang")
    language.set(qn("w:val"), "zh-CN")
    language.set(qn("w:eastAsia"), "zh-CN")
    properties.extend((fonts, color, underline, size, language))
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def _native_toc(doc: FormalDocument) -> None:
    """Backward-compatible local name; the implementation is shared by all services."""

    add_native_toc(doc)


def _source_share_donut(labels: list[str], counts: list[int], *, total: int, title: str) -> BytesIO:
    """Render a complete-composition donut: named top sites plus an ``其他`` slice."""

    width, height = 1500, 820
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    regular = ImageFont.truetype(str(FONT_REGULAR), 28)
    small = ImageFont.truetype(str(FONT_REGULAR), 24)
    bold = ImageFont.truetype(str(FONT_BOLD), 39)
    center_bold = ImageFont.truetype(str(FONT_BOLD), 42)
    draw.text((70, 34), title, fill="#1F2D3D", font=bold)

    total = max(int(total), 0)
    top_total = sum(max(int(value), 0) for value in counts)
    other = max(0, total - top_total)
    chart_labels = [*labels, "其他网站"]
    chart_counts = [*[max(int(value), 0) for value in counts], other]
    nonzero = [
        (label, count) for label, count in zip(chart_labels, chart_counts, strict=True) if count > 0
    ]
    if not nonzero:
        nonzero = [("暂无引用", 1)]
        denominator = 1
    else:
        denominator = max(total, sum(count for _, count in nonzero), 1)

    box = (95, 145, 735, 785)
    start = -90.0
    for index, (_, count) in enumerate(nonzero):
        sweep = count / denominator * 360.0
        color = f"#{SOURCE_SHARE_COLORS[min(index, len(SOURCE_SHARE_COLORS) - 1)]}"
        draw.pieslice(box, start=start, end=start + sweep, fill=color, outline="#FFFFFF", width=4)
        start += sweep
    draw.ellipse((270, 320, 560, 610), fill=f"#{WHITE}")
    center_text = f"{total:,}\n条全部引用"
    center_box = draw.multiline_textbbox(
        (0, 0), center_text, font=center_bold, spacing=8, align="center"
    )
    center_width = center_box[2] - center_box[0]
    center_height = center_box[3] - center_box[1]
    draw.multiline_text(
        (415 - center_width / 2, 465 - center_height / 2),
        center_text,
        fill="#12355B",
        font=center_bold,
        spacing=8,
        align="center",
    )

    legend_x, legend_y = 800, 150
    for index, (label, count) in enumerate(nonzero):
        y = legend_y + index * 76
        color = f"#{SOURCE_SHARE_COLORS[min(index, len(SOURCE_SHARE_COLORS) - 1)]}"
        draw.rounded_rectangle((legend_x, y + 6, legend_x + 30, y + 36), radius=5, fill=color)
        display = label if len(label) <= 24 else f"{label[:23]}…"
        draw.text((legend_x + 48, y), display, fill="#1F2D3D", font=regular)
        percent = count / denominator * 100
        draw.text(
            (legend_x + 48, y + 36),
            f"{count:,} 条 · {percent:.2f}%",
            fill=f"#{MUTED}",
            font=small,
        )

    output = BytesIO()
    image.save(output, format="PNG", optimize=True)
    output.seek(0)
    return output


def _add_source_share_donut(doc: FormalDocument, rows: list[dict[str, Any]], *, total: int) -> None:
    top_rows = rows[:7]
    picture = _source_share_donut(
        [str(row.get("sitename") or "（未知网站）") for row in top_rows],
        [int(row.get("count") or 0) for row in top_rows],
        total=total,
        title="全部引用的网站来源占比",
    )
    paragraph = doc.document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(picture, width=Cm(16.7))


def _anchor_boxes(anchor: dict[str, Any] | None) -> list[tuple[int, int, int, int]]:
    if not isinstance(anchor, dict):
        return []
    raw_boxes = anchor.get("bboxes")
    if not isinstance(raw_boxes, list):
        raw_boxes = [anchor.get("bbox")]
    boxes: list[tuple[int, int, int, int]] = []
    for raw in raw_boxes:
        if not isinstance(raw, list) or len(raw) != 4:
            continue
        try:
            x, y, width, height = (int(value) for value in raw)
        except (TypeError, ValueError):
            continue
        if width > 0 and height > 0:
            boxes.append((x, y, width, height))
    return boxes


def _mention_view(
    payload: bytes,
    platform: str,
    *,
    image_kind: str,
    anchor: dict[str, Any] | None,
) -> tuple[BytesIO, str, bool]:
    """Build one readable image with context and optional reviewed mention boxes.

    Official share images are clean customer-facing exports and therefore retain their
    full horizontal canvas.  Runtime screenshots are a fallback only; their navigation
    sidebar is removed before presentation.  A reviewed anchor controls only the vertical
    context and red boxes—absence of an anchor never creates a synthetic frame.
    """

    with Image.open(BytesIO(payload)) as source:
        image = source.convert("RGB")
    width, height = image.size
    privacy_left = 0
    if image_kind not in {"share_image", "answer_excerpt_screenshot"}:
        privacy_left = {
            "doubao": min(200, max(0, width // 7)),
            "yiyan": min(220, max(0, width // 6)),
            "deepseek": 0,
        }.get(platform, 0)
        if isinstance(anchor, dict) and isinstance(anchor.get("privacy_left"), int):
            privacy_left = max(0, min(int(anchor["privacy_left"]), width - 1))

    boxes = []
    for x, y, box_width, box_height in _anchor_boxes(anchor):
        x0 = max(privacy_left, min(x, width - 1))
        y0 = max(0, min(y, height - 1))
        x1 = max(x0 + 1, min(x + box_width, width))
        y1 = max(y0 + 1, min(y + box_height, height))
        boxes.append((x0, y0, x1, y1))

    if boxes:
        top_hit = min(box[1] for box in boxes)
        bottom_hit = max(box[3] for box in boxes)
        vertical_margin = max(150, int((bottom_hit - top_hit) * 1.7))
        crop_top = max(0, top_hit - vertical_margin)
        crop_bottom = min(height, bottom_hit + vertical_margin)
        minimum_height = min(height, 620)
        if crop_bottom - crop_top < minimum_height:
            center = (top_hit + bottom_hit) // 2
            crop_top = max(0, center - minimum_height // 2)
            crop_bottom = min(height, crop_top + minimum_height)
            crop_top = max(0, crop_bottom - minimum_height)
        method = str((anchor or {}).get("method") or "")
        if method.startswith("dom_"):
            label = "被提及位置（采集时 DOM 坐标）"
        elif method.startswith("ocr_"):
            label = "被提及位置（采集时 OCR 坐标）"
        else:
            label = "被提及位置（人工复核坐标）"
    else:
        bands = {
            "doubao": (0.13, 0.42),
            "yiyan": (0.12, 0.82),
            "deepseek": (0.66, 0.98),
        }
        top, bottom = bands.get(platform, (0.18, 0.78))
        crop_top = int(height * top)
        crop_bottom = max(crop_top + 1, int(height * bottom))
        label = "回答关键局部（当前无可复核像素坐标）"

    # A long share image squeezed to page height is not readable at 100% zoom.
    # Bound the vertical slice relative to its usable width; the byte-for-byte
    # long image remains in the evidence package and is never discarded.
    usable_width = max(1, width - privacy_left)
    maximum_readable_height = max(620, int(usable_width * 0.9))
    if crop_bottom - crop_top > maximum_readable_height:
        center = (crop_top + crop_bottom) // 2
        if boxes:
            center = (min(box[1] for box in boxes) + max(box[3] for box in boxes)) // 2
        crop_top = max(0, center - maximum_readable_height // 2)
        crop_bottom = min(height, crop_top + maximum_readable_height)
        crop_top = max(0, crop_bottom - maximum_readable_height)

    crop = image.crop((privacy_left, crop_top, width, crop_bottom))
    if boxes:
        draw = ImageDraw.Draw(crop)
        stroke = max(5, min(crop.size) // 115)
        for x0, y0, x1, y1 in boxes:
            draw.rectangle(
                (x0 - privacy_left, y0 - crop_top, x1 - privacy_left, y1 - crop_top),
                outline="#DC2626",
                width=stroke,
            )
    stream = BytesIO()
    crop.save(stream, format="PNG", optimize=True)
    stream.seek(0)
    return stream, label, bool(boxes)


def _add_screenshot_panel(
    doc: FormalDocument,
    payload: bytes,
    *,
    platform: str,
    number: int,
    image_kind: str,
    anchor: dict[str, Any] | None,
    figure_prefix: str = "6",
) -> None:
    stream, label, anchored = _mention_view(
        payload,
        platform,
        image_kind=image_kind,
        anchor=anchor,
    )
    with Image.open(stream) as view:
        width, height = view.size
    stream.seek(0)
    max_width_cm, max_height_cm = 16.4, 15.2
    width_cm = min(max_width_cm, max_height_cm * width / max(height, 1))
    paragraph = doc.document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(stream, width=Cm(max(5.5, width_cm)))
    source_label = (
        "平台官方分享图片"
        if image_kind == "share_image"
        else "干净回答证据图"
        if image_kind == "answer_excerpt_screenshot"
        else "运行页截图（历史退级）"
    )
    if anchored:
        coordinate_source = label.removeprefix("被提及位置（").removesuffix("）")
        anchor_note = f"红框依据{coordinate_source}"
    else:
        anchor_note = "未绘制未经复核的定位框"
    _caption(
        doc,
        f"图 {figure_prefix}-{number}  {source_label} · {label}；{anchor_note}。"
        "本页仅展示可读局部，完整长图见证据包；红框不改变原始文字内容。",
    )


def _toc(doc: FormalDocument) -> None:
    _native_toc(doc)
    doc.page_break()


def _group_insight(group: dict[str, Any], snapshot: dict[str, Any]) -> str:
    target = snapshot.get("target") or {}
    rate = float(target.get("appearance_rate") or 0)
    avg_rank = target.get("avg_rank")
    if rate >= 80:
        state = "该场景下品牌覆盖较强"
    elif rate >= 50:
        state = "该场景已有基础覆盖，但仍存在空白问题或平台"
    else:
        state = "该场景是当前主要可见性短板"
    rank_text = f"，提及时平均位次为 {avg_rank}" if avg_rank is not None else ""
    return f"{group['title']}：{state}{rank_text}。"


def _metrics_explanation(doc: FormalDocument, *, target_brand: str) -> None:
    doc.heading("2.3 指标口径", level=2)
    doc.table(
        ["指标", "计算方式", "客户应如何理解"],
        [
            (
                "品牌提及率",
                f"提及{target_brand}的回答数 ÷ 全部主样本回答数",
                "回答中是否出现品牌，不等同于排在前列",
            ),
            (
                "综合品牌排名",
                "按出现次数 ÷ 提及时平均位次形成综合分后排序",
                "兼顾覆盖和位置；分数只用于本批品牌间排序",
            ),
            (
                "平均推荐位次",
                "仅在已提及样本中，对品牌 1-based 顺序求平均",
                "数值越小越靠前；未提及回答不硬塞入平均值",
            ),
            (
                "Top1/3/5 出现率",
                "位次不超过 N 的回答数 ÷ 全部主样本回答数",
                "本报告主表统一以全部回答为分母",
            ),
            (
                "带引用回答覆盖",
                "至少捕获 1 条信源 URL 的回答数 ÷ 全部回答数",
                "只表示平台展示了引用，不代表引用支持每个结论",
            ),
        ],
        widths=(31, 75, 66),
        font_size=8,
    )


def _repetition_copy(current: int, required: int, formal_answer_count: int) -> dict[str, str]:
    """Return internally consistent sampling copy for complete/incomplete matrices."""

    if current >= required:
        return {
            "callout": (
                f"已按冻结矩阵完成报价要求的 {required} 次独立重复；"
                f"正式签发前复核弱问题在 {current} 次观测中是否稳定复现。"
            ),
            "region_finding": (
                f"当前每单元已完成 {current} 次独立观测；该差异仅描述本窗口，不外推为长期稳定差异。"
            ),
            "design_requirement": f"本窗口已完成 {current} 次独立重复",
            "sample_requirement": f"已达到报价主样本 {formal_answer_count} 条",
            "region_review": (
                f"基于每单元 {current} 次独立观测，该差异仅作为本窗口结果，不外推为长期稳定差异。"
            ),
            "signoff_intro": (
                "本节不是向客户内容运营团队分派任务，而是列出评测执行方在正式报告"
                "签发前必须完成的完整性确认、异常复核与一致性检查。"
                "具体内容优化方案应在正式评测结果冻结后另行制定。"
            ),
            "stability_current": f"每个问题×平台×地域均已完成 {current} 次独立观测",
            "stability_action": f"按冻结矩阵核对 {current} 次观测完整性，并复核异常差异",
            "weak_question_action": (
                f"逐条复核 {current} 次独立观测，并在后续窗口保持原题不变以监测稳定性"
            ),
        }
    return {
        "callout": ("正式签发前由评测执行方按冻结矩阵补齐重复采样，并复核弱问题是否稳定复现。"),
        "region_finding": f"当前每单元仅 {current} 次观测，该差异只能作为补采关注点。",
        "design_requirement": f"正式应完成 {required} 次独立重复",
        "sample_requirement": f"按报价补采完成后应为 {formal_answer_count} 条",
        "region_review": (
            f"由于当前每单元只有 {current} 次观测，应在正式重复采样后再判断差异是否稳定。"
        ),
        "signoff_intro": (
            "本节不是向客户内容运营团队分派任务，而是列出评测执行方在正式报告签发前"
            "必须完成的补采、复核与一致性检查。具体内容优化方案应在正式评测结果冻结后"
            "另行制定。"
        ),
        "stability_current": f"当前每个问题×平台×地域只有 {current} 次观测",
        "stability_action": f"由评测执行方补齐至 {required} 次独立重复，并按同一矩阵重算",
        "weak_question_action": "原题保持不变完成复测，并复核同义改写是否造成稳定差异",
    }


def _add_representative_overview(
    doc: FormalDocument,
    rows: list[dict[str, Any]],
    screenshots: dict[str, bytes],
) -> None:
    """Render the overview with real thumbnails instead of evidence-status prose."""

    doc.table(
        ["组", "业务问题", "平台/地域", "品牌位次", "引用", "代表图片"],
        [
            (
                row["display_number"],
                row["group_title"],
                f"{row['platform_label']}/{row['region']}",
                f"第 {row['target_rank']} 位" if row["target_rank"] is not None else "未提及",
                f"{row['citation_count']} 条",
                "图片待补",
            )
            for row in rows
        ],
        widths=(9, 42, 31, 23, 17, 50),
        font_size=7.4,
    )
    table = doc.document.tables[-1]
    for row_index, item in enumerate(rows, 1):
        cell = table.cell(row_index, 5)
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        payload = screenshots.get(str(item["answer_pub_id"]))
        if not payload:
            run = paragraph.add_run("图片待补")
            _set_font(run, size=7)
            run.font.color.rgb = RGBColor.from_string(MUTED)
            continue
        image_kind = str(item.get("preferred_image_kind") or "answer_screenshot")
        stream, _, _ = _mention_view(
            payload,
            str(item["platform"]),
            image_kind=image_kind,
            anchor=item.get("answer_anchor"),
        )
        with Image.open(stream) as view:
            image_width, image_height = view.size
        stream.seek(0)
        width_cm = min(4.3, 3.2 * image_width / max(image_height, 1))
        paragraph.add_run().add_picture(stream, width=Cm(max(2.7, width_cm)))
        label_paragraph = cell.add_paragraph()
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_paragraph.paragraph_format.space_after = Pt(0)
        label = (
            "官方分享图片"
            if image_kind == "share_image"
            else "干净回答证据图"
            if image_kind == "answer_excerpt_screenshot"
            else "运行页截图（历史退级）"
        )
        label_run = label_paragraph.add_run(label)
        _set_font(label_run, size=6.4)
        label_run.font.color.rgb = RGBColor.from_string(MUTED)


def _answer_evidence_page(
    doc: FormalDocument,
    item: dict[str, Any],
    screenshot: bytes | None,
    *,
    target_brand: str,
    start_new_page: bool = True,
) -> None:
    if start_new_page:
        doc.page_break()
    number = int(item["display_number"])
    target_rank = item.get("target_rank")
    brand_performance = (
        f"回答品牌序列中提及 {target_brand}，推荐位次第 {target_rank} 位"
        if target_rank is not None
        else f"回答品牌序列中未提及 {target_brand}"
    )
    image_kind = str(item.get("preferred_image_kind") or "missing")
    image_status_label = (
        "官方分享图片"
        if image_kind == "share_image"
        else "干净回答证据图"
        if image_kind == "answer_excerpt_screenshot"
        else "运行页截图（历史退级）"
    )
    doc.heading(f"6.{number} {item['group_title']} · 代表回答")
    doc.table(
        ["字段", "实测内容"],
        [
            ("问题", item["question"]),
            (
                "采样环境",
                f"{item['platform_label']} · {item['region']} · 深度思考 · "
                f"{_fmt_datetime(item['capture_time'])}",
            ),
            ("品牌表现", brand_performance),
            (
                "证据状态",
                f"代表图片：{'已载入' if screenshot else '当前未载入'}（"
                f"{image_status_label}）；"
                f"信源 URL：{item['citation_count']} 条；正文：{item['response_chars']} 字符",
            ),
        ],
        widths=(28, 144),
        font_size=8.2,
    )
    if screenshot:
        _add_screenshot_panel(
            doc,
            screenshot,
            platform=str(item["platform"]),
            number=number,
            image_kind=str(item.get("preferred_image_kind") or "answer_screenshot"),
            anchor=item.get("answer_anchor"),
        )
    else:
        doc.callout(
            "截图状态", "该代表样本未能从证据存储载入截图，正式签发前须补齐。", kind="warning"
        )

    doc.heading("回答原文摘录", level=2)
    doc.paragraph(str(item["answer_excerpt"]))
    doc.callout(
        "读图说明",
        f"本页可直接核对平台、问题文本，以及回答中是否出现 {target_brand}。"
        "品牌顺序来自完整回答的结构化抽取；代表图优先使用平台官方分享图片，"
        "其后使用干净回答证据图；只有前两者缺失时才使用已去除会话侧栏的历史运行页截图。",
        kind="success",
    )
    if item.get("citations"):
        doc.heading("回答内前置信源 URL", level=2)
        citation_rows = [
            (
                citation.get("ordinal") or "—",
                citation.get("host") or "（未知）",
                citation.get("title") or "（未捕获标题）",
                "打开原网页" if citation.get("url") else "（未捕获 URL）",
            )
            for citation in item["citations"][:3]
        ]
        doc.table(
            ["序", "网站域名", "标题", "网页链接"],
            citation_rows,
            widths=(9, 31, 54, 78),
            font_size=6.8,
        )
        citation_table = doc.document.tables[-1]
        for row_index, citation in enumerate(item["citations"][:3], 1):
            url = str(citation.get("url") or "")
            if not url:
                continue
            paragraph = citation_table.cell(row_index, 3).paragraphs[0]
            paragraph.clear()
            _hyperlink(paragraph, url)
        doc.numbered(
            [
                f"本页为可读性展示前 {len(citation_rows)} 条；该回答共捕获 "
                f"{item['citation_count']} 条信源 URL，全部链接见附录 C.{number}。",
            ]
        )
    else:
        doc.callout(
            "引用披露",
            "该回答未捕获信源 URL。报告不会据此推断其事实来源。",
            kind="warning",
        )


def render_service1_v2_docx(
    facts: dict[str, Any], *, screenshots: dict[str, bytes] | None = None
) -> bytes:
    """Render the complete service-1 V2 review report from real measurement facts."""

    service1 = facts["service1"]
    delivery = service1.get("delivery_v2")
    if not isinstance(delivery, dict):
        raise ValueError("service1.delivery_v2_missing")
    screenshots = screenshots or {}
    overall = service1["overall"]
    target = overall.get("target") or {}
    scope = delivery["scope"]
    target_brand = str(facts["target_brand"])
    model_labels = {"doubao": "豆包", "deepseek": "DeepSeek", "yiyan": "文心一言"}
    selected_groups = [
        group for group in service1["candidate_groups"] if group["selected_for_main_report"]
    ]
    selected_groups.sort(key=lambda group: int(group["index"]))
    group_results = [
        (group, service1["by_group"][group["id"]].get("target") or {}) for group in selected_groups
    ]
    strongest_group = max(
        group_results, key=lambda item: float(item[1].get("appearance_rate") or 0)
    )
    weakest_group = min(group_results, key=lambda item: float(item[1].get("appearance_rate") or 0))
    model_results = [
        (model, snapshot.get("target") or {}) for model, snapshot in service1["by_model"].items()
    ]
    region_results = list(delivery["by_region"].items())
    required_repetitions = int(service1["quotation_required_repetitions_per_cell"])
    current_repetitions = int(scope["current_repetitions"])
    formal_answer_count = int(
        service1.get("expected_formal_answers")
        or (
            scope["answers"] * required_repetitions / current_repetitions
            if current_repetitions
            else 0
        )
    )
    repetition_copy = _repetition_copy(
        current_repetitions,
        required_repetitions,
        formal_answer_count,
    )

    doc = FormalDocument(
        title="品牌 GEO 推荐结果评测报告",
        subtitle="服务 1 · AI 推荐可见性、竞品格局与证据明细",
        facts=facts,
    )
    doc.cover(report_code=build_report_code(facts, service_number=1, version="V2"))
    _toc(doc)

    doc.heading("1. 执行摘要")
    doc.kpis(
        [
            (
                "品牌提及率",
                _fmt_percent(target.get("appearance_rate")),
                f"{target.get('mentions', 0)}/{overall['answers']} 条",
            ),
            (
                "综合品牌排名",
                f"第 {target.get('overall_rank') or '—'} 位",
                f"本批共观察 {scope['brands_observed']} 个品牌",
            ),
            ("平均推荐位次", str(target.get("avg_rank") or "—"), "仅在提及样本内"),
            (
                "Top 3 出现率",
                _fmt_percent((target.get("top_rates") or {}).get("3", {}).get("of_total")),
                "分母=全部主样本",
            ),
        ]
    )
    doc.paragraph(
        f"本次审阅口径覆盖 {scope['selected_groups']} 组业务问题、{scope['questions']} 个问题文本、"
        f"{scope['platforms']} 个 AI 平台和 {scope['regions']} 个地域，"
        f"共形成 {scope['answers']} 条平衡主样本。"
        f"{target_brand} 在其中被提及 {target.get('mentions', 0)} 次，提及率为 "
        f"{float(target.get('appearance_rate') or 0):.2f}%，综合品牌排名第 "
        f"{target.get('overall_rank') or '—'} 位。"
    )
    doc.callout(
        "一句话结论",
        f"{target_brand} 当前综合品牌排名第 {target.get('overall_rank') or '—'} 位；"
        f"{strongest_group[0]['title']}提及率最高（"
        f"{_fmt_percent(strongest_group[1].get('appearance_rate'))}），"
        f"{weakest_group[0]['title']}最低（"
        f"{_fmt_percent(weakest_group[1].get('appearance_rate'))}）。" + repetition_copy["callout"],
        kind="success",
    )
    doc.table(
        ["当前数据状态", "数值", "审阅判断"],
        [
            (
                "品牌抽取成功",
                _fmt_ratio(scope["extract_ok"], scope["answers"]),
                "可用于本批排名统计",
            ),
            (
                "回答截图留存",
                _fmt_ratio(scope["answer_screenshots"], scope["answers"]),
                (
                    "主样本截图覆盖完整"
                    if scope["answer_screenshots"] == scope["answers"]
                    else f"仍缺 {scope['answers'] - scope['answer_screenshots']} 条"
                ),
            ),
            (
                "带引用回答",
                _fmt_ratio(scope["answers_with_citation"], scope["answers"]),
                f"共捕获 {scope['citation_references']} 条引用记录",
            ),
            (
                "重复次数",
                f"{scope['current_repetitions']} 次/单元",
                (
                    f"报价要求 {required_repetitions} 次；"
                    f"本窗口实际为 {current_repetitions} 次/单元"
                    if is_formal_document(facts)
                    else f"报价要求 {required_repetitions} 次；当前仍为预正式审阅"
                ),
            ),
        ],
        widths=(42, 34, 96),
        font_size=8.4,
    )

    doc.heading("1.1 核心发现", level=2)
    findings = [
        _group_insight(group, service1["by_group"][group["id"]]) for group in selected_groups
    ]
    model_finding = (
        "平台差异："
        + "；".join(
            f"{model_labels.get(model, model)}提及率 "
            f"{_fmt_percent(result.get('appearance_rate'))}"
            + (
                f"、提及时平均位次 {result.get('avg_rank')}"
                if result.get("avg_rank") is not None
                else "、本批未形成提及位次"
            )
            for model, result in model_results
        )
        + "。"
    )
    region_rates = [float(row.get("appearance_rate") or 0) for _, row in region_results]
    region_gap = max(region_rates) - min(region_rates) if region_rates else 0
    region_finding = (
        "地域差异："
        + "；".join(
            f"{region}提及率 {_fmt_percent(row.get('appearance_rate'))}"
            for region, row in region_results
        )
        + f"；本批最高与最低相差 {region_gap:.2f} 个百分点。"
        + repetition_copy["region_finding"]
    )
    findings.extend([model_finding, region_finding])
    doc.bullets(findings)
    doc.page_break()

    doc.heading("2. 评测设计与范围")
    doc.table(
        ["维度", "本次审阅口径", "正式交付要求"],
        [
            (
                "业务问题",
                f"{len(service1['candidate_groups'])} 组候选中选取证据最完整的 "
                f"{len(selected_groups)} 组；主文共 {scope['questions']} 个问题文本",
                "保留全部候选与选择记录",
            ),
            (
                "AI 平台",
                "、".join(model_labels.get(model, model) for model in service1["primary_models"])
                + f"，模式={service1['primary_mode']}",
                "固定平台、模式和账号",
            ),
            (
                "地域",
                "、".join(service1["primary_regions"]) + "浏览器地域采样",
                "保留账号/出口地域审计台账",
            ),
            (
                "重复",
                f"当前 {current_repetitions} 次/问题×平台×地域单元",
                repetition_copy["design_requirement"],
            ),
            (
                "主样本",
                f"{scope['questions']}×{scope['platforms']}×{scope['regions']}×"
                f"{current_repetitions} = {scope['answers']} 条",
                repetition_copy["sample_requirement"],
            ),
            (
                "数据窗口",
                f"{facts['window']['start']} 至 {facts['window']['end']}",
                "签发时锁定正式采集窗口",
            ),
        ],
        widths=(29, 84, 59),
        font_size=8.2,
    )

    doc.heading("2.1 候选问题组与选择记录", level=2)
    doc.paragraph(
        "为避免只展示对目标品牌有利的问题，候选组选择只使用证据完整度：45% 单元覆盖、"
        "20% 品牌抽取覆盖、20% 带引用回答覆盖、10% 回答完整度、5% 平台/地域广度。"
        "评分不读取目标品牌提及、位次或竞品结果。"
    )
    doc.table(
        ["名次", "候选问题组", "完整度", "有效单元", "带引用", "用途"],
        [
            (
                group["selection_rank"],
                group["title"],
                f"{group['selection_score']:.2f}",
                _fmt_ratio(group["observed_cells"], group["expected_cells"]),
                _fmt_ratio(group["answers_with_citation"], group["observed_cells"]),
                "主文" if group["selected_for_main_report"] else "附录",
            )
            for group in sorted(
                service1["candidate_groups"], key=lambda value: value["selection_rank"]
            )
        ],
        widths=(13, 62, 23, 25, 25, 24),
        font_size=8.1,
    )

    doc.heading("2.2 主文问题清单", level=2)
    question_list_rows = []
    for group in selected_groups:
        for question_index, question in enumerate(group["questions"], 1):
            question_list_rows.append(
                (f"{group['index']}-{question_index}", group["title"], question)
            )
    doc.table(
        ["编号", "业务问题组", "实际提问文本"],
        question_list_rows,
        widths=(14, 49, 109),
        font_size=8,
    )
    _metrics_explanation(doc, target_brand=target_brand)

    doc.heading("3. 品牌可见性详细结果")
    doc.heading("3.1 目标品牌与主要竞品", level=2)
    competitor_rows = [(target_brand, target)] + [
        (row.get("brand_input") or row.get("brand"), row)
        for row in overall.get("competitors") or []
        if row
    ]
    doc.chart(
        [str(name) for name, _ in competitor_rows],
        [float(row.get("appearance_rate") or 0) for _, row in competitor_rows],
        title=f"目标品牌与重点对比品牌提及率（{scope['answers']} 条主样本）",
    )
    doc.table(
        ["品牌", "综合位次", "提及", "提及率", "平均位次", "Top1", "Top3", "Top5"],
        [
            (
                name,
                row.get("overall_rank") or "—",
                row.get("mentions") or 0,
                _fmt_percent(row.get("appearance_rate")),
                row.get("avg_rank") or "—",
                _fmt_percent((row.get("top_rates") or {}).get("1", {}).get("of_total")),
                _fmt_percent((row.get("top_rates") or {}).get("3", {}).get("of_total")),
                _fmt_percent((row.get("top_rates") or {}).get("5", {}).get("of_total")),
            )
            for name, row in competitor_rows
        ],
        widths=(27, 18, 16, 23, 23, 21, 21, 21),
        font_size=7.5,
    )
    full_ranking = delivery["full_brand_ranking"]
    target_position = next(
        (index for index, row in enumerate(full_ranking) if row["is_target"]), None
    )
    if target_position is None:
        competitor_readout = f"{target_brand} 未进入本批综合品牌排名。"
    elif target_position == 0:
        competitor_readout = f"{target_brand} 当前位列本批综合品牌排名第 1 位。"
    else:
        above = full_ranking[target_position - 1]
        competitor_readout = (
            f"{target_brand} 综合排名第 {target.get('overall_rank')} 位；紧邻其上的品牌为"
            f"{above['brand']}（提及率 {_fmt_percent(above['appearance_rate'])}）。"
            f"{target_brand} 的 Top1 出现率为 "
            f"{_fmt_percent((target.get('top_rates') or {}).get('1', {}).get('of_total'))}。"
        )
    doc.numbered([competitor_readout])

    doc.heading("3.2 分平台表现", level=2)
    doc.table(
        ["平台", "样本", "提及", "提及率", "平均位次", "Top1", "Top3", "带引用"],
        [
            (
                model_labels.get(model, model),
                snapshot["answers"],
                (snapshot.get("target") or {}).get("mentions") or 0,
                _fmt_percent((snapshot.get("target") or {}).get("appearance_rate")),
                (snapshot.get("target") or {}).get("avg_rank") or "—",
                _fmt_percent(
                    ((snapshot.get("target") or {}).get("top_rates") or {})
                    .get("1", {})
                    .get("of_total")
                ),
                _fmt_percent(
                    ((snapshot.get("target") or {}).get("top_rates") or {})
                    .get("3", {})
                    .get("of_total")
                ),
                _fmt_ratio(snapshot["answers_with_citation"], snapshot["answers"]),
            )
            for model, snapshot in service1["by_model"].items()
        ],
        widths=(27, 17, 17, 24, 25, 21, 21, 20),
        font_size=7.7,
    )

    doc.heading("3.3 分地域表现", level=2)
    doc.table(
        ["地域", "样本", "提及", "提及率", "平均位次", "Top1", "Top3", "引用覆盖"],
        [
            (
                region,
                row["answers"],
                row["mentions"],
                _fmt_percent(row["appearance_rate"]),
                row["avg_rank"] or "—",
                _fmt_percent((row["top_rates"].get("1") or {}).get("of_total")),
                _fmt_percent((row["top_rates"].get("3") or {}).get("of_total")),
                _fmt_ratio(row["answers_with_citation"], row["answers"]),
            )
            for region, row in delivery["by_region"].items()
        ],
        widths=(25, 18, 18, 24, 25, 21, 21, 20),
        font_size=7.8,
    )
    doc.numbered(
        [
            "；".join(
                f"{region} {row['answers']} 条、提及率 {_fmt_percent(row['appearance_rate'])}"
                for region, row in region_results
            )
            + "。",
            f"最高与最低相差 {region_gap:.2f} 个百分点；" + repetition_copy["region_review"],
        ]
    )

    doc.heading("3.4 分业务问题组表现", level=2)
    group_rows = []
    for group in selected_groups:
        snapshot = service1["by_group"][group["id"]]
        group_target = snapshot.get("target") or {}
        group_rows.append(
            (
                group["title"],
                snapshot["answers"],
                group_target.get("mentions") or 0,
                _fmt_percent(group_target.get("appearance_rate")),
                group_target.get("avg_rank") or "—",
                _fmt_percent((group_target.get("top_rates") or {}).get("1", {}).get("of_total")),
                _fmt_percent((group_target.get("top_rates") or {}).get("3", {}).get("of_total")),
                _fmt_ratio(snapshot["answers_with_citation"], snapshot["answers"]),
            )
        )
    doc.table(
        ["业务问题组", "样本", "提及", "提及率", "平均位次", "Top1", "Top3", "带引用"],
        group_rows,
        widths=(47, 17, 17, 23, 23, 20, 20, 20),
        font_size=7.6,
    )
    doc.chart(
        [group["title"] for group in selected_groups],
        [
            float(
                (service1["by_group"][group["id"]].get("target") or {}).get("appearance_rate") or 0
            )
            for group in selected_groups
        ],
        title="不同业务问题组的品牌提及率",
    )

    doc.heading("3.5 逐题结果", level=2)
    doc.paragraph(
        f"下表将每个实际问题的 {scope['platforms'] * scope['regions'] * current_repetitions} "
        f"条回答（{scope['platforms']} 平台×{scope['regions']} 地域×"
        f"{current_repetitions} 次）单独计算。"
        "它比总体平均更能定位品牌在哪种自然语言表达中被召回或遗漏。"
    )
    doc.table(
        ["题号", "实际问题", "提及", "提及率", "平均位次", "Top1/3/5", "引用覆盖"],
        [
            (
                f"{row['group_index']}-{row['question_index']}",
                row["question"],
                _fmt_ratio(row["mentions"], row["answers"]),
                _fmt_percent(row["appearance_rate"]),
                row["avg_rank"] or "—",
                f"{row['top1']}/{row['top3']}/{row['top5']}",
                _fmt_ratio(row["answers_with_citation"], row["answers"]),
            )
            for row in delivery["question_rows"]
        ],
        widths=(13, 76, 18, 23, 21, 22, 20),
        font_size=7.1,
    )
    zero_questions = [row["question"] for row in delivery["question_rows"] if row["mentions"] == 0]
    strongest = max(delivery["question_rows"], key=lambda row: float(row["appearance_rate"]))
    doc.numbered(
        [
            f"覆盖最完整的题目之一是“{strongest['question']}”"
            f"（{_fmt_percent(strongest['appearance_rate'])}）。",
            (
                "当前零提及问题包括：“"
                + "”；“".join(zero_questions)
                + "”。这些问题应作为内容补强和正式复测的 P0 清单。"
                if zero_questions
                else "本批主问题没有零提及项。"
            ),
        ]
    )

    doc.heading("3.6 推荐位次分布", level=2)
    doc.chart(
        [str(row["label"]) for row in delivery["rank_distribution"]],
        [float(row["count"]) for row in delivery["rank_distribution"]],
        title=f"{target_brand}在 {scope['answers']} 条主样本中的位次分布",
        suffix="条",
    )
    mentioned_answers = sum(
        int(row["count"]) for row in delivery["rank_distribution"] if row["label"] != "未提及"
    )

    def distribution_insight(row: dict[str, Any]) -> str:
        count = int(row["count"])
        if count == 0:
            return "本批没有落入该区间的回答。"
        label = str(row["label"])
        if label == "第 1 位":
            return "回答把品牌置于首位；这是本表唯一能直接证明“首推”的区间。"
        if label == "第 2–3 位":
            return "已进入优先比较范围，但仍被至少一个品牌排在前面。"
        if label == "第 4–5 位":
            return "进入候选清单中段，有可见性但不属于优先推荐。"
        if label == "第 6–10 位":
            return "仅处于长名单后段；能被检索到，但决策优先级较弱。"
        if label == "第 11 位以后":
            return "虽被提及但位置很靠后，主要反映基础收录而非有效推荐。"
        return (
            f"完全未进入品牌序列，是当前覆盖缺口；相当于全部样本的 "
            f"{count / max(scope['answers'], 1) * 100:.2f}%。"
        )

    doc.table(
        ["位次区间", "回答数", "占全部样本", "解读"],
        [
            (
                row["label"],
                row["count"],
                _fmt_percent(float(row["count"]) / scope["answers"] * 100),
                distribution_insight(row),
            )
            for row in delivery["rank_distribution"]
        ],
        widths=(42, 25, 38, 67),
        font_size=8.2,
    )
    doc.numbered(
        [
            f"共 {mentioned_answers}/{scope['answers']} 条回答提及 {target_brand}；"
            "位次分布不能把“提及”自动解释成“推荐”。",
            "第 1 位表示直接首推；第 2–5 位表示进入优先或中段候选；第 6 位以后更接近长名单曝光。",
            "“未提及”与“提及但靠后”是两种不同短板：前者要解决召回，后者要解决排序与证据竞争力。",
        ]
    )

    doc.page_break()
    doc.heading("4. 竞品品牌格局")
    doc.paragraph(
        f"主样本中共观察到 {scope['brands_observed']} 个规范化品牌。"
        "综合排名按“出现次数 ÷ 提及时平均位次”排序，"
        "因此覆盖广且位置靠前的品牌会取得更高名次。下表展示 Top 20；完整排名在附录 A。"
    )
    top20 = delivery["full_brand_ranking"][:20]
    doc.table(
        ["名次", "品牌", "提及", "提及率", "平均位次", "综合分"],
        [
            (
                row["rank"],
                row["brand"],
                row["occurrences"],
                _fmt_percent(row["appearance_rate"]),
                row["avg_rank"],
                f"{float(row['score']):.3f}",
            )
            for row in top20
        ],
        widths=(15, 57, 20, 27, 28, 25),
        font_size=7.5,
    )
    target_row_index = next((index for index, row in enumerate(top20, 1) if row["is_target"]), None)
    if target_row_index is not None:
        # The most recently added table is the Top-20 table; retain the target cue.
        _highlight_row(doc.document.tables[-1], target_row_index)

    doc.heading("5. 信源结构分析")
    sources = delivery["sources"]
    doc.kpis(
        [
            ("引用记录", str(sources["total"]), f"来自 {scope['answers_with_citation']} 条回答"),
            ("唯一 URL", str(sources["unique_urls"]), "按规范化 URL 去重"),
            ("Top3 集中度", _fmt_percent(sources["top3_concentration"]), "按信源记录数计算"),
            (
                "无引用回答",
                str(scope["answers"] - scope["answers_with_citation"]),
                "不推断其信息来源",
            ),
        ]
    )
    source_top = list(sources.get("sources") or [])[:12]
    _add_source_share_donut(doc, source_top, total=int(sources["total"]))
    doc.table(
        ["名次", "引用网站（域名）", "引用次数", "占全部引用", "平均引用位置", "首位引用"],
        [
            (
                row["rank"],
                row["sitename"],
                row["count"],
                _fmt_percent(row["percent"]),
                f"{float(row['avg_ordinal']):.2f}" if row.get("avg_ordinal") is not None else "—",
                row.get("first_position_count") or 0,
            )
            for row in source_top
        ],
        widths=(12, 58, 25, 29, 28, 20),
        font_size=7.8,
    )
    doc.numbered(
        [
            "饼图按全部引用记录计算：单列引用量最高的 7 个网站，其余网站合并为“其他网站”，"
            "各扇区合计为 100%。",
            "“引用网站（域名）”指 AI 回答列出的 URL 所属网站，例如 example.com；"
            "该名称直接说明统计对象，也不把抓取文档数混入回答引用统计。",
            "引用次数按回答中的 URL 引用条目累计；同一 URL 被不同回答重复引用时会重复计数。",
            "平均引用位置是该网站各次引用序号的算术平均，数值越小表示通常越靠前；"
            "不再使用跨回答累加的倒数序号总分，因此不会随样本量增加而产生难以解释的数值。",
            "首位引用表示该网站在多少条引用记录中位于第 1 位。上述位置指标只描述"
            "列表顺序，不代表权威性、事实正确性或内容已被回答采纳。",
        ]
    )

    doc.heading("6. 代表回答与证据链")
    representative_count = len(delivery["representative_answers"])
    representative_platforms = "、".join(
        dict.fromkeys(str(row["platform_label"]) for row in delivery["representative_answers"])
    )
    doc.paragraph(
        f"以下 {representative_count} 组各展示一条真实回答。为兼顾证据可读性和平台多样性，"
        f"代表页实际覆盖 {representative_platforms}；每组优先选择官方分享图片，"
        "分享图片缺失时再按运行页截图可用性、引用和正文完整度选择。"
        "选择不读取品牌是否提及或位次。"
        f"代表页不参与指标计算，其余 {scope['answers'] - representative_count} 条样本"
        "仍进入统计和附录。"
    )
    _add_representative_overview(
        doc,
        list(delivery["representative_answers"]),
        screenshots,
    )
    doc.numbered(
        [
            delivery["evidence_policy"],
            "报告只展示上下文充分的“被提及位置”；红框用于强调人工复核的品牌文字，"
            "不会用整图边框冒充命中标注。",
        ]
    )
    for item in delivery["representative_answers"]:
        _answer_evidence_page(
            doc,
            item,
            screenshots.get(str(item["answer_pub_id"])),
            target_brand=target_brand,
            start_new_page=True,
        )
    doc.page_break()

    doc.heading("7. 评测结论与正式复测要求")
    doc.heading("7.1 已由当前数据直接支持的结论", level=2)
    top_rates = target.get("top_rates") or {}
    group_conclusion = "；".join(
        f"{group['title']} {_fmt_percent(result.get('appearance_rate'))}"
        for group, result in group_results
    )
    model_conclusion = "；".join(
        f"{model_labels.get(model, model)}提及率 "
        f"{_fmt_percent(result.get('appearance_rate'))}、平均位次 "
        f"{result.get('avg_rank') if result.get('avg_rank') is not None else '未形成'}"
        for model, result in model_results
    )
    zero_question_rows = [row for row in delivery["question_rows"] if int(row["mentions"]) == 0]
    zero_question_conclusion = (
        f"逐题：本批有 {len(zero_question_rows)} 个零提及问题："
        + "；".join(f"“{row['question']}”" for row in zero_question_rows)
        + "。"
        if zero_question_rows
        else "逐题：本批没有零提及问题；仍应关注提及率最低的问题。"
    )
    doc.bullets(
        [
            f"总体：{target_brand} 在 {scope['answers']} 条主样本中提及 "
            f"{target.get('mentions', 0)} 次，提及率 "
            f"{_fmt_percent(target.get('appearance_rate'))}，综合排名第 "
            f"{target.get('overall_rank') or '—'} 位；Top1/Top3/Top5 出现率分别为 "
            f"{_fmt_percent((top_rates.get('1') or {}).get('of_total'))} / "
            f"{_fmt_percent((top_rates.get('3') or {}).get('of_total'))} / "
            f"{_fmt_percent((top_rates.get('5') or {}).get('of_total'))}。",
            f"业务场景：{group_conclusion}。不同问题组的结果不应合并解释为单一品牌能力。",
            f"平台：{model_conclusion}。",
            zero_question_conclusion,
            f"证据：{scope['answer_screenshots']}/{scope['answers']} 条主样本有回答截图，"
            f"{scope['answers_with_citation']}/{scope['answers']} 条带引用，"
            f"共捕获 {scope['citation_references']} 条引用记录。",
        ]
    )
    doc.heading("7.2 正式复测与签发检查", level=2)
    doc.paragraph(repetition_copy["signoff_intro"])
    question_rows_by_rate = sorted(
        delivery["question_rows"], key=lambda row: (float(row["appearance_rate"]), row["question"])
    )
    weakest_question = question_rows_by_rate[0]
    lagging_models = [
        (model, result) for model, result in model_results if result.get("avg_rank") is not None
    ]
    lagging_model, lagging_model_result = (
        max(lagging_models, key=lambda item: float(item[1]["avg_rank"]))
        if lagging_models
        else model_results[0]
    )
    lagging_rank_text = lagging_model_result.get("avg_rank")
    doc.table(
        ["检查项", "当前试采发现", "正式评测处理", "签发判定"],
        [
            (
                "样本稳定性",
                repetition_copy["stability_current"],
                repetition_copy["stability_action"],
                "主样本单元完整；同时报告均值、离散程度与异常复核结果",
            ),
            (
                "弱问题复核",
                f"最低覆盖问题为“{weakest_question['question']}”；"
                f"当前提及 {weakest_question['mentions']}/{weakest_question['answers']}，"
                f"提及率 {_fmt_percent(weakest_question['appearance_rate'])}",
                repetition_copy["weak_question_action"],
                "区分可重复短板与单次随机波动；不以一次零提及直接下定论",
            ),
            (
                "问题组差异",
                f"最高组“{strongest_group[0]['title']}”与最低组“{weakest_group[0]['title']}”"
                "存在提及率差异",
                "分别报告三组结果，不用总体均值覆盖问题意图差异",
                "每组样本数、提及率、平均位次和 Top1/3/5 均可追溯",
            ),
            (
                "平台差异",
                (
                    f"{model_labels.get(lagging_model, lagging_model)} 提及时平均位次 "
                    f"{lagging_rank_text}，为本批平台中的最大值"
                    if lagging_rank_text is not None
                    else "本批各平台均未形成目标品牌提及位次"
                ),
                "固定问题、模式、账号与地域条件完成同题跨平台对照",
                "平台差异表与对应样本证据一致，不从平台内部机制做无证据推断",
            ),
            (
                "证据完整性",
                f"当前分享图片 {scope['share_images']}/{scope['answers']}；运行页截图 "
                f"{scope['answer_screenshots']}/{scope['answers']}；带引用回答 "
                f"{scope['answers_with_citation']}/{scope['answers']}",
                "逐样本保存回答正文、采集时间、分享图片；分享图缺失时保存去侧栏截图及引用 URL",
                "代表图可阅读、全量样本登记完整、代表回答全部信源可在附录核对",
            ),
        ],
        widths=(28, 50, 51, 43),
        font_size=7.3,
    )
    doc.numbered(
        [
            "正式复测必须沿用已冻结的问题矩阵；如需换题，应作为新批次单独报告。",
            "评测结论只回答品牌是否被提及、处于什么位次、与哪些竞品共同出现及列出哪些信源。",
            "官网内容是否被实际采纳、信源正文是否准确支持回答，分别由服务 3 和服务 2 核查，"
            "不在服务 1 中越界下结论。",
        ]
    )
    doc.page_break()

    doc.heading("附录说明")
    doc.numbered(
        [
            "附录 A 按综合名次连续列出全部品牌；跨页时自动重复表头，不按固定行数人为拆表。",
            f"附录 B 按问题组、问题、平台、地域连续登记全部 {scope['answers']} 条主样本；"
            "跨页时自动重复表头。",
            "附录 C 按三条代表回答分别成表，因为每张表对应一条独立的回答—信源证据链；"
            "表内列出该回答捕获的全部信源 URL。",
            "附录 D 按候选问题组分别成表，因为问题组是评测设计中的独立语义单元；"
            "签发限制单列为最后一节。",
        ]
    )
    doc.page_break()

    doc.heading("附录 A · 完整品牌排名")
    doc.paragraph(
        f"下表披露当前 {scope['answers']} 条主样本中观察到的全部 "
        f"{scope['brands_observed']} 个规范化品牌。"
        "完整披露用于避免只挑选少数竞品造成误读。"
    )
    ranking_rows = delivery["full_brand_ranking"]
    doc.table(
        ["名次", "品牌", "提及", "提及率", "平均位次", "综合分"],
        [
            (
                row["rank"],
                row["brand"],
                row["occurrences"],
                _fmt_percent(row["appearance_rate"]),
                row["avg_rank"],
                f"{float(row['score']):.3f}",
            )
            for row in ranking_rows
        ],
        widths=(15, 57, 20, 27, 28, 25),
        font_size=7,
    )
    ranking_table = doc.document.tables[-1]
    for row_index, row in enumerate(ranking_rows, 1):
        if row["is_target"]:
            _highlight_row(ranking_table, row_index)
    doc.page_break()

    doc.heading(f"附录 B · {scope['answers']} 条主样本登记")
    doc.paragraph("本表按问题组、问题、平台和地域逐条登记用于统计的主样本。")
    registry = delivery["sample_registry"]
    doc.table(
        ["序", "组", "问题", "平台", "地域", "提及/位次", "引用", "截图"],
        [
            (
                row["display_number"],
                row["group_title"],
                row["question"],
                row["platform_label"],
                row["region"],
                f"是/第{row['target_rank']}位" if row["mentioned"] else "否",
                row["citation_count"],
                "有" if row["has_answer_screenshot"] else "无",
            )
            for row in registry
        ],
        widths=(9, 30, 56, 20, 14, 21, 11, 11),
        font_size=6.2,
    )
    doc.page_break()

    doc.heading("附录 C · 代表回答完整信源列表")
    doc.paragraph(
        "本附录逐条披露正文三张代表回答所捕获的全部信源 URL；各表以回答为边界，"
        "不把不同回答的引用合并为一条证据链。"
    )
    for representative_index, item in enumerate(delivery["representative_answers"]):
        number = int(item["display_number"])
        if representative_index:
            doc.page_break()
        doc.heading(
            f"C.{number} {item['group_title']} · {item['platform_label']}/{item['region']}"
            f"（{item['citation_count']} 条）",
            level=2,
        )
        citation_rows = list(item.get("citations") or [])
        if not citation_rows:
            doc.paragraph("该代表回答未捕获信源 URL。")
            continue
        doc.table(
            ["序", "网站域名", "页面标题", "网页链接"],
            [
                (
                    citation.get("ordinal") or "—",
                    citation.get("host") or "（未知）",
                    citation.get("title") or "（未捕获标题）",
                    "打开原网页" if citation.get("url") else "（未捕获 URL）",
                )
                for citation in citation_rows
            ],
            widths=(10, 42, 88, 32),
            font_size=6.5,
        )
        citation_table = doc.document.tables[-1]
        for row_index, citation in enumerate(citation_rows, 1):
            url = str(citation.get("url") or "")
            if not url:
                continue
            link_paragraph = citation_table.cell(row_index, 3).paragraphs[0]
            link_paragraph.clear()
            _hyperlink(link_paragraph, url)
    doc.page_break()

    doc.heading("附录 D · 候选问题与签发限制")
    for group in service1["candidate_groups"]:
        marker = "主文选用" if group["selected_for_main_report"] else "未入主文，保留审计"
        doc.heading(f"D.{group['index']} {group['title']}（{marker}）", level=2)
        doc.table(
            ["序", "问题文本"],
            [(index, question) for index, question in enumerate(group["questions"], 1)],
            widths=(11, 161),
            font_size=8,
        )
    doc.heading("D.5 限制与正式签发前检查", level=2)
    doc.bullets(
        [
            (
                "本报告基于已冻结的正式评估窗口事实生成；结论仅适用于披露的窗口与样本口径。"
                if is_formal_document(facts)
                else "当前数据为联调/试采样，不是客户签收的正式运行数据。"
            ),
            f"正式报价口径要求每个问题×平台×地域单元完成 {required_repetitions} "
            f"次独立重复；当前为 {current_repetitions} 次。",
            "地域为浏览器地域采样标签；正式采集须补充账号、出口和地域台账。",
            "品牌顺序来自结构化抽取，代表截图用于人工复核；正式签发前须抽查异常高/低位次。",
            "信源分析只表明回答捕获了哪些 URL，不自动证明网页内容被答案准确采用。",
            "候选问题组按证据完整度自动选择，未使用目标品牌结果；未选组仍完整披露。",
        ]
    )
    return doc.save()
