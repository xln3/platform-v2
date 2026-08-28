"""DOCX renderer for a pre-validated Metrics V2 snapshot projection.

This module deliberately has no answer-analysis or metric-engine dependency.  It
formats immutable values and query-contribution rows supplied by the application
boundary; it cannot classify answers or calculate a project KPI.
"""

from __future__ import annotations

from collections.abc import Mapping, Sequence
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from domain.reporting.metric_snapshot_binding import MetricSnapshotSetBinding

_METRIC_LABELS = {
    "ai_impression_neutral_spontaneous_association_rate_v2": "品牌中性 AI 印象自发联想率",
    "ai_recommendation_organic_mention_rate_v2": "中性 AI 推荐自然提及率",
    "ai_recommendation_organic_recommendation_rate_v2": "中性 AI 推荐自然推荐率",
    "ai_recommendation_rankable_response_rate_v2": "中性 AI 推荐可排序回答率",
    "ai_recommendation_organic_top1_visibility_rate_v2": "中性 AI 推荐 Top1 可见率",
    "ai_recommendation_organic_top3_visibility_rate_v2": "中性 AI 推荐 Top3 可见率",
    "ai_recommendation_organic_top5_visibility_rate_v2": "中性 AI 推荐 Top5 可见率",
    "ai_recommendation_organic_top1_given_rankable_rate_v2": "可排序回答内 Top1 率",
    "ai_recommendation_organic_top3_given_rankable_rate_v2": "可排序回答内 Top3 率",
    "ai_recommendation_organic_top5_given_rankable_rate_v2": "可排序回答内 Top5 率",
    "ai_recommendation_mean_rank_given_target_ranked_v2": "目标有推荐排名时的平均名次",
    "ai_recommendation_entity_share_v2": "中性推荐实体份额",
    "prompted_recommendation_positive_rate_v2": "焦点品牌点名后正向推荐率",
    "prompted_recommendation_conditional_rate_v2": "焦点品牌点名后有条件推荐率",
    "prompted_recommendation_negative_rate_v2": "焦点品牌点名后负向推荐率",
    "prompted_recommendation_neutral_rate_v2": "焦点品牌点名后中性推荐率",
    "competitor_anchored_target_bring_in_rate_v2": "其他品牌点名后焦点品牌带出率",
    "competitor_anchored_target_alternative_rate_v2": "其他品牌点名后替代推荐率",
    "multibrand_pairwise_win_rate_v2": "多品牌同问两两胜出率",
    "multibrand_pairwise_tie_rate_v2": "多品牌同问两两持平率",
    "multibrand_pairwise_loss_rate_v2": "多品牌同问两两落后率",
    "multibrand_corecommendation_rate_v2": "多品牌同问共同推荐率",
}


def _metric_label(name: str) -> str:
    label = _METRIC_LABELS.get(name)
    return f"{label}（{name}）" if label else name


def _display_value(value: str | None, *, state: str) -> str:
    if value is None:
        return f"不形成正式数值（{state}）"
    return value


def _display_coverage(value: str | None) -> str:
    if value is None:
        return "—"
    try:
        numeric = float(value)
    except ValueError:
        return value
    return f"{numeric:.2%}"


def _text(value: object, *, fallback: str = "—") -> str:
    if value is None or value == "":
        return fallback
    if isinstance(value, Sequence) and not isinstance(value, str | bytes):
        return "、".join(str(item) for item in value) or fallback
    return str(value)


def render_bound_metric_snapshot_docx(
    *,
    title: str,
    service_number: int,
    binding: MetricSnapshotSetBinding,
    query_contributions: Sequence[Mapping[str, Any]],
    document_status: str,
    governance: Mapping[str, Any],
) -> bytes:
    """Render values exactly as frozen; no formula or fallback exists here."""

    document = Document()
    document.core_properties.title = title
    document.core_properties.subject = "Metrics V2 冻结快照正式报告"
    document.core_properties.author = str(governance.get("prepared_by") or "GEO 项目组")

    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        f"服务 {service_number} · {document_status} · 查询等权（query_macro）"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("报告数据绑定", level=1)
    for label, value in (
        ("快照集 ID", binding.snapshot_set_pub_id),
        ("快照集哈希", binding.snapshot_set_hash),
        ("快照集状态", binding.snapshot_set_state),
        ("成员依赖哈希", binding.dependency_hash),
        ("项目 ID", binding.project_pub_id),
        ("统计窗口", f"{binding.window_start.isoformat()} 至 {binding.window_end.isoformat()}"),
        ("过滤条件", _text(dict(binding.filters))),
    ):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}：").bold = True
        paragraph.add_run(str(value))

    document.add_paragraph(
        "本报告仅格式化上述不可变 Metrics V2 快照及贡献明细；报告生成过程不读取原始回答，"
        "不执行分类、指标公式或模型判定。"
    )

    document.add_heading("指标摘要", level=1)
    table = document.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    headers = (
        "指标",
        "焦点实体",
        "版本",
        "状态",
        "正式值",
        "分子",
        "分母",
        "查询数",
        "语义覆盖",
        "成员快照 ID",
    )
    for cell, label in zip(table.rows[0].cells, headers, strict=True):
        cell.text = label
    for metric in binding.snapshots:
        cells = table.add_row().cells
        metric_values = (
            _metric_label(metric.metric_name),
            metric.focal_entity_id,
            metric.metric_version,
            metric.state,
            _display_value(metric.value, state=metric.state),
            metric.raw_numerator,
            metric.raw_denominator,
            metric.unique_query_count,
            _display_coverage(metric.semantic_coverage),
            metric.snapshot_pub_id,
        )
        for cell, metric_value in zip(cells, metric_values, strict=True):
            cell.text = str(metric_value)

    for metric in (item for item in binding.snapshots if item.state == "limited"):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f"范围限制：{_metric_label(metric.metric_name)}仅描述在已配置的"
            f" {metric.unique_query_count} 个查询中观察到的结果，不外推行业总体或市场概率。"
        )
        run.bold = True
    if any(
        metric.state in {"insufficient", "experimental", "failed"} for metric in binding.snapshots
    ):
        document.add_paragraph(
            "insufficient、experimental 或 failed 成员仅披露状态，不形成正式数值或结论。"
        )

    document.add_heading("查询级贡献附录", level=1)
    document.add_paragraph(
        "下表逐行引用冻结贡献；完整逐回答、判定、事件、排除项、设计单元与哈希见配套 XLSX。"
    )
    query_table = document.add_table(rows=1, cols=9)
    query_table.style = "Table Grid"
    query_headers = (
        "成员快照 ID",
        "query_key",
        "查询文本",
        "查询权重",
        "分子",
        "分母",
        "值",
        "未知权重",
        "贡献哈希",
    )
    for cell, label in zip(query_table.rows[0].cells, query_headers, strict=True):
        cell.text = label
    for row in query_contributions:
        cells = query_table.add_row().cells
        query_values = (
            row.get("snapshot_pub_id"),
            row.get("query_key"),
            row.get("query_text"),
            row.get("query_weight"),
            row.get("query_numerator", row.get("numerator")),
            row.get("query_denominator", row.get("denominator")),
            row.get("query_value", row.get("value")),
            row.get("unknown_weight"),
            row.get("contribution_hash"),
        )
        for cell, value in zip(cells, query_values, strict=True):
            cell.text = _text(value)

    footer = document.add_paragraph()
    footer_run = footer.add_run(
        f"校验：set={binding.snapshot_set_pub_id} · hash={binding.snapshot_set_hash}"
    )
    footer_run.font.size = Pt(8)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


__all__ = ["render_bound_metric_snapshot_docx"]
