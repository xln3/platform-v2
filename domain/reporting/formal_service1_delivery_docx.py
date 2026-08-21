"""Compact customer narrative for governed Service-1 delivery artifacts.

V15 restructure (client-perspective content review CR-01..CR-40): the body now
leads with factual conclusions, separates measurement reliability from brand KPIs,
classifies query intent before interpreting zero mentions, and moves methods,
aliases, full per-question results, competitor scope and version governance into
self-contained appendices.  Engineering vocabulary (frozen matrix, manifest,
worksheet names, internal codes) is rejected at render time by _CLIENT_FORBIDDEN.
"""

from __future__ import annotations

import re
from typing import Any

from docx.shared import Pt

from .formal_review_docx import (
    FormalDocument,
    _fmt_datetime,
    _set_font,
    add_native_toc,
    build_report_code,
)
from .formal_review_service1_docx import _add_screenshot_panel, _hyperlink
from .service1_governance import release_state_label
from .service1_metrics import QUERY_INTENT_LABELS

_PLATFORM_LABELS = {"doubao": "豆包", "deepseek": "DeepSeek", "yiyan": "文心一言"}
_INTENT_ORDER = ("recommend", "selection", "knowledge")
_INTENT_EXPLANATION = {
    "recommend": "直接要求列出厂商或服务商（如“有哪些厂商”“找谁做”）",
    "selection": "询问如何选择或评估厂商；回答可能列出厂商，也可能只给方法",
    "knowledge": "询问评估标准或概念；回答通常不需要列出厂商，品牌不出现不属于短板",
}
_CLIENT_FORBIDDEN = (
    "本次审阅口径",
    "审阅判断",
    "当前试采",
    "正式复测与签发检查",
    "冻结矩阵",
    "mode=",
    "[特殊字符]",
    "发布门禁未通过",
    "不得对外交付",
    "阻断项已逐项",
    "缺失时不得批准签发",
    "不作出的承诺",
    # V15 client-language gates: keep engineering vocabulary out of the body.
    "分子/分母",
    "规范实体",
    "1-based",
    "引用核验",
    "共现",
    "manifest",
    "工作表",
    "客户优先行动",
    "重复一致性",
    "样本索引",
    "待办项",
)


def _to_float(value: object) -> float | None:
    if value is None:
        return None
    try:
        return float(str(value))
    except ValueError:
        return None


def _pct(value: object) -> str:
    return f"{(_to_float(value) or 0.0):.1f}%"


def _rank(value: object) -> str:
    number = _to_float(value)
    return "—" if number is None else f"{number:.1f}"


def _fraction(row: dict[str, Any], key: str = "mentions") -> str:
    return f"{int(row.get(key) or 0)}/{int(row.get('answers') or 0)}"


def _clip(value: object, limit: int = 26) -> str:
    text = " ".join(str(value or "").split())
    return text if len(text) <= limit else f"{text[: limit - 1]}…"


def _clean_answer(value: object, limit: int = 1400) -> str:
    text = str(value or "")
    # Preserve the raw capture in the evidence package, but repair known UTF-8-as-
    # Windows-1252 display artefacts in the customer-readable excerpt.
    replacements = {
        "â†’": "→",
        "Â·": "·",
        "Â ": " ",
        "â€œ": "“",
        "â€\u009d": "”",
        "â€™": "’",
        "â€“": "–",
        "â€”": "—",
        "â€¦": "…",
    }
    for source, replacement in replacements.items():
        text = text.replace(source, replacement)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"(?m)^\s{0,3}#{1,6}\s*", "", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    text = re.sub(r"(?m)^\s*[-*_]{3,}\s*$", "", text)
    text = re.sub(r"(?m)^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*$", "", text)
    text = re.sub(r"(?m)^\s*[-*+]\s+", "• ", text)
    text = re.sub(
        r"(?m)^\s*\|(.+)\|\s*$",
        lambda match: " ｜ ".join(
            cell.strip() for cell in match.group(1).split("|") if cell.strip()
        ),
        text,
    )
    text = re.sub(r"[\u0000-\u0008\u000b\u000c\u000e-\u001f\u007f]", "", text)
    text = re.sub(r"[\u200b-\u200f\ufeff]", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text if len(text) <= limit else text[: limit - 1].rstrip() + "…"


def _page(doc: FormalDocument, title: str) -> None:
    doc.page_break()
    doc.heading(title)


def _ai_disclaimer(doc: FormalDocument) -> None:
    doc.callout(
        "重要说明",
        "以下为 AI 生成原文，未经事实核验，不代表评测方结论。报告仅评测品牌是否出现、"
        "出现顺序及回答所列链接。",
        kind="warning",
    )


def _section_readout(doc: FormalDocument, conclusion: str, usage: str) -> None:
    """Close every main chapter with a plain conclusion and a usage note (CR-40)."""

    doc.numbered([f"本节结论：{conclusion}", f"建议如何使用：{usage}"])


def _add_visible_url(doc: FormalDocument, *, ordinal: object, title: object, url: str) -> None:
    lead = doc.document.add_paragraph()
    lead.paragraph_format.space_after = Pt(1)
    run = lead.add_run(f"{ordinal or '—'}. {title or '（未捕获标题）'}")
    _set_font(run, size=8)
    run.bold = True
    link = doc.document.add_paragraph()
    link.paragraph_format.space_after = Pt(5)
    _hyperlink(link, url, text=url)
    for run in link.runs:
        _set_font(run, size=7.2)


def _key_excerpt(text: str, target_brand: str, *, limit: int = 1100) -> str:
    """Return a brand-centred window of a longer answer, with explicit ellipses.

    Deep-thinking answers often open with tool/search narration that can contain
    the brand name ("让我深入读取盛邦安全的方法论…").  Anchoring the excerpt on
    that first token would show the narration instead of the actual answer, so an
    occurrence inside the first tenth of the text yields to the next one.
    """

    compact = text.strip()
    if len(compact) <= limit:
        return compact
    position = compact.find(target_brand)
    if position >= 0 and position < max(1, len(compact) // 10):
        later = compact.find(target_brand, max(1, len(compact) // 10))
        if later >= 0:
            position = later
    if position < 0:
        return f"{compact[: limit - 1].rstrip()}…"
    before = max(0, position - limit // 3)
    after = min(len(compact), before + limit)
    before = max(0, after - limit)
    # Snap the window to line boundaries so the excerpt never starts mid-row.
    if before > 0:
        line_start = compact.find("\n", before, position)
        if line_start >= 0:
            before = line_start + 1
    if after < len(compact):
        line_end = compact.rfind("\n", position, after)
        if line_end > position:
            after = line_end
    prefix = "…" if before else ""
    suffix = "…" if after < len(compact) else ""
    excerpt = f"{prefix}{compact[before:after].strip()}{suffix}"
    # Drop a trailing reference-list tail; the full links are listed right below.
    excerpt = re.split(r"(?m)^\s*(参考来源|参考资料|参考文献|来源)[:：]?\s*$", excerpt)[0].rstrip()
    return excerpt


def _citation_table(doc: FormalDocument, citations: list[dict[str, Any]]) -> None:
    """List every cited link; show the full address for the first entries.

    LibreOffice only preserves PDF link annotations for hyperlinks whose visible
    label is the URL itself, so the first entries render as visible-URL lines
    (these are exactly the URLs the publication QA expects to be clickable).
    Remaining entries stay in a compact host/title table; their full addresses
    ship in the machine-readable delivery files.
    """

    if not citations:
        doc.paragraph("该回答没有列出链接；报告不据此推断其信息来源。")
        return
    visible = citations[:3]
    for citation in visible:
        url = str(citation.get("url") or "")
        if url:
            _add_visible_url(
                doc,
                ordinal=citation.get("ordinal"),
                title=citation.get("title"),
                url=url,
            )
    rest = citations[len(visible) :]
    if not rest:
        return
    doc.paragraph(f"其余 {len(rest)} 条链接的完整地址见随附样本明细文件与证据包：")
    doc.table(
        ["序", "网站域名", "页面标题"],
        [
            (
                citation.get("ordinal") or "—",
                citation.get("host") or "（未知）",
                _clip(citation.get("title") or "（未捕获标题）", 42),
            )
            for citation in rest
        ],
        widths=(10, 52, 110),
        font_size=6.2,
    )


def _representative_pages(
    doc: FormalDocument,
    item: dict[str, Any],
    sample: dict[str, Any],
    screenshot: bytes | None,
    *,
    target_brand: str,
    start_new_page: bool = True,
) -> None:
    number = int(item["display_number"])
    heading = f"8.{number} {item['platform_label']}代表回答 · {item['group_title']}"
    if start_new_page:
        _page(doc, heading)
    else:
        doc.heading(heading)
    # Exactly one in-place notice per representative sample (CR-33).
    _ai_disclaimer(doc)
    doc.table(
        ["字段", "实测内容"],
        [
            ("实际问题", item["question"]),
            ("平台与地域标签", f"{item['platform_label']} · {item['region']}"),
            ("采集时间", _fmt_datetime(item.get("capture_time"))),
            (
                "目标品牌表现",
                f"提到，品牌在回答所列品牌中第 {item['target_rank']} 个出现"
                if item.get("target_rank") is not None
                else "未提到",
            ),
            ("回答所列链接", f"{item.get('citation_count') or 0} 条（链接内容未经事实核验）"),
            ("样本编号", str(sample.get("sample_id") or "—")),
        ],
        widths=(31, 141),
        font_size=8.1,
    )
    full_text = _clean_answer(sample.get("response_text"), limit=20000)
    excerpt = (
        _key_excerpt(full_text, target_brand, limit=420)
        if full_text
        else _clean_answer(item.get("answer_excerpt"), limit=420)
    )
    if excerpt:
        doc.heading(
            "品牌相关片段" if item.get("target_rank") is not None else "回答片段（未提到目标品牌）",
            level=2,
        )
        doc.paragraph(excerpt)
    doc.paragraph(
        f"该回答的关键片段与全部所列链接见附录 D.{number}；完整原文与完整长图见随附证据包。"
    )
    if screenshot:
        # Keep the evidence image and its caption together on a dedicated page.
        doc.page_break()
        _add_screenshot_panel(
            doc,
            screenshot,
            platform=str(item["platform"]),
            number=number,
            image_kind=str(item.get("preferred_image_kind") or "answer_screenshot"),
            anchor=item.get("answer_anchor"),
            figure_prefix="8",
        )
    else:
        doc.callout("图片状态", "该代表样本没有可载入的回答图片。", kind="warning")


def _short_question(value: object, limit: int = 18) -> str:
    """Clip a question without dangling half an ASCII token or parenthesis."""

    text = _clip(value, limit)
    if text.endswith("…"):
        body = text[:-1].rstrip()
        while body and (body[-1].isascii() or body[-1] in "（("):
            body = body[:-1].rstrip()
        text = (body or text[:-1]) + "…"
    return text


def _sample_gap_text(delivery: dict[str, Any]) -> str:
    cells = list(delivery.get("incomplete_cells") or [])
    if not cells:
        return "全部组合均已完成计划次数。"
    parts = [
        f"“{_short_question(cell['question'])}”在{cell['platform_label']}·{cell['region']}"
        f"仅有 {cell['observed']}/{cell['required']} 次"
        for cell in cells[:3]
    ]
    suffix = "。" if len(cells) <= 3 else f"；另有 {len(cells) - 3} 个组合类似。"
    return "；".join(parts) + suffix


def render_service1_delivery_docx(
    facts: dict[str, Any],
    *,
    screenshots: dict[str, bytes] | None = None,
    service_number: int = 1,
    report_title: str | None = None,
    report_subtitle: str = "服务 1 · 品牌 AI 可见性与竞品表现",
) -> bytes:
    """Render the customer-readable service-1 report; bulk audit data stay in sidecars."""

    service1 = facts["service1"]
    delivery = service1.get("delivery_v3")
    if not isinstance(delivery, dict):
        raise ValueError("service1.delivery_v3_missing")
    screenshots = screenshots or {}
    target_brand = str(facts["target_brand"])
    target = delivery["target"]
    scope = delivery["scope"]
    quotation_catalog = facts.get("service_catalog_version") == "quotation_services_v2"
    risk_service_number = 3 if quotation_catalog else 2
    official_service_number = 4 if quotation_catalog else 3
    scope_label = str(scope.get("scope_label") or "本次三组已测业务场景")
    title = report_title or f"{scope_label}品牌 GEO 推荐结果评测报告"
    facts = {**facts, "report_title": title}
    doc = FormalDocument(
        title=title,
        subtitle=report_subtitle,
        facts=facts,
    )
    version = str((facts.get("document_governance") or {}).get("version") or "V1.0")
    doc.cover(report_code=build_report_code(facts, service_number=service_number, version=version))
    add_native_toc(doc, heading_levels="1-2")

    question_count = int(scope["questions"])
    answers = int(scope["answers"])
    expected_answers = int(scope.get("expected_answers") or 0)
    top_counts = target["top_counts"]
    platform_rows = list(delivery["by_platform"].items())
    group_rows = list(delivery["by_group"].items())
    strongest_platform_slug, strongest_platform = max(
        platform_rows, key=lambda item: item[1]["mention_rate"]
    )
    weakest_platform_slug, weakest_platform = min(
        platform_rows, key=lambda item: item[1]["mention_rate"]
    )
    strongest_group_name, strongest_group = max(
        group_rows, key=lambda item: item[1]["mention_rate"]
    )
    weakest_group_name, weakest_group = min(group_rows, key=lambda item: item[1]["mention_rate"])
    strongest_platform_label = _PLATFORM_LABELS.get(
        strongest_platform_slug, strongest_platform_slug
    )
    weakest_platform_label = _PLATFORM_LABELS.get(weakest_platform_slug, weakest_platform_slug)
    intent_breakdown = delivery.get("intent_breakdown") or {}
    knowledge_questions = int((intent_breakdown.get("knowledge") or {}).get("questions") or 0)
    observed_aliases = [
        str(value) for value in (delivery.get("target_aliases") or {}).get("observed") or []
    ]
    alias_observation = (
        f"本批回答中品牌仅以“{observed_aliases[0]}”一种写法出现"
        if len(observed_aliases) == 1
        else (
            f"本批回答中品牌出现 {len(observed_aliases)} 种写法"
            f"（{'、'.join(observed_aliases[:4])}）"
            if observed_aliases
            else "本批回答未观测到品牌写法"
        )
    )

    # ------------------------------------------------------------------ 1. 摘要
    _page(doc, "1. 执行摘要")
    doc.kpis(
        [
            (
                "品牌提及率",
                _pct(target["mention_rate"]),
                target["mention_rate_fraction"] + " 条回答",
            ),
            (
                "平均首次出现顺序",
                f"第 {_rank(target['avg_rank'])} 位" if target.get("avg_rank") else "—",
                "仅在提到品牌的回答内",
            ),
            (
                "前三出现率",
                _pct(target["top_rates"]["3"]),
                f"{top_counts['3']}/{target['answers']} 条回答",
            ),
            (
                "主样本回答",
                f"{answers} 条",
                (
                    f"计划 {expected_answers} 条，当前 {answers} 条"
                    if expected_answers and answers != expected_answers
                    else "计划样本已全部完成"
                ),
            ),
        ]
    )
    doc.paragraph(
        f"本次测试围绕{scope_label}，覆盖 {scope['selected_groups']} 组业务问题、"
        f"{question_count} 个问题文本、{scope['platforms']} 个 AI 平台和 "
        f"{scope['regions']} 个地域标签，共形成 {answers} 条有效回答。"
        f"{target_brand}在其中被提到 {target['mentions']} 次，提及率 "
        f"{_pct(target['mention_rate'])}；提到品牌的回答中，品牌平均第 "
        f"{_rank(target['avg_rank'])} 个出现。"
    )
    doc.heading("1.1 三个事实结论", level=2)
    doc.table(
        ["事实结论", "数据依据"],
        [
            (
                f"{target_brand}在本次问题集中的出现不足一半",
                f"提到 {target['mentions']}/{answers}，提及率 {_pct(target['mention_rate'])}；"
                f"前三出现率 {_pct(target['top_rates']['3'])}"
                f"（{top_counts['3']}/{answers}）",
            ),
            (
                "三类场景覆盖不均衡",
                f"最高：{strongest_group_name} {_fraction(strongest_group)}（"
                f"{_pct(strongest_group['mention_rate'])}）；最低：{weakest_group_name} "
                f"{_fraction(weakest_group)}（{_pct(weakest_group['mention_rate'])}）",
            ),
            (
                "三个平台覆盖不均衡",
                f"{strongest_platform_label} {_fraction(strongest_platform)}（"
                f"{_pct(strongest_platform['mention_rate'])}）；{weakest_platform_label} "
                f"{_fraction(weakest_platform)}（{_pct(weakest_platform['mention_rate'])}）",
            ),
        ],
        widths=(52, 120),
        font_size=8,
    )
    doc.heading("1.2 对客户的含义", level=2)
    doc.bullets(
        [
            f"品牌在三类场景中已有出现，但覆盖不均：“{weakest_group_name}”是当前最明显的空白场景。",
            f"同一问题集下，品牌在{strongest_platform_label}的提及率明显高于"
            f"{weakest_platform_label}；该差异只对应本批问题和窗口，不推断平台机制。",
            (
                f"本批有 {knowledge_questions} 个知识型问题（询问评估标准），其回答本就不需要"
                "列出厂商，品牌在其中不出现不计为品牌短板。"
                if knowledge_questions
                else "本批问题均合理期待品牌出现，零提及问题可直接视为覆盖缺口。"
            ),
        ]
    )
    doc.heading("1.3 建议的下一步", level=2)
    doc.bullets(
        [
            "先盘点、后投入：盘点“" + weakest_group_name + "”场景下品牌公开材料是否存在、"
            "是否清楚说明品牌与产品能力，再决定内容投入（具体建议见第 7 章）。",
            "统一品牌表达：公开材料中统一品牌简称、公司全称与产品归属写法，"
            "再用同一问题集复测观察提及变化。",
            (
                f"补齐本批数据：计划 {expected_answers} 条、当前 {answers} 条；"
                "补齐缺失测试与采样台账后重算，再确认上述场景与平台差异是否稳定。"
                if expected_answers and answers != expected_answers
                else "在后续窗口沿用同一问题集复测，确认上述场景与平台差异是否稳定。"
            ),
        ]
    )
    doc.paragraph(
        f"阅读口径：本报告结论只对应{scope_label}、本报告披露的 {question_count} 个问题文本、"
        "三个平台、两个地域标签和本次采集窗口；不代表品牌在所有业务场景或整体市场中的"
        "表现。更广范围的判断需扩展问题集后另行评估。"
    )

    # ------------------------------------------------------------ 2. 测试范围
    _page(doc, "2. 本次测试回答什么问题")
    doc.table(
        ["范围维度", "本次口径", "说明"],
        [
            (
                "业务场景",
                scope_label,
                "报价单按业务线组织核心问题，本次服务 1 选取其中三个资产治理相关场景，逐组列于 2.1",
            ),
            (
                "问题",
                f"{scope['selected_groups']} 组、{question_count} 个问题文本",
                "每组 1 个客户原始问题加 3 个语义变体；测试开始前问题已确定，"
                "测试过程中未根据结果更换",
            ),
            (
                "平台",
                "、".join(
                    _PLATFORM_LABELS.get(model, model)
                    for model in service1.get("primary_models") or []
                )
                + "（深度思考模式）",
                "同一问题在同一平台内比较",
            ),
            (
                "地域",
                "、".join(service1.get("primary_regions") or []) + "（采样地域标签）",
                "本批账号与出口台账尚未补齐，本报告不作地域差异结论",
            ),
            (
                "每格测试次数",
                "计划每个“问题×平台×地域”组合 "
                f"{service1['quotation_required_repetitions_per_cell']} 次",
                _sample_gap_text(delivery),
            ),
            (
                "样本规模",
                f"计划 {expected_answers} 条，当前 {answers} 条有效回答",
                "有效回答指成功采集且完成品牌抽取的回答；本批无剔除",
            ),
            (
                "数据窗口",
                f"{facts['window']['start']} 至 {facts['window']['end']}",
                "结论仅对应该窗口",
            ),
        ],
        widths=(26, 66, 80),
        font_size=7.8,
    )
    if answers != expected_answers:
        doc.callout(
            "样本缺口",
            f"计划 {expected_answers} 条回答，当前 {answers} 条。"
            + _sample_gap_text(delivery)
            + "缺口补齐并重算前，本报告保持内部审核稿状态。",
            kind="warning",
        )

    doc.heading("2.1 三个业务场景与十二个问题", level=2)
    for position, group in enumerate(delivery["selected_groups"], 1):
        doc.heading(f"场景{('一二三')[position - 1]}：{group['title']}", level=3)
        doc.table(
            ["变体", "实际提问文本", "提问意图"],
            [
                (
                    "原题" if index == 1 else chr(63 + index),
                    question,
                    QUERY_INTENT_LABELS[
                        next(
                            (
                                row["query_intent"]
                                for row in delivery["question_rows"]
                                if row["question"] == question
                            ),
                            "recommend",
                        )
                    ],
                )
                for index, question in enumerate(group["questions"], 1)
            ],
            widths=(16, 118, 38),
            font_size=7.8,
        )
    doc.heading("2.2 提问意图与阅读方式", level=2)
    doc.table(
        ["提问意图", "含义", "本批问题数"],
        [
            (
                QUERY_INTENT_LABELS[intent],
                _INTENT_EXPLANATION[intent],
                (intent_breakdown.get(intent) or {}).get("questions") or 0,
            )
            for intent in _INTENT_ORDER
            if intent in intent_breakdown
        ],
        widths=(24, 116, 32),
        font_size=8,
    )
    _section_readout(
        doc,
        f"本次测试用 {question_count} 个已确定的问题，在三个平台、两个地域标签下观察"
        f"{target_brand}是否被 AI 回答提到。",
        "先看第 3 章总体结果，再按第 4 章定位具体场景或平台；解释单个问题的结果前，"
        "先确认它的提问意图。",
    )

    # ------------------------------------------------------------ 3. 总体结果
    _page(doc, "3. 品牌总体结果")
    doc.chart(
        ["提及率", "第一出现率", "前三出现率", "前五出现率"],
        [
            target["mention_rate"],
            target["top_rates"]["1"],
            target["top_rates"]["3"],
            target["top_rates"]["5"],
        ],
        title=f"{target_brand}在全部 {answers} 条有效回答中的可见性",
    )
    doc.table(
        ["指标", "计算方式", "结果", "客户应如何理解"],
        [
            (
                "品牌提及率",
                f"提到{target_brand}或已确认品牌别名的回答数 ÷ 全部有效回答数",
                f"{target['mention_rate_fraction']}＝{_pct(target['mention_rate'])}",
                "回答是否提到品牌，不等同于排在前列",
            ),
            (
                "第一出现率",
                "品牌第一个出现的回答数 ÷ 全部有效回答数",
                f"{top_counts['1']}/{answers}＝{_pct(target['top_rates']['1'])}",
                "品牌是回答中最先出现的品牌",
            ),
            (
                "前三出现率",
                "品牌属于最先出现三个品牌之一的回答数 ÷ 全部有效回答数",
                f"{top_counts['3']}/{answers}＝{_pct(target['top_rates']['3'])}",
                "品牌进入回答的优先比较范围",
            ),
            (
                "前五出现率",
                "品牌属于最先出现五个品牌之一的回答数 ÷ 全部有效回答数",
                f"{top_counts['5']}/{answers}＝{_pct(target['top_rates']['5'])}",
                "品牌进入回答的较长候选清单",
            ),
            (
                "平均首次出现顺序",
                "在提到品牌的回答中，品牌第几个出现的平均值",
                f"第 {_rank(target['avg_rank'])} 位" if target.get("avg_rank") else "—",
                "数值越小越靠前；未提到的回答不参与平均",
            ),
        ],
        widths=(26, 56, 34, 56),
        font_size=7.4,
    )
    doc.bullets(
        [
            "品牌的不同写法已合并：回答提到" + target_brand + "或已确认的品牌别名，"
            "均计为一次提到；同一回答重复出现不重复计数。别名清单见附录 A。",
            "“首次出现顺序”按品牌在回答文字中第一次出现的位置计算；叙述性回答中的出现"
            "顺序不等于平台的明确推荐排名。",
            f"本批 {answers} 条回答全部通过品牌抽取，无剔除；两次测试的波动情况见附录 A。",
        ]
    )
    doc.heading("3.1 首次出现顺序分布", level=2)
    doc.chart(
        [str(row["label"]) for row in delivery["rank_distribution"]],
        [float(row["count"]) for row in delivery["rank_distribution"]],
        title=f"{target_brand}在 {answers} 条有效回答中的首次出现顺序分布",
        suffix="条",
    )
    doc.table(
        ["区间", "回答数", "占全部样本"],
        [
            (
                row["label"],
                row["count"],
                _pct(float(row["count"]) / max(answers, 1) * 100),
            )
            for row in delivery["rank_distribution"]
        ],
        widths=(48, 40, 84),
        font_size=8.2,
    )
    _section_readout(
        doc,
        f"{target_brand}在 {answers} 条回答中被提到 {target['mentions']} 次；提到时多数位于"
        f"前列（前三出现率 {_pct(target['top_rates']['3'])}）。",
        "把提及率当作“是否被想到”，把首次出现顺序当作“被想到时大致排第几”；"
        "两个指标都要结合提问意图阅读。",
    )

    # ------------------------------------------------- 4. 分场景与分平台结果
    _page(doc, "4. 分场景与分平台结果")
    doc.heading("4.1 分场景表现", level=2)
    doc.table(
        ["业务场景", "回答数", "提及", "提及率", "前三出现率", "平均首次出现顺序"],
        [
            (
                name,
                row["answers"],
                row["mentions"],
                _pct(row["mention_rate"]),
                f"{row['top_counts']['3']}/{row['answers']}＝{_pct(row['top_rates']['3'])}",
                f"第 {_rank(row['avg_rank'])} 位" if row.get("avg_rank") else "—",
            )
            for name, row in group_rows
        ],
        widths=(52, 18, 18, 26, 34, 24),
        font_size=7.6,
    )
    doc.chart(
        [name for name, _ in group_rows],
        [float(row["mention_rate"]) for _, row in group_rows],
        title="不同业务场景的品牌提及率",
    )
    doc.heading("4.2 按提问意图表现", level=2)
    doc.table(
        ["提问意图", "问题数", "回答数", "提及", "提及率"],
        [
            (
                QUERY_INTENT_LABELS[intent],
                bucket["questions"],
                bucket["answers"],
                bucket["mentions"],
                f"{bucket['mention_rate_fraction']}＝{_pct(bucket['mention_rate'])}",
            )
            for intent in _INTENT_ORDER
            if intent in intent_breakdown
            for bucket in [intent_breakdown[intent]]
        ],
        widths=(26, 22, 24, 22, 78),
        font_size=8,
    )
    doc.paragraph(
        "知识型问题（如“评估产品应关注哪些指标”）本来就不期待回答列出厂商，"
        "其零提及用于核对问题设计，不计为品牌短板；推荐型问题的零提及才是覆盖缺口。"
    )
    doc.heading("4.3 分平台表现", level=2)
    doc.table(
        ["平台", "回答数", "提及", "提及率", "第一出现率", "前三出现率", "平均首次出现顺序"],
        [
            (
                _PLATFORM_LABELS.get(platform, platform),
                row["answers"],
                row["mentions"],
                _pct(row["mention_rate"]),
                _pct(row["top_rates"]["1"]),
                _pct(row["top_rates"]["3"]),
                f"第 {_rank(row['avg_rank'])} 位" if row.get("avg_rank") else "—",
            )
            for platform, row in platform_rows
        ],
        widths=(28, 18, 18, 26, 26, 26, 30),
        font_size=7.6,
    )
    doc.paragraph(
        "平台差异只描述本批问题与窗口下的观察，用于定位复测重点；不对平台内部机制做推断。"
    )
    _section_readout(
        doc,
        f"场景维度：“{weakest_group_name}”提及率最低；平台维度：{weakest_platform_label}最低。"
        "知识型问题的低提及不属于短板。",
        "按“场景×平台”定位下一批内容盘点与复测的优先级；单一组合样本很小，不要据单个问题下结论。",
    )

    # ---------------------------------------------------- 5. 竞品表现（范围内）
    _page(doc, "5. 本次问题范围内的品牌出现对比")
    comparison = delivery["competitor_comparison"]
    competitor_rows = [comparison["target"], *comparison["competitors"]]
    doc.table(
        ["统一品牌名称", "身份", "提及", "提及率", "前三出现率", "平均首次出现顺序"],
        [
            (
                row["canonical_name"],
                "目标品牌" if row["canonical_name"] == target_brand else "对比品牌",
                row["mention_rate_fraction"],
                _pct(row["mention_rate"]),
                f"{row['top_counts']['3']}/{row['answers']}＝{_pct(row['top_rates']['3'])}",
                f"第 {_rank(row['avg_rank'])} 位" if row.get("avg_rank") else "—",
            )
            for row in competitor_rows
        ],
        widths=(40, 22, 26, 27, 34, 23),
        font_size=7.6,
    )
    doc.bullets(
        [
            "同一品牌的不同写法已合并；本表只比较网络安全厂商品牌，开源工具、机构名称"
            "和尚未确认的名称不进入对比（口径见附录 C）。",
            "对比品牌按本批回答中出现次数自动选取前 5 个厂商品牌；该名单尚未经客户确认，"
            "确认后可替换为指定竞品重算。",
            "本表只对应本次 12 个问题，不是行业排名，也不推断各品牌整体市场地位。",
        ]
    )
    _section_readout(
        doc,
        f"在本次问题范围内，{target_brand}提及率 "
        f"{_pct(target['mention_rate'])}，与对比品牌的同题逐项差值见附录 C。",
        "把本表当作“同一批问题下谁更常被 AI 想到”的对照，不要当作市场份额或行业排名引用。",
    )

    # ------------------------------------------------------------ 6. 三个发现
    _page(doc, "6. 三个最重要发现")
    representatives = list(delivery["representative_answers"])
    weakest_rep = next(
        (item for item in representatives if item["group_title"] == weakest_group_name),
        representatives[0] if representatives else None,
    )
    recommend_rates: dict[str, tuple[int, int]] = {}
    for row in delivery["question_rows"]:
        if str(row.get("query_intent")) != "recommend":
            continue
        prev_mentions, prev_answers = recommend_rates.get(str(row["group_title"]), (0, 0))
        recommend_rates[str(row["group_title"])] = (
            prev_mentions + int(row["mentions"]),
            prev_answers + int(row["answers"]),
        )
    weakest_rate = recommend_rates.get(weakest_group_name)
    weakest_still_lowest = bool(
        weakest_rate
        and weakest_rate[1]
        and all(
            weakest_rate[0] / weakest_rate[1] <= (mentions / answers if answers else 0)
            for mentions, answers in recommend_rates.values()
        )
    )
    findings = [
        (
            "发现一：品牌出现不足一半，且集中在部分场景",
            f"在 {answers} 条有效回答中，{target_brand}被提到 {target['mentions']} 次"
            f"（{_pct(target['mention_rate'])}）；提到品牌的回答中，品牌平均第 "
            f"{_rank(target['avg_rank'])} 个出现，前三出现率 "
            f"{_pct(target['top_rates']['3'])}。",
            "品牌已进入 AI 回答的候选范围，但还不是稳定出现；当前更接近“部分场景可见”，"
            "而非“普遍被推荐”。",
        ),
        (
            f"发现二：“{weakest_group_name}”是最明显的空白场景",
            f"该场景提及 {_fraction(weakest_group)}（{_pct(weakest_group['mention_rate'])}），"
            f"明显低于“{strongest_group_name}”的 {_fraction(strongest_group)}"
            f"（{_pct(strongest_group['mention_rate'])}）。",
            "其中知识型问题不期待厂商名单；"
            + ("即便只看推荐型问题，该场景提及率仍为最低；" if weakest_still_lowest else "")
            + "值得优先盘点品牌在该主题的公开材料。"
            + (
                f"该场景的代表回答见第 8.{int(weakest_rep['display_number'])} 节。"
                if weakest_rep
                else ""
            ),
        ),
        (
            "发现三：平台之间差异明显",
            f"{strongest_platform_label}提及 {_fraction(strongest_platform)}"
            f"（{_pct(strongest_platform['mention_rate'])}），{weakest_platform_label}仅 "
            f"{_fraction(weakest_platform)}（{_pct(weakest_platform['mention_rate'])}）。",
            "差异只对应本批问题与窗口。建议先核查各平台回答所列链接的来源结构"
            f"（服务 {risk_service_number}、服务 {official_service_number}），"
            "再决定是否按平台分别投入。",
        ),
    ]
    for heading, data_text, explanation in findings:
        doc.heading(heading, level=2)
        doc.paragraph(data_text, bold_lead="数据：")
        doc.paragraph(explanation, bold_lead="解释：")
    _section_readout(
        doc,
        "三个发现互相对应：总体覆盖不足一半，缺口集中在个别场景与个别平台。",
        "把发现作为第 7 章行动的观察依据；每个发现都可以在样本明细中逐条复核。",
    )

    # ------------------------------------------------------------ 7. 建议行动
    _page(doc, "7. 建议行动")
    doc.paragraph(
        "以下每条建议都写明本批观察依据、建议动作、责任方和验证方式。这些建议是待验证的"
        "工作假设：本服务证明的是品牌是否出现，内容投入与提及变化之间的因果关系需要"
        "用同一问题集的复测来检验。"
    )
    doc.table(
        ["本批观察依据", "建议动作", "责任方", "如何验证"],
        [
            (
                f"“{weakest_group_name}”提及率最低（{_fraction(weakest_group)}）",
                "先盘点该场景下品牌公开材料是否存在、是否清楚说明品牌定位与产品能力，"
                "再决定内容建设投入",
                "客户内容团队，评测方提供清单",
                "沿用同一问题集复测，观察提及率与前三出现率变化",
            ),
            (
                alias_observation,
                "在官网与公开页面统一品牌简称、公司全称与产品归属表达",
                "客户内容团队",
                "复测时按统一品牌名称口径比较提及与顺序变化",
            ),
            (
                f"{weakest_platform_label}提及率明显低于{strongest_platform_label}",
                f"暂不直接按平台投入；先结合服务 {risk_service_number} 的被拉踩内容核查和"
                f"服务 {official_service_number} 的官网引用能效定位原因",
                "评测方",
                f"服务 {risk_service_number}/{official_service_number} 结果出来后，"
                "再决定是否需要平台级动作",
            ),
            (
                f"{scope['answers_with_citation']}/{answers} 条回答列出了链接",
                "持续维护可访问、标题明确的公开说明页面，保持可被 AI 检索引用",
                "客户内容团队",
                "复测时观察回答所列链接中品牌相关页面的出现情况",
            ),
        ],
        widths=(44, 56, 30, 42),
        font_size=7.3,
    )
    doc.callout(
        "服务边界",
        "服务 1 只回答“品牌是否出现、以什么顺序出现、回答列出了哪些链接”。"
        f"被拉踩内容由服务 {risk_service_number} 核查，官网内容是否被 AI 实际采纳由"
        f"服务 {official_service_number} 评估；"
        "本报告不越界下结论。",
        kind="info",
    )

    # ------------------------------------------------------------ 8. 代表回答
    _page(doc, "8. 代表回答与截图")
    doc.paragraph(
        f"以下 {len(representatives)} 条代表回答按固定平台轮换选取，覆盖三个平台；"
        "选择只看证据完整度，不看品牌是否提到。代表回答不参与指标计算；"
        f"全部 {answers} 条回答的明细见随附样本明细文件。"
    )
    sample_by_answer = {str(row.get("answer_pub_id")): row for row in delivery["sample_registry"]}
    doc.table(
        ["业务场景", "平台/地域标签", "目标品牌", "所列链接", "样本编号"],
        [
            (
                item["group_title"],
                f"{item['platform_label']}/{item['region']}",
                (
                    f"第 {item['target_rank']} 个出现"
                    if item.get("target_rank") is not None
                    else "未提到"
                ),
                item["citation_count"],
                str(
                    sample_by_answer.get(str(item.get("answer_pub_id") or ""), {}).get("sample_id")
                    or "—"
                ),
            )
            for item in representatives
        ],
        widths=(56, 34, 30, 20, 32),
        font_size=7.6,
    )
    if not delivery.get("representative_platforms_complete"):
        doc.callout(
            "代表证据缺口", "三个平台的代表回答未齐全，出版检查会阻断交付候选状态。", kind="warning"
        )
    for representative_index, item in enumerate(representatives):
        answer_id = str(item.get("answer_pub_id") or "")
        _representative_pages(
            doc,
            item,
            sample_by_answer.get(answer_id, {}),
            screenshots.get(answer_id),
            target_brand=target_brand,
            start_new_page=representative_index > 0,
        )

    # -------------------------------------------------- 9. 使用说明与交付物
    _page(doc, "9. 结果使用说明与交付物")
    doc.bullets(
        [
            f"适用范围：结论只对应{scope_label}和本次采集窗口；扩展到其他业务场景需要另行评估。",
            "证据口径：每条结论都可回溯到具体回答、截图和样本编号；AI 回答原文均未经事实核验。",
            "统计口径：所有比例都给出分子与分母；本批为固定问题集，结果不能外推到整体市场。",
            "品牌口径：品牌写法合并与对比品牌入选规则见附录 A、附录 C。",
        ]
    )
    doc.table(
        ["交付物", "内容", "使用方式"],
        [
            ("主报告（DOCX/PDF）", "全部结论、计算方式、附录与代表证据", "本文档"),
            (
                "样本明细（电子表格）",
                f"{answers} 条回答的逐条明细、品牌清单与同题差值",
                "按样本编号检索",
            ),
            (
                "原始证据包（ZIP）",
                "完整回答原文、回答截图、分享图片与网页快照",
                "按样本编号目录查看",
            ),
            (
                "文件校验清单（JSON）",
                "各交付文件的大小与 SHA-256 校验值",
                "用于核对文件未被改动",
            ),
        ],
        widths=(42, 82, 48),
        font_size=7.8,
    )

    # ------------------------------------------------------------ 附录 A
    _page(doc, "附录 A · 计算方式、品牌别名与方法附注")
    doc.heading("A.1 指标计算方式", level=2)
    doc.table(
        ["指标", "客户可理解的含义", "计算方式"],
        [
            (
                "品牌提及率",
                "回答是否提到" + target_brand + "或已确认的品牌别名",
                "提到品牌的回答数 ÷ 本次全部有效回答数",
            ),
            (
                "首次出现顺序",
                "品牌在回答所列品牌中第几个出现（第一个出现记为第 1 位）",
                "仅在提到品牌的回答中计算平均值",
            ),
            (
                "前三出现率",
                "品牌属于回答中最先出现的三个品牌之一",
                "符合条件的回答数 ÷ 本次全部有效回答数",
            ),
            (
                "带链接的回答",
                "回答中列出了至少一个网页链接",
                "带链接回答数 ÷ 本次全部有效回答数；链接内容未经事实核验",
            ),
        ],
        widths=(28, 74, 70),
        font_size=7.8,
    )
    doc.heading("A.2 目标品牌别名", level=2)
    aliases = delivery.get("target_aliases") or {}
    registered_aliases = [str(value) for value in aliases.get("registered") or []]
    observed_aliases = [str(value) for value in aliases.get("observed") or []]
    doc.bullets(
        [
            "计数规则：回答提到" + target_brand + "或下列已确认别名，均计为一次提到；"
            "同一回答中多个别名只计一次。",
            "已确认别名清单："
            + ("、".join(registered_aliases) if registered_aliases else "（本批未登记别名）"),
            "本批回答中实际出现的写法："
            + ("、".join(observed_aliases) if observed_aliases else "（未观测到写法）"),
        ]
    )
    doc.heading("A.3 测试结果的波动情况", level=2)
    consistency = delivery["repeat_consistency"]
    complete_pairs = int(consistency["complete_pairs"] or 0)
    expected_pairs = int(consistency["expected_pairs"] or 0)
    agreement_pairs = int(consistency["mention_agreement_pairs"] or 0)
    both_mentioned = int(consistency["both_mentioned_pairs"] or 0)
    changed_pairs = complete_pairs - agreement_pairs
    doc.paragraph(
        "大模型对同一问题的两次回答可能不同，因此每个“问题×平台×地域”组合计划测试两次，"
        "用于判断结果是稳定出现还是单次随机。该指标只描述测试本身的稳定性，"
        "不用于评价品牌优劣。"
    )
    doc.table(
        ["波动观察项", "数值", "说明"],
        [
            (
                "完成两次测试的组合",
                f"{complete_pairs}/{expected_pairs}",
                "每个组合应为 2 次独立测试",
            ),
            (
                "两次结果一致的组合",
                f"{agreement_pairs}/{complete_pairs}"
                + (
                    f"（{_pct(consistency['mention_agreement_rate'])}）"
                    if consistency.get("mention_agreement_rate") is not None
                    else ""
                ),
                "两次都提到或两次都未提到",
            ),
            (
                "提及状态发生变化的组合",
                f"{changed_pairs} 组",
                "一次提到、一次未提到，说明单次结果不能直接定论",
            ),
            (
                "两次都提到的组合的平均位次差",
                (
                    f"{_rank(consistency['mean_absolute_rank_delta'])} 位"
                    f"（仅 {both_mentioned} 组参与）"
                    if consistency.get("mean_absolute_rank_delta") is not None
                    else "—"
                ),
                "只基于两次都提到的组合，不代表全部组合的波动",
            ),
        ],
        widths=(52, 50, 70),
        font_size=7.8,
    )
    doc.paragraph(
        "注意：本批缺少账号、浏览器与出口地域台账，两次测试的独立性尚未得到正式证明；"
        "以上数值仅用于了解测量波动，不作为稳定性背书。"
    )
    interval = target.get("mention_rate_wilson_95") or []
    if len(interval) == 2:
        doc.heading("A.4 提及率的不确定区间（方法附注）", level=2)
        doc.paragraph(
            f"若把本批回答近似看作独立同分布样本，提及率的 95% Wilson 区间为 "
            f"{interval[0]:.1f}%–{interval[1]:.1f}%。本批实际是固定问题、同一问题重复"
            "观测的相关样本，该区间只描述本批测量的不确定程度，不能外推到整体市场或"
            "其他问题集。"
        )
    doc.heading("A.5 地域标签分布（不作结论）", level=2)
    doc.table(
        ["地域标签", "回答数", "提及", "提及率"],
        [
            (
                region,
                row["answers"],
                row["mentions"],
                f"{row['mention_rate_fraction']}＝{_pct(row['mention_rate'])}",
            )
            for region, row in delivery["by_region"].items()
        ],
        widths=(36, 34, 34, 68),
        font_size=8,
    )
    doc.paragraph(
        "地域为浏览器采样标签；本批账号、浏览器与出口地域台账尚未补齐，上表仅登记分布，"
        "不构成地域差异结论。"
    )

    # ------------------------------------------------------------ 附录 B
    _page(doc, "附录 B · 十二个问题完整结果")
    for group in delivery["selected_groups"]:
        rows = [row for row in delivery["question_rows"] if row["group_title"] == group["title"]]
        doc.heading(f"B.{group['index']} {group['title']}", level=2)
        doc.table(
            ["变体", "实际问题", "意图", "提及", "提及率", "平均顺序", "前三", "带链接"],
            [
                (
                    "原题" if row["question_index"] == 1 else chr(63 + row["question_index"]),
                    row["question"],
                    QUERY_INTENT_LABELS[str(row["query_intent"])],
                    row["mention_rate_fraction"],
                    _pct(row["mention_rate"]),
                    _rank(row["avg_rank"]),
                    f"{row['top_counts']['3']}/{row['answers']}",
                    f"{row['answers_with_citation']}/{row['answers']}",
                )
                for row in rows
            ],
            widths=(13, 68, 16, 20, 20, 17, 18, 18),
            font_size=7,
        )
    doc.paragraph(
        "“带链接”指回答列出了至少一个网页链接；链接是否支持回答内容未经核验。"
        "逐条回答明细见随附样本明细文件。"
    )

    # ------------------------------------------------------------ 附录 C
    _page(doc, "附录 C · 品牌对比口径与同题差值明细")
    doc.heading("C.1 对比品牌入选规则与别名", level=2)
    entity_rows = {str(row["canonical_name"]): row for row in delivery.get("entity_ranking") or []}
    doc.bullets(
        [
            "入选规则：本批全部回答中，按出现次数取前 5 个已确认为厂商品牌的名称；"
            "名单由程序生成，尚未经客户确认。",
            "写法合并：同一品牌的不同写法合并为一个统一品牌名称后再计数；下表列出"
            "本批实际观测到的写法。",
        ]
    )
    doc.table(
        ["统一品牌名称", "本批观测到的写法", "身份"],
        [
            (
                row["canonical_name"],
                "、".join(
                    str(value)
                    for value in (entity_rows.get(str(row["canonical_name"])) or {}).get(
                        "raw_aliases"
                    )
                    or [row["canonical_name"]]
                ),
                "目标品牌" if row["canonical_name"] == target_brand else "对比品牌",
            )
            for row in competitor_rows
        ],
        widths=(36, 106, 30),
        font_size=7.6,
    )
    type_counts = delivery.get("entity_type_counts") or {}
    doc.heading("C.2 未进入对比的名称类别", level=2)
    doc.paragraph(
        f"本批回答中共识别出 {scope['entity_rows_audited']} 个不同名称："
        f"{type_counts.get('company', 0)} 个厂商品牌和 {type_counts.get('product', 0)} 个产品品牌"
        "进入候选对比（产品品牌单独成行计数），"
        f"{type_counts.get('tool', 0)} 个开源工具、"
        f"{type_counts.get('institution', 0)} 个机构类名称和 "
        f"{type_counts.get('unknown', 0)} 个待确认名称不进入品牌对比。"
        "待确认名称完成分类前，对比名单可能调整。"
    )
    doc.heading("C.3 同题、同平台逐项差值", level=2)
    all_combos = list(comparison.get("same_question_platform") or [])
    by_competitor: dict[str, dict[str, Any]] = {}
    for row in all_combos:
        bucket = by_competitor.setdefault(
            str(row["competitor"]),
            {"combos": 0, "target_higher": 0, "competitor_higher": 0, "even": 0, "max_gap": 0.0},
        )
        bucket["combos"] += 1
        gap = float(row["mention_rate_gap_pp"])
        if gap > 0:
            bucket["target_higher"] += 1
        elif gap < 0:
            bucket["competitor_higher"] += 1
        else:
            bucket["even"] += 1
        bucket["max_gap"] = max(bucket["max_gap"], abs(gap))
    doc.paragraph("先给出全部组合的全景分布，再列出差值最大的组合明细；不挑选极端值单独展示。")
    doc.table(
        ["对比品牌", "组合数", "目标更高", "对比品牌更高", "持平", "最大单项差值"],
        [
            (
                competitor,
                bucket["combos"],
                bucket["target_higher"],
                bucket["competitor_higher"],
                bucket["even"],
                f"{bucket['max_gap']:.1f} 个百分点",
            )
            for competitor, bucket in sorted(
                by_competitor.items(), key=lambda item: -item[1]["max_gap"]
            )
        ],
        widths=(34, 24, 26, 30, 20, 38),
        font_size=7.6,
    )
    same_scope = sorted(
        (
            row
            for row in all_combos
            if int(row.get("target_mentions") or 0) > 0
            or int(row.get("competitor_mentions") or 0) > 0
        ),
        key=lambda row: -abs(float(row["mention_rate_gap_pp"])),
    )
    silent_combos = len(all_combos) - len(same_scope)
    doc.paragraph(
        f"下表为绝对差值最大的 20 个组合；全部 {len(all_combos)} 个组合（含双方均未提及的 "
        f"{silent_combos} 个）及每一项的分子、分母见随附样本明细文件。每个组合仅 4 条回答"
        "（2 个地域标签 × 2 次测试），差值波动大，只用于定位，不代表稳定差距。"
    )
    doc.table(
        ["平台", "问题（截短）", "对比品牌", "目标提及", "对比提及", "提及率差"],
        [
            (
                _PLATFORM_LABELS.get(row["platform"], row["platform"]),
                _short_question(row["question"], 16),
                row["competitor"],
                f"{row['target_mentions']}/{row['answers']}",
                f"{row['competitor_mentions']}/{row['answers']}",
                f"{float(row['mention_rate_gap_pp']):+.1f}",
            )
            for row in same_scope[:20]
        ],
        widths=(19, 57, 27, 17, 17, 15),
        font_size=6.2,
    )
    doc.paragraph("差值＝目标品牌 − 对比品牌，单位为百分点；正值表示目标品牌提及率更高。")

    # ------------------------------------------------------------ 附录 D
    _page(doc, "附录 D · 代表回答关键片段与所列链接")
    doc.paragraph(
        "以下内容均为 AI 生成原文，未经事实核验，不代表评测方结论。"
        "回答中提及的行业份额、客户案例等内容来自 AI，不应作为事实引用；"
        "关键片段围绕品牌出现位置截取，逐字完整原文与原始截图保存在随附证据包中。"
    )
    for representative_index, item in enumerate(representatives):
        number = int(item["display_number"])
        if representative_index:
            doc.page_break()
        doc.heading(
            f"D.{number} {item['platform_label']}代表回答 · {item['group_title']}",
            level=2,
        )
        answer_id = str(item.get("answer_pub_id") or "")
        sample = sample_by_answer.get(answer_id, {})
        doc.table(
            ["字段", "内容"],
            [
                ("实际问题", item["question"]),
                ("平台与地域标签", f"{item['platform_label']} · {item['region']}"),
                ("采集时间", _fmt_datetime(item.get("capture_time"))),
                ("样本编号", str(sample.get("sample_id") or "—")),
            ],
            widths=(31, 141),
            font_size=8,
        )
        full_text = _clean_answer(sample.get("response_text"), limit=20000)
        excerpt = _key_excerpt(full_text, target_brand, limit=1100)
        if excerpt:
            doc.heading("关键片段", level=3)
            doc.paragraph(excerpt)
        citations = list(item.get("citations") or [])
        doc.heading(f"D.{number}.1 回答所列全部链接（{len(citations)} 条）", level=3)
        _citation_table(doc, citations)

    # ------------------------------------------------------------ 附录 E
    _page(doc, "附录 E · 版本、审批与本批限制")
    governance = facts.get("document_governance") or {}
    doc.table(
        ["治理字段", "记录"],
        [
            (
                "项目与服务",
                f"{facts.get('project_name') or '—'} · 服务1 · 品牌GEO推荐结果评测",
            ),
            (
                "版本与状态",
                f"{governance.get('version') or 'V1.0'} · "
                f"{release_state_label(str(facts.get('document_status') or ''))}",
            ),
            (
                "编制",
                f"{governance.get('prepared_by') or 'GEO 项目组'} · "
                f"{governance.get('prepared_date') or str(facts['generated_at'])[:10]}",
            ),
            (
                "复核",
                f"{governance.get('reviewed_by') or '待复核'} · "
                f"{governance.get('reviewed_date') or '待定'}",
            ),
            (
                "批准",
                f"{governance.get('approved_by') or '待批准'} · "
                f"{governance.get('approved_date') or '待定'}",
            ),
            ("保密级别", "客户机密—仅限指定项目组"),
        ],
        widths=(36, 136),
        font_size=7.8,
    )
    doc.heading("E.1 本批限制", level=2)
    limitations = []
    if answers != expected_answers:
        limitations.append(
            f"样本：计划 {expected_answers} 条回答，当前 {answers} 条；"
            + _sample_gap_text(delivery)
            + "补齐后按同一问题集重算。"
        )
    limitations.extend(
        [
            "选题登记：本批问题在测试开始前未形成书面登记记录；本报告按报价单问题顺序"
            "还原测试范围，后续批次将在采样前完成书面登记。",
            "地域与独立性：本批缺少账号、浏览器与出口地域台账，两次测试的独立性未得到"
            "正式证明；地域仅作采集标签，不作结论。",
            "对比品牌名单：由程序按出现频次生成，尚未经客户确认；回答中另有 "
            f"{scope['unclassified_entities']} 个待确认名称未进入对比。",
            "文档状态：本报告为内部审核稿；完成上述补齐并经人工复核、批准后，"
            "方可升级为客户交付候选稿。",
        ]
    )
    doc.bullets(limitations)
    doc.heading("E.2 修订记录", level=2)
    doc.table(
        ["版本", "日期", "修订说明", "状态"],
        [
            (
                governance.get("version") or "V1.0",
                governance.get("prepared_date") or str(facts["generated_at"])[:10],
                "客户视角重构：事实结论先行，意图分类解读，方法、别名与全量结果入附录",
                release_state_label(str(facts.get("document_status") or "")),
            )
        ],
        widths=(24, 31, 89, 28),
        font_size=8,
    )
    doc.paragraph(
        "版权与使用：本报告及其证据包仅供指定客户项目组在本项目范围内使用。未经书面许可，"
        "不得向无关第三方传播。"
    )

    payload = bytes(doc.save())
    # Renderer-level defence.  OOXML is compressed, so the publication QA repeats
    # this check on extracted document/PDF text after conversion.
    visible_values = " ".join(
        [paragraph.text for paragraph in doc.document.paragraphs]
        + [cell.text for table in doc.document.tables for row in table.rows for cell in row.cells]
    )
    found = [value for value in _CLIENT_FORBIDDEN if value in visible_values]
    if found:
        raise ValueError("customer_report_internal_language:" + ",".join(found))
    return payload


__all__ = ["render_service1_delivery_docx"]
