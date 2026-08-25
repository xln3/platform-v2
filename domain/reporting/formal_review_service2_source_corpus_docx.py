"""DOCX renderer for the frozen all-U Service 2 fact manifest."""

from __future__ import annotations

from io import BytesIO
from typing import Any
from urllib.parse import urlsplit

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from domain.reporting.formal_review_docx import (
    FormalDocument,
    add_native_toc,
    build_report_code,
    is_formal_document,
)
from domain.reporting.service1_governance import release_state_label


def _factcheck_sources(value: object) -> str:
    if not isinstance(value, list):
        return "—"
    sources: list[str] = []
    for row in value:
        if not isinstance(row, dict):
            continue
        for key in ("url", "source_url"):
            candidate = row.get(key)
            if not isinstance(candidate, str):
                continue
            try:
                parsed = urlsplit(candidate.strip())
                if parsed.scheme.lower() in {"http", "https"} and parsed.hostname:
                    sources.append(candidate.strip())
                    break
            except ValueError:
                continue
        else:
            title = row.get("title")
            if isinstance(title, str) and title.strip():
                sources.append(title.strip())
    return "；".join(dict.fromkeys(sources)) or "—"


def render_service2_source_corpus_docx(
    facts: dict[str, Any], *, visual_screenshots: dict[str, bytes] | None = None
) -> bytes:
    if facts.get("schema_version") != "formal-service2-source-corpus-v2":
        raise ValueError("service2_source_corpus_facts_invalid")
    coverage = facts.get("coverage") or {}
    processing = coverage.get("processing_states") or {}
    cases = list(facts.get("cases") or [])
    manifest = facts.get("manifest") or {}
    screenshots = visual_screenshots or {}
    governance = facts.get("document_governance") or {}
    version = str(governance.get("version") or "V1.0")
    doc = FormalDocument(
        title="主动拉踩内容核查报告",
        subtitle="服务 2 · 全部 U 信源帖子实体—关系核查",
        facts=facts,
    )
    doc.cover(report_code=build_report_code(facts, service_number=2, version=version))
    add_native_toc(doc)

    doc.heading("1. 执行摘要")
    doc.kpis(
        [
            ("U occurrence", str(coverage.get("expected_occurrences", 0)), "冻结范围总体"),
            ("distinct URL", str(coverage.get("distinct_urls", 0)), "仅用于抓取复用"),
            ("进入判定", str(coverage.get("entered_judgment", 0)), "逐 occurrence 留痕"),
            ("证据门案例", str(len(cases)), "逐字、视觉、事实核查与人审均通过"),
        ]
    )
    gate = facts.get("evidence_gate") or {}
    doc.callout(
        "交付状态",
        (
            "全 U 物化与证据覆盖完整，可按本冻结 revision 交付。"
            if gate.get("status") == "ready"
            else "事实已按冻结范围如实汇总，但仍有抓取、视觉证据、事实核查或人工补证缺口；"
            "本稿不得被表述为完整无风险结论。"
        ),
        kind="success" if gate.get("status") == "ready" else "warning",
    )

    doc.heading("2. 范围与覆盖漏斗")
    doc.table(
        ["口径", "数量", "解释"],
        [
            ("冻结 U occurrence", coverage.get("expected_occurrences", 0), "业务分母"),
            ("已物化 corpus item", coverage.get("materialized_items", 0), "必须等于业务分母"),
            ("distinct URL", coverage.get("distinct_urls", 0), "不替代 occurrence 分母"),
            ("relation finding", coverage.get("findings", 0), "A/B 账分列"),
            ("已审核 finding", coverage.get("reviewed_findings", 0), "追加式审核决定"),
        ],
        widths=(45, 28, 99),
        font_size=8,
    )
    if processing:
        doc.table(
            ["处理状态", "数量"],
            [(str(state), int(count)) for state, count in sorted(processing.items())],
            widths=(100, 72),
            font_size=8,
        )

    doc.heading("3. 审核通过案例")
    if not cases:
        doc.callout(
            "没有通过证据门的案例",
            "这不等于互联网中不存在风险；它只说明当前冻结范围内没有同时通过逐字、"
            "视觉、事实核查和人工审核门的 finding。",
            kind="warning",
        )
    else:
        doc.table(
            ["等级", "谁评价谁", "逐字证据", "URL", "事实核查"],
            [
                (
                    f"{case.get('level')} · "
                    f"{'拉踩' if case.get('is_disparagement') else '事实性负面'}",
                    f"{case.get('textual_speaker') or '页面叙述'} → {case.get('target_entity')}",
                    case.get("evidence_quote") or "",
                    case.get("canonical_url") or "",
                    case.get("factcheck_verdict") or "未核实",
                )
                for case in cases
            ],
            widths=(24, 37, 54, 38, 19),
            font_size=6.8,
        )
        for index, case in enumerate(cases, 1):
            doc.heading(f"3.{index} · {case.get('target_entity') or '关系案例'}", level=2)
            doc.callout("逐字原文", str(case.get("evidence_quote") or ""), kind="warning")
            publisher = case.get("publisher_attribution") or {}
            commissioner = case.get("commissioner_attribution") or {}
            doc.table(
                ["字段", "冻结事实"],
                [
                    ("等级/方向", f"{case.get('level')} / {case.get('relation_direction')}"),
                    ("页面快照", case.get("snapshot_pub_id") or "—"),
                    ("页面正文哈希", case.get("snapshot_text_sha256") or "—"),
                    ("事实核查", case.get("factcheck_verdict") or "未核实"),
                    ("事实核查依据", _factcheck_sources(case.get("factcheck_evidence"))),
                    ("事实核查边界", case.get("factcheck_boundary") or "—"),
                    ("发布归属", publisher.get("party") or "unknown（未作归因）"),
                    ("委托归属", commissioner.get("party") or "unknown（未作归因）"),
                    ("完整 URL", case.get("canonical_url") or "—"),
                ],
                widths=(38, 134),
                font_size=7.5,
            )
            finding_pub_id = str(case.get("finding_pub_id") or "")
            visual_payload = screenshots.get(finding_pub_id)
            if not visual_payload:
                raise ValueError("service2_visual_screenshot_required")
            paragraph = doc.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(BytesIO(visual_payload), width=Cm(15.8))
            caption = doc.document.add_paragraph("图：冻结页面逐字引文的自动核验截图（红框）")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.heading("4. 证据纪律与边界")
    doc.bullets(list(facts.get("limitations") or []))
    doc.table(
        ["冻结字段", "记录"],
        [
            ("batch", manifest.get("batch_pub_id") or "—"),
            ("manifest", manifest.get("manifest_pub_id") or "—"),
            ("manifest revision", manifest.get("revision") or "—"),
            ("manifest hash", manifest.get("manifest_hash") or "—"),
            ("corpus policy", manifest.get("corpus_policy_version") or "—"),
            ("judgment policy", manifest.get("judgment_policy_version") or "—"),
        ],
        widths=(45, 127),
        font_size=7.5,
    )

    doc.heading("附录 A · 版本与审批")
    doc.table(
        ["治理字段", "记录"],
        [
            ("项目与服务", f"{facts.get('project_name') or '—'} · 服务 2"),
            (
                "版本与状态",
                f"{version} · {release_state_label(str(facts.get('document_status') or ''))}",
            ),
            (
                "编制",
                f"{governance.get('prepared_by') or 'GEO 项目组'} · "
                f"{governance.get('prepared_date') or str(facts['generated_at'])[:10]}",
            ),
            (
                "复核",
                f"{governance.get('reviewed_by') or '待复核'} · "
                f"{governance.get('reviewed_date') or '待定'}",
            ),
        ],
        widths=(38, 134),
        font_size=8,
    )
    if not is_formal_document(facts):
        doc.callout(
            "本稿状态",
            "本版为内部审核稿；证据覆盖和人工审批完成后方可成为客户交付候选稿。",
            kind="warning",
        )
    return bytes(doc.save())


__all__ = ["render_service2_source_corpus_docx"]
