"""Customer-readable Service-2 V2 DOCX renderer with visual evidence."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from domain.reporting.formal_review_docx import (
    FONT_BOLD,
    MUTED,
    FormalDocument,
    _fmt_datetime,
    _set_font,
    add_native_toc,
    build_report_code,
    is_formal_document,
)
from domain.reporting.service1_governance import release_state_label

_MODE_LABELS = {"deep_think": "深度思考", "normal": "快速", "web": "联网检索"}
_CLIENT_FORBIDDEN = (
    "执行行",
    "采集批次",
    "口径修正",
    "系统缺陷",
    "运营复核清单",
    "deep_think",
    "manifest",
    "工作表",
)


def _caption(doc: FormalDocument, text: str) -> None:
    paragraph = doc.document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    _set_font(run, size=7.8)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(MUTED)


def _hyperlink(paragraph: Any, url: str, text: str | None = None) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text or url
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def _image_stream(image: Image.Image) -> BytesIO:
    stream = BytesIO()
    image.save(stream, format="PNG", optimize=True)
    stream.seek(0)
    return stream


def _answer_views(
    payload: bytes, anchor: dict[str, Any] | None
) -> tuple[BytesIO, BytesIO | None, str]:
    with Image.open(BytesIO(payload)) as source:
        full = source.convert("RGB")
    full_stream = _image_stream(full)
    if not anchor or not isinstance(anchor.get("bbox"), list):
        return full_stream, None, "原图中没有经复核的可视命中坐标"

    raw_bbox = [int(value) for value in anchor["bbox"]]
    if len(raw_bbox) != 4:
        return full_stream, None, "历史标注坐标无效"
    x, y, width, height = raw_bbox
    x0 = max(0, min(x, full.width - 1))
    y0 = max(0, min(y, full.height - 1))
    x1 = max(x0 + 1, min(x + width, full.width))
    y1 = max(y0 + 1, min(y + height, full.height))

    horizontal_margin = max(90, int((x1 - x0) * 0.45))
    vertical_margin = max(150, int((y1 - y0) * 1.6))
    crop_left = max(0, x0 - horizontal_margin)
    crop_top = max(0, y0 - vertical_margin)
    crop_right = min(full.width, x1 + horizontal_margin)
    crop_bottom = min(full.height, y1 + vertical_margin)
    if crop_right - crop_left < 850:
        center = (x0 + x1) // 2
        crop_left = max(0, center - 520)
        crop_right = min(full.width, center + 520)
    crop = full.crop((crop_left, crop_top, crop_right, crop_bottom))
    draw = ImageDraw.Draw(crop)
    stroke = max(5, min(crop.size) // 120)
    draw.rectangle(
        (
            x0 - crop_left,
            y0 - crop_top,
            x1 - crop_left,
            y1 - crop_top,
        ),
        outline="#DC2626",
        width=stroke,
    )
    label = str(anchor.get("label") or "AI 回答命中表述")
    font = ImageFont.truetype(str(FONT_BOLD), max(24, crop.width // 36))
    padding = 12
    label_box = draw.textbbox((0, 0), label, font=font)
    label_width = min(crop.width - 24, label_box[2] - label_box[0] + padding * 2)
    label_height = label_box[3] - label_box[1] + padding * 2
    badge_y = max(8, y0 - crop_top - label_height - 12)
    draw.rounded_rectangle(
        (8, badge_y, 8 + label_width, badge_y + label_height),
        radius=8,
        fill="#FFFFFF",
        outline="#DC2626",
        width=max(3, stroke // 2),
    )
    draw.text((8 + padding, badge_y + padding // 2), label, font=font, fill="#991B1B")
    method = str(anchor.get("method") or "")
    if method.startswith("dom_") or method.startswith("ocr_"):
        anchor_note = "红框按采集时保存的文本位置绘制"
    else:
        anchor_note = "红框按人工复核的位置绘制"
    return full_stream, _image_stream(crop), f"{anchor_note}；红框仅标命中原句"


def _add_answer_screenshot(
    doc: FormalDocument,
    case: dict[str, Any],
    payload: bytes | None,
) -> None:
    if not payload:
        doc.callout(
            "回答截图缺口",
            "采集记录指向回答截图，但本次生成未能加载图像。正式签发前必须恢复资产。",
            kind="warning",
        )
        return
    full, crop, note = _answer_views(payload, case.get("answer_anchor"))
    if crop is None:
        doc.callout(
            "回答图未展示",
            "本条没有可复核的命中坐标。报告保留上方逐字原句和上下文，但不放入"
            "缩小后不可读、又无法准确标注的整页截图。原始图仍在受控证据库中留存。",
            kind="warning",
        )
        return

    del full
    with Image.open(crop) as crop_image:
        crop_ratio = crop_image.width / max(crop_image.height, 1)
    crop.seek(0)
    crop_width = min(16.2, 12.6 * crop_ratio)
    paragraph = doc.document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(crop, width=Cm(max(7.0, crop_width)))
    _caption(
        doc,
        f"AI 回答命中区域及必要上下文；红框仅强调判定原句。{note}。",
    )


def _source_capture_status(row: dict[str, Any] | None) -> str:
    if not row:
        return "未执行网页快照"
    if row.get("capture_status") != "captured":
        return f"抓取失败：{row.get('error') or '未知错误'}"
    if row.get("content_status") != "ok":
        return (
            f"页面不可作为核查证据：content_status={row.get('content_status')}，"
            f"HTTP={row.get('http_status') or '—'}"
        )
    if not row.get("matched_terms"):
        return "页面已抓取，但未找到可逐字标注的核查锚点"
    return "核查页面已抓取且锚点可见"


def _add_source_capture(
    doc: FormalDocument,
    source: dict[str, Any],
    capture: dict[str, Any] | None,
    *,
    number: str,
    max_width_cm: float = 14.8,
) -> bool:
    url = str(source.get("url") or "")
    status = _source_capture_status(capture)
    if (
        capture is None
        or capture.get("capture_status") != "captured"
        or capture.get("content_status") != "ok"
        or not capture.get("matched_terms")
        or not isinstance(capture.get("payload"), bytes)
    ):
        # Broken, 404 or unanchored URLs remain in the operator capture manifest, not
        # in a customer evidence card.  Listing an unusable URL adds no evidence and
        # can be mistaken for support merely because it is visible in the report.
        return False
    paragraph = doc.document.add_paragraph()
    lead = paragraph.add_run("原网址：")
    _set_font(lead, size=8.5)
    lead.bold = True
    _hyperlink(paragraph, url)
    doc.paragraph(status, bold_lead="网页快照状态：")
    payload = capture.get("payload")
    assert isinstance(payload, bytes)
    paragraph = doc.document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(BytesIO(payload), width=Cm(max_width_cm))
    fallback = capture.get("transport_fallback")
    fallback_note = "；HTTPS 拒绝连接后改用同主机 HTTP 页面，已在审计清单记录" if fallback else ""
    terms = "、".join(str(value) for value in capture.get("matched_terms") or [])
    _caption(
        doc,
        f"图 {number}  事实核查网页可见区域；红框/黄底为网页中逐字锚点：{terms}{fallback_note}。",
    )
    role = str(source.get("role") or "")
    if "反证" in role:
        doc.callout(
            "如何使用本图",
            "本图可核对公开页面中的具体数字和口径；引用时应注意网页口径与回答口径"
            "是否一致，不能把不同口径的数字直接互相替换。",
            kind="success",
        )
    else:
        doc.callout(
            "如何使用本图",
            "本图只能证明目标品牌官网存在相应产品/能力描述；它不能证明 AI 回答中的"
            "“弱”、高低顺序、价格优劣或“不如”结论，因此本条仍为证据不足。",
        )
    return True


def _toc(doc: FormalDocument, *, answer_cases: int, source_cases: int) -> None:
    del answer_cases, source_cases
    add_native_toc(doc)
    doc.page_break()


def _case_page(
    doc: FormalDocument,
    case: dict[str, Any],
    *,
    answer_screenshot: bytes | None,
    source_captures: dict[str, dict[str, Any]],
    page_break: bool = True,
    section_label: str | None = None,
) -> None:
    if page_break:
        doc.page_break()
    heading_label = section_label or f"5.{str(case['case_id']).removeprefix('A-')}"
    doc.heading(f"{heading_label} {case['case_id']} · {case['direction']}")
    if case.get("customer_scope_note"):
        doc.callout(
            "补充案例边界",
            str(case["customer_scope_note"]),
            kind="info",
        )
    verdict_kind = "warning"
    doc.callout(
        str(case["expression_verdict"]),
        str(case["customer_conclusion"]),
        kind=verdict_kind,
    )
    doc.table(
        ["字段", "实测内容"],
        [
            ("表述方向", case["direction"]),
            (
                "采集环境",
                f"{case['platform_label']} · {case['region']} · "
                f"{_MODE_LABELS.get(str(case['mode']), str(case['mode'] or '—'))} · "
                f"{_fmt_datetime(case['capture_time'])}",
            ),
            ("原始问题", case["question"]),
            ("表述类型", case["statement_type"]),
            ("表达判定", case["expression_verdict"]),
            ("事实判定", case["fact_verdict"]),
        ],
        widths=(29, 143),
        font_size=8.1,
    )
    doc.paragraph(str(case["attribution"]), bold_lead="归因边界：")
    doc.heading("原回答与命中表述", level=2)
    doc.callout("AI 回答原句", str(case["evidence_quote"]), kind="warning")
    doc.paragraph(str(case["answer_context"]), bold_lead="上下文摘录：")
    paragraph = doc.document.add_paragraph()
    lead = paragraph.add_run("平台入口：")
    _set_font(lead, size=8.5)
    lead.bold = True
    platform_url = str(case.get("platform_entry_url") or "")
    if platform_url:
        _hyperlink(paragraph, platform_url)
    note = paragraph.add_run(f"  （{case['platform_url_note']}）")
    _set_font(note, size=8)
    note.font.color.rgb = RGBColor.from_string(MUTED)
    _add_answer_screenshot(doc, case, answer_screenshot)

    doc.heading("事实核查解释", level=2)
    doc.paragraph(str(case["why"]), bold_lead="为什么这样判：")
    doc.paragraph(str(case["customer_conclusion"]), bold_lead="客户可用结论：")
    sources = list(case.get("factcheck_sources") or [])
    if not sources:
        doc.callout(
            "公开来源缺口",
            "本条没有形成公开核查 URL，不能进入确定风险清单。",
            kind="warning",
        )
    visible_sources = []
    for source in sources:
        capture = source_captures.get(str(source.get("url") or ""))
        if (
            capture
            and capture.get("capture_status") == "captured"
            and capture.get("content_status") == "ok"
            and capture.get("matched_terms")
            and isinstance(capture.get("payload"), bytes)
        ):
            visible_sources.append((source, capture))
    if sources and not visible_sources:
        doc.callout(
            "公开核查页缺口",
            "本案未取得可逐字定位的有效公开核查页；404、抓取失败和未命中锚点的"
            "链接不在本报告展示，也不作为支持或反证。",
            kind="warning",
        )
    for source_index, (source, capture) in enumerate(visible_sources, 1):
        doc.heading(f"核查来源 {source_index}", level=3)
        _add_source_capture(
            doc,
            source,
            capture,
            number=f"5-{case['case_id']}-{source_index}",
            max_width_cm=13.0 if len(sources) > 1 else 14.8,
        )


def _source_case_page(
    doc: FormalDocument,
    case: dict[str, Any],
    screenshot_payload: bytes | None,
) -> None:
    doc.heading(f"5.{case['case_id']} · 公开信源正文风险", level=2)
    doc.callout(
        str(case["expression_verdict"]),
        str(case["customer_conclusion"]),
        kind="warning",
    )
    doc.table(
        ["字段", "实测内容"],
        [
            ("表述方向", case["direction"]),
            ("信源网站", case["platform_label"]),
            ("命中原句", case["evidence_quote"]),
            ("表达判定", case["expression_verdict"]),
            ("事实判定", case["fact_verdict"]),
        ],
        widths=(29, 143),
        font_size=8.1,
    )
    url = str(case.get("source_url") or "")
    if url:
        paragraph = doc.document.add_paragraph()
        lead = paragraph.add_run("原信源网页：")
        _set_font(lead, size=8.5)
        lead.bold = True
        _hyperlink(paragraph, url, "打开原网页")
    descriptor = case.get("source_screenshot") or {}
    if screenshot_payload:
        anchor = {
            "bbox": descriptor.get("bbox"),
            "label": f"{case.get('target_brand') or '目标品牌'}所在段落",
        }
        full, crop, note = _answer_views(screenshot_payload, anchor)
        image = crop or full
        paragraph = doc.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(image, width=Cm(15.8))
        _caption(
            doc,
            f"信源网页中目标品牌所在段落的可视证据；{note}。完整原始网页截图继续保留在证据存储中。",
        )
    else:
        doc.callout(
            "信源截图缺口",
            "该信源案例尚未载入目标品牌段落截图；正式签发前必须补齐 DOM 锚点和可读上下文图。",
            kind="warning",
        )
    doc.numbered(
        [
            "该原句来自已成功抽取的信源正文，而不是 AI 回答文本。",
            "系统先确认正文提及目标品牌，再检查品牌所在上下文段落是否存在贬低别人或被贬低的表述。",
            str(case["attribution"]),
        ]
    )


def render_service2_v2_docx(
    facts: dict[str, Any],
    *,
    answer_screenshots: dict[str, bytes] | None = None,
    source_captures: dict[str, dict[str, Any]] | None = None,
    source_case_screenshots: dict[str, bytes] | None = None,
) -> bytes:
    """Render corrected Service-2 V2 facts and real visual evidence."""

    service2 = facts["service2"]
    delivery = service2.get("delivery_v2")
    if not isinstance(delivery, dict):
        raise ValueError("service2.delivery_v2_missing")
    answer_screenshots = answer_screenshots or {}
    source_captures = source_captures or {}
    source_case_screenshots = source_case_screenshots or {}
    citations = delivery["citation_funnel"]
    fetch = delivery["source_fetch"]
    judgments = delivery["judgment_funnel"]
    cases = list(delivery["cases"])
    source_cases = list(delivery.get("source_cases") or [])
    source_audit = delivery["source_content_audit"]
    verdicts = delivery["case_verdict_counts"]

    doc = FormalDocument(
        title="品牌 GEO 内容生态风险核查报告",
        subtitle="服务 2 · AI 拉踩表述、公开事实核查与可视证据",
        facts=facts,
    )
    version = str((facts.get("document_governance") or {}).get("version") or "V1.0")
    doc.cover(report_code=build_report_code(facts, service_number=2, version=version))
    _toc(doc, answer_cases=len(cases), source_cases=len(source_cases))

    doc.heading("1. 执行摘要")
    doc.kpis(
        [
            ("独立合格回答", str(citations["eligible_answers"]), "按独立回答去重"),
            (
                "带引用回答",
                str(citations["answers_with_citation"]),
                f"共 {citations['citation_references']:,} 条引用记录",
            ),
            (
                "去重风险线索",
                str(judgments["unique_cases"]),
                f"回答 {len(cases)} 项 / 信源正文 {len(source_cases)} 项",
            ),
            ("公开证据冲突", str(verdicts.get("refuted", 0)), "其余为证据不足/未定案"),
        ]
    )
    total_cases = len(cases) + len(source_cases)
    if total_cases:
        core_text = (
            f"本报告只保留直接涉及{facts['target_brand']}的线索："
            f"{len(cases)} 项来自 AI 回答，{len(source_cases)} 项来自公开信源正文。"
            f"{verdicts.get('refuted', 0)} 项存在公开数据冲突，"
            f"{verdicts.get('unverifiable', 0)} 项缺少同口径公开证据。"
            "可以确认的是 AI 回答出现了无来源排序、贬低性比较或不完整负向标签；"
            "没有证据证明这些内容由竞品或第三方撰写、投放。"
        )
    else:
        core_text = (
            f"本批未发现直接涉及{facts['target_brand']}的拉踩式或负向比较线索："
            "AI 回答 0 项、公开信源正文 0 项。"
            "该结论只对应本批实际核查范围：完成表述判定的回答 "
            f"{judgments['ok_distinct_answers']}/{citations['eligible_answers']} 条、"
            f"完成正文核查的网页 {fetch['ok']}/{citations['unique_canonical_urls']:,} 个；"
            "未覆盖部分不构成“没有风险”的证明，也不外推到全网。"
        )
    doc.callout("核心结论", core_text, kind="warning" if total_cases else "success")
    doc.paragraph(
        "本报告把三件事分开回答：第一，AI 回答是否出现拉踩式或负向比较表达；"
        f"第二，被引用的公开信源是否提及{facts['target_brand']}，以及品牌所在段落是否"
        "存在贬低别人或被贬低的表述；第三，具体比较是否有公开事实支持。"
        "‘表达成立、事实无法核验’并不矛盾，"
        "也不等于该表述为真。"
    )
    doc.table(
        ["客户应采取的动作", "适用线索", "含义"],
        [
            ("立即纠正/平台反馈", "公开证据冲突", "原回答数字无来源且与公开数据冲突"),
            ("保留监测并补证", "拉踩表达成立、事实不足", "先保留截图与原句，不宣称虚假事实已坐实"),
            ("无需立即动作", "原句不完整、暂不定性", "由评测方补全上下文并复核后再判断"),
        ],
        widths=(40, 56, 76),
        font_size=8.2,
    )

    doc.page_break()
    doc.heading("2. 本次核查的数据范围")
    doc.paragraph(
        "本章说明风险线索的数据来源：多少条独立 AI 回答、其中多少条列出了公开链接、"
        "多少网页完成了正文核查。所有数量均按独立回答、独立网页去重；同一回答经多轮"
        "复核时只计一次。"
    )
    doc.table(
        ["阶段", "实际数量", "口径", "客户应如何理解"],
        [
            ("合格 AI 回答", citations["eligible_answers"], "独立回答", "本批核查的对象"),
            (
                "带引用回答",
                citations["answers_with_citation"],
                "独立回答",
                "回答中至少列出 1 个链接",
            ),
            (
                "回答引用记录",
                f"{citations['citation_references']:,}",
                "按最新一轮分析结果",
                f"{citations['unique_canonical_urls']:,} 个不同链接地址",
            ),
            (
                "网页正文核查",
                fetch["documents"],
                "独立网页",
                f"成功读取正文 {fetch['ok']} 份；并非每条链接都核查了正文",
            ),
            (
                "回答表述判定",
                judgments["ok_distinct_answers"],
                "独立回答",
                "完成拉踩表述检查的回答数",
            ),
            (
                "信源正文判定",
                judgments["ok_distinct_source_documents"],
                "独立网页",
                f"其中目标品牌段落风险线索 {len(source_cases)} 项",
            ),
            (
                "风险线索",
                judgments["unique_cases"],
                "按命中原句去重",
                f"回答 {len(cases)} 项 / 信源正文 {len(source_cases)} 项",
            ),
        ],
        widths=(35, 28, 53, 56),
        font_size=7.8,
    )
    doc.numbered(
        [
            f"表述判定覆盖 {judgments['ok_distinct_answers']}/{citations['eligible_answers']} "
            "条合格回答；未判定的回答不参与本报告结论，也不推断为无风险。",
        ]
    )
    doc.heading("2.1 为什么回答中的链接多、实际核查的网页少", level=2)
    doc.paragraph(
        f"分析环节完整保留了全部链接：{citations['answers_with_citation']} 条带引用回答共 "
        f"{citations['citation_references']:,} 条引用，平均每条带引用回答 "
        f"{citations['avg_refs_cited_answers']:.2f} 条、"
        f"最多 {citations['max_refs_one_answer']} 条。"
        "正文核查则需要逐网页访问读取，本批只覆盖了其中一部分。"
    )
    doc.table(
        ["来源正文核查环节", "当前值", "含义"],
        [
            (
                "回答中列出的链接",
                f"{citations['citation_references']:,} 条 / "
                f"{citations['unique_canonical_urls']:,} 个不同地址",
                "分析环节完整保留",
            ),
            (
                "进入正文核查的网页",
                f"{fetch['documents']} 份",
                "本批实际读取了正文的网页",
            ),
            ("正文读取成功", f"{fetch['ok']} 份", "可进入品牌提及与段落风险检查"),
            (
                "回答与网页的对应关系",
                f"{fetch['answer_document_relations']} 条 / "
                f"{fetch['answers_with_planned_documents']} 条回答",
                "同一网页只读取一次，但会关联到所有列出它的回答",
            ),
        ],
        widths=(39, 63, 70),
        font_size=8,
    )
    doc.numbered(
        [
            "本批正文核查只覆盖了部分链接；未核查的网页不构成“无风险”结论。",
            "正文核查按回答逐一规划：同一网页去重读取，并关联到所有列出它的回答。",
            "各级覆盖数量以本表披露为准；未覆盖部分在附录的限制说明中列出。",
        ]
    )

    doc.heading("2.2 信源正文如何检查拉踩", level=2)
    doc.table(
        ["检查阶段", "当前数量", "判定规则"],
        [
            ("正文读取成功", source_audit["successful_documents"], "网页正文可读取"),
            (
                "目标品牌可视提及",
                source_audit["documents_with_target_brand_visual_anchor"],
                f"正文中逐字出现{facts['target_brand']}，并保存了页面位置标注",
            ),
            (
                "完成信源段落判定",
                source_audit["judged_distinct_documents"],
                "只检查目标品牌所在上下文段落",
            ),
            (
                "形成客户风险线索",
                source_audit["flagged_target_brand_cases"],
                "段落中存在目标品牌贬低别人或被贬低的可核对原句",
            ),
        ],
        widths=(42, 29, 101),
        font_size=7.8,
    )
    doc.numbered(
        [
            source_audit["method"],
            f"不直接涉及{facts['target_brand']}的竞品间比较不进入本报告。",
            "本批只核查了部分被引用网页的正文，因此当前“没有信源正文风险线索”"
            "不能解释为全网没有风险。",
        ]
    )

    doc.heading("2.3 截图与命中位置覆盖", level=2)
    doc.paragraph(
        "每条客户案例都应同时具备可核对原句、回答截图和命中位置标注；信源正文案例还应"
        "具备网页截图及命中段落。缺少任一环节时，本报告明确披露证据缺口，不自动"
        "推测命中位置，也不把读取失败的网页当作证据。"
    )

    doc.page_break()
    doc.heading("3. 判定方法与证据边界")
    doc.table(
        ["层级", "回答的问题", "本报告输出"],
        [
            ("表达判定", "是否存在高低排序、贬低、不如、弱等拉踩式表达", "成立 / 线索命中未定性"),
            ("事实核查", "公开资料能否支持或推翻具体比较", "公开冲突 / 有支持 / 无法核验"),
            ("主体归因", "谁生成了这句话、是否能认定竞品投放", "仅归因 AI 回答；无作者证据不外推"),
            ("可视证据", "原回答与核查网页能否逐字定位", "原图、红框、URL、抓取状态与证据边界"),
        ],
        widths=(31, 78, 63),
        font_size=8,
    )
    doc.heading("3.1 结论词的直白解释", level=2)
    doc.bullets(
        [
            "公开证据冲突：已有页面提供了相反数字或明确口径；仍需避免跨口径替换。",
            "无法核验：没有同指标、同场景、同时间、同样本的可靠比较；不代表原话真实。",
            "拉踩式表达成立：句子形式存在无依据的高低比较；不自动等同于事实虚假。",
            "暂不定性：原句或表头不完整，无法说明负向词对应的维度，需在后续批次补采后再判断。",
        ]
    )

    doc.heading("4. 风险线索总表")
    doc.heading("4.1 AI 回答中直接涉及目标品牌的线索", level=2)
    if cases:
        doc.table(
            ["案例", "对象与比较方向", "命中原句", "表达判定", "事实判定"],
            [
                (
                    case["case_id"],
                    case["direction"],
                    str(case["evidence_quote"])[:78]
                    + ("…" if len(str(case["evidence_quote"])) > 78 else ""),
                    case["expression_verdict"],
                    case["fact_verdict"],
                )
                for case in cases
            ],
            widths=(14, 39, 57, 32, 30),
            font_size=6.9,
        )
    else:
        doc.callout(
            "本批未发现回答侧线索",
            f"在完成表述判定的回答中，没有出现直接涉及{facts['target_brand']}的拉踩式或"
            "负向比较原句。",
            kind="success",
        )
    doc.numbered(
        [
            f"本报告只呈现直接涉及{facts['target_brand']}的线索；"
            f"另有 {judgments['excluded_competitor_only_cases']} 项不直接涉及目标品牌的"
            "比较类线索留在内部复核，不在本报告展示，也不计入风险数量。",
            "同一回答、同一原句经多轮复核时，本报告合并为一个案例。",
        ]
    )

    doc.heading("4.2 被引用信源正文中的目标品牌段落线索", level=2)
    if source_cases:
        doc.table(
            ["案例", "来源网站", "命中原句", "表达判定", "事实判定"],
            [
                (
                    case["case_id"],
                    case["platform_label"],
                    str(case["evidence_quote"])[:86]
                    + ("…" if len(str(case["evidence_quote"])) > 86 else ""),
                    case["expression_verdict"],
                    case["fact_verdict"],
                )
                for case in source_cases
            ],
            widths=(14, 35, 61, 32, 30),
            font_size=6.9,
        )
    else:
        doc.callout(
            "当前没有可签发的信源正文案例",
            "这只表示本批已完成正文核查的网页中未形成目标品牌段落风险线索；"
            "由于正文核查只覆盖了部分被引用网页，不能据此声称全部引用网页都没有风险。",
            kind="warning",
        )

    doc.heading("5. 逐案可视证据")
    if cases or source_cases:
        doc.paragraph(
            "每个案例均按同一顺序展示：谁对谁、问题与采集环境、AI 原句、原回答截图、"
            "公开核查网址、网页截图/标注、证据为何充分或不足、客户可用结论。"
        )
    else:
        doc.paragraph(
            "本批没有进入逐案展示的案例。后续批次出现线索时，将按统一版式逐案展示："
            "谁对谁、问题与采集环境、AI 原句、原回答截图、公开核查网址、网页截图/标注、"
            "证据为何充分或不足、客户可用结论。"
        )
    for case in cases:
        _case_page(
            doc,
            case,
            answer_screenshot=answer_screenshots.get(str(case["answer_pub_id"])),
            source_captures=source_captures,
            # Let Word flow the cases naturally.  An explicit break immediately
            # after a full evidence page can produce a blank page in LibreOffice.
            page_break=False,
        )
    for case in source_cases:
        descriptor = case.get("source_screenshot") or {}
        _source_case_page(
            doc,
            case,
            source_case_screenshots.get(str(descriptor.get("pub_id") or "")),
        )

    doc.page_break()
    doc.heading("附录 A · 限制说明与版本记录")
    doc.bullets(delivery["limitations"])
    doc.heading("A.1 版本与审批", level=2)
    governance = facts.get("document_governance") or {}
    doc.table(
        ["治理字段", "记录"],
        [
            (
                "项目与服务",
                f"{facts.get('project_name') or '—'} · 服务2 · 品牌GEO内容生态风险核查",
            ),
            (
                "版本与状态",
                f"{governance.get('version') or 'V1.0'} · "
                f"{release_state_label(str(facts.get('document_status') or ''))}",
            ),
            (
                "编制",
                f"{governance.get('prepared_by') or 'GEO 项目组'} · "
                f"{governance.get('prepared_date') or str(facts['generated_at'])[:10]}",
            ),
            (
                "复核",
                f"{governance.get('reviewed_by') or '待复核'} · "
                f"{governance.get('reviewed_date') or '待定'}",
            ),
            (
                "批准",
                f"{governance.get('approved_by') or '待批准'} · "
                f"{governance.get('approved_date') or '待定'}",
            ),
            ("保密级别", "客户机密—仅限指定项目组"),
        ],
        widths=(36, 136),
        font_size=7.8,
    )
    if is_formal_document(facts):
        doc.callout(
            "报告状态",
            "本报告基于已冻结的正式评估窗口事实签发。来源核查与截图标注覆盖仍以"
            "正文披露的实际数量为准，不将未覆盖页面解释为无风险。",
        )
    else:
        doc.callout(
            "本稿状态",
            "本版为内部审核稿，用于完成内容、数据、证据和版式复核。"
            "在网页正文核查覆盖与正式重复采样补齐前，本报告不作为全网风险完备性结论；"
            "数据补齐与人工复核、批准完成后，方可成为客户交付候选稿。",
            kind="warning",
        )
    payload = bytes(doc.save())
    visible_values = " ".join(
        [paragraph.text for paragraph in doc.document.paragraphs]
        + [cell.text for table in doc.document.tables for row in table.rows for cell in row.cells]
    )
    found = [value for value in _CLIENT_FORBIDDEN if value in visible_values]
    if found:
        raise ValueError("customer_report_internal_language:" + ",".join(found))
    return payload


__all__ = ["render_service2_v2_docx"]
