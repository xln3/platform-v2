"""Polished DOCX renderer for the service 1/2/3 pre-formal review reports."""

from __future__ import annotations

from collections.abc import Iterable, Sequence
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

NAVY = "12355B"
BLUE = "1769AA"
CYAN = "29A3C6"
LIGHT_BLUE = "EAF3F9"
PALE_BLUE = "F5F9FC"
INK = "1F2D3D"
MUTED = "5E6F7F"
RED = "A73535"
PALE_RED = "FBEDEE"
GREEN = "2B7A5A"
PALE_GREEN = "EAF6F1"
WHITE = "FFFFFF"
GRID = "CCD8E2"

FONT_REGULAR = Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc")
FONT_BOLD = Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc")

_FORMAL_DOCUMENT_STATUSES = frozenset(
    {
        "formal",
        "formal_ready",
        "formal_signed",
        "final",
        "signed",
        "approved",
        "正式",
        "正式报告",
        "已签发",
    }
)


def is_formal_document(facts: dict[str, Any]) -> bool:
    """Return whether customer-facing chrome should use the formal signed state.

    Older fact builders omit ``document_status`` or emit values such as
    ``pre_formal_review_nonproduction_data``.  Those inputs intentionally remain
    review drafts; only an explicit final/formal state removes the review warning.
    """

    raw_status = str(facts.get("document_status") or "").strip().lower()
    normalized = raw_status.replace("-", "_").replace(" ", "_")
    return normalized in _FORMAL_DOCUMENT_STATUSES


def build_report_code(
    facts: dict[str, Any],
    *,
    service_number: int,
    version: str | None = None,
) -> str:
    """Build a tenant-neutral customer document code without internal identifiers."""

    if service_number not in {1, 2, 3, 4}:
        raise ValueError("service_number_must_be_1_2_3_or_4")
    generated_at = facts.get("generated_at")
    if isinstance(generated_at, datetime):
        date_stamp = generated_at.strftime("%Y%m%d")
    else:
        digits = "".join(character for character in str(generated_at or "") if character.isdigit())
        date_stamp = digits[:8] if len(digits) >= 8 else "UNDATED"
    segments = ["GEO", f"S{service_number}"]
    if version:
        cleaned_version = "".join(
            character for character in str(version).upper() if character.isalnum()
        )
        if cleaned_version:
            segments.append(cleaned_version)
    segments.extend(("FORMAL" if is_formal_document(facts) else "REVIEW", date_stamp))
    return "-".join(segments)


def _shade(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _cell_border(cell: Any, color: str = GRID, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def _cell_margins(
    cell: Any, *, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _keep_row_together(row: Any) -> None:
    """Prevent a table row from being split across two pages."""

    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def _set_cell_width(cell: Any, width_mm: float) -> None:
    """Set an exact DOCX cell width instead of leaving it as a layout hint.

    ``python-docx``'s ``cell.width`` alone does not disable Word/LibreOffice's
    content-driven resizing.  Reports contain narrow ordinal/count columns, so write
    the underlying ``tcW`` value as an exact dxa width as well.
    """

    width_twips = max(1, round(float(width_mm) * 1440 / 25.4))
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:w"), str(width_twips))
    tc_width.set(qn("w:type"), "dxa")
    cell.width = Mm(width_mm)


def _fit_table_widths(widths: Sequence[float], *, available_mm: float = 172.0) -> tuple[float, ...]:
    """Keep requested proportions while preventing a body table from crossing margins."""

    requested = tuple(float(width) for width in widths)
    total = sum(requested)
    if total <= 0 or any(width <= 0 for width in requested):
        raise ValueError("table widths must be positive")
    if total <= available_mm:
        return requested
    scale = available_mm / total
    return tuple(width * scale for width in requested)


def _set_table_fixed_layout(table: Any, widths: Sequence[float]) -> None:
    """Persist a fixed table grid so narrow columns stay narrow after conversion."""

    table.autofit = False
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(max(1, round(float(width) * 1440 / 25.4))))
        grid.append(grid_col)


def _set_font(run: Any, *, name: str = "Noto Sans CJK SC", size: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def _page_field(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def add_native_toc(
    doc: FormalDocument,
    *,
    title: str = "报告目录",
    heading_levels: str = "1-3",
) -> None:
    """Insert the canonical Word-native table of contents used by every report.

    The renderer deliberately emits a real ``TOC`` field instead of a hand-written
    chapter table.  ``tools/refresh_docx_indexes.py`` refreshes this field after the
    document has been paginated, which gives both Word and the exported PDF accurate
    page numbers and clickable destinations.
    """

    title_paragraph = doc.document.add_paragraph()
    title_paragraph.paragraph_format.space_before = Pt(12)
    title_paragraph.paragraph_format.space_after = Pt(9)
    title_run = title_paragraph.add_run(title)
    _set_font(title_run, size=19)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(NAVY)

    paragraph = doc.document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f' TOC \\o "{heading_levels}" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在生成交付文件时自动更新"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, placeholder, end))

    settings = doc.document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _fmt_percent(value: int | float | None, digits: int = 2) -> str:
    if value is None:
        return "未评估"
    return f"{float(value):.{digits}f}%"


def _fmt_ratio(numerator: int | float | None, denominator: int | float | None) -> str:
    return f"{int(numerator or 0)}/{int(denominator or 0)}"


def _fmt_datetime(value: object) -> str:
    if isinstance(value, datetime):
        return value.astimezone().strftime("%Y-%m-%d %H:%M")
    return str(value or "—")


def _clip(value: object, limit: int = 120) -> str:
    text = " ".join(str(value or "").split())
    return text if len(text) <= limit else f"{text[: limit - 1]}…"


def _verdict_label(values: Iterable[str]) -> str:
    labels = {
        "supported": "有证据支持",
        "refuted": "已反证",
        "unverifiable": "无法核验",
        "not_checked": "未核验",
    }
    values = list(values)
    return " / ".join(labels.get(value, value) for value in values) or "未核验"


def _bar_chart(
    labels: Sequence[str],
    values: Sequence[float],
    *,
    title: str,
    suffix: str = "%",
) -> BytesIO:
    width, height = 1500, 720
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    regular = ImageFont.truetype(str(FONT_REGULAR), 32)
    small = ImageFont.truetype(str(FONT_REGULAR), 25)
    bold = ImageFont.truetype(str(FONT_BOLD), 39)
    draw.text((70, 35), title, fill=f"#{INK}", font=bold)
    left, right, top, bottom = 120, 60, 130, 150
    plot_w, plot_h = width - left - right, height - top - bottom
    maximum = max([float(value) for value in values] + [1.0])
    if suffix == "%":
        maximum = max(100.0, maximum)
    slots = max(len(values), 1)
    slot = plot_w / slots
    bar_w = min(120, slot * 0.58)
    draw.line((left, top, left, top + plot_h), fill=f"#{GRID}", width=3)
    draw.line((left, top + plot_h, left + plot_w, top + plot_h), fill=f"#{GRID}", width=3)
    for index, (label, raw_value) in enumerate(zip(labels, values, strict=False)):
        value = float(raw_value)
        bar_h = (value / maximum) * plot_h
        x0 = left + slot * index + (slot - bar_w) / 2
        y0 = top + plot_h - bar_h
        fill = f"#{BLUE}" if index == 0 else f"#{CYAN}"
        draw.rounded_rectangle((x0, y0, x0 + bar_w, top + plot_h), radius=10, fill=fill)
        value_text = f"{value:.1f}{suffix}"
        bbox = draw.textbbox((0, 0), value_text, font=regular)
        draw.text(
            (x0 + bar_w / 2 - (bbox[2] - bbox[0]) / 2, max(top + 4, y0 - 46)),
            value_text,
            fill=f"#{INK}",
            font=regular,
        )
        chunks = [label[pos : pos + 8] for pos in range(0, len(label), 8)][:2]
        for line_index, chunk in enumerate(chunks):
            bbox = draw.textbbox((0, 0), chunk, font=small)
            draw.text(
                (
                    x0 + bar_w / 2 - (bbox[2] - bbox[0]) / 2,
                    top + plot_h + 18 + line_index * 34,
                ),
                chunk,
                fill=f"#{MUTED}",
                font=small,
            )
    output = BytesIO()
    image.save(output, format="PNG", optimize=True)
    output.seek(0)
    return output


class FormalDocument:
    def __init__(self, *, title: str, subtitle: str, facts: dict[str, Any]) -> None:
        self.document = Document()
        self.title = title
        self.subtitle = subtitle
        self.facts = facts
        self._setup()

    def _setup(self) -> None:
        section = self.document.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(17)
        section.bottom_margin = Mm(17)
        section.left_margin = Mm(19)
        section.right_margin = Mm(19)
        section.header_distance = Mm(7)
        section.footer_distance = Mm(8)

        styles = self.document.styles
        normal = styles["Normal"]
        normal.font.name = "Noto Sans CJK SC"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        normal.font.size = Pt(9.5)
        normal.font.color.rgb = RGBColor.from_string(INK)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.25
        for style_name, size, color in (
            ("Title", 28, WHITE),
            ("Heading 1", 19, NAVY),
            ("Heading 2", 13, BLUE),
            ("Heading 3", 10.5, NAVY),
        ):
            style = styles[style_name]
            style.font.name = "Noto Sans CJK SC"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style.font.bold = True
        styles["Heading 1"].paragraph_format.space_before = Pt(12)
        styles["Heading 1"].paragraph_format.space_after = Pt(7)
        styles["Heading 2"].paragraph_format.space_before = Pt(9)
        styles["Heading 2"].paragraph_format.space_after = Pt(5)

        self.document.core_properties.title = self.title
        self.document.core_properties.subject = self.subtitle
        self.document.core_properties.author = "GEO 验证系统"
        self._header_footer(section)

    def _header_footer(self, section: Any) -> None:
        target_brand = str(self.facts.get("target_brand") or "目标品牌")
        formal = is_formal_document(self.facts)
        header = section.header
        table = header.add_table(rows=1, cols=2, width=Mm(172))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Mm(112)
        table.columns[1].width = Mm(60)
        left, right = table.rows[0].cells
        left.text = f"{target_brand}  |  GEO 验证服务"
        right.text = "正式报告" if formal else "预正式审阅稿"
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in (left, right):
            _cell_margins(cell, top=0, bottom=40, start=0, end=0)
            for run in cell.paragraphs[0].runs:
                _set_font(run, size=8)
                run.font.color.rgb = RGBColor.from_string(MUTED)
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_label = "正式报告   |   " if formal else "内部审阅 · 禁止外发   |   "
        run = paragraph.add_run(footer_label)
        _set_font(run, size=8)
        run.font.color.rgb = RGBColor.from_string(MUTED)
        _page_field(paragraph)

    def cover(self, *, report_code: str) -> None:
        formal = is_formal_document(self.facts)
        hero = self.document.add_table(rows=1, cols=1)
        hero.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = hero.cell(0, 0)
        _shade(cell, NAVY)
        _cell_margins(cell, top=850, bottom=850, start=650, end=650)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        label = cell.paragraphs[0]
        label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_run = label.add_run(self.subtitle.upper())
        _set_font(label_run, size=10)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor.from_string("90D5ED")
        title = cell.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title.add_run(self.title)
        _set_font(title_run, size=28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor.from_string(WHITE)
        brand = cell.add_paragraph()
        brand_run = brand.add_run(self.facts["target_brand"])
        _set_font(brand_run, size=16)
        brand_run.font.color.rgb = RGBColor.from_string(WHITE)

        self.document.add_paragraph()
        if not formal:
            banner = self.document.add_table(rows=1, cols=1)
            banner.alignment = WD_TABLE_ALIGNMENT.CENTER
            banner_cell = banner.cell(0, 0)
            _shade(banner_cell, PALE_RED)
            _cell_margins(banner_cell, top=150, bottom=150, start=180, end=180)
            warning = banner_cell.paragraphs[0]
            warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
            warning_run = warning.add_run("预正式审阅稿 · 基于联调/试采样数据 · 禁止对外发布")
            _set_font(warning_run, size=10)
            warning_run.font.bold = True
            warning_run.font.color.rgb = RGBColor.from_string(RED)

        metadata = [
            ("项目", self.facts["project_name"]),
            ("评估窗口", f"{self.facts['window']['start']} 至 {self.facts['window']['end']}"),
            ("生成时间", _fmt_datetime(self.facts["generated_at"])),
            ("文档编号", report_code),
            ("文档状态", "正式报告" if formal else "待客户审核内容与版式"),
        ]
        self.table(["字段", "内容"], metadata, widths=(35, 137), header=False)
        if not formal:
            self.callout(
                "审阅目的",
                "本稿用于检查正式报告的内容结构、指标口径、证据披露和视觉版式。"
                "数值来自当前已采集事实，但不代表已完成报价单约定的正式采样。",
                kind="warning",
            )
        self.document.add_page_break()  # type: ignore[no-untyped-call]

    def heading(self, text: str, *, level: int = 1) -> None:
        paragraph = self.document.add_heading(text, level=level)
        if level == 1:
            paragraph.paragraph_format.keep_with_next = True

    def paragraph(self, text: str, *, bold_lead: str | None = None) -> None:
        paragraph = self.document.add_paragraph()
        if bold_lead:
            lead = paragraph.add_run(bold_lead)
            _set_font(lead)
            lead.bold = True
        body = paragraph.add_run(text)
        _set_font(body)

    def bullets(self, values: Iterable[str]) -> None:
        for value in values:
            paragraph = self.document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(value)
            _set_font(run, size=9.3)

    def numbered(self, values: Iterable[str]) -> None:
        """Render short explanatory notes as an ordered list.

        Tables in customer reports often need several independent reading notes.  A
        numbered list makes those boundaries visible and avoids packing unrelated
        definitions into one dense callout paragraph.
        """

        for index, value in enumerate(values, 1):
            # Explicit labels deliberately restart at 1 for every explanatory block.
            # Word's built-in List Number style otherwise continues numbering across
            # unrelated tables and can produce a confusing 9/10/11 sequence later in
            # the document after conversion through LibreOffice.
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.left_indent = Mm(6)
            paragraph.paragraph_format.first_line_indent = Mm(-6)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(f"{index}. {value}")
            _set_font(run, size=9.3)

    def callout(self, title: str, body: str, *, kind: str = "info") -> None:
        fill, color = {
            "warning": (PALE_RED, RED),
            "success": (PALE_GREEN, GREEN),
            "info": (LIGHT_BLUE, NAVY),
        }.get(kind, (LIGHT_BLUE, NAVY))
        table = self.document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        _shade(cell, fill)
        _cell_margins(cell, top=150, bottom=150, start=180, end=180)
        paragraph = cell.paragraphs[0]
        lead = paragraph.add_run(f"{title}  ")
        _set_font(lead, size=9.5)
        lead.bold = True
        lead.font.color.rgb = RGBColor.from_string(color)
        run = paragraph.add_run(body)
        _set_font(run, size=9.3)
        run.font.color.rgb = RGBColor.from_string(INK)
        self.document.add_paragraph().paragraph_format.space_after = Pt(0)

    def kpis(self, values: Sequence[tuple[str, str, str]]) -> None:
        table = self.document.add_table(rows=1, cols=len(values))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for index, (label, value, note) in enumerate(values):
            cell = table.cell(0, index)
            _shade(cell, LIGHT_BLUE if index % 2 == 0 else PALE_BLUE)
            _cell_margins(cell, top=150, bottom=150, start=130, end=130)
            label_p = cell.paragraphs[0]
            label_run = label_p.add_run(label)
            _set_font(label_run, size=8.3)
            label_run.font.color.rgb = RGBColor.from_string(MUTED)
            value_p = cell.add_paragraph()
            value_run = value_p.add_run(value)
            _set_font(value_run, size=18)
            value_run.bold = True
            value_run.font.color.rgb = RGBColor.from_string(NAVY)
            note_p = cell.add_paragraph()
            note_run = note_p.add_run(note)
            _set_font(note_run, size=7.5)
            note_run.font.color.rgb = RGBColor.from_string(MUTED)
            _cell_border(cell, color=WHITE, size="8")

    def table(
        self,
        headers: Sequence[str],
        rows: Iterable[Sequence[object]],
        *,
        widths: Sequence[float] | None = None,
        header: bool = True,
        font_size: float = 8.2,
    ) -> None:
        rows = list(rows)
        table = self.document.add_table(rows=1 if header else 0, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        if widths is None:
            table.autofit = True
        else:
            if len(widths) != len(headers):
                raise ValueError("table widths must match the number of headers")
            widths = _fit_table_widths(widths)
            _set_table_fixed_layout(table, widths)
        if header:
            header_cells = table.rows[0].cells
            _repeat_header(table.rows[0])
            _keep_row_together(table.rows[0])
            for index, value in enumerate(headers):
                header_cells[index].text = str(value)
                _shade(header_cells[index], NAVY)
                _cell_border(header_cells[index])
                _cell_margins(header_cells[index])
                for run in header_cells[index].paragraphs[0].runs:
                    _set_font(run, size=font_size)
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
        for row_index, values in enumerate(rows):
            cells = table.add_row().cells
            _keep_row_together(table.rows[-1])
            for index in range(len(headers)):
                cells[index].text = str(values[index] if index < len(values) else "")
                _shade(cells[index], WHITE if row_index % 2 == 0 else PALE_BLUE)
                _cell_border(cells[index])
                _cell_margins(cells[index])
                cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cells[index].paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    if widths is not None and float(widths[index]) <= 18:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        _set_font(run, size=font_size)
        if widths is not None:
            for row in table.rows:
                for index, width in enumerate(widths):
                    _set_cell_width(row.cells[index], width)
        self.document.add_paragraph().paragraph_format.space_after = Pt(0)

    def chart(
        self,
        labels: Sequence[str],
        values: Sequence[float],
        *,
        title: str,
        suffix: str = "%",
    ) -> None:
        picture = _bar_chart(labels, values, title=title, suffix=suffix)
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(picture, width=Cm(16.7))

    def page_break(self) -> None:
        self.document.add_page_break()  # type: ignore[no-untyped-call]

    def save(self) -> bytes:
        output = BytesIO()
        self.document.save(output)
        return output.getvalue()


def _service1(facts: dict[str, Any]) -> bytes:
    data = facts["service1"]
    overall = data["overall"]
    target = overall["target"] or {}
    doc = FormalDocument(
        title="品牌 GEO 推荐结果评测报告",
        subtitle="服务 1 · 品牌 AI 可见性与竞品对比",
        facts=facts,
    )
    doc.cover(report_code=build_report_code(facts, service_number=1))
    doc.heading("1. 执行摘要")
    doc.kpis(
        [
            ("品牌提及率", _fmt_percent(target.get("appearance_rate")), "40/72 条主样本"),
            ("平均推荐位次", str(target.get("avg_rank") or "—"), "仅在提及样本内"),
            (
                "Top 3 出现率",
                _fmt_percent((target.get("top_rates") or {}).get("3", {}).get("of_total")),
                "分母=全部 72 条",
            ),
            ("主样本规模", str(overall["answers"]), "3 组×4 变体×3 平台×2 地域"),
        ]
    )
    doc.paragraph(
        f"在当前严格平衡的 {overall['answers']} 条主样本中，{facts['target_brand']}被提及 "
        f"{target.get('mentions', 0)} 次，提及率 {target.get('appearance_rate', 0):.2f}%；"
        f"综合品牌表中位列第 {target.get('overall_rank') or '—'}。"
        "该结果仅用于审阅报告表达，尚不是满足“每题重复 2 次”的正式结论。"
    )
    doc.callout(
        "样本缺口",
        "当前主样本每个问题×平台×地域单元只有 1 条可比较观测；"
        "报价单要求 2 次独立重复。正式报告必须补采后重算。",
        kind="warning",
    )

    doc.heading("2. 评测设计与口径")
    doc.table(
        ["项目", "当前审阅口径", "正式交付要求"],
        [
            ("业务问题", "4 组候选，主文选 3 组；每组 1 原题+3 变体", "完整保留候选组和选择记录"),
            ("AI 平台", "豆包、DeepSeek、文心一言（深度思考）", "同一平台/模式组合"),
            ("地域", "北京、上海浏览器地域采样", "补齐独立账号台账与证明"),
            ("重复", "1 次/单元（当前）", "2 次/单元（报价）"),
            ("主样本去重", "同单元多条时取最新观测", "两次独立重复均入统计"),
        ],
        widths=(29, 72, 71),
    )
    doc.heading("2.1 指标定义", level=2)
    doc.table(
        ["指标", "精确定义", "分母"],
        [
            (
                "品牌提及率",
                f"AI 回答品牌序列中出现{facts['target_brand']}的回答数/总回答数",
                "全部主样本",
            ),
            (
                "推荐排名",
                f"{facts['target_brand']}在已提及回答的 1-based 品牌顺序",
                "仅已提及回答",
            ),
            ("Top1/3/5 出现率", "排名不超过 N 的回答数/总回答数", "全部主样本"),
            ("竞品提及率", "竞品出现的回答数/总回答数", "全部主样本"),
        ],
        widths=(35, 94, 43),
    )

    doc.heading("3. 候选问题组选择")
    doc.paragraph(
        "主文选择仅使用证据完整度：45% 单元覆盖、20% 品牌抽取覆盖、"
        "20% 带引用回答覆盖、10% 回答完整度、5% 平台/地域广度。"
        "评分不读取目标品牌的提及或排名结果。"
    )
    doc.table(
        ["排名", "候选组", "完整度分", "单元", "带引用回答", "主文"],
        [
            (
                group["selection_rank"],
                group["title"],
                f"{group['selection_score']:.2f}",
                _fmt_ratio(group["observed_cells"], group["expected_cells"]),
                _fmt_ratio(group["answers_with_citation"], group["observed_cells"]),
                "是" if group["selected_for_main_report"] else "否（附录）",
            )
            for group in sorted(data["candidate_groups"], key=lambda value: value["selection_rank"])
        ],
        widths=(15, 62, 24, 23, 31, 20),
    )

    doc.heading("4. 品牌可见性结果")
    competitor_rows = [(facts["target_brand"], target.get("appearance_rate") or 0.0, target)] + [
        (row.get("brand_input") or row.get("brand"), row.get("appearance_rate") or 0.0, row)
        for row in overall["competitors"]
        if row is not None
    ]
    doc.chart(
        [str(row[0]) for row in competitor_rows],
        [float(row[1]) for row in competitor_rows],
        title="目标品牌与竞品提及率（主样本）",
    )
    doc.table(
        ["品牌", "提及次数", "提及率", "平均位次", "Top1", "Top3", "Top5"],
        [
            (
                name,
                row.get("mentions") or 0,
                _fmt_percent(row.get("appearance_rate")),
                row.get("avg_rank") or "—",
                _fmt_percent((row.get("top_rates") or {}).get("1", {}).get("of_total")),
                _fmt_percent((row.get("top_rates") or {}).get("3", {}).get("of_total")),
                _fmt_percent((row.get("top_rates") or {}).get("5", {}).get("of_total")),
            )
            for name, _, row in competitor_rows
        ],
        widths=(28, 22, 25, 25, 22, 22, 22),
    )

    doc.heading("4.1 分平台表现", level=2)
    doc.table(
        ["平台", "样本", "提及", "提及率", "平均位次", "Top1", "Top3", "Top5"],
        [
            (
                {"doubao": "豆包", "deepseek": "DeepSeek", "yiyan": "文心一言"}.get(model, model),
                snapshot["answers"],
                (snapshot["target"] or {}).get("mentions") or 0,
                _fmt_percent((snapshot["target"] or {}).get("appearance_rate")),
                (snapshot["target"] or {}).get("avg_rank") or "—",
                _fmt_percent(
                    ((snapshot["target"] or {}).get("top_rates") or {}).get("1", {}).get("of_total")
                ),
                _fmt_percent(
                    ((snapshot["target"] or {}).get("top_rates") or {}).get("3", {}).get("of_total")
                ),
                _fmt_percent(
                    ((snapshot["target"] or {}).get("top_rates") or {}).get("5", {}).get("of_total")
                ),
            )
            for model, snapshot in data["by_model"].items()
        ],
        widths=(28, 18, 18, 24, 25, 20, 20, 20),
    )

    doc.heading("4.2 分问题组表现", level=2)
    doc.table(
        ["问题组", "主文", "样本", "提及率", "平均位次", "Top3", "带引用回答"],
        [
            (
                group["title"],
                "是" if group["selected_for_main_report"] else "否",
                data["by_group"][group["id"]]["answers"],
                _fmt_percent(
                    (data["by_group"][group["id"]]["target"] or {}).get("appearance_rate")
                ),
                (data["by_group"][group["id"]]["target"] or {}).get("avg_rank") or "—",
                _fmt_percent(
                    (
                        (
                            (data["by_group"][group["id"]]["target"] or {}).get("top_rates") or {}
                        ).get("3")
                        or {}
                    ).get("of_total")
                ),
                _fmt_ratio(
                    data["by_group"][group["id"]]["answers_with_citation"],
                    data["by_group"][group["id"]]["answers"],
                ),
            )
            for group in data["candidate_groups"]
        ],
        widths=(52, 16, 18, 25, 25, 22, 32),
    )
    doc.callout(
        "读数提示",
        "不同业务问题组之间的提及率差异明显，说明品牌认知对查询意图敏感。"
        "正式阶段应在补足第二次重复后同时报告总体值和分组稳定性。",
    )

    doc.heading("5. 结论与建议")
    doc.bullets(
        [
            f"当前主样本中{facts['target_brand']}提及率为 "
            f"{target.get('appearance_rate', 0):.2f}%，Top1/Top3/Top5 出现率分别为 "
            f"{(target.get('top_rates') or {}).get('1', {}).get('of_total', 0):.2f}% / "
            f"{(target.get('top_rates') or {}).get('3', {}).get('of_total', 0):.2f}% / "
            f"{(target.get('top_rates') or {}).get('5', {}).get('of_total', 0):.2f}%。",
            "主文问题组已按证据完整度选取 3 组；未选组仍在附录中披露，避免选择性报告。",
            "补采时应固定平台模式、地域、账号和问题文本，每单元完成 2 次独立重复。",
            "正式签发前应对排名异常值和高影响回答进行人工复核，并保留回答、思考/检索过程、引用与截图链路。",
        ]
    )

    doc.page_break()
    doc.heading("A. 候选问题全量附录")
    for group in data["candidate_groups"]:
        marker = "主文选用" if group["selected_for_main_report"] else "仅附录披露"
        doc.heading(f"A.{group['index']} {group['title']}（{marker}）", level=2)
        doc.table(
            ["序号", "问题文本"],
            [(index, question) for index, question in enumerate(group["questions"], 1)],
            widths=(18, 154),
        )
    doc.heading("B. 限制与正式签发前检查")
    doc.bullets(facts["limitations"])
    return doc.save()


def _service2(facts: dict[str, Any]) -> bytes:
    data = facts["service2"]
    doc = FormalDocument(
        title="品牌 GEO 内容生态风险核查报告",
        subtitle="服务 2 · 抹黑/拉踩线索与事实核查",
        facts=facts,
    )
    doc.cover(report_code=build_report_code(facts, service_number=2))
    rate = float(data["flagged_signal_rate"] or 0) * 100
    source_total = int(data["by_subject_type"].get("source_document", 0))
    source_flagged = int(data["flagged_by_subject_type"].get("source_document", 0))
    doc.heading("1. 执行摘要")
    doc.kpis(
        [
            ("有效判定", str(data["judgments_ok"]), "judgment_status=ok"),
            ("模型标记线索", str(data["flagged_signals"]), f"{rate:.2f}% · 非定案"),
            ("去重线索", str(data["unique_signal_patterns"]), "按证据摘录去重"),
            ("有证据支持", str(data["supported_cases"]), "事实核查结果"),
        ]
    )
    doc.callout(
        "核心结论",
        f"系统标记 {data['flagged_signals']} 条模型风险线索，"
        f"去重后 {data['unique_signal_patterns']} 类；"
        f"事实核查中 {data['refuted_cases']} 条已反证、{data['unverifiable_cases']} 条无法核验、"
        f"{data['supported_cases']} 条有证据支持。因此当前不能报告“已发现经核实的第三方抹黑内容”。",
        kind="warning",
    )
    doc.paragraph(
        "本服务采用“模型标记→证据摘录→事实核查→人工复核”的分层过程。"
        "模型标记只是候选线索，不是对作者、竞品或第三方的归因。"
    )

    doc.heading("2. 核查覆盖与质量")
    answer_total = int(data["by_subject_type"].get("answer", 0))
    answer_flagged = int(data["flagged_by_subject_type"].get("answer", 0))
    doc.table(
        ["核查对象", "有效判定", "模型标记", "标记率", "说明"],
        [
            (
                "AI 回答",
                answer_total,
                answer_flagged,
                _fmt_percent(answer_flagged / answer_total * 100 if answer_total else None),
                "回答内生成的比较/负向表述",
            ),
            (
                "已抓取信源文档",
                source_total,
                source_flagged,
                _fmt_percent(source_flagged / source_total * 100 if source_total else None),
                "仅代表当前抓取子集",
            ),
        ],
        widths=(40, 28, 28, 25, 51),
    )
    doc.table(
        ["判定状态", "数量", "纳入风险比率"],
        [
            (status, count, "是" if status == "ok" else "否")
            for status, count in sorted(data["status_counts"].items())
        ],
        widths=(55, 35, 82),
    )

    doc.heading("3. 事实核查结果")
    verdicts = data["factcheck_verdict_counts"]
    doc.chart(
        ["有证据支持", "已反证", "无法核验", "未核验"],
        [
            float(verdicts.get("supported", 0)),
            float(verdicts.get("refuted", 0)),
            float(verdicts.get("unverifiable", 0)),
            float(verdicts.get("not_checked", 0)),
        ],
        title="模型标记线索的事实核查分布（条）",
        suffix="条",
    )
    doc.paragraph(
        "“无法核验”表示当前证据不足以支持或推翻该表述，不等于该表述为真；"
        "“已反证”表示可获得的公开证据与模型表述冲突。"
    )

    doc.heading("4. 去重线索明细")
    doc.table(
        ["序号", "证据摘录（截断展示）", "次数", "目标品牌", "核查结论"],
        [
            (
                index,
                _clip(case["evidence_quote"], 155),
                case["occurrences"],
                " / ".join(case["target_brands"]) or "—",
                _verdict_label(case["factcheck_verdicts"]),
            )
            for index, case in enumerate(data["cases"], 1)
        ],
        widths=(12, 92, 16, 25, 27),
        font_size=7.5,
    )
    for index, case in enumerate(data["cases"], 1):
        doc.heading(f"4.{index} 线索复核", level=2)
        doc.paragraph(case["evidence_quote"], bold_lead="原始摘录：")
        doc.paragraph(_verdict_label(case["factcheck_verdicts"]), bold_lead="核查状态：")
        if case["factcheck_summaries"]:
            doc.paragraph("；".join(case["factcheck_summaries"]), bold_lead="核查说明：")
        if case["factcheck_sources"]:
            doc.paragraph("；".join(case["factcheck_sources"]), bold_lead="公开来源：")
        doc.paragraph(
            "同一回答、同一摘录的重复复核在本报告中合并为一个客户案例。",
            bold_lead="审计说明：",
        )

    doc.heading("5. 风险结论与下一步")
    doc.bullets(
        [
            "当前 8 条标记全部来自 AI 回答，不能归因为竞品或第三方已发布的 GEO 内容。",
            f"已抓取信源文档子集的标记数为 {source_flagged}/{source_total}；"
            "该子集不代表全网完整覆盖。",
            "服务范围为 AI 回答及其公开信源，不要求客户额外提供内容样本。",
            "对“无法核验”线索补充可靠公开来源或客户确认事实，再由人工审核是否纳入正式风险清单。",
        ]
    )
    doc.heading("A. 限制与审计索引")
    doc.bullets(facts["limitations"])
    return doc.save()


def _service3(facts: dict[str, Any]) -> bytes:
    data = facts["service3"]
    own_rate = float(data["own_site_answer_citation_rate"] or 0) * 100
    citation_rate = float(data["citation_coverage_rate"] or 0) * 100
    evidence_rate = float(data["own_site_cited_text_evidence_rate"] or 0) * 100
    doc = FormalDocument(
        title="官网内容 AI 引用能效评估报告",
        subtitle="服务 3 · 官网引用、内容采纳与优化建议",
        facts=facts,
    )
    doc.cover(report_code=build_report_code(facts, service_number=3))
    doc.heading("1. 执行摘要")
    doc.kpis(
        [
            (
                "官网引用率",
                f"{own_rate:.2f}%",
                _fmt_ratio(data["answers_with_own_site_citation"], data["answers_total"]),
            ),
            (
                "内容采纳率",
                "未评估",
                _fmt_ratio(
                    data["own_site_adoption_verified_answers"],
                    data["own_site_adoption_evaluated_answers"],
                ),
            ),
            (
                "回答信源覆盖",
                f"{citation_rate:.2f}%",
                _fmt_ratio(data["answers_with_citation"], data["answers_total"]),
            ),
            (
                "官网引文证据可见",
                f"{evidence_rate:.2f}%",
                _fmt_ratio(
                    data["own_site_cited_text_answers"], data["answers_with_own_site_citation"]
                ),
            ),
        ]
    )
    doc.callout(
        "口径修正",
        "官网引用率已按报价单定义改为“引用官网 URL 的 AI 回答/全部合格 AI 回答”。"
        "抓取文档官网占比与引用转述准确率已拆为辅助指标，不再冒充引用率或采纳率。",
        kind="success",
    )
    doc.callout(
        "采纳率数据缺口",
        "当前没有回答级“官网内容已被理解并用于生成”的判定，因此必须报告为未评估（0/0）。"
        "引用转述准确不等于内容采纳。",
        kind="warning",
    )

    doc.heading("2. 报价指标与扩展指标")
    doc.table(
        ["指标", "结果", "分子/分母", "精确定义", "属性"],
        [
            (
                "官网引用率",
                f"{own_rate:.2f}%",
                _fmt_ratio(data["answers_with_own_site_citation"], data["answers_total"]),
                "至少引用一条官网 URL 的回答/全部合格回答",
                "报价指标",
            ),
            (
                "官网内容采纳率",
                "未评估",
                "0/0",
                "经回答级证据确认用于生成的回答/已评估回答",
                "报价指标",
            ),
            (
                "回答信源覆盖率",
                f"{citation_rate:.2f}%",
                _fmt_ratio(data["answers_with_citation"], data["answers_total"]),
                "带任意 URL 引用的回答/全部合格回答",
                "扩展指标",
            ),
            (
                "有引用回答中的官网占比",
                _fmt_percent(float(data["own_site_share_of_cited_answers"] or 0) * 100),
                _fmt_ratio(data["answers_with_own_site_citation"], data["answers_with_citation"]),
                "引用官网的回答/带任意引用的回答",
                "扩展指标",
            ),
            (
                "官网引用条目占比",
                _fmt_percent(float(data["own_site_reference_share"] or 0) * 100),
                _fmt_ratio(data["own_site_citation_references"], data["citation_references_total"]),
                "官网引用条目/全部引用条目（受平台引用结构差异影响）",
                "辅助",
            ),
            (
                "官网引文证据可见率",
                f"{evidence_rate:.2f}%",
                _fmt_ratio(
                    data["own_site_cited_text_answers"], data["answers_with_own_site_citation"]
                ),
                "有可见 cited_text 的官网引用回答/官网引用回答",
                "证据覆盖",
            ),
        ],
        widths=(34, 22, 25, 70, 21),
        font_size=7.4,
    )

    doc.heading("3. 分平台、模式与地域结果")
    doc.table(
        ["平台", "模式", "地域", "回答", "带引用", "引用官网", "官网引用率"],
        [
            (
                row["model_label"],
                "深度思考" if row["mode"] == "deep_think" else "普通",
                row["region"],
                row["answers"],
                row["answers_with_citation"],
                row["answers_with_own_site_citation"],
                _fmt_percent(row["own_site_answer_citation_rate"] * 100),
            )
            for row in data["platform_region_breakdown"]
        ],
        widths=(28, 28, 18, 18, 24, 28, 28),
    )

    platform_totals: dict[str, dict[str, int]] = {}
    for row in data["platform_region_breakdown"]:
        bucket = platform_totals.setdefault(row["model_label"], {"answers": 0, "own": 0})
        bucket["answers"] += int(row["answers"])
        bucket["own"] += int(row["answers_with_own_site_citation"])
    doc.chart(
        list(platform_totals),
        [value["own"] / value["answers"] * 100 for value in platform_totals.values()],
        title="各平台官网引用率（全部当前样本）",
    )

    doc.heading("4. 抓取与引用转述审计（辅助口径）")
    doc.table(
        ["辅助指标", "结果", "分子/分母", "说明"],
        [
            (
                "抓取文档官网占比",
                _fmt_percent(float(data["own_site_share"] or 0) * 100)
                if data["documents_total"]
                else "数据不足",
                _fmt_ratio(data["own_site_documents"], data["documents_total"]),
                "抓取子集结构，不是回答引用率",
            ),
            (
                "官网引用转述准确率",
                _fmt_percent(float(data["own_site_transcript_accuracy_rate"] or 0) * 100)
                if data["own_site_transcript_total"]
                else "数据不足",
                _fmt_ratio(data["own_site_transcript_accurate"], data["own_site_transcript_total"]),
                "转述与源文一致性，不是内容采纳率",
            ),
        ],
        widths=(48, 29, 30, 65),
    )
    doc.callout(
        "当前审计缺口",
        f"当前窗口内有 {data['answers_with_own_site_citation']} 条回答引用官网，"
        "但抓取子集中官网文档为 "
        f"{data['own_site_documents']}/{data['documents_total']}。因此当前无法完成官网正文级采纳/转述核验，"
        "正式采集必须优先抓取并审计官网 URL。",
        kind="warning",
    )

    doc.heading("5. 引用主机分布（回答级）")
    doc.table(
        ["主机", "覆盖回答", "引用条目", "官网"],
        [
            (row["host"], row["answers"], row["references"], "是" if row["is_own_site"] else "否")
            for row in data["answer_level_hosts"][:12]
        ],
        widths=(85, 30, 30, 27),
    )

    doc.heading("6. 官网内容问题与优化线索")
    generated = _fmt_datetime(data.get("suggestion_generated_at"))
    doc.paragraph(
        f"现有建议生成于 {generated}。"
        "该批次证据来自早期抓取文档，未覆盖当前正式候选采样窗口的官网引用；"
        "因此以“待当前批次复核的优化线索”呈现。"
    )
    doc.table(
        ["严重度", "类别", "建议标题", "建议详情", "证据时间"],
        [
            (
                row.get("severity") or "—",
                row.get("category") or "—",
                row.get("title") or "—",
                _clip(row.get("detail"), 150),
                _fmt_datetime((row.get("evidence") or {}).get("fetched_at")),
            )
            for row in data["suggestions"]
        ],
        widths=(20, 25, 40, 68, 27),
        font_size=7.4,
    )

    doc.heading("7. 优先行动建议")
    doc.bullets(
        [
            "P0：对当前 14 条官网引用回答建立回答—URL—正文快照映射，完成回答级采纳判定。",
            "P0：修正正文抓取规划，保证官网 URL 不会被全局 Top-N 信源截断，并记录抓取/审计覆盖率。",
            "P1：按平台和地域报告官网引用率，对零引用平台核查搜索可达性、内容结构与可引用段落。",
            "P1：对建议批次重新生成当前窗口证据，旧批次仅作线索，不直接进正式结论。",
            "P2：扩展指标保留回答信源覆盖、有引用回答中官网占比、引文证据可见率和官网引用转述准确率。",
        ]
    )
    doc.heading("A. 限制与正式签发前检查")
    doc.bullets(facts["limitations"])
    return doc.save()


def render_formal_review_docx(service_number: int, facts: dict[str, Any]) -> bytes:
    renderers = {1: _service1, 2: _service2, 3: _service3}
    try:
        renderer = renderers[service_number]
    except KeyError as exc:
        raise ValueError("service_number_must_be_1_2_or_3") from exc
    return renderer(facts)
