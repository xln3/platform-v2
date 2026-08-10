# ruff: noqa: E501
from __future__ import annotations

import os
from collections.abc import Mapping, Sequence
from functools import partial
from html import escape
from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.shared import Pt
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.shapes import Drawing
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

# 组件类型全集见 migrations/versions/s02_0001_intelligence_plane.py 的 CHECK 约束
# （kpi/chart/section/evidence/recommendation）；未知类型一律降级为「标题+正文」，不崩。
_COMPONENT_TYPES = {"kpi", "chart", "section", "evidence", "recommendation"}

# chart payload 惯例：series=[{date, value}]；点数上限防止自由 JSONB 撑爆版面。
_CHART_MAX_POINTS = 64

_NO_CHART_DATA_NOTE = "图表数据不可用（series 为空或格式异常）"


class RenderingError(ValueError):
    """报告产物渲染失败（fail-loud，绝不静默输出不可读内容）。"""


def render_html(title: str, sections: Sequence[Mapping[str, object]]) -> bytes:
    blocks = [f"<h1>{escape(title)}</h1>"]
    for component in sections:
        component_type = str(component.get("component_type", "section"))
        component_title = escape(str(component.get("title", "")))
        body = escape(str(component.get("body", "")))
        if component_type == "kpi":
            trace_token = escape(str(component.get("trace_token", "")))
            trace = f"<p class='kpi-trace'>trace: {trace_token}</p>" if trace_token else ""
            blocks.append(
                f"<div class='kpi-card'><h3>{component_title}</h3>"
                f"<p class='kpi-value'>{body}</p>{trace}</div>"
            )
        elif component_type == "chart":
            blocks.append(f"<figure class='chart'><figcaption>{component_title}</figcaption>")
            if body:
                blocks.append(f"<p class='chart-desc'>{body}</p>")
            blocks.append(_html_chart_svg(component.get("series")))
            blocks.append("</figure>")
        elif component_type in {"evidence", "recommendation"}:
            label = "证据" if component_type == "evidence" else "建议"
            blocks.append(
                f"<div class='tag-block tag-{component_type}'>"
                f"<strong class='tag-label'>【{label}】{component_title}</strong>"
                f"<p>{body}</p></div>"
            )
        else:
            # section 与未知类型的诚实降级：标题 + 正文段落。
            blocks.append(f"<section><h2>{component_title}</h2><p>{body}</p></section>")
    css = (
        "body{font-family:system-ui,-apple-system,'Segoe UI','PingFang SC','Microsoft YaHei',"
        "sans-serif;margin:2rem auto;max-width:52rem;padding:0 1rem;line-height:1.6;color:#1f2933}"
        "h1{border-bottom:2px solid #2563eb;padding-bottom:.4rem}"
        "section{margin:1.2rem 0}"
        ".kpi-card{display:inline-block;border:1px solid #d9e2ec;border-radius:8px;"
        "padding:.8rem 1.2rem;margin:.4rem .4rem .4rem 0;background:#f8fafc}"
        ".kpi-card h3{margin:0;font-size:.95rem;color:#52606d}"
        ".kpi-value{font-size:1.8rem;font-weight:700;margin:.2rem 0;color:#102a43}"
        ".kpi-trace{font-size:.72rem;color:#9aa5b1;margin:0}"
        ".chart{border:1px solid #d9e2ec;border-radius:8px;padding:.8rem;margin:1rem 0}"
        ".chart figcaption{font-weight:600}"
        ".chart-desc{color:#52606d;font-size:.9rem}"
        ".chart-empty{color:#9aa5b1;font-size:.85rem}"
        ".tag-block{border-left:4px solid #2563eb;background:#f0f4f8;padding:.5rem .9rem;margin:.8rem 0}"
        ".tag-recommendation{border-left-color:#0e9f6e}"
        ".tag-label{color:#102a43}"
    )
    return (
        "<!doctype html><html lang='zh-CN'><head><meta charset='utf-8'>"
        f"<title>{escape(title)}</title><style>{css}</style></head>"
        f"<body>{''.join(blocks)}</body></html>"
    ).encode()


def _html_chart_svg(raw_series: object) -> str:
    points, raw_points = _extract_series(raw_series)
    if raw_points and len(points) != len(raw_points):
        rows = "".join(
            f"<tr><td>{escape(date)}</td><td>{escape(value)}</td></tr>"
            for date, value in raw_points
        )
        return (
            "<table class='chart-fallback'><thead><tr><th>date</th><th>value</th></tr></thead>"
            f"<tbody>{rows}</tbody></table>"
            f"<p class='chart-empty'>{escape(_NO_CHART_DATA_NOTE)}，已按原始数据点列表呈现。</p>"
        )
    if not points:
        return f"<p class='chart-empty'>{escape(_NO_CHART_DATA_NOTE)}。</p>"
    width, height, pad_left, pad_bottom, pad_top = 640, 240, 48, 40, 16
    plot_width = width - pad_left - 16
    plot_height = height - pad_top - pad_bottom
    max_value = max(value for _, value in points) or 1.0
    slot = plot_width / len(points)
    bar_width = max(2.0, slot * 0.6)
    shapes = [
        f"<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 {width} {height}' "
        "role='img' width='100%' height='auto'>",
        f"<line x1='{pad_left}' y1='{pad_top}' x2='{pad_left}' y2='{height - pad_bottom}' "
        "stroke='#9aa5b1' stroke-width='1'/>",
        f"<line x1='{pad_left}' y1='{height - pad_bottom}' x2='{width - 16}' "
        f"y2='{height - pad_bottom}' stroke='#9aa5b1' stroke-width='1'/>",
        f"<text x='4' y='{pad_top + 10}' font-size='10' fill='#52606d'>"
        f"{escape(_format_tick(max_value))}</text>",
    ]
    for index, (date, value) in enumerate(points):
        bar_height = plot_height * (value / max_value) if max_value else 0.0
        x = pad_left + slot * index + (slot - bar_width) / 2
        y = height - pad_bottom - bar_height
        shapes.append(
            f"<rect x='{x:.1f}' y='{y:.1f}' width='{bar_width:.1f}' height='{bar_height:.1f}' "
            "fill='#2563eb'/>"
        )
        label = date if len(points) <= 12 or index % max(1, len(points) // 12) == 0 else ""
        if label:
            shapes.append(
                f"<text x='{x + bar_width / 2:.1f}' y='{height - pad_bottom + 14}' font-size='10' "
                f"fill='#52606d' text-anchor='middle'>{escape(label)}</text>"
            )
        shapes.append(
            f"<text x='{x + bar_width / 2:.1f}' y='{y - 4:.1f}' font-size='10' "
            f"fill='#102a43' text-anchor='middle'>{escape(_format_tick(value))}</text>"
        )
    shapes.append("</svg>")
    return "".join(shapes)


def render_docx(title: str, sections: Sequence[Mapping[str, object]]) -> bytes:
    document = Document()
    document.add_heading(title, level=0)
    for component in sections:
        component_type = str(component.get("component_type", "section"))
        component_title = str(component.get("title", ""))
        body = str(component.get("body", ""))
        if component_type == "kpi":
            heading = document.add_paragraph()
            heading.add_run(component_title).bold = True
            value_paragraph = document.add_paragraph()
            value_run = value_paragraph.add_run(body)
            value_run.font.size = Pt(24)
            value_run.bold = True
            trace_token = str(component.get("trace_token", ""))
            if trace_token:
                trace_paragraph = document.add_paragraph()
                trace_run = trace_paragraph.add_run(f"trace: {trace_token}")
                trace_run.font.size = Pt(8)
        elif component_type == "chart":
            # 诚实降级：python-docx 无法生成原生 OOXML 图表，
            # chart 组件以「标题行 + date/value 数据点表格」落地，不编造图形。
            document.add_paragraph(f"{component_title}（图表，数据点表格）")
            if body:
                document.add_paragraph(body)
            points, raw_points = _extract_series(component.get("series"))
            rows = [(date, _format_tick(value)) for date, value in points] or raw_points
            if rows:
                table = document.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                table.rows[0].cells[0].text = "date"
                table.rows[0].cells[1].text = "value"
                for date, value in rows:
                    cells = table.add_row().cells
                    cells[0].text = date
                    cells[1].text = value
            else:
                document.add_paragraph(f"（{_NO_CHART_DATA_NOTE}）")
        elif component_type in {"evidence", "recommendation"}:
            label = "证据" if component_type == "evidence" else "建议"
            paragraph = document.add_paragraph()
            label_run = paragraph.add_run(f"【{label}】{component_title}")
            label_run.bold = True
            if body:
                document.add_paragraph(body)
        else:
            # section 与未知类型的诚实降级：标题 + 正文段落。
            document.add_heading(component_title, level=1)
            document.add_paragraph(body)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def render_xlsx(rows: Sequence[Mapping[str, object]]) -> bytes:
    columns = sorted({key for row in rows for key in row})
    all_rows: list[list[object]] = [
        list(columns),
        *[[row.get(column, "") for column in columns] for row in rows],
    ]
    sheet_rows: list[str] = []
    for row_number, values in enumerate(all_rows, 1):
        cells = "".join(
            f'<c r="{_column_name(index)}{row_number}" t="inlineStr">'
            f"<is><t>{_xml(str(value))}</t></is></c>"
            for index, value in enumerate(values, 1)
        )
        sheet_rows.append(f'<row r="{row_number}">{cells}</row>')
    sheet = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f"<sheetData>{''.join(sheet_rows)}</sheetData></worksheet>"
    )
    return _zip(
        {
            "[Content_Types].xml": (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/xl/workbook.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
                '<Override PartName="/xl/worksheets/sheet1.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
                "</Types>"
            ),
            "_rels/.rels": (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
                'Target="xl/workbook.xml"/></Relationships>'
            ),
            "xl/workbook.xml": (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
                'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
                '<sheets><sheet name="Metrics" sheetId="1" r:id="rId1"/></sheets></workbook>'
            ),
            "xl/_rels/workbook.xml.rels": (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
                'Target="worksheets/sheet1.xml"/></Relationships>'
            ),
            "xl/worksheets/sheet1.xml": sheet,
        }
    )


# 中文字体候选：env GEO_REPORT_PDF_FONT_PATH 优先，其后按覆盖率/常见度排序；
# 第一个存在且 TTFont 注册成功的获胜。Noto CJK 是 .ttc/CFF 轮廓，TTFont 不支持，勿加。
_PDF_FONT_ENV = "GEO_REPORT_PDF_FONT_PATH"
_PDF_FONT_FALLBACKS = (
    "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
    "/home/xln/.fonts/SimHei.ttf",
    "/home/xln/.fonts/SimSun.ttf",
)
_PDF_CJK_FONT_NAME = "GeoReportCJK"
# 模块级缓存：None=未解析；""=已解析但无可用字体；非空 str=已注册字体名。
# 测试可 monkeypatch 该缓存与候选清单来覆盖无字体分支。
_pdf_font_resolution: str | None = None


def _resolve_pdf_cjk_font() -> str:
    """返回已注册的 CJK 字体名；无可用字体返回 ""（只解析一次，模块级缓存）。"""
    global _pdf_font_resolution
    if _pdf_font_resolution is not None:
        return _pdf_font_resolution
    candidates = [
        *([os.environ[_PDF_FONT_ENV]] if os.environ.get(_PDF_FONT_ENV) else []),
        *_PDF_FONT_FALLBACKS,
    ]
    resolved = ""
    for candidate in candidates:
        if not os.path.exists(candidate):
            continue
        try:
            pdfmetrics.registerFont(TTFont(_PDF_CJK_FONT_NAME, candidate))
        except Exception:
            continue
        resolved = _PDF_CJK_FONT_NAME
        break
    _pdf_font_resolution = resolved
    return resolved


def render_pdf(title: str, sections: Sequence[Mapping[str, object]]) -> bytes:
    font_name = _resolve_pdf_cjk_font()
    if not font_name and _has_non_latin1(title, sections):
        raise RenderingError(
            "report PDF contains non-Latin-1 text but no usable CJK font was found; "
            f"set {_PDF_FONT_ENV} or install one of: {', '.join(_PDF_FONT_FALLBACKS)}"
        )
    base_font = font_name or "Helvetica"
    styles = {
        "title": ParagraphStyle(
            "GeoTitle", fontName=base_font, fontSize=20, leading=26, spaceAfter=10 * mm
        ),
        "heading": ParagraphStyle(
            "GeoHeading",
            fontName=base_font,
            fontSize=14,
            leading=18,
            spaceBefore=6 * mm,
            spaceAfter=2 * mm,
        ),
        "body": ParagraphStyle(
            "GeoBody", fontName=base_font, fontSize=10.5, leading=15, spaceAfter=2 * mm
        ),
        "kpi_value": ParagraphStyle(
            "GeoKpiValue",
            fontName=base_font,
            fontSize=24,
            leading=28,
            textColor=colors.HexColor("#102a43"),
        ),
        "trace": ParagraphStyle(
            "GeoTrace",
            fontName=base_font,
            fontSize=7.5,
            leading=10,
            textColor=colors.HexColor("#9aa5b1"),
            spaceAfter=4 * mm,
        ),
        "note": ParagraphStyle(
            "GeoNote",
            fontName=base_font,
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#9aa5b1"),
        ),
    }
    story: list[object] = [Paragraph(_para_text(title), styles["title"])]
    for component in sections:
        component_type = str(component.get("component_type", "section"))
        component_title = str(component.get("title", ""))
        body = str(component.get("body", ""))
        if component_type == "kpi":
            story.append(Paragraph(_para_text(component_title), styles["heading"]))
            story.append(Paragraph(_para_text(body), styles["kpi_value"]))
            trace_token = str(component.get("trace_token", ""))
            if trace_token:
                story.append(Paragraph(_para_text(f"trace: {trace_token}"), styles["trace"]))
            else:
                story.append(Spacer(1, 4 * mm))
        elif component_type == "chart":
            story.append(Paragraph(_para_text(component_title), styles["heading"]))
            if body:
                story.append(Paragraph(_para_text(body), styles["body"]))
            story.extend(_pdf_chart_flowables(component.get("series"), styles, base_font))
        elif component_type in {"evidence", "recommendation"}:
            label = "证据" if component_type == "evidence" else "建议"
            story.append(Paragraph(_para_text(f"【{label}】{component_title}"), styles["heading"]))
            if body:
                story.append(Paragraph(_para_text(body), styles["body"]))
        else:
            # section 与未知类型的诚实降级：标题 + 正文段落。
            story.append(Paragraph(_para_text(component_title), styles["heading"]))
            story.append(Paragraph(_para_text(body), styles["body"]))
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        title=title,
        author="GEO Platform",
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    # 既有消费方断言 %PDF-1.7 文件头（tests/integration/test_s02_evidence_service.py）。
    document.build(story, canvasmaker=partial(canvas.Canvas, pdfVersion=(1, 7)))
    return output.getvalue()


def _pdf_chart_flowables(
    raw_series: object, styles: Mapping[str, ParagraphStyle], base_font: str
) -> list[object]:
    points, raw_points = _extract_series(raw_series)
    if raw_points and len(points) != len(raw_points):
        # 数值无法全部解析：不编造图形，以数据点表格 + 说明诚实降级。
        table = Table(
            [["date", "value"], *[[date, value] for date, value in raw_points]],
            colWidths=[60 * mm, 40 * mm],
        )
        table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), base_font),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d9e2ec")),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f4f8")),
                ]
            )
        )
        return [
            table,
            Paragraph(
                _para_text(f"{_NO_CHART_DATA_NOTE}，已按原始数据点列表呈现。"), styles["note"]
            ),
            Spacer(1, 2 * mm),
        ]
    if not points:
        return [
            Paragraph(_para_text(f"{_NO_CHART_DATA_NOTE}。"), styles["note"]),
            Spacer(1, 2 * mm),
        ]
    drawing = Drawing(170 * mm, 60 * mm)
    chart = VerticalBarChart()
    chart.x = 12 * mm
    chart.y = 10 * mm
    chart.width = 150 * mm
    chart.height = 42 * mm
    chart.data = [[value for _, value in points]]
    chart.categoryAxis.categoryNames = [date for date, _ in points]
    chart.categoryAxis.labels.fontName = base_font
    chart.categoryAxis.labels.fontSize = 7
    chart.valueAxis.labels.fontName = base_font
    chart.valueAxis.labels.fontSize = 7
    chart.bars[0].fillColor = colors.HexColor("#2563eb")
    drawing.add(chart)
    return [drawing, Spacer(1, 4 * mm)]


def _extract_series(raw_series: object) -> tuple[list[tuple[str, float]], list[tuple[str, str]]]:
    """从自由 JSONB series 提取数据点。

    返回 (数值点, 原始点)：数值点全部可解析为 float 时供图表使用；
    原始点保留字符串原值供降级表格使用，两点集不一致时调用方应降级而非截断。
    """
    points: list[tuple[str, float]] = []
    raw_points: list[tuple[str, str]] = []
    if not isinstance(raw_series, Sequence) or isinstance(raw_series, str | bytes):
        return points, raw_points
    for item in list(raw_series)[:_CHART_MAX_POINTS]:
        if not isinstance(item, Mapping):
            continue
        date = str(item.get("date", ""))
        value = item.get("value")
        raw_points.append((date, str(value)))
        try:
            points.append((date, float(value)))  # type: ignore[arg-type]
        except (TypeError, ValueError):
            pass
    return points, raw_points


def _has_non_latin1(title: str, sections: Sequence[Mapping[str, object]]) -> bool:
    def non_latin1(text: str) -> bool:
        return any(ord(char) > 0xFF for char in text)

    if non_latin1(title):
        return True
    for component in sections:
        for key in ("title", "body", "trace_token"):
            if non_latin1(str(component.get(key, ""))):
                return True
    return False


def _para_text(value: str) -> str:
    """Paragraph 是 XML 标记文本：转义后把换行落成 <br/>。"""
    return escape(value).replace("\n", "<br/>")


def _format_tick(value: float) -> str:
    return f"{value:g}"


def _zip(files: Mapping[str, str]) -> bytes:
    output = BytesIO()
    with ZipFile(output, "w", compression=ZIP_DEFLATED) as archive:
        for name, content in files.items():
            archive.writestr(name, content)
    return output.getvalue()


def _xml(value: str) -> str:
    return escape(value, quote=True)


def _column_name(index: int) -> str:
    name = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        name = chr(65 + remainder) + name
    return name
