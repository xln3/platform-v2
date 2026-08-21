"""Client-oriented DOCX renderer for Service 3 website citation/adoption V2."""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageStat

from domain.reporting.formal_review_docx import (
    MUTED,
    FormalDocument,
    _fmt_datetime,
    _set_font,
    add_native_toc,
    build_report_code,
    is_formal_document,
)
from domain.reporting.service1_governance import release_state_label

STATUS_LABELS = {
    "confirmed": "可确认直接内容复用（采纳下界）",
    "weak": "弱证据，待人工复核",
    "no_direct_evidence": "未见直接文本复用",
    "not_evaluated": "未覆盖（缺官网正文快照）",
}
_MODE_LABELS = {"deep_think": "深度思考", "normal": "快速", "web": "联网检索"}
_CLIENT_FORBIDDEN = (
    "分子/分母",
    "live_valid",
    "opened_pages",
    "未扇出",
    "历史证据退级",
    "manifest",
    "工作表",
)


def _ratio(numerator: object, denominator: object) -> str:
    return f"{_to_int(numerator)}/{_to_int(denominator)}"


def _clip(value: object, limit: int) -> str:
    text = " ".join(str(value or "").split())
    return text if len(text) <= limit else f"{text[: limit - 1]}…"


def _clean_excerpt(value: object) -> str:
    """Strip answer markup so excerpts never carry raw markdown into the report."""

    text = str(value or "")
    text = re.sub(r"(?m)^\s{0,3}#{1,6}\s*", "", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    text = re.sub(r"(?m)^\s*[-*+]\s+", "• ", text)
    text = re.sub(r"(?m)^\s*[-*_]{3,}\s*$", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def _pct_ratio(value: object) -> str:
    return "未评估" if value is None else f"{_to_float(value) * 100:.2f}%"


def _to_int(value: object) -> int:
    try:
        return int(str(value or 0))
    except (TypeError, ValueError):
        return 0


def _to_float(value: object) -> float:
    try:
        return float(str(value or 0))
    except (TypeError, ValueError):
        return 0.0


def _caption(doc: FormalDocument, text: str) -> None:
    paragraph = doc.document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    _set_font(run, size=7.8)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(MUTED)


def _hyperlink(paragraph: Any, url: str, text: str = "打开原网页") -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def _image_stream(payload: bytes) -> tuple[BytesIO, int, int]:
    with Image.open(BytesIO(payload)) as source:
        image = source.convert("RGB")
    stream = BytesIO()
    image.save(stream, format="PNG", optimize=True)
    stream.seek(0)
    return stream, image.width, image.height


def _is_visually_blank(payload: bytes) -> bool:
    """Reject legacy screenshot files that contain an effectively empty white page."""

    with Image.open(BytesIO(payload)) as source:
        preview = source.convert("L")
        preview.thumbnail((192, 192))
        stats = ImageStat.Stat(preview)
    return stats.mean[0] >= 250 and stats.stddev[0] <= 5


def _add_image(
    doc: FormalDocument,
    payload: bytes,
    *,
    caption: str,
    max_width_cm: float = 16.2,
    max_height_cm: float = 15.0,
) -> None:
    stream, width, height = _image_stream(payload)
    width_cm = min(max_width_cm, max_height_cm * width / max(height, 1))
    paragraph = doc.document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(stream, width=Cm(width_cm))
    _caption(doc, caption)


def _source_url_table(doc: FormalDocument, sources: list[dict[str, Any]]) -> None:
    """Render bounded URL rows; preserve full destinations as hyperlinks, not raw text."""

    doc.table(
        ["序", "网站域名", "类型", "网页链接"],
        [
            (
                source["ordinal"],
                source["host"] or "未解析",
                "官网" if source["is_own_site"] else "第三方",
                "打开原网页",
            )
            for source in sources
        ],
        widths=(12, 85, 25, 50),
        font_size=7.2,
    )
    table = doc.document.tables[-1]
    for row_index, source in enumerate(sources, 1):
        url = str(source.get("url") or "")
        if not url:
            continue
        paragraph = table.cell(row_index, 3).paragraphs[0]
        paragraph.clear()
        _hyperlink(paragraph, url)


def _toc(doc: FormalDocument, facts: dict[str, Any]) -> None:
    del facts
    add_native_toc(doc)
    doc.callout(
        "关键边界",
        "报告中的‘可确认直接内容复用’是可见证据下的保守下界；"
        "不等同于可以观测 AI 的隐藏语义采纳或内部取舍过程。",
        kind="warning",
    )
    doc.page_break()


def _evidence_case(
    doc: FormalDocument,
    case: dict[str, Any],
    assets: dict[str, bytes],
    official_captures: dict[str, dict[str, Any]],
    *,
    number: int,
) -> None:
    doc.page_break()
    doc.heading(f"6.{number} {case['query']}")
    doc.table(
        ["字段", "实测事实"],
        [
            (
                "AI 平台",
                f"{case['model_label']} · {case['region']} · "
                f"{_MODE_LABELS.get(str(case['mode']), str(case['mode'] or '—'))}",
            ),
            ("采集时间", _fmt_datetime(case["capture_time"])),
            (
                "全部信源",
                f"{case['all_source_count']} 条 URL，其中官网 {len(case['official_sources'])} 条",
            ),
            (
                "官网页面",
                ("见下方“打开原网页”链接" if case["best_official_url"] else "当前未捕获官网 URL"),
            ),
            ("官网页标题", case["best_official_title"] or "未捕获"),
            (
                "证据关系",
                {
                    "direct": "已直接绑定到该回答",
                    "same_url_current_window_reuse": (
                        "复用同网址在当前窗口的网页快照（采集时未直接绑定到该回答）"
                    ),
                    "missing": "缺当前窗口官网快照",
                }.get(case["snapshot_relation"], case["snapshot_relation"]),
            ),
            ("判定", STATUS_LABELS[case["status"]]),
            ("判定依据", case["status_basis"]),
        ],
        widths=(33, 139),
        font_size=7.8,
    )

    doc.heading("AI 回答主文证据", level=2)
    doc.paragraph(_clean_excerpt(case.get("answer_excerpt")) or "当前无可展示的回答摘录。")
    answer_asset = case.get("answer_screenshot") or {}
    answer_payload = assets.get(str(answer_asset.get("pub_id") or ""))
    if answer_payload:
        answer_image_kind = str(case.get("answer_screenshot_kind") or "answer_screenshot")
        _add_image(
            doc,
            answer_payload,
            caption=(
                f"图 6-{number}-A  "
                + (
                    "AI 回答正文干净证据图"
                    if answer_image_kind == "answer_excerpt_screenshot"
                    else "AI 回答运行页截图（较早批次的存证图）"
                )
            ),
            max_height_cm=14.2,
        )
    else:
        doc.callout("回答截图", "当前证据关系中未能载入回答截图。", kind="warning")

    doc.heading("官网正文证据", level=2)
    doc.paragraph(_clean_excerpt(case.get("source_excerpt")) or "当前缺官网正文摘录。")
    if case.get("best_official_url"):
        paragraph = doc.document.add_paragraph()
        lead = paragraph.add_run("官网原网页：")
        _set_font(lead, size=8.5)
        lead.bold = True
        _hyperlink(paragraph, str(case["best_official_url"]))
    if case.get("matched_phrase"):
        doc.callout(
            "可比对片段",
            f"归一化后最长连续重合 {case['match_length']} 字符："
            f"“{case['matched_phrase']}”。此结果已排除回答末尾的参考来源列表。",
            kind="success" if case["status"] == "confirmed" else "info",
        )
    official_asset = case.get("official_screenshot") or {}
    official_payload = assets.get(str(official_asset.get("pub_id") or ""))
    official_url = str(case.get("best_official_url") or "")
    live_capture = official_captures.get(official_url) or {}
    live_payload = live_capture.get("payload")
    live_capture_is_usable = (
        live_capture.get("capture_status") == "captured"
        and live_capture.get("content_status") == "ok"
        and bool(live_capture.get("matched_terms"))
        and isinstance(live_payload, bytes)
    )
    if live_capture_is_usable:
        assert isinstance(live_payload, bytes)
        matched_terms = "、".join(str(value) for value in live_capture.get("matched_terms") or [])
        _add_image(
            doc,
            live_payload,
            caption=(
                f"图 6-{number}-B  官网页面相关部分；红框/黄底为与回答可比对的"
                f"网页原句：{matched_terms}。"
            ),
            max_width_cm=16.2,
            max_height_cm=10.5,
        )
        doc.numbered(
            [
                "本图是重新打开原网页后定位到对应句子的当前可读视口，不是把整页"
                "缩成不可读的存证图。",
                "完整整页截图仍保留在证据存储中，但不在客户正文重复占用版面。",
            ]
        )
    elif official_payload and not _is_visually_blank(official_payload):
        doc.callout(
            "官网可视证据边界",
            "当前历史资产只有整页官网截图，没有与上述正文摘录对应的句子级像素坐标。"
            "整页画框没有信息增量且会把正文缩小到不可读，因此本稿不再嵌入该图；"
            "本次重新打开页面也未取得可逐字定位的有效视口图；以可读正文摘录和原网页"
            "链接呈现，待复采形成段落级锚点后再展示红框截图。",
            kind="warning",
        )
    elif official_payload or case.get("official_screenshot_status") == "blank_or_low_information":
        doc.callout(
            "官网截图质量",
            "该 URL 已有当前窗口截图资产，但画面实测为空白，未把空白图冒充视觉证据；"
            "本条正文摘录来自同窗口、完整性校验通过的结构化页面快照。",
            kind="warning",
        )
    else:
        doc.callout(
            "官网截图",
            "该 URL 在当前窗口未找到可校验的官网截图，本条不进入采纳率分母。",
            kind="warning",
        )

    doc.heading("全部回答信源 URL", level=2)
    _source_url_table(doc, list(case["all_sources"]))
    doc.numbered(
        [
            "表内不再铺开不可换行的原始 URL；网站域名用于识别来源，完整地址保留在"
            "“打开原网页”超链接中。",
            "“官网/第三方”按项目确认官网域名匹配；只表示 URL 归属，不表示内容已被准确采纳。",
        ]
    )
    doc.heading("平台公开思考/检索摘要", level=2)
    if case.get("surface_reasoning"):
        doc.paragraph(_clean_excerpt(case["surface_reasoning"]))
    else:
        doc.callout(
            "证据边界",
            "当前未存证平台公开返回的思考/检索摘要；不推断隐藏推理。",
            kind="warning",
        )


def render_service3_v2_docx(
    facts: dict[str, Any],
    *,
    evidence_assets: dict[str, bytes] | None = None,
    official_captures: dict[str, dict[str, Any]] | None = None,
    service_number: int = 3,
    report_title: str = "官网内容 AI 引用能效评估报告",
    report_subtitle: str = "服务 3 · 回答—URL—官网正文证据链",
) -> bytes:
    """Render a client-readable Service 3 V2 report and evidence appendix."""

    evidence_assets = evidence_assets or {}
    official_captures = official_captures or {}
    metrics = facts["metrics"]
    doc = FormalDocument(
        title=report_title,
        subtitle=report_subtitle,
        facts=facts,
    )
    version = str((facts.get("document_governance") or {}).get("version") or "V1.0")
    doc.cover(report_code=build_report_code(facts, service_number=service_number, version=version))
    _toc(doc, facts)

    doc.heading("1. 执行摘要")
    doc.kpis(
        [
            (
                "官网引用率",
                _pct_ratio(metrics["own_site_answer_citation_rate"]),
                _ratio(metrics["answers_with_own_site_citation"], metrics["answers_total"]),
            ),
            (
                "采纳评价覆盖",
                _pct_ratio(metrics["adoption_evaluation_coverage_rate"]),
                _ratio(
                    metrics["adoption_evaluated_answers"],
                    metrics["answers_with_own_site_citation"],
                ),
            ),
            (
                "可确认直接复用",
                _pct_ratio(metrics["conservative_adoption_rate"]),
                _ratio(
                    metrics["adoption_verified_answers"],
                    metrics["adoption_evaluated_answers"],
                ),
            ),
            (
                "官网引用 URL",
                str(metrics["own_site_citation_references"]),
                f"全部引用 {metrics['citation_references_total']} 条",
            ),
        ]
    )
    doc.callout(
        "当前结论",
        f"当前 {metrics['answers_total']} 条合格 AI 回答中，"
        f"{metrics['answers_with_own_site_citation']} 条至少引用 1 个官网 URL，"
        f"官网引用率为 {_pct_ratio(metrics['own_site_answer_citation_rate'])}。"
        f"其中只有 {metrics['adoption_evaluated_answers']} 条具备当前窗口官网正文快照；"
        f"保守规则下 {metrics['adoption_verified_answers']} 条可确认直接内容复用。",
        kind="success" if metrics["adoption_verified_answers"] else "warning",
    )
    doc.callout(
        "不能过度解读",
        f"尚有 {metrics['not_evaluated_answers']} 条官网引用回答因缺当前窗口正文快照未评价；"
        f"{metrics['weak_evidence_answers']} 条只有弱证据。"
        "因此当前比率只是‘可见的直接内容复用下界’，不是 AI 隐藏语义采纳的完整比率。",
        kind="warning",
    )

    doc.heading("2. 指标、判定方法与证据边界")
    doc.table(
        ["指标", "结果", "计算方式", "定义"],
        [
            (
                "官网引用率",
                _pct_ratio(metrics["own_site_answer_citation_rate"]),
                _ratio(metrics["answers_with_own_site_citation"], metrics["answers_total"]),
                "至少引用 1 条官网 URL 的回答/全部合格回答",
            ),
            (
                "任意信源覆盖率",
                _pct_ratio(metrics["citation_coverage_rate"]),
                _ratio(metrics["answers_with_citation"], metrics["answers_total"]),
                "带任意 URL 引用的回答/全部合格回答",
            ),
            (
                "有引用回答中的官网覆盖",
                _pct_ratio(metrics["own_site_share_of_cited_answers"]),
                _ratio(metrics["answers_with_own_site_citation"], metrics["answers_with_citation"]),
                "引用官网的回答/带任意引用的回答",
            ),
            (
                "采纳评价覆盖率",
                _pct_ratio(metrics["adoption_evaluation_coverage_rate"]),
                _ratio(
                    metrics["adoption_evaluated_answers"],
                    metrics["answers_with_own_site_citation"],
                ),
                "具有当前窗口官网正文快照的官网引用回答/全部官网引用回答",
            ),
            (
                "可确认直接内容复用率",
                _pct_ratio(metrics["conservative_adoption_rate"]),
                _ratio(metrics["adoption_verified_answers"], metrics["adoption_evaluated_answers"]),
                "回答主文与官网正文有≥20字符连续重合的回答/已完成正文评价的回答",
            ),
        ],
        widths=(40, 25, 30, 77),
        font_size=7.2,
    )
    doc.numbered(
        [
            f"分母：{facts['adoption_method']['denominator']}。",
            f"可确认规则：{facts['adoption_method']['confirmed_rule']}。",
            f"弱证据规则：{facts['adoption_method']['weak_rule']}。",
            facts["adoption_method"]["boundary"],
        ]
    )

    doc.heading("3. 分平台、模式与地域结果")
    doc.table(
        ["平台", "模式", "地域", "回答", "带引用", "引用官网", "官网引用率"],
        [
            (
                row["model_label"],
                "深度思考" if row["mode"] == "deep_think" else "普通",
                row["region"],
                row["answers"],
                row["answers_with_citation"],
                row["answers_with_own_site_citation"],
                _pct_ratio(row["own_site_answer_citation_rate"]),
            )
            for row in facts["platform_region_breakdown"]
        ],
        widths=(27, 25, 18, 19, 24, 28, 31),
        font_size=7.5,
    )
    doc.heading("3.1 当前能观察到哪些检索阶段", level=2)
    doc.table(
        ["平台", "回答", "有检索摘要", "候选结果可见", "打开页面可见", "最终引用可见"],
        [
            (
                row["model_label"],
                row["answers"],
                row["trace_available"],
                row["candidate_stage_observed"],
                row["opened_stage_observed"],
                row["final_citation_stage_observed"],
            )
            for row in facts["retrieval_observability_by_platform"]
        ],
        widths=(30, 22, 28, 33, 31, 28),
        font_size=7.5,
    )
    observability_notes = [
        "最终引用阶段可以直接区分“平台返回了引用”与“没有返回可解析引用”，也能"
        "判断最终引用中是否包含官网。",
        "只有保存了搜索候选 URL 的回答，才能判断官网是否曾进入候选；只有保存了"
        "页面打开事件的回答，才能继续判断候选是否被打开。",
    ]
    for row in facts["retrieval_observability_by_platform"]:
        model = str(row.get("model") or "")
        if not model:
            model = {"DeepSeek": "deepseek", "豆包": "doubao", "文心一言": "yiyan"}.get(
                str(row.get("model_label") or ""), ""
            )
        current_window = (
            f"{row['model_label']}：本窗口 {row['candidate_stage_observed']}/"
            f"{row['answers']} 条保存候选阶段、{row['opened_stage_observed']}/"
            f"{row['answers']} 条保存页面打开阶段、{row['final_citation_stage_observed']}/"
            f"{row['answers']} 条保存最终引用阶段。{row['boundary']}。"
        )
        if model == "deepseek":
            if int(row["opened_stage_observed"]) < int(row["answers"]):
                current_window += (
                    "本批部分回答未保存页面打开事件，这些回答不能判断当时是否打开官网。"
                )
        elif model == "doubao":
            if int(row["opened_stage_observed"]) < int(row["answers"]):
                current_window += "该平台当前没有可观察的页面打开事件，不能区分候选后是否实际打开。"
        elif model == "yiyan":
            if int(row["candidate_stage_observed"]) < int(row["answers"]):
                current_window += (
                    "该平台当前主要能观察到最终引用，不能完整区分“未进候选”和“候选后未选用”。"
                )
        observability_notes.append(current_window)
    observability_notes.append(
        "因此后续补采可以验证 DeepSeek 的候选→打开→最终引用链，以及豆包的候选→"
        "最终引用链；平台未暴露的阶段无法凭空补齐。"
    )
    doc.numbered(observability_notes)

    probe = facts.get("latest_cross_platform_probe")
    if isinstance(probe, dict) and probe.get("rows"):
        doc.heading("3.2 同题补充实测：实际能区分到哪一层", level=2)
        doc.callout(
            "已完成实测",
            f"{probe['scope']}。采集时间为 {_fmt_datetime(probe['capture_start'])} 至 "
            f"{_fmt_datetime(probe['capture_end'])}；三个回答均成功采集且质检合格。",
            kind="success",
        )

        def stage_text(row: dict[str, Any], stage: str) -> str:
            if not row[f"{stage}_stage_observed"]:
                return "平台未暴露/未保存"
            count = int(row[f"{stage}_urls"])
            official = bool(row[f"official_{stage}_observed"])
            return f"可见 {count} 个 URL；{'含官网' if official else '未见官网'}"

        def probe_conclusion(row: dict[str, Any]) -> str:
            if row["opened_stage_observed"]:
                return "可区分候选→打开→最终引用"
            if row["candidate_stage_observed"]:
                return "可区分候选→最终引用；不能确认打开阶段"
            return "只能确认最终引用；候选和打开阶段未知"

        doc.paragraph(str(probe["query"]), bold_lead="同一试点问题：")
        doc.table(
            ["平台", "候选结果", "页面打开", "最终引用", "本次可下结论"],
            [
                (
                    row["model_label"],
                    stage_text(row, "candidate"),
                    stage_text(row, "opened"),
                    f"{row['final_citations']} 条；"
                    f"{'含官网' if row['official_final_citation'] else '未见官网'}",
                    probe_conclusion(row),
                )
                for row in probe["rows"]
            ],
            widths=(24, 38, 38, 31, 41),
            font_size=7.1,
        )
        probe_by_model = {str(row["model"]): row for row in probe["rows"]}
        deepseek_probe = probe_by_model.get("deepseek", {})
        doubao_probe = probe_by_model.get("doubao", {})
        yiyan_probe = probe_by_model.get("yiyan", {})
        doc.numbered(
            [
                f"DeepSeek 本次实测保存了 {deepseek_probe.get('candidate_urls', 0)} 个"
                f"候选 URL、{deepseek_probe.get('opened_urls', 0)} 个打开页和 "
                f"{deepseek_probe.get('final_citations', 0)} 个最终引用；官网在三个阶段"
                "均出现，因此三层可以逐层核对。",
                f"豆包本次保存了 {doubao_probe.get('candidate_urls', 0)} 个候选 URL 和 "
                f"{doubao_probe.get('final_citations', 0)} 个最终引用，但没有稳定的页面"
                "打开事件；不能把“候选后未引用”进一步解释成“打开后未采用”。",
                f"文心本次保存了 {yiyan_probe.get('final_citations', 0)} 个最终引用；"
                "即使提示词要求实际打开网页，也不能据此声称系统观察到了完整候选或"
                "打开动作。",
            ]
        )

    doc.heading("3.3 零官网引用单元：逐项可证实结论", level=2)
    zero_groups = list(facts["zero_citation_groups"])
    if zero_groups:
        doc.table(
            ["平台/模式/地域", "回答", "带任意引用", "官网引用"],
            [
                (
                    f"{row['model_label']} / "
                    f"{'深度思考' if row['mode'] == 'deep_think' else '普通'} / "
                    f"{row['region']}",
                    row["answers"],
                    row["answers_with_citation"],
                    row["answers_with_own_site_citation"],
                )
                for row in zero_groups
            ],
            widths=(82, 28, 32, 30),
            font_size=7.5,
        )
        for row in zero_groups:
            doc.heading(
                f"{row['model_label']} / "
                f"{'深度思考' if row['mode'] == 'deep_think' else '普通'} / "
                f"{row['region']}",
                level=3,
            )
            doc.numbered(row["diagnosis_items"])
    else:
        doc.numbered(
            [
                "当前聚合窗口内，每个平台/模式/地域单元都至少出现过 1 条官网引用；"
                "这不表示每条回答都引用官网，也不表示官网内容均被准确采用。"
            ]
        )

    doc.heading("4. AI 回答引用的网站来源")
    doc.paragraph(
        "‘网站来源’是 AI 回答中 URL 的域名。覆盖回答数是至少引用该站一次的回答数；"
        "URL 引用条目是全部引用次数。该表不是抓取文档数，也不表示站点权威性。"
    )
    doc.table(
        ["排名", "网站域名", "覆盖回答", "URL 引用条目", "是否官网"],
        [
            (
                index,
                row["host"],
                row["answers"],
                row["references"],
                "官网" if row["is_own_site"] else "第三方",
            )
            for index, row in enumerate(facts["answer_source_domains"], 1)
        ],
        widths=(14, 76, 28, 31, 23),
        font_size=7.6,
    )

    doc.heading("5. 官网引用回答证据总表")
    doc.table(
        ["序", "问题", "平台/地域", "全部URL", "官网URL", "快照", "判定"],
        [
            (
                index,
                _clip(row["query"], 42),
                f"{row['model_label']}/{row['region']}",
                row["all_source_count"],
                len(row["official_sources"]),
                (
                    "有"
                    if any(item["has_current_text_snapshot"] for item in row["official_sources"])
                    else "缺"
                ),
                STATUS_LABELS[row["status"]],
            )
            for index, row in enumerate(facts["evaluations"], 1)
        ],
        widths=(10, 55, 29, 17, 18, 14, 29),
        font_size=6.6,
    )
    doc.callout(
        "快照覆盖差距",
        f"当前只有 {metrics['direct_snapshot_bound_answers']} 条回答直接绑定官网快照；"
        f"按同 URL 复用当前窗口快照后，可覆盖 {metrics['same_url_snapshot_covered_answers']} 条。"
        f"其中 {metrics.get('usable_screenshot_covered_answers', 0)} 条有非空白官网截图。"
        "这限制了回答级证据链的完整呈现，相关差距在附录限制中如实披露。",
        kind="warning",
    )

    doc.heading("6. 回答—URL—官网正文证据卡")
    doc.paragraph(
        "以下优先展示不同判定类型中证据最完整的真实回答。"
        "每张卡都披露全部信源 URL，不只展示官网 URL。"
    )
    for number, case in enumerate(facts["selected_evidence_cases"], 1):
        _evidence_case(
            doc,
            case,
            evidence_assets,
            official_captures,
            number=number,
        )

    doc.page_break()
    doc.heading("7. 官网内容问题与优化动作")
    doc.paragraph(
        "下表只根据当前回答—URL—官网正文证据链产生，不沿用旧建议批次，"
        "也不把系统开发缺口写成客户的官网责任。"
    )
    doc.table(
        ["优先级", "当前事实", "建议动作", "责任方"],
        [
            (row["priority"], row["fact"], row["action"], row["owner"])
            for row in facts["client_actions"]
        ],
        widths=(20, 55, 75, 22),
        font_size=7.5,
    )

    doc.heading("附录 A · 全部官网 URL 与证据状态")
    appendix_rows = []
    appendix_urls: list[str] = []
    for answer_index, row in enumerate(facts["evaluations"], 1):
        for source in row["official_sources"]:
            appendix_rows.append(
                (
                    answer_index,
                    source["ordinal"],
                    "打开官网页面",
                    "有" if source["has_cited_text"] else "无",
                    (
                        "正文+截图"
                        if source["has_current_text_snapshot"] and source["has_current_screenshot"]
                        else "正文有/截图空白"
                        if source["has_current_text_snapshot"]
                        and source.get("screenshot_status") == "blank_or_low_information"
                        else "仅正文"
                        if source["has_current_text_snapshot"]
                        else "缺失"
                    ),
                    "直接" if source["direct_answer_relation"] else "未直接绑定",
                    STATUS_LABELS[row["status"]],
                )
            )
            appendix_urls.append(str(source["url"] or ""))
    doc.table(
        ["回答", "URL序", "官网页面", "引文", "快照", "关系", "回答判定"],
        appendix_rows,
        widths=(13, 13, 45, 14, 24, 26, 37),
        font_size=6.1,
    )
    appendix_table = doc.document.tables[-1]
    for row_index, url in enumerate(appendix_urls, 1):
        if not url:
            continue
        paragraph = appendix_table.cell(row_index, 2).paragraphs[0]
        paragraph.clear()
        _hyperlink(paragraph, url, "打开官网页面")
    doc.heading("附录 B · 限制说明与版本记录")
    limitations = list(facts["limitations"])
    if is_formal_document(facts):
        limitations = [
            item for item in limitations if "联调/试采样数据，不是正式运行签发结论" not in str(item)
        ]
        limitations.insert(
            0,
            "本报告基于已冻结的正式评估窗口事实生成；结论仅适用于披露的窗口与证据范围。",
        )
    doc.bullets(limitations)
    doc.heading("B.1 版本与审批", level=2)
    governance = facts.get("document_governance") or {}
    doc.table(
        ["治理字段", "记录"],
        [
            (
                "项目与服务",
                f"{facts.get('project_name') or '—'} · 服务{service_number} · {report_title}",
            ),
            (
                "版本与状态",
                f"{governance.get('version') or 'V1.0'} · "
                f"{release_state_label(str(facts.get('document_status') or ''))}",
            ),
            (
                "编制",
                f"{governance.get('prepared_by') or 'GEO 项目组'} · "
                f"{governance.get('prepared_date') or str(facts['generated_at'])[:10]}",
            ),
            (
                "复核",
                f"{governance.get('reviewed_by') or '待复核'} · "
                f"{governance.get('reviewed_date') or '待定'}",
            ),
            (
                "批准",
                f"{governance.get('approved_by') or '待批准'} · "
                f"{governance.get('approved_date') or '待定'}",
            ),
            ("保密级别", "客户机密—仅限指定项目组"),
        ],
        widths=(36, 136),
        font_size=7.8,
    )
    payload = bytes(doc.save())
    visible_values = " ".join(
        [paragraph.text for paragraph in doc.document.paragraphs]
        + [cell.text for table in doc.document.tables for row in table.rows for cell in row.cells]
    )
    found = [value for value in _CLIENT_FORBIDDEN if value in visible_values]
    if found:
        raise ValueError("customer_report_internal_language:" + ",".join(found))
    return payload
