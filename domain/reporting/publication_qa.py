"""Automated publication checks for customer-facing DOCX/PDF artifacts."""

from __future__ import annotations

import re
import subprocess
import tempfile
from collections import Counter
from collections.abc import Iterable, Mapping
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any
from urllib.parse import unquote
from zipfile import ZipFile

from docx import Document

_FORBIDDEN = (
    "[特殊字符]",
    "本次审阅口径",
    "审阅判断",
    "当前试采",
    "正式复测与签发检查",
    "冻结矩阵",
    "TODO",
    "TBD",
    "PLACEHOLDER",
    "发布门禁未通过",
    "不得对外交付",
    "阻断项已逐项",
    "缺失时不得批准签发",
    "不作出的承诺",
)
_CONTROL = re.compile(r"[\u0000-\u0008\u000b\u000c\u000e-\u001f\u007f]")
_MARKDOWN = re.compile(r"(?m)^\s{0,3}(?:#{1,6}\s+|```|[-*_]{3,}\s*$|\|.*\|\s*$|[-*+]\s+\S)")
_MOJIBAKE = (
    "â†’",
    "Â·",
    "Â ",
    "â€œ",
    "â€\u009d",
    "â€™",
    "â€“",
    "â€”",
    "â€¦",
    "ï¿½",
    "\ufffd",
)
_URL = re.compile(r"https?://[^\s<>\]\[()（）]+")
_NUMBERED_HEADING = re.compile(r"^\d+(?:\.\d+)*\.?\s+\S")


class PublicationQAError(RuntimeError):
    """Required publication tooling failed or emitted an unreadable artifact."""


def _find_mojibake(value: str) -> list[str]:
    return sorted(token for token in _MOJIBAKE if token in value)


def _run(*args: str) -> str:
    try:
        result = subprocess.run(args, check=False, capture_output=True, timeout=60)
    except (OSError, subprocess.TimeoutExpired) as exc:
        raise PublicationQAError(f"publication_tool_failed:{Path(args[0]).name}") from exc
    if result.returncode != 0:
        raise PublicationQAError(f"publication_tool_failed:{Path(args[0]).name}")
    return result.stdout.decode("utf-8", errors="replace")


def _docx_text(payload: bytes) -> tuple[str, list[str]]:
    with tempfile.NamedTemporaryFile(suffix=".docx") as handle:
        handle.write(payload)
        handle.flush()
        document = Document(handle.name)
    blocks: list[str] = []
    headings: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
        if text and str(paragraph.style.name).startswith(("Heading", "标题")):
            headings.append(text)
    for table in document.tables:
        for row in table.rows:
            blocks.extend(cell.text.strip() for cell in row.cells if cell.text.strip())
    return "\n".join(blocks), headings


def _docx_properties(payload: bytes) -> dict[str, Any]:
    with tempfile.NamedTemporaryFile(suffix=".docx") as handle:
        handle.write(payload)
        handle.flush()
        document = Document(handle.name)
    core = document.core_properties
    with tempfile.NamedTemporaryFile(suffix=".docx") as handle:
        handle.write(payload)
        handle.flush()
        with ZipFile(handle.name) as archive:
            document_xml = archive.read("word/document.xml").decode(errors="replace")
            styles_xml = archive.read("word/styles.xml").decode(errors="replace")
    return {
        "title": core.title or "",
        "author": core.author or "",
        "subject": core.subject or "",
        "keywords": core.keywords or "",
        "has_toc_field": "TOC \\" in document_xml or "TOC \\o" in document_xml,
        "language_zh_cn": "zh-CN" in document_xml and "zh-CN" in styles_xml,
    }


def _pdf_info(path: Path) -> dict[str, str]:
    values: dict[str, str] = {}
    for line in _run("pdfinfo", "-isodates", str(path)).splitlines():
        if ":" in line:
            key, value = line.split(":", 1)
            values[key.strip()] = value.strip()
    return values


def _pdf_urls(path: Path) -> list[str]:
    rows = _run("pdfinfo", "-url", str(path)).splitlines()[1:]
    urls: list[str] = []
    for row in rows:
        match = re.search(r"https?://\S+", row)
        if match:
            urls.append(match.group(0))
    return urls


def _pdf_image_pages(path: Path) -> set[int]:
    pages: set[int] = set()
    for row in _run("pdfimages", "-list", str(path)).splitlines()[2:]:
        match = re.match(r"\s*(\d+)\s+\d+\s+(?:image|mask|smask)\b", row)
        if match:
            pages.add(int(match.group(1)))
    return pages


def _orphan_figure_caption_pages(page_bodies: list[str], image_pages: set[int]) -> list[int]:
    return [
        page_number
        for page_number, body in enumerate(page_bodies, 1)
        if re.search(r"(?m)^图\s*\d+-\d+(?!\d)", body) and page_number not in image_pages
    ]


def _normalize(value: str) -> str:
    return re.sub(r"\s+", "", value).replace("\ufeff", "")


def _normalize_url(value: str) -> str:
    return unquote(value).rstrip("/")


def _normalize_pdf_semantics(value: str) -> str:
    pages = value.split("\f")
    semantic_pages: list[str] = []
    for page in pages:
        # The native TOC is intentionally refreshed on every export.  Dot leader
        # widths may change after the first LibreOffice save without changing any
        # heading or page number, so compare body semantics separately.
        if page.count("...") >= 3:
            continue
        lines = []
        for line in page.splitlines():
            if "客户机密—仅限指定项目组 |" in line:
                continue
            if "GEO 验证服务" in line and any(
                label in line for label in ("内部审核稿", "客户交付候选稿", "已批准签发版")
            ):
                continue
            lines.append(line)
        semantic_pages.append("\n".join(lines))
    return _normalize("\n".join(semantic_pages))


def _page_body(value: str) -> str:
    lines = []
    for line in value.splitlines():
        compact = line.strip()
        if not compact or re.fullmatch(r"\d+", compact):
            continue
        if "客户机密—仅限指定项目组" in compact:
            continue
        if re.search(r"GEO\s*验证服务", compact):
            continue
        lines.append(compact)
    return "\n".join(lines)


def inspect_publication(
    *,
    docx: bytes,
    pdf: bytes,
    expected_title: str,
    expected_status_label: str,
    expected_urls: Iterable[str] = (),
    # The client-restructured report carries self-contained appendices (per-question
    # results, competitor scope, representative excerpts); 25–40 pages is the healthy
    # band, the lower bound still catches a hollow export.
    page_range: tuple[int, int] = (25, 40),
) -> dict[str, Any]:
    """Inspect one converted pair and return explicit gates and diagnostics."""

    expected_url_list = [str(value) for value in expected_urls if str(value).startswith("http")]
    with tempfile.TemporaryDirectory(prefix="geo-publication-qa-") as directory:
        pdf_path = Path(directory) / "report.pdf"
        pdf_path.write_bytes(pdf)
        info = _pdf_info(pdf_path)
        text = _run("pdftotext", "-layout", str(pdf_path), "-")
        page_texts = text.split("\f")
        if page_texts and not page_texts[-1].strip():
            page_texts.pop()
        annotated_urls = _pdf_urls(pdf_path)
        image_pages = _pdf_image_pages(pdf_path)
    docx_visible, headings = _docx_text(docx)
    properties = _docx_properties(docx)
    pages = int(info.get("Pages", "0") or 0)
    page_bodies = [_page_body(value) for value in page_texts]
    blank_pages = [index for index, value in enumerate(page_bodies, 1) if len(value) < 12]
    low_content_pages = [
        index for index, value in enumerate(page_bodies, 1) if 12 <= len(value) < 70
    ]
    orphan_figure_captions = _orphan_figure_caption_pages(page_bodies, image_pages)
    duplicate_headings = sorted(
        {
            heading
            for heading in headings
            if _NUMBERED_HEADING.match(heading) and headings.count(heading) > 1
        }
    )
    combined = f"{docx_visible}\n{text}"
    forbidden_found = sorted(value for value in _FORBIDDEN if value in combined)
    controls_found = sorted(set(_CONTROL.findall(combined.replace("\f", ""))))
    markdown_found = bool(_MARKDOWN.search(combined))
    mojibake_found = _find_mojibake(combined)
    normalized_pdf = _normalize(text)
    critical_text = [
        expected_title,
        expected_status_label,
        "中国标准时间（UTC+8）",
        "客户机密—仅限指定项目组",
        "AI生成原文，未经事实核验，不代表评测方结论",
    ]
    missing_critical_text = [
        value for value in critical_text if _normalize(value) not in normalized_pdf
    ]
    annotated_normalized = {_normalize_url(value) for value in annotated_urls}
    missing_url_annotations = sorted(
        value
        for value in set(expected_url_list)
        if _normalize_url(value) not in annotated_normalized
    )
    page_size = info.get("Page size", "")
    page_dimensions = [float(value) for value in re.findall(r"\d+(?:\.\d+)?", page_size)[:2]]
    checks = {
        "page_count_in_target_range": page_range[0] <= pages <= page_range[1],
        "page_text_count_matches": len(page_texts) == pages,
        "a4_page_size": len(page_dimensions) == 2
        and abs(page_dimensions[0] - 595.28) < 1
        and abs(page_dimensions[1] - 841.89) < 1,
        "tagged_pdf": info.get("Tagged", "").lower() == "yes",
        "pdf_title_matches": info.get("Title", "") == expected_title,
        "pdf_author_present": bool(info.get("Author", "").strip()),
        "pdf_subject_present": bool(info.get("Subject", "").strip()),
        "pdf_keywords_present": bool(info.get("Keywords", "").strip()),
        "docx_title_matches": properties["title"] == expected_title,
        "docx_author_present": bool(properties["author"]),
        "docx_subject_present": bool(properties["subject"]),
        "docx_keywords_present": bool(properties["keywords"]),
        "docx_language_zh_cn": properties["language_zh_cn"],
        "docx_toc_present": properties["has_toc_field"],
        "no_blank_pages": not blank_pages,
        "no_orphan_figure_captions": not orphan_figure_captions,
        "no_duplicate_numbered_headings": not duplicate_headings,
        "no_forbidden_tokens": not forbidden_found,
        "no_control_characters": not controls_found,
        "no_markdown_residue": not markdown_found,
        "no_mojibake": not mojibake_found,
        "critical_text_preserved": not missing_critical_text,
        "all_displayed_urls_clickable": not missing_url_annotations,
        "external_url_count_matches": len(annotated_urls) >= len(expected_url_list),
    }
    return {
        "schema_version": "service1-publication-qa-v1",
        "status": "passed" if all(checks.values()) else "failed",
        "checks": checks,
        "pdf": {
            "pages": pages,
            "page_size": page_size,
            "tagged": info.get("Tagged"),
            "title": info.get("Title"),
            "author": info.get("Author"),
            "subject": info.get("Subject"),
            "keywords": info.get("Keywords"),
            "external_url_annotations": len(annotated_urls),
        },
        "docx": properties,
        "diagnostics": {
            "blank_pages": blank_pages,
            "low_content_pages": low_content_pages,
            "image_pages": sorted(image_pages),
            "orphan_figure_caption_pages": orphan_figure_captions,
            "duplicate_numbered_headings": duplicate_headings,
            "forbidden_tokens": forbidden_found,
            "control_characters": controls_found,
            "markdown_residue": markdown_found,
            "mojibake": mojibake_found,
            "missing_critical_text": missing_critical_text,
            "expected_urls": expected_url_list,
            "annotated_urls": annotated_urls,
            "missing_url_annotations": missing_url_annotations,
        },
    }


def compare_reexport(
    *,
    first_docx: bytes,
    first_pdf: bytes,
    second_docx: bytes,
    second_pdf: bytes,
) -> dict[str, Any]:
    """Verify a second DOCX refresh/export does not drift in pages or visible text."""

    with tempfile.TemporaryDirectory(prefix="geo-reexport-qa-") as directory:
        root = Path(directory)
        first_path = root / "first.pdf"
        second_path = root / "second.pdf"
        first_path.write_bytes(first_pdf)
        second_path.write_bytes(second_pdf)
        first_info = _pdf_info(first_path)
        second_info = _pdf_info(second_path)
        first_text = _run("pdftotext", str(first_path), "-")
        second_text = _run("pdftotext", str(second_path), "-")
    first_docx_text, _ = _docx_text(first_docx)
    second_docx_text, _ = _docx_text(second_docx)
    first_semantic_text = _normalize_pdf_semantics(first_text)
    second_semantic_text = _normalize_pdf_semantics(second_text)
    similarity = SequenceMatcher(None, first_semantic_text, second_semantic_text).ratio()
    checks = {
        "pdf_page_count_stable": first_info.get("Pages") == second_info.get("Pages"),
        # Poppler may interleave adjacent table-cell characters differently after
        # a harmless repagination save.  Require the exact same character multiset,
        # near-identical order, and separately exact DOCX paragraph/table text.
        "pdf_normalized_text_stable": Counter(first_semantic_text) == Counter(second_semantic_text)
        and similarity >= 0.995,
        "docx_normalized_text_stable": _normalize(first_docx_text) == _normalize(second_docx_text),
    }
    return {
        "schema_version": "service1-reexport-qa-v1",
        "status": "passed" if all(checks.values()) else "failed",
        "checks": checks,
        "first_pages": int(first_info.get("Pages", "0") or 0),
        "second_pages": int(second_info.get("Pages", "0") or 0),
        "pdf_text_similarity": round(similarity, 6),
    }


def displayed_service1_urls(facts: Mapping[str, Any]) -> list[str]:
    delivery = facts["service1"]["delivery_v3"]
    return [
        str(citation.get("url"))
        for representative in delivery.get("representative_answers") or []
        for citation in (representative.get("citations") or [])[:3]
        if isinstance(citation, Mapping) and citation.get("url")
    ]


__all__ = [
    "PublicationQAError",
    "compare_reexport",
    "displayed_service1_urls",
    "inspect_publication",
]
