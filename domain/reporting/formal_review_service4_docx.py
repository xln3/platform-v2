"""DOCX renderer for the Service-4 GEO pilot and effect review."""

from __future__ import annotations

from typing import Any

from docx.shared import RGBColor

from domain.reporting.formal_review_docx import (
    MUTED,
    FormalDocument,
    _fmt_datetime,
    _set_font,
    add_native_toc,
    build_report_code,
)

MODEL_LABELS = {
    "doubao": "豆包",
    "deepseek": "DeepSeek",
    "yiyan": "文心一言",
    "tongyi": "通义千问",
    "yuanbao": "腾讯元宝",
}


def _percent(value: Any) -> str:
    if value is None:
        return "证据不足"
    return f"{float(value):.2f}%"


def _metric_value(row: dict[str, Any], key: str) -> str:
    value = row.get(key)
    if value is None:
        return "证据不足"
    if row.get("unit") == "rank":
        return f"{float(value):.2f} 位"
    return _percent(value)


def _interval(value: object, *, unit: str) -> str:
    if not isinstance(value, list | tuple) or len(value) != 2:
        return "区间不足"
    suffix = " 位" if unit == "rank" else "%"
    return f"{float(value[0]):.2f}–{float(value[1]):.2f}{suffix}"


def _absolute(row: dict[str, Any]) -> str:
    value = row.get("absolute_change")
    if value is None:
        return "不计算"
    suffix = " 位" if row.get("unit") == "rank" else " 个百分点"
    return f"{float(value):+.2f}{suffix}"


def _relative(value: Any) -> str:
    if value is None:
        return "不适用"
    return f"{float(value):+.2f}%"


def _status_label(value: object) -> str:
    return {
        "matched": "通过",
        "mismatched": "不一致",
        "unverifiable": "无法核验",
        "comparable": "可比",
        "not_comparable": "不可直接比较",
    }.get(str(value), str(value or "—"))


def _dynamic_header(doc: FormalDocument, facts: dict[str, Any]) -> None:
    """Keep the shared layout while replacing its legacy project-specific label."""

    header = doc.document.sections[0].header
    if not header.tables:
        return
    table = header.tables[-1]
    left = table.cell(0, 0)
    left.text = f"{facts['target_brand']}  |  GEO 试点与效果验证"
    for run in left.paragraphs[0].runs:
        _set_font(run, size=8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def _metric_table(doc: FormalDocument, rows: list[dict[str, Any]]) -> None:
    if not rows:
        doc.callout(
            "证据不足",
            "品牌抽取、候选问题组或样本读取未通过证据闸门，本报告不输出伪零值和伪变化。",
            kind="warning",
        )
        return
    doc.table(
        ["指标", "before", "after", "绝对变化", "相对变化", "样本量", "稳定性"],
        [
            (
                row["label"],
                f"{_metric_value(row, 'before')}\n95%区间 "
                f"{_interval(row.get('before_interval_95'), unit=str(row.get('unit')))}",
                f"{_metric_value(row, 'after')}\n95%区间 "
                f"{_interval(row.get('after_interval_95'), unit=str(row.get('unit')))}",
                _absolute(row),
                _relative(row.get("relative_change_percent")),
                (
                    f"before {row.get('before_numerator', '—')}/{row['before_n']}\n"
                    f"after {row.get('after_numerator', '—')}/{row['after_n']}"
                ),
                row["stability"],
            )
            for row in rows
        ],
        widths=(28, 23, 23, 20, 19, 25, 34),
        font_size=6.5,
    )
    doc.numbered(
        [
            "绝对变化按 after−before 计算；比例指标使用百分点，平均推荐位次使用位次差。",
            "相对变化以前测值的绝对值为分母；前测为 0 或任一臂无值时不计算。",
            "平均推荐位次越小通常越靠前；其样本量分子为品牌被提及次数，分母为全部合格回答。",
            "95% 区间表达抽样不确定性，不等于显著性检验，也不单独证明优化因果。",
        ]
    )


def render_service4_review_docx(facts: dict[str, Any]) -> bytes:
    """Render one report with two explicit parts sharing the same frozen facts."""

    required = {
        "target_brand",
        "project_name",
        "generated_at",
        "windows",
        "candidate_groups",
        "comparability",
        "evidence_gate",
        "pilot_plan",
    }
    missing = sorted(required - set(facts))
    if missing:
        raise ValueError(f"service4_facts_missing:{','.join(missing)}")
    doc = FormalDocument(
        title="GEO 试点与效果验证报告 V2",
        subtitle="服务 4 · 优化试点方案与同矩阵前后对比",
        facts=facts,
    )
    _dynamic_header(doc, facts)
    doc.cover(report_code=build_report_code(facts, service_number=4, version="V2"))
    add_native_toc(doc)

    before_arm = facts.get("arms", {}).get("before", {})
    after_arm = facts.get("arms", {}).get("after", {})
    checks = list(facts["comparability"].get("checks") or [])
    passed_checks = sum(bool(row.get("passed")) for row in checks)
    selected_groups = [
        row for row in facts["candidate_groups"] if row.get("selected_for_main_report")
    ]

    doc.heading("1. 执行摘要")
    doc.kpis(
        [
            (
                "证据闸门",
                "可描述" if facts["metrics"] else "证据不足",
                str(facts["evidence_gate"]["status"]),
            ),
            (
                "可比性检查",
                f"{passed_checks}/{len(checks)}",
                _status_label(facts["comparability"]["status"]),
            ),
            (
                "正文问题组",
                f"{len(selected_groups)}/3",
                "仅按证据完整度选择",
            ),
            (
                "双臂主样本",
                f"{before_arm.get('answers_selected_groups', 0)} / "
                f"{after_arm.get('answers_selected_groups', 0)}",
                "before / after",
            ),
        ]
    )
    doc.callout(
        "签发结论",
        str(facts["evidence_gate"]["conclusion"]),
        kind=(
            "success"
            if facts["comparability"]["status"] == "comparable" and facts["metrics"]
            else "warning"
        ),
    )
    doc.paragraph(
        "本报告把“GEO 优化试点方案”和“优化前后效果对比”放在同一制品中。"
        "两部分共用同一份冻结事实、候选组选择记录、指标版本和证据闸门，"
        "因此方案边界与效果口径不会在两个文件之间漂移。"
    )
    doc.table(
        ["臂", "窗口", "全部候选回答", "正文回答", "抽取成功", "带引用", "有图片"],
        [
            (
                "before",
                f"{facts['windows']['before']['start']} 至 {facts['windows']['before']['end']}",
                before_arm.get("answers_all_candidates", 0),
                before_arm.get("answers_selected_groups", 0),
                before_arm.get("extract_ok_selected", 0),
                before_arm.get("answers_with_citation_selected", 0),
                before_arm.get("answers_with_visual_selected", 0),
            ),
            (
                "after",
                f"{facts['windows']['after']['start']} 至 {facts['windows']['after']['end']}",
                after_arm.get("answers_all_candidates", 0),
                after_arm.get("answers_selected_groups", 0),
                after_arm.get("extract_ok_selected", 0),
                after_arm.get("answers_with_citation_selected", 0),
                after_arm.get("answers_with_visual_selected", 0),
            ),
        ],
        widths=(18, 48, 25, 23, 22, 18, 18),
        font_size=7.1,
    )
    doc.page_break()

    plan = facts["pilot_plan"]
    doc.heading("第一部分 GEO 优化试点方案")
    doc.callout(
        "方案状态",
        "本节是待执行方案，不是已经完成的优化记录。只有发布、版本和页面快照留证完成后，"
        "相应动作才可转为已执行。",
        kind="warning",
    )
    doc.heading("2. 试点目标、对象与依据", level=2)
    doc.paragraph(str(plan["objective"]), bold_lead="试点目标：")
    doc.paragraph(str(plan["scope"]), bold_lead="问题范围：")
    doc.paragraph(str(plan["baseline_evidence"]), bold_lead="基线依据：")
    doc.heading("3. 外部信息源与内容优化动作", level=2)
    doc.table(
        ["序号", "责任方", "优化对象", "证据依据", "执行内容", "执行边界"],
        [
            (
                row["number"],
                row["owner"],
                row["object"],
                row["basis"],
                row["work"],
                row["boundary"],
            )
            for row in plan["actions"]
        ],
        widths=(12, 27, 28, 31, 43, 31),
        font_size=6.7,
    )
    doc.numbered(
        [
            "外部信息源建设只使用公开、可核验、可保存版本的页面，不把付费曝光或虚假背书当作证据。",
            "官网内容以事实可引用性为目标，不能为了指标加入无法证明的绝对化宣传。",
            "复测由评测执行方按冻结矩阵完成，客户内容审核与评测执行责任分开留痕。",
        ]
    )
    doc.heading("4. 验收指标、复测矩阵与停止条件", level=2)
    configured = facts["design"]["before_configured"]
    doc.table(
        ["矩阵维度", "冻结值"],
        [
            (
                "正文业务问题组",
                "、".join(str(row["title"]) for row in selected_groups) or "不足 3 组",
            ),
            ("每组问题文本", "基准问题及语义变体；完整文本见附录 A"),
            (
                "AI 平台",
                "、".join(MODEL_LABELS.get(value, value) for value in configured["models"]),
            ),
            ("模式", "、".join(str(value) for value in configured["modes"])),
            ("地域", "、".join(str(value) for value in configured["regions"])),
            ("逐单元重复", f"{facts['design']['required_repetitions_per_cell']} 次"),
            ("指标版本", str(facts["design"]["metric_version"])),
            ("冻结规则", str(facts["design"]["freeze_rule"])),
        ],
        widths=(42, 130),
        font_size=8,
    )
    doc.paragraph("；".join(plan["acceptance_metrics"]), bold_lead="验收指标：")
    doc.numbered([f"停止条件：{value}" for value in plan["stop_conditions"]])
    doc.page_break()

    doc.heading("第二部分 优化前后效果对比")
    doc.heading("5. 双臂可比性披露", level=2)
    doc.callout(
        "可比性总判定",
        (
            "全部检查通过，可以按同口径报告变化；但没有干预执行台账时，仍不能单独证明因果。"
            if facts["comparability"]["status"] == "comparable"
            else "至少一项检查不一致或无法核验。下列变化最多是描述性时间窗差异，不得归因优化。"
        ),
        kind="success" if facts["comparability"]["status"] == "comparable" else "warning",
    )
    doc.table(
        ["序号", "检查项", "before", "after", "状态", "披露"],
        [
            (
                index,
                row["label"],
                row["before"],
                row["after"],
                _status_label(row["status"]),
                row["disclosure"],
            )
            for index, row in enumerate(checks, 1)
        ],
        widths=(12, 30, 34, 34, 20, 42),
        font_size=6.6,
    )
    doc.numbered(
        [
            "问题、平台、模式、地域、逐单元重复、账号策略、指标/抽取版本和冻结规则均进入闸门。",
            "“无法核验”不等于相同；它与明确不一致一样阻断优化效果归因。",
            "账号与浏览器配置按策略计数核验；表中披露双臂是否一致及缺口。",
        ]
    )

    doc.heading("6. 核心指标前后对比", level=2)
    _metric_table(doc, list(facts.get("metrics") or []))

    doc.heading("6.1 三组业务问题结果", level=3)
    group_results = list(facts.get("group_results") or [])
    if group_results:
        for index, group in enumerate(group_results, 1):
            doc.heading(f"{index}. {group['title']}", level=3)
            _metric_table(doc, list(group.get("metrics") or []))
    else:
        doc.callout("未形成分组结果", "证据闸门未通过，不展示伪造的分组差异。", kind="warning")

    doc.heading("7. 竞品品牌格局", level=2)
    landscape = list(facts.get("competitor_landscape") or [])
    if landscape:
        doc.table(
            ["序号", "品牌", "对象", "before 提及率", "after 提及率", "绝对变化", "平均位次前/后"],
            [
                (
                    index,
                    row["brand"],
                    "目标品牌" if row["is_target"] else "竞品观察",
                    f"{_percent(row['before_appearance_rate'])} ({row['before_mentions']} 次)",
                    f"{_percent(row['after_appearance_rate'])} ({row['after_mentions']} 次)",
                    f"{float(row['absolute_change']):+.2f} 个百分点",
                    f"{row.get('before_avg_rank') or '—'} / {row.get('after_avg_rank') or '—'}",
                )
                for index, row in enumerate(landscape[:15], 1)
            ],
            widths=(12, 32, 24, 31, 31, 23, 19),
            font_size=6.8,
        )
        doc.numbered(
            [
                "除目标品牌外，所有被观察品牌均作为竞品格局对象。",
                "提及率分母为对应臂正文合格回答数；未出现按真实 0 次披露，不代表市场份额为 0。",
                "表内最多展示 15 个高覆盖品牌；完整动态品牌格局见附录 C。",
            ]
        )
    else:
        doc.callout("证据不足", "未通过品牌抽取证据闸门，不输出竞品格局。", kind="warning")

    doc.heading("8. AI 回答引用的网站结构", level=2)
    source_rows = list(facts.get("source_structure") or [])
    if source_rows:
        doc.table(
            ["序号", "引用网站（域名）", "before 条目/占比", "after 条目/占比", "占比绝对变化"],
            [
                (
                    index,
                    row["site"],
                    f"{row['before_count']} / {_percent(row['before_share'])}",
                    f"{row['after_count']} / {_percent(row['after_share'])}",
                    f"{float(row['absolute_change']):+.2f} 个百分点",
                )
                for index, row in enumerate(source_rows[:15], 1)
            ],
            widths=(14, 67, 34, 34, 23),
            font_size=7.1,
        )
        doc.numbered(
            [
                "占比使用对应臂全部最终引用条目为分母，各臂完整网站占比合计为 100%。",
                "网站顺序只描述最终引用结构，不代表权威性、页面已打开或事实正确。",
                "正文展示前 15 个网站；完整动态列表见附录 C。",
            ]
        )
    else:
        doc.callout(
            "本批未形成网站结构",
            "所选回答没有可解析最终引用，不能用已抓取文档或候选 URL 冒充回答引用。",
            kind="warning",
        )

    doc.heading("8.1 官网引用扩展指标", level=3)
    own = facts.get("own_site_extension") or {}
    if own.get("status") == "evaluated":
        before = own["before"]
        after = own["after"]
        doc.table(
            ["指标", "before", "after", "绝对变化", "相对变化", "证据边界"],
            [
                (
                    "回答级官网引用率",
                    f"{_percent(before['answer_citation_rate'])} "
                    f"({before['answers_with_own_site']}/{before['answers']})",
                    f"{_percent(after['answer_citation_rate'])} "
                    f"({after['answers_with_own_site']}/{after['answers']})",
                    f"{float(own['absolute_change']):+.2f} 个百分点",
                    _relative(own.get("relative_change_percent")),
                    own["boundary"],
                )
            ],
            widths=(34, 28, 28, 23, 22, 37),
            font_size=7,
        )
        doc.numbered([str(own["evidence_basis"]), str(own["stability"])])
    else:
        doc.callout("扩展指标不足", str(own.get("boundary") or "证据不足。"), kind="warning")

    doc.heading("9. 稳定性、归因边界与复测结论", level=2)
    doc.numbered([str(value) for value in facts.get("limitations") or []])
    doc.callout(
        "因果边界",
        str(facts["evidence_gate"]["causal_claim_blocker"])
        + "；因此本报告即使通过矩阵可比性，也把结果表述为同口径观察变化，而不是单因果证明。",
        kind="warning",
    )

    doc.page_break()
    doc.heading("附录 A：全部候选问题组与选择评分")
    doc.paragraph(str(facts["design"]["selection_policy"]))
    doc.table(
        ["序号", "候选问题组", "问题数", "before 覆盖", "after 覆盖", "完整度", "用途"],
        [
            (
                row["index"],
                row["title"],
                len(row["questions"]),
                f"{row['before_evidence']['covered_observations']}/"
                f"{row['before_evidence']['expected_observations']}",
                f"{row['after_evidence']['covered_observations']}/"
                f"{row['after_evidence']['expected_observations']}",
                f"{float(row['selection_score']):.2f}",
                "正文" if row["selected_for_main_report"] else "附录",
            )
            for row in facts["candidate_groups"]
        ],
        widths=(13, 51, 18, 27, 27, 19, 17),
        font_size=7.2,
    )
    doc.numbered(
        [
            str(facts["candidate_groups"][0]["selection_basis"])
            if facts["candidate_groups"]
            else "未读取到可评分候选组。",
            "所有未入选组仍保留在本附录和冻结事实快照；不因 after 表现较差而剔除。",
        ]
    )
    for group in facts["candidate_groups"]:
        doc.heading(f"A.{group['index']} {group['title']}", level=2)
        doc.table(
            ["问题序号", "实际问题文本"],
            [(index, question) for index, question in enumerate(group["questions"], 1)],
            widths=(22, 150),
            font_size=8,
        )

    doc.heading("附录 B：逐单元样本量登记")
    registry = list(facts["comparability"].get("cell_registry") or [])
    if registry:
        doc.table(
            ["序号", "问题", "平台", "模式", "地域", "before n", "after n", "一致"],
            [
                (
                    index,
                    row["question"],
                    MODEL_LABELS.get(row["platform"], row["platform"]),
                    row["mode"],
                    row["region"],
                    row["before_n"],
                    row["after_n"],
                    "是" if row["matched"] else "否",
                )
                for index, row in enumerate(registry, 1)
            ],
            widths=(12, 67, 23, 20, 17, 13, 13, 7),
            font_size=6.5,
        )
    else:
        doc.callout("无可登记单元", "候选组或答案证据不足。", kind="warning")

    doc.heading("附录 C：完整品牌格局与网站结构")
    if landscape:
        doc.heading("C.1 完整动态品牌格局", level=2)
        doc.table(
            ["序号", "品牌", "对象", "before 次数/率", "after 次数/率", "绝对变化"],
            [
                (
                    index,
                    row["brand"],
                    "目标品牌" if row["is_target"] else "竞品观察",
                    f"{row['before_mentions']} / {_percent(row['before_appearance_rate'])}",
                    f"{row['after_mentions']} / {_percent(row['after_appearance_rate'])}",
                    f"{float(row['absolute_change']):+.2f} 个百分点",
                )
                for index, row in enumerate(landscape, 1)
            ],
            widths=(13, 40, 25, 35, 35, 24),
            font_size=6.9,
        )
    if source_rows:
        doc.heading("C.2 完整动态引用网站结构", level=2)
        doc.table(
            ["序号", "引用网站（域名）", "before 次数/占比", "after 次数/占比", "占比变化"],
            [
                (
                    index,
                    row["site"],
                    f"{row['before_count']} / {_percent(row['before_share'])}",
                    f"{row['after_count']} / {_percent(row['after_share'])}",
                    f"{float(row['absolute_change']):+.2f} 个百分点",
                )
                for index, row in enumerate(source_rows, 1)
            ],
            widths=(13, 69, 34, 34, 22),
            font_size=7,
        )
    if not landscape and not source_rows:
        doc.callout("证据不足", "没有可列示的品牌格局或最终引用网站结构。", kind="warning")

    doc.heading("附录 D：签发限制")
    doc.numbered(
        [
            f"报告生成时间：{_fmt_datetime(facts['generated_at'])}；生成后新增数据不自动进入本次冻结事实。",
            f"证据状态：{facts['evidence_gate']['status']}；"
            f"可比性状态：{facts['comparability']['status']}。",
            *[str(value) for value in facts.get("limitations") or []],
        ]
    )
    return doc.save()
